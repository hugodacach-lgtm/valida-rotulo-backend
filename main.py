import os, base64, json, io, asyncio, re
try:
    from pypdf import PdfReader
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False
import httpx
from fastapi import FastAPI, UploadFile, File, Form, Request
from fastapi.responses import StreamingResponse, JSONResponse, HTMLResponse
from fastapi.middleware.cors import CORSMiddleware
import html
from datetime import datetime, timezone, timedelta

# ═══════════════════════════════════════════════════════════════════════════════
# SENTRY — Monitoramento de erros em produção
# Env: SENTRY_DSN (obter em sentry.io → Project → DSN)
# Falha silenciosamente se não configurado
# ═══════════════════════════════════════════════════════════════════════════════
try:
    import sentry_sdk
    from sentry_sdk.integrations.fastapi import FastApiIntegration
    from sentry_sdk.integrations.httpx import HttpxIntegration
    _SENTRY_DSN = os.environ.get("SENTRY_DSN", "")
    SENTRY_DSN = _SENTRY_DSN  # alias compat — algumas funções usam sem underscore
    if _SENTRY_DSN:
        sentry_sdk.init(
            dsn=_SENTRY_DSN,
            integrations=[FastApiIntegration(), HttpxIntegration()],
            traces_sample_rate=0.1,   # 10% das requests — não sobrecarrega
            profiles_sample_rate=0.1,
            environment=os.environ.get("ENVIRONMENT", "production"),
            release=os.environ.get("RENDER_GIT_COMMIT", "unknown"),
        )
except ImportError:
    pass  # sentry-sdk não instalado — sem monitoramento, sem crash

# Garantia: SENTRY_DSN sempre definido no escopo global (mesmo se sentry_sdk falhou)
if "SENTRY_DSN" not in dir():
    SENTRY_DSN = ""
if "_SENTRY_DSN" not in dir():
    _SENTRY_DSN = ""

app = FastAPI(title="ValidaRótulo IA v7 — Cobertura Universal: POA + Vegetais + Bebidas + Suplementos")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")

# Modelo Anthropic — controlável via env var ANTHROPIC_MODEL no Render.
# Default: Sonnet 4.6 (recomendação Anthropic para novas integrações).
# Para rollback rápido: setar ANTHROPIC_MODEL=claude-sonnet-4-5 no Render.
ANTHROPIC_MODEL = os.environ.get("ANTHROPIC_MODEL", "claude-sonnet-4-6")

# ═══════════════════════════════════════════════════════════════════════════════
# KNOWLEDGE BASE — URLs dos PDFs oficiais MAPA/ANVISA/INMETRO
# 50 categorias cobrindo toda a cadeia POA
# ═══════════════════════════════════════════════════════════════════════════════
MAPA_URLS = {
    # ── CARNES E EMBUTIDOS ────────────────────────────────────────────────
    "embutidos":          ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN042000salsichamortadelalinguia.pdf"],
    "salame":             ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN222000RTsalamesalaminhocopaprescrupresparmalingcolpepperoni.pdf"],
    "presunto":           ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/Port7652023RTIQpresunto.pdf"],
    "hamburguer":         ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/Port7242022RThamburguer1.pdf"],
    "bacon":              ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/copy_of_Port7482023RTbacon.pdf"],
    "carne_moida":        ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/Port6642022RTIQcarnemoda1.pdf"],
    "charque":            ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN922020RTCharqueCarneSalgadaMidoSalgado.pdf"],
    "fiambre":            ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/Port7062022RTIQfiambre.pdf"],
    "carne_maturada":     ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/Port7232022RTcarnematuradabovino.pdf"],
    "carnes_temperadas":  ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN172018RTcrneostemperados.pdf"],
    "almondega_kibe":     ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN202000RTcrneosalmondegakibe.pdf"],
    "gelatina_colageno":  ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/Port3842021Gelatinagelatinahidrolisadaecolgeno.pdf"],
    "corned_beef":        ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN832003RTcornedbeefbovinoconserva.pdf"],
    "paleta_salgada":     ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN62001RTcarneospaletasalgadosempanadospresserranopratopronto.pdf"],

    # ── LATICÍNIOS ────────────────────────────────────────────────────────
    "laticinios_geral":       ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/Port1461996RTqueijomanteigacremedeleitegorduralctealeitefluido.pdf"],
    "queijo_coalho_manteiga": ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN302001RTmanteigatemaoumanteigadegarrafaqueijodecoalhoequeijodemanteiga.pdf"],
    "mussarela":              ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/Port3661997RTMassaqueijomussarela.pdf"],
    "requeijao":              ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/Port3591997RTrequeijaoouqueijofundido.pdf"],
    "leite_uht":              ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/Port3701997RTleiteUHTUAT.pdf"],
    "leite_pasteurizado":     ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN762018RTleitecrurefrleitepasteuhomogleitefluido.pdf"],
    "leite_fermentado":       ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN462007RTleitesfermentados.pdf"],
    "doce_de_leite":          ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/Port3541997RTdocedeleite.pdf"],
    "soro_de_leite":          ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN942020RTSorodeleite.pdf"],
    "queijo_prato":           ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/Port3581997RTQueijoprato.pdf"],
    "queijo_processado":      ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/Port3561997RTQueijoprocessado.pdf"],

    # ── PESCADO ───────────────────────────────────────────────────────────
    "pescado_fresco":         ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/Port1851997RTpeixefresco.pdf"],
    "pescado_congelado":      ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN212017RTpeixecongelado.pdf"],
    "pescado_salgado":        ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN012019RTpeixesalgadopeixesalgadoseco.pdf"],
    "camarao":                ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN232019RTcamaraofrescoresfcongeladodescongelado.pdf"],
    "lagosta":                ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN242019RTlagostafrescacongelada.pdf"],
    "moluscos_cefalopodes":   ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/Port10222024RTmoluscocefal%C3%B3pode.pdf"],
    "conserva_sardinha":      ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN222011RTconservassardinhas.pdf"],
    "conserva_atum":          ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN462011RTconservasatuns.pdf"],
    "conserva_peixe":         ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN452011RTconservaspeixe.pdf"],

    # ── MEL E APÍCOLA ─────────────────────────────────────────────────────
    "mel":                    ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/Port7952023RTIQmel.pdf"],

    # ── OVOS ──────────────────────────────────────────────────────────────
    "ovos":                   ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/Port012020RTovosdemesadeovosparasementeovoscaipiras.pdf"],
    "ovos_derivados":         ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/Port7282022RTovointegralpasteurizadodesidratado.pdf"],

    # ── LATICÍNIOS EXTRAS ────────────────────────────────────────────────
    "leite_em_po":            ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN532018RTleiteempoMercosul.pdf"],
    "leite_condensado":       ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN472018RTleitecondensado.pdf"],
    "nata":                   ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN232012RTnata.pdf"],
    "queijo_provolone":       ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN732020RTqueijoprovolone.pdf"],
    "queijo_parmesao":        ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/Port3531997RTqueijoparmesao.pdf"],

    # ── NORMAS ANVISA NA KB ──────────────────────────────────────────────
    "alergenos_rdc727":       ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/rdc727_rotulagem%20geral.pdf"],
    "nutricional_in75":       ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/INSTRU%C3%87%C3%83O%20NORMATIVA-IN%20N%C2%BA%2075%2C%20DE%208%20DE%20OUTUBRO%20DE%202020%20-%20INSTRU%C3%87%C3%83O%20NORMATIVA-IN%20N%C2%BA%2075%2C%20DE%208%20DE%20OUTUBRO%20DE%202020%20-%20DOU%20-%20Imprensa%20Nacional.pdf"],
    "aditivos_rdc778":        ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/RDC%20778%20-%20Aditivos%20alimentares%20%28geral%29.pdf"],

    # ── NORMAS GERAIS ─────────────────────────────────────────────────────
    "rotulagem_geral":        [
        "https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN222005Regulamentortuloregistroproduto.pdf",
        "https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/Port2402021AlteraIN222005rotulagem.pdf",
        "https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/Port4492022AlteraIN222005rotulagem.pdf",
    ],
    "nomenclatura":           ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/Port14852025PadrocategoriaenomeclaturaPOA.pdf"],

    # ── CÁRNEOS ADICIONAIS ─────────────────────────────────────────────────
    "apresuntado":            ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/Port7012022RTIQapresuntado.pdf"],
    "paleta_empanados":       ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN62001RTcarneospaletasalgadosempanadospresserranopratopronto.pdf"],
    "corned_beef":            ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN832003RTcornedbeefbovinoconserva.pdf"],

    # ── PESCADO ADICIONAIS ──────────────────────────────────────────────────
    "camarao":                ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN232019RTcamaraofrescoresfriradocongelado.pdf"],
    "conserva_sardinhas":     ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN222011RTconservasardinhas.pdf"],
    "conserva_peixes":        ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN452011RTconservadepeixe.pdf"],
    "conserva_atuns":         ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN462011RTconservadeatuns.pdf"],
    "lagosta":                ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN242019RTlagostafrescocongelado.pdf"],

    # ── LATICÍNIOS ADICIONAIS ───────────────────────────────────────────────
    "leite_em_po":            ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN532018RTleiteempoMercosul.pdf"],
    "leite_condensado":       ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN472018RTleitecondensado.pdf"],
    "bebida_lactea":          ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN162005RTbebidalactea.pdf"],
    "soro_leite":             ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN942020RTsoroleite.pdf"],

    # ── MEL E APÍCOLA ADICIONAIS ────────────────────────────────────────────
    "mel_qualidade":          ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN112000RTmeldequal.pdf"],
    "apicola_derivados":      ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN032001RTapitoxinacerageleiarealpropolis.pdf"],

    # ── CÁRNEOS — GAPS FECHADOS ───────────────────────────────────────────
    "pate":                   ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN202000RTcrneosalmondegakibe.pdf"],
    "prato_pronto_poa":       ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN62001RTcarneospaletasalgadosempanadospresserranopratopronto.pdf"],
    "funcional_poa":          [
        "https://www.in.gov.br/en/web/dou/-/resolucao-rdc-n-267-de-22-de-setembro-de-2005-13651025",
        "https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN762018RTleitecrurefrleitepasteuhomogleitefluido.pdf",
    ],
    "jerked_beef":            ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN222000RTjerkedbeef.pdf"],
    "carne_sol":              ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN922020RTCharqueCarneSalgadaMidoSalgado.pdf"],
    "frango_inteiro":         ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/Port14852025PadrocategoriaenomeclaturaPOA.pdf"],

    # ── LATICÍNIOS — GAPS FECHADOS ────────────────────────────────────────
    "queijo_suico_gruyere":   ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/Port1461996RTqueijomanteigacremedeleitegorduralctealeitefluido.pdf"],
    "queijo_gouda_edam":      ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/Port1461996RTqueijomanteigacremedeleitegorduralctealeitefluido.pdf"],
    "creme_azedo":            ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN232012RTnata.pdf"],
    "miudos_visceras":        ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/Port14852025PadrocategoriaenomeclaturaPOA.pdf"],
    "peru_pato":              ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/Port14852025PadrocategoriaenomeclaturaPOA.pdf"],
    "leite_cabra":            ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/Port1461996RTqueijomanteigacremedeleitegorduralctealeitefluido.pdf"],
    "ovo_codorna":            ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/Port012020RTovosdemesadeovosparasementeovoscaipiras.pdf"],
    "queijo_provolone":       ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN732020RTqueijoprovolone.pdf"],
    "queijo_brie_camembert":  ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/Port1461996RTqueijomanteigacremedeleitegorduralctealeitefluido.pdf"],
    "bebida_lactea":          ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN162005RTbebidalactea.pdf"],
    "composto_lacteo":        ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN282007RTcompostolacteo.pdf"],

    # ── OVOS ADICIONAIS ─────────────────────────────────────────────────────
    "ovos_pasteurizados":     ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/Port7282022RTovointegralpasteurizadodesidratado.pdf"],

    # ── MAPA — RIISPOA ATUALIZADO ────────────────────────────────────────────
    "riispoa_atualizado":     ["https://www.planalto.gov.br/ccivil_03/_ato2019-2022/2020/decreto/d10468.htm"],

    # ── INMETRO — CONTEÚDO LÍQUIDO COMPLETO ──────────────────────────────────
    "inmetro_conteudo_liq":   [
        "https://www.inmetro.gov.br/legislacao/rtac/pdf/RTAC002775.pdf",
        "https://www.inmetro.gov.br/legislacao/rtac/pdf/RTAC002776.pdf",
    ],
    "inmetro_carnes_queijos": ["https://www.inmetro.gov.br/legislacao/rtac/pdf/RTAC003070.pdf"],

    # ── ANVISA — NOVA FÓRMULA (2024) ──────────────────────────────────────────
    "nova_formula_rdc902":    ["https://www.in.gov.br/en/web/dou/-/resolucao-rdc-n-902-de-6-de-setembro-de-2024-583517765"],

    # ── ANVISA — REGULARIZAÇÃO SNVS (2024) ───────────────────────────────────
    "regularizacao_rdc843":   ["https://www.in.gov.br/en/web/dou/-/resolucao-rdc-n-843-de-22-de-fevereiro-de-2024-543791616"],

    # ── ANVISA — ALÉRGENOS (obrigatório para TODOS os produtos) ─────────────
    "alergenos_rdc727":       ["https://www.in.gov.br/en/web/dou/-/resolucao-rdc-n-727-de-1-de-julho-de-2022-413249279"],

    # ── LATICÍNIOS — ÚLTIMAS 3 ────────────────────────────────────────────
    "leite_aromatizado":      ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN262007RTleitearomatizado.pdf"],
    "composto_lacteo":        ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN282007RTcompostolacteo.pdf"],
    "nata":                   ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN232012RTnata.pdf"],

    # ── OVOS — NOMENCLATURA ATUALIZADA ───────────────────────────────────
    "ovos_nomenclatura":      ["https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/Port11792024nomenclaturaovos.pdf"],

    # ── ANVISA — ROTULAGEM GERAL ──────────────────────────────────────────
    "rotulagem_anvisa_rdc715": ["https://www.in.gov.br/en/web/dou/-/resolucao-rdc-n-715-de-1-de-julho-de-2022-413249117"],

    # ── ANVISA — AROMATIZANTES ────────────────────────────────────────────
    # RDC 725/2022 — PDF scanned, sem texto extraível
    # Conteúdo embutido diretamente como texto (resumo da norma)
    "aromatizantes_rdc725":   ["https://www.in.gov.br/en/web/dou/-/resolucao-rdc-n-725-de-1-de-julho-de-2022-413249198"],

    # ── ANVISA — ADITIVOS IT 70/2016 (alegações aditivos em embutidos/laticínios)
    "aditivos_it70":          ["https://www.gov.br/anvisa/pt-br/assuntos/noticias-anvisa/2016/arquivos/2016_informe_tecnico_70.pdf"],

    # ════════════════════════════════════════════════════════════════════════
    # NORMAS GERAIS — APLICÁVEIS A TODOS OS ALIMENTOS (POA e não-POA)
    # ════════════════════════════════════════════════════════════════════════

    # ── INMETRO — CONTEÚDO LÍQUIDO COMPLEMENTARES ─────────────────────────
    "inmetro_251_requisitos": ["https://www.inmetro.gov.br/legislacao/rtac/pdf/RTAC002776.pdf"],
    "inmetro_265_medidas":    ["https://www.inmetro.gov.br/legislacao/rtac/pdf/RTAC003057.pdf"],
    "inmetro_248_verificacao":["https://www.inmetro.gov.br/legislacao/rtac/pdf/RTAC001781.pdf"],

    # ── ANVISA — IRRADIAÇÃO ────────────────────────────────────────────────
    "irradiacao_rdc21":       ["https://www.in.gov.br/en/web/dou/-/resolucao-rdc-n-21-de-26-de-janeiro-de-2001-396649938"],

    # ── ANVISA — REGULARIZAÇÃO CATEGORIAS (IN 281/2024) ───────────────────
    "regularizacao_in281":    ["https://www.in.gov.br/en/web/dou/-/instrucao-normativa-n-281-de-22-de-fevereiro-de-2024-547762772"],

    # ── ANVISA — SUPLEMENTOS ALIMENTARES ──────────────────────────────────
    "suplementos_rdc243":     ["https://www.in.gov.br/en/web/dou/-/resolucao-rdc-n-243-de-26-de-julho-de-2018-41232201"],
    "suplementos_in28":       ["https://www.in.gov.br/en/web/dou/-/instrucao-normativa-n-28-de-26-de-julho-de-2018-41232253"],

    # ── ANVISA — CEREAIS, FARINHAS, MASSAS ────────────────────────────────
    "cereais_rdc711":         ["https://www.in.gov.br/en/web/dou/-/resolucao-rdc-n-711-de-1-de-julho-de-2022-413249064"],
    "cereais_integrais_rdc712":["https://www.in.gov.br/en/web/dou/-/resolucao-rdc-n-712-de-1-de-julho-de-2022-413249083"],

    # ── LEGISLAÇÕES FEDERAIS GERAIS ────────────────────────────────────────
    "dec_lei_986_1969":       ["https://www.planalto.gov.br/ccivil_03/decreto-lei/del0986.htm"],
    "lei_8078_cdc":           ["https://www.planalto.gov.br/ccivil_03/leis/l8078compilado.htm"],
    "lei_6437_infracoes":     ["https://www.planalto.gov.br/ccivil_03/leis/l6437.htm"],
    "lei_12849_latex":        ["https://www.planalto.gov.br/ccivil_03/_ato2011-2014/2013/lei/l12849.htm"],
    "port_392_2021_alteracao":["https://www.in.gov.br/en/web/dou/-/portaria-n-392-de-29-de-setembro-de-2021-350099068"],

    # ── MAPA — TRANSGÊNICOS COMPLEMENTARES ────────────────────────────────
    "transgenicos_ini1_2004": ["https://www.planalto.gov.br/ccivil_03/_Ato2004-2006/2004/Instrucao_normativa/INI-1-2004.htm"],
    "transgenicos_port2658":  ["https://www.planalto.gov.br/ccivil_03/portaria/P_2658.htm"],

    # ── MAPA — BEBIDAS ─────────────────────────────────────────────────────
    "bebidas_dec6871":        ["https://www.planalto.gov.br/ccivil_03/_ato2007-2010/2009/decreto/d6871.htm"],

    # ════════════════════════════════════════════════════════════════════════
    # CATEGORIAS NÃO-POA — ANVISA + MAPA (FASE 2)
    # ════════════════════════════════════════════════════════════════════════

    # ── CEREAIS, PÃES, MASSAS, BISCOITOS ─────────────────────────────────
    "cereais_pao_massa":      ["https://www.in.gov.br/en/web/dou/-/resolucao-rdc-n-711-de-1-de-julho-de-2022-413249064"],
    "alimentos_integrais":    ["https://www.in.gov.br/en/web/dou/-/resolucao-rdc-n-712-de-1-de-julho-de-2022-413249083"],
    "farinhas_amidos":        ["https://www.in.gov.br/en/web/dou/-/resolucao-rdc-n-711-de-1-de-julho-de-2022-413249064"],
    "biscoito_bolacha":       ["https://www.in.gov.br/en/web/dou/-/resolucao-rdc-n-711-de-1-de-julho-de-2022-413249064"],
    "macarrao_massa":         ["https://www.in.gov.br/en/web/dou/-/resolucao-rdc-n-711-de-1-de-julho-de-2022-413249064"],

    # ── ÓLEOS E GORDURAS ─────────────────────────────────────────────────
    "oleos_gorduras":         ["https://antigo.anvisa.gov.br/documents/33916/392655/RDC+270+de+22+de+setembro+de+2005.pdf"],
    "azeite_oliva":           ["https://antigo.anvisa.gov.br/documents/33916/392655/RDC+270+de+22+de+setembro+de+2005.pdf"],
    "margarina":              ["https://antigo.anvisa.gov.br/documents/33916/392655/RDC+270+de+22+de+setembro+de+2005.pdf"],

    # ── AÇÚCAR, MEL VEGETAL E PRODUTOS AÇUCARADOS ───────────────────────
    "acucar_derivados":       ["https://www.planalto.gov.br/ccivil_03/decreto-lei/del0986.htm"],
    "chocolate_cacau":        ["https://antigo.anvisa.gov.br/documents/33916/392655/RDC+264_2005.pdf"],
    "doces_geleias":          ["https://antigo.anvisa.gov.br/documents/33916/392655/RDC+272_2005.pdf"],
    "sorvete_gelado":         ["https://antigo.anvisa.gov.br/documents/33916/392655/RDC+266_2005.pdf"],

    # ── FRUTAS, HORTALIÇAS E CONSERVAS ───────────────────────────────────
    "conservas_vegetais":     ["https://antigo.anvisa.gov.br/documents/33916/392655/RDC+272_2005.pdf"],
    "frutas_processadas":     ["https://antigo.anvisa.gov.br/documents/33916/392655/RDC+272_2005.pdf"],
    "cogumelos":              ["https://antigo.anvisa.gov.br/documents/33916/392655/RDC+272_2005.pdf"],

    # ── CONDIMENTOS, MOLHOS E TEMPEROS ───────────────────────────────────
    "condimentos_temperos":   ["https://antigo.anvisa.gov.br/documents/33916/392655/RDC+276_2005.pdf"],
    "molhos_ketchup":         ["https://antigo.anvisa.gov.br/documents/33916/392655/RDC+276_2005.pdf"],
    "vinagre":                ["https://antigo.anvisa.gov.br/documents/33916/392655/RDC+278_2005.pdf"],
    "cafe_cevada_cha":        ["https://antigo.anvisa.gov.br/documents/33916/392655/RDC+277_2005.pdf"],

    # ── LEGUMINOSAS E GRÃOS ──────────────────────────────────────────────
    "leguminosas_graos":      ["https://www.planalto.gov.br/ccivil_03/decreto-lei/del0986.htm"],
    "feijao_ervilha_graos":   ["https://www.planalto.gov.br/ccivil_03/decreto-lei/del0986.htm"],

    # ── BEBIDAS (expansão) ────────────────────────────────────────────────
    "suco_néctar":            ["https://www.planalto.gov.br/ccivil_03/_ato2007-2010/2009/decreto/d6871.htm"],
    "refrigerante":           ["https://www.planalto.gov.br/ccivil_03/_ato2007-2010/2009/decreto/d6871.htm"],
    "agua_mineral":           ["https://www.planalto.gov.br/ccivil_03/leis/l9433.htm"],
    "cerveja":                ["https://www.planalto.gov.br/ccivil_03/_ato2007-2010/2009/decreto/d6871.htm"],
    "vinho":                  ["https://www.planalto.gov.br/ccivil_03/leis/l7678.htm"],
    "cachaça_destilados":     ["https://www.planalto.gov.br/ccivil_03/_ato2007-2010/2009/decreto/d6871.htm"],
    "bebida_energetica":      ["https://www.in.gov.br/en/web/dou/-/resolucao-rdc-n-727-de-1-de-julho-de-2022-413249279"],
    "isotonica":              ["https://www.in.gov.br/en/web/dou/-/resolucao-rdc-n-727-de-1-de-julho-de-2022-413249279"],

    # ── SUPLEMENTOS ALIMENTARES ──────────────────────────────────────────
    "suplementos_rdc243_v2":  ["https://www.in.gov.br/en/web/dou/-/resolucao-rdc-n-243-de-26-de-julho-de-2018-41232201"],
    "suplementos_in28_v2":    ["https://www.in.gov.br/en/web/dou/-/instrucao-normativa-n-28-de-26-de-julho-de-2018-41232253"],
    "proteina_po":            ["https://www.in.gov.br/en/web/dou/-/resolucao-rdc-n-243-de-26-de-julho-de-2018-41232201"],
    "vitaminas_minerais":     ["https://www.in.gov.br/en/web/dou/-/resolucao-rdc-n-243-de-26-de-julho-de-2018-41232201"],

    # ════════════════════════════════════════════════════════════════════════
    # MAPA VEGETAL — BEBIDAS (Cartilhão de Bebidas — normas base acessíveis)
    # Fonte: planalto.gov.br (HTML estático, sempre acessível)
    # ════════════════════════════════════════════════════════════════════════

    # Lei base de bebidas
    "lei_8918_bebidas":       ["https://www.planalto.gov.br/ccivil_03/leis/l8918.htm"],
    "dec_6871_bebidas_tipos": ["https://www.planalto.gov.br/ccivil_03/_ato2007-2010/2009/decreto/d6871.htm"],
    # Vinho e derivados
    "lei_7678_vinho":         ["https://www.planalto.gov.br/ccivil_03/leis/l7678.htm"],
    "dec_8198_vinho":         ["https://www.planalto.gov.br/ccivil_03/_ato2011-2014/2014/decreto/D8198.htm"],
    # Suco de fruta — MAPA
    "in_mapa_49_2018_suco":   ["https://www.in.gov.br/en/web/dou/-/instrucao-normativa-n-49-de-26-de-setembro-de-2018-40018720"],
    # Bebidas não alcoólicas — nova regra DQI
    "port_mapa_123_2021":     ["https://www.in.gov.br/en/web/dou/-/portaria-n-123-de-14-de-maio-de-2021-319484390"],
    # Kombucha
    "in_mapa_41_2019_kombucha": ["https://www.in.gov.br/en/web/dou/-/instrucao-normativa-n-41-de-17-de-setembro-de-2019-215441545"],
    # Cachaça e aguardente
    "dec_6871_cachaca":       ["https://www.planalto.gov.br/ccivil_03/_ato2007-2010/2009/decreto/d6871.htm"],
    # Polpa e suco artesanal
    "lei_13648_2018_polpa":   ["https://www.planalto.gov.br/ccivil_03/_ato2015-2018/2018/lei/l13648.htm"],

    # ════════════════════════════════════════════════════════════════════════
    # MAPA VEGETAL — PRODUTOS VEGETAIS (Qualidade Vegetal — planalto)
    # ════════════════════════════════════════════════════════════════════════

    # Café
    "in_mapa_cafe_torrado":   ["https://antigo.anvisa.gov.br/documents/33916/392655/RDC+277_2005.pdf"],
    # Açúcar e derivados
    "dec_6268_classificacao": ["https://www.planalto.gov.br/ccivil_03/_ato2007-2010/2007/decreto/d6268.htm"],
    "lei_9972_classificacao": ["https://www.planalto.gov.br/ccivil_03/leis/l9972.htm"],
    # Frutas e hortaliças
    "rdc_714_frutas_hortalicas": ["https://www.in.gov.br/en/web/dou/-/resolucao-rdc-n-714-de-1-de-julho-de-2022-413249103"],
    # Orgânicos
    "lei_10831_organicos":    ["https://www.planalto.gov.br/ccivil_03/leis/2003/l10831.htm"],
    "dec_6323_organicos":     ["https://www.planalto.gov.br/ccivil_03/_ato2007-2010/2007/decreto/d6323.htm"],

    # ════════════════════════════════════════════════════════════════════════
    # ANVISA BIBLIOTECA ALIMENTOS — RDCs principais (PDF direto antigo.anvisa)
    # ════════════════════════════════════════════════════════════════════════
    "rdc_429_rotulagem_nutri": ["https://www.in.gov.br/en/web/dou/-/resolucao-rdc-n-429-de-8-de-outubro-de-2020-283614641"],
    "in_75_porcoes":          ["https://www.in.gov.br/en/web/dou/-/instrucao-normativa-n-75-de-8-de-outubro-de-2020-283615832"],
    "rdc_266_sorvetes":       ["https://antigo.anvisa.gov.br/documents/33916/392655/RDC+266_2005.pdf"],
    "rdc_268_proteina_vegetal": ["https://antigo.anvisa.gov.br/documents/33916/392655/RDC+268_2005.pdf"],
    "rdc_243_alimentos_funcionais": ["https://www.in.gov.br/en/web/dou/-/resolucao-rdc-n-243-de-26-de-julho-de-2018-41232201"],
    "rdc_277_cafe_cha":       ["https://antigo.anvisa.gov.br/documents/33916/392655/RDC+277_2005.pdf"],
    "rdc_264_chocolate":      ["https://antigo.anvisa.gov.br/documents/33916/392655/RDC+264_2005.pdf"],
    "rdc_270_oleos":          ["https://antigo.anvisa.gov.br/documents/33916/392655/RDC+270+de+22+de+setembro+de+2005.pdf"],
    "rdc_276_condimentos":    ["https://antigo.anvisa.gov.br/documents/33916/392655/RDC+276_2005.pdf"],
    "rdc_716_oleos_2022":     ["https://www.in.gov.br/en/web/dou/-/resolucao-rdc-n-716-de-1-de-julho-de-2022-413249130"],
    "rdc_713_acucar_2022":    ["https://www.in.gov.br/en/web/dou/-/resolucao-rdc-n-713-de-1-de-julho-de-2022-413249084"],

    # ── ALIMENTOS PARA FINS ESPECIAIS ────────────────────────────────────
    "dieta_diabetico":        ["https://antigo.anvisa.gov.br/documents/33916/392655/resolucao+29-98.pdf"],
    "formula_infantil":       ["https://www.in.gov.br/en/web/dou/-/resolucao-rdc-n-241-de-26-de-julho-de-2018-41232198"],
    "alimento_celíaco":       ["https://www.planalto.gov.br/ccivil_03/leis/2003/l10.674.htm"],

    # ── SNACKS, SALGADINHOS E INDUSTRIALIZADOS MISTOS ───────────────────
    "salgadinhos_snacks":     ["https://www.in.gov.br/en/web/dou/-/resolucao-rdc-n-711-de-1-de-julho-de-2022-413249064"],
    "barrinha_cereal":        ["https://antigo.anvisa.gov.br/documents/33916/392655/RDC+263_2005.pdf"],
    "amendoim_castanhas":     ["https://antigo.anvisa.gov.br/documents/33916/392655/RDC+276_2005.pdf"],
}
_kb_cache: dict = {}
_startup_ready: bool = False  # T1: True após warmup completo
_startup_ts: str = ""          # T1: timestamp do último startup bem-sucedido

# ═══════════════════════════════════════════════════════════════════════════════
# DETECÇÃO DE CATEGORIA — palavras-chave por produto
# ═══════════════════════════════════════════════════════════════════════════════
CATEGORIA_KEYWORDS = {
    # Embutidos
    "embutidos":        ["salsicha", "linguiça", "linguica", "mortadela", "apresuntado"],
    "salame":           ["salame", "salaminho", "copa", "coppa", "pepperoni", "presunto parma", "presunto serrano"],
    "presunto":         ["presunto", "presunto cozido", "presunto defumado"],
    "hamburguer":       ["hambúrguer", "hamburguer", "burger"],
    "bacon":            ["bacon", "panceta", "pancetta"],
    "carne_moida":      ["carne moída", "carne moida"],
    "charque":          ["charque", "carne salgada", "miúdo salgado", "jabá", "jaba"],
    "fiambre":          ["fiambre"],
    "carne_maturada":   ["carne maturada", "dry aged"],
    "carnes_temperadas":["carne temperada", "frango temperado", "carne marinada"],
    "almondega_kibe":   ["almôndega", "almondega", "kibe", "quibe"],
    "gelatina_colageno":["gelatina", "colágeno", "colageno"],

    # Laticínios
    "laticinios_geral": ["queijo", "queijinho", "cheese"],
    "queijo_coalho_manteiga": ["queijo coalho", "queijo de manteiga", "manteiga de garrafa", "manteiga da terra"],
    "mussarela":        ["mussarela", "mozzarela", "mozzarella", "muçarela"],
    "leite_uht":        ["leite uht", "leite longa vida", "leite uat"],
    "leite_pasteurizado":["leite pasteurizado", "leite integral", "leite desnatado", "leite semidesnatado"],
    "doce_de_leite":    ["doce de leite"],
    "leite_fermentado": ["iogurte", "yogurte", "leite fermentado", "kefir", "bebida láctea", "bebida lactea"],
    "requeijao":        ["requeijão", "requeijao", "queijo fundido", "cream cheese"],

    # Pescado
    "pescado_fresco":   ["peixe", "pescado", "tilápia", "tilapia", "salmão", "salmao", "atum", "sardinha", "merluza", "badejo"],
    "pescado_congelado":["peixe congelado", "pescado congelado", "filé congelado", "file congelado"],
    "pescado_salgado":  ["bacalhau", "peixe salgado", "pescado salgado"],
    "camarao":          ["camarão", "camarao", "lagosta", "lula", "polvo"],

    # Mel e apícola
    "mel":              ["mel", "própolis", "propolis", "geleia real", "pólen", "pollen", "apicola", "apícola"],

    # Ovos
    "ovos":             ["ovo", "ovos", "ovo caipira", "ovo codorna"],

    # Aves
    "aves_geral":       ["frango", "frango inteiro", "peito de frango", "coxa", "sobrecoxa", "pato", "peru", "aves"],
    # Carnes adicionais
    "apresuntado":      ["apresuntado"],
    "paleta_empanados": ["paleta cozida", "empanado", "presunto serrano", "presunto parma", "prato pronto"],
    "corned_beef":      ["corned beef", "carne em conserva", "carne enlatada", "carne bovina em conserva"],
    # Pescado adicional
    "camarao":          ["camarão", "camarao", "camaroes"],
    "conserva_sardinhas":["sardinha", "conserva de sardinha", "sardinha em óleo", "sardinha em molho"],
    "conserva_peixes":  ["peixe em conserva", "peixe enlatado", "atum em lata", "conserva de peixe"],
    "conserva_atuns":   ["atum", "bonito", "conserva de atum", "atum em óleo", "atum ao natural"],
    "lagosta":          ["lagosta", "caranguejo"],
    # Laticínios adicionais
    "leite_em_po":      ["leite em pó", "leite em po", "leite desidratado", "leite integral em pó"],
    "leite_condensado": ["leite condensado", "leite condensado adoçado"],
    "bebida_lactea":    ["bebida láctea", "bebida lactea"],
    "soro_leite":       ["soro de leite", "soro lácteo", "whey"],
    # Mel adicionais
    "mel_qualidade":    ["mel", "mel puro", "mel orgânico", "mel silvestre", "mel florada"],
    "poa_caprino_ovino": ["leite de cabra", "queijo de cabra", "iogurte caprino",
                         "carne ovina", "cordeiro", "carneiro", "ovino", "caprino",
                         "chèvre", "chevre", "queijo caprino"],
    "poa_organico":     ["frango orgânico", "carne orgânica", "suíno orgânico", "bovino orgânico",
                         "ovos orgânicos", "leite orgânico", "queijo orgânico", "iogurte orgânico",
                         "sisorg", "certificado orgânico", "produto orgânico poa"],
    "apicola_derivados":["própolis", "propolis", "geleia real", "pólen apícola", "pollen", "cera de abelha", "apitoxina", "extrato de própolis"],
    # Ovos adicionais
    "ovos_pasteurizados":["ovo pasteurizado", "ovo desidratado", "ovo em pó", "ovo integral"],
    "ovos_nomenclatura": ["ovos caipira", "ovos de granja", "ovo orgânico", "ovo cage free"],
    # Laticínios finais
    "leite_aromatizado": ["leite aromatizado", "leite com chocolate", "leite com morango", "leite sabor"],
    "composto_lacteo":   ["composto lácteo", "composto lacteo", "alimento lácteo composto"],
    "nata":              ["nata", "creme nata"],
    # ANVISA geral e MAPA geral
    "rotulagem_anvisa_rdc715": ["rotulagem", "rótulo alimento embalado"],
    "aromatizantes_rdc725":    ["aromatizante", "aroma natural", "flavorizante"],
    "aditivos_it70":           ["aditivo alimentar", "conservante", "corante", "estabilizante", "espessante"],
    "riispoa_atualizado":      ["inspeção industrial", "produto de origem animal", "registro produto"],
    "inmetro_carnes_queijos":  ["produto cárneo", "queijo", "requeijão", "conteúdo líquido"],
    "nova_formula_rdc902":     ["nova fórmula", "alteração composição"],
    "regularizacao_rdc843":    ["regularização", "notificação produto"],
    # INMETRO complementares
    "inmetro_251_requisitos":  ["conteúdo líquido", "volume nominal", "embalagem"],
    "inmetro_265_medidas":     ["conteúdo líquido", "indicação quantitativa", "massa volume"],
    "inmetro_248_verificacao": ["verificação conteúdo", "produto pré-medido"],
    # ANVISA gerais
    "irradiacao_rdc21":        ["irradiado", "irradiação", "tratamento radiação"],
    "regularizacao_in281":     ["notificação", "isenção registro", "categoria alimento"],
    "suplementos_rdc243":      ["suplemento alimentar", "suplemento nutricional", "whey protein", "proteína em pó"],
    "suplementos_in28":        ["suplemento alimentar", "BCAA", "creatina", "termogênico"],
    "cereais_rdc711":          ["cereal", "farinha", "amido", "massa alimentícia", "biscoito", "pão", "farelo"],
    "cereais_integrais_rdc712":["integral", "cereal integral", "grão integral", "farinha integral"],
    # Legislações federais
    "dec_lei_986_1969":        ["alimento", "denominação", "fraude alimentar"],
    "lei_8078_cdc":            ["consumidor", "direito do consumidor", "código defesa"],
    "lei_6437_infracoes":      ["infração sanitária", "autuação", "penalidade"],
    "lei_12849_latex":         ["látex", "latex natural", "alergia látex"],
    "port_392_2021_alteracao": ["redução de conteúdo", "downsizing", "embalagem menor"],
    # Transgênicos
    "transgenicos_ini1_2004":  ["transgênico", "ogm", "organismo geneticamente modificado"],
    "transgenicos_port2658":   ["símbolo transgênico", "símbolo T", "símbolo ogm"],
    # Bebidas
    "bebidas_dec6871":         ["bebida", "suco", "néctar", "refrigerante", "cerveja", "vinho", "cachaça"],
    # Pescado — categorias novas
    "camarao":          ["camarão", "camarao", "camarões", "camaroes"],
    "lagosta":          ["lagosta"],
    "moluscos_cefalopodes": ["lula", "polvo", "pota", "cefalópode", "cefalopode"],
    "conserva_peixe_sardinha": ["sardinha", "conserva de sardinha", "sardinha em lata"],
    "conserva_atum":    ["atum", "atum em conserva", "atum enlatado", "bonito", "conserva de atum"],

    # Laticínios extras
    "leite_em_po":      ["leite em pó", "leite em po", "leite desidratado"],
    "leite_condensado": ["leite condensado", "leite condençado"],
    "nata_creme":       ["nata", "creme de leite"],
    "queijo_prato":     ["queijo prato"],
    "queijo_parmesao":  ["queijo parmesão", "queijo parmesao", "parmesano", "reggiano"],
    "queijo_provolone": ["queijo provolone", "provolone"],
    "queijo_fundido":   ["queijo fundido", "queijo processado", "queijo cremoso"],

    # ════════════════════════════════════════════════════════════════════
    # CATEGORIAS NÃO-POA (Fase 2)
    # ════════════════════════════════════════════════════════════════════

    # ── CEREAIS, PÃES, MASSAS, BISCOITOS ─────────────────────────────────
    "cereais_pao_massa":   ["pão", "pao", "bisnaguinha", "baguete", "macarrão", "macarrao",
                            "massa", "espaguete", "lasanha", "arroz", "aveia", "granola",
                            "muesli", "cereal matinal", "biscoito", "bolacha", "wafer",
                            "crackers", "torrada"],
    "farinhas_amidos":     ["farinha de trigo", "farinha de milho", "farinha de arroz",
                            "amido de milho", "maisena", "fécula", "polvilho", "tapioca",
                            "fubá", "farelo", "flocos"],
    "biscoito_bolacha":    ["biscoito", "bolacha", "wafer", "cookie", "crackers"],
    "macarrao_massa":      ["macarrão", "macarrao", "massa", "espaguete", "penne",
                            "fusilli", "lasanha", "talharim", "nhoque"],
    "alimentos_integrais": ["integral", "grão inteiro", "cereal integral", "pão integral",
                            "farinha integral", "macarrão integral"],

    # ── ÓLEOS, GORDURAS E MARGARINAS ─────────────────────────────────────
    "oleos_gorduras":      ["óleo", "oleo", "gordura vegetal", "óleo de soja",
                            "óleo de girassol", "óleo de canola", "óleo de milho",
                            "óleo de palma", "óleo de coco"],
    "azeite_oliva":        ["azeite", "azeite de oliva", "olive oil", "azeite extra virgem"],
    "margarina":           ["margarina", "creme vegetal", "gordura vegetal culinária"],

    # ── AÇÚCAR E PRODUTOS AÇUCARADOS ─────────────────────────────────────
    "acucar_derivados":    ["açúcar", "acucar", "açúcar refinado", "açúcar cristal",
                            "açúcar mascavo", "açúcar demerara", "rapadura", "melado"],
    "chocolate_cacau":     ["chocolate", "cacau", "achocolatado", "creme de avelã",
                            "trufas", "bombom"],
    "doces_geleias":       ["geleia", "compota", "doce de frutas", "marmelada",
                            "goiabada", "pasta de amendoim"],
    "sorvete_gelado":      ["sorvete", "picolé", "picole", "gelato", "frozen", "sorbet"],

    # ── FRUTAS, HORTALIÇAS E CONSERVAS ───────────────────────────────────
    "conservas_vegetais":  ["conserva", "picles", "ervilha enlatada", "milho enlatado",
                            "tomate pelado", "extrato de tomate", "polpa de tomate",
                            "palmito", "azeitona"],
    "frutas_processadas":  ["polpa de fruta", "purê de fruta", "fruta desidratada",
                            "fruta seca", "uva passa", "banana passa"],
    "cogumelos":           ["cogumelo", "shitake", "champignon", "funghi"],

    # ── CONDIMENTOS, MOLHOS E TEMPEROS ───────────────────────────────────
    "condimentos_temperos":["tempero", "condimento", "pimenta", "colorau", "urucum",
                            "alho em pó", "cebola em pó", "curry", "curcuma", "canela",
                            "cominho", "páprica", "sal temperado", "caldo em cubo",
                            "caldo em pó", "sazon"],
    "molhos_ketchup":      ["molho", "ketchup", "maionese", "mostarda", "shoyu",
                            "molho de soja", "molho inglês", "molho barbecue", "pesto"],
    "vinagre":             ["vinagre", "aceto", "vinagre de maçã", "vinagre balsâmico"],
    "cafe_cevada_cha":     ["café", "cafe", "nescafé", "espresso", "cevada",
                            "chá", "cha", "erva mate", "mate", "camomila"],

    # ── LEGUMINOSAS E GRÃOS ──────────────────────────────────────────────
    "leguminosas_graos":   ["feijão", "feijao", "lentilha", "grão de bico", "ervilha",
                            "fava", "quinoa", "trigo em grão"],
    "feijao_ervilha_graos":["feijão carioca", "feijão preto", "feijão branco",
                            "grão de bico", "lentilha vermelha"],

    # ── BEBIDAS ──────────────────────────────────────────────────────────
    "suco_néctar":         ["suco", "néctar", "nectar", "suco de laranja", "suco de uva",
                            "suco integral", "néctar de fruta", "refresco", "polpa de açaí"],
    "refrigerante":        ["refrigerante", "coca-cola", "pepsi", "guaraná", "fanta",
                            "sprite", "água tônica", "cola"],
    "agua_mineral":        ["água mineral", "agua mineral", "água com gás",
                            "água saborizada", "água aromatizada"],
    "cerveja":             ["cerveja", "chope", "lager", "ale", "pilsen", "ipa", "stout",
                            "weiss", "cerveja sem álcool", "cerveja zero"],
    "vinho":               ["vinho", "espumante", "champagne", "prosecco",
                            "vinho tinto", "vinho branco", "vinho rosé"],
    "cachaça_destilados":  ["cachaça", "cachaca", "pinga", "aguardente", "vodka",
                            "rum", "whisky", "gin", "tequila", "conhaque", "licor"],
    "bebida_energetica":   ["energético", "energetico", "red bull", "monster",
                            "energy drink", "bebida energizante"],
    "isotonica":           ["isotônico", "isotonico", "gatorade", "powerade",
                            "bebida esportiva", "bebida hidratante"],

    # ── SUPLEMENTOS ALIMENTARES ──────────────────────────────────────────
    "proteina_po":         ["whey", "proteína em pó", "proteina em po", "proteína isolada",
                            "proteína concentrada", "caseína", "albumina",
                            "proteína vegetal", "proteína de ervilha"],
    "vitaminas_minerais":  ["vitamina c", "vitamina d", "vitamina b12", "complexo b",
                            "multivitamínico", "zinco", "magnésio", "ômega 3",
                            "fish oil", "colágeno hidrolisado"],
    "suplementos_rdc243_v2":["suplemento", "bcaa", "creatina", "glutamina",
                              "termogênico", "pre-treino", "hipercalórico"],
    "dieta_diabetico":     ["diet", "zero açúcar", "sem açúcar adicionado", "diabético"],
    "formula_infantil":    ["fórmula infantil", "formula infantil", "leite maternizado"],

    # ── SNACKS E INDUSTRIALIZADOS MISTOS ────────────────────────────────
    "salgadinhos_snacks":  ["salgadinho", "snack", "chips", "doritos", "cheetos",
                            "batata frita", "amendoim crocante"],
    "barrinha_cereal":     ["barra de cereal", "barrinha", "granola bar",
                            "barra de proteína"],
    "amendoim_castanhas":  ["amendoim", "castanha", "nozes", "amêndoa", "amendoa",
                            "macadâmia", "avelã", "pistache", "mix de nuts"],

    # Carnes extras (mantido)
    "pate_bacon_copa":     ["patê", "pate", "copa"],
    "paleta_empanados_presunto_serrano": ["paleta cozida", "empanado", "nugget", "presunto serrano"],
    "corned_beef_conserva": ["corned beef", "carne enlatada", "carne em conserva"],

    # Ovos derivados
    "ovos_pasteurizados": ["ovo pasteurizado", "ovo desidratado", "ovo integral pasteurizado"],

    # ANVISA
    "alergenos_rdc727":   ["alérgenos", "alergenos", "alergênico", "alergenico"],
    "nutricional_in75":   ["tabela nutricional", "informação nutricional"],
    "aditivos_rdc778":    ["aditivo", "conservante", "corante", "estabilizante"],

    # Pescado extras
    "conserva_sardinha":  ["sardinha", "sardinha em lata", "conserva de sardinha"],
    "conserva_atum":      ["atum", "atum em conserva", "atum enlatado", "bonito"],
    "conserva_peixe":     ["conserva de peixe", "peixe em lata", "peixe enlatado"],

    # Laticínios extras
    "leite_em_po":        ["leite em pó", "leite em po", "leite desidratado", "leite pó"],
    "leite_condensado":   ["leite condensado"],
    "nata":               ["nata", "creme de leite"],
    "queijo_provolone":   ["provolone", "queijo provolone"],
    "queijo_parmesao":    ["parmesão", "parmesao", "parmesano", "queijo parmesão"],

    # Ovos derivados
    "ovos_derivados":     ["ovo pasteurizado", "ovo desidratado", "ovo integral pasteurizado", "ovoproduto"],

    # ── CARNES — GAPS FECHADOS ─────────────────────────────────────────────
    "pate":               ["patê", "pate", "patê de frango", "patê de fígado", "patê de presunto"],
    "prato_pronto_poa":   ["empanado", "nugget", "steak de frango", "bife empanado", "lasanha",
                           "lasanha com carne", "croquete", "prato pronto", "semiprontos",
                           "produto cárneo elaborado", "hambúrguer recheado", "cordon bleu"],
    "funcional_poa":      ["leite com cálcio", "leite enriquecido", "leite com vitaminas",
                           "leite com ômega", "ovo enriquecido", "ovo com ômega", "probiótico",
                           "prebiótico", "lactobacillus", "bifidobacterium", "leite funcional",
                           "iogurte funcional", "alimento funcional", "reduzido em gordura"],
    "jerked_beef":        ["jerked beef", "jerk beef", "carne bovina salgada", "carne curada"],
    "carne_sol":          ["carne de sol", "carne do sol", "carne sol"],
    "frango_inteiro":     ["frango inteiro", "frango resfriado", "frango congelado", "galeto", "frango caipira"],
    "miudos_visceras":    ["fígado", "figado", "coração bovino", "coracao bovino", "rim bovino", "língua bovina",
                           "lingua bovina", "bucho", "mocotó", "mocoto", "tutano", "rabo bovino",
                           "moela", "coração de frango", "fígado de frango", "figado de frango", "miúdo", "miudo"],
    "peru_pato":          ["peru", "peito de peru", "coxa de peru", "pato", "carne de pato", "pato inteiro"],
    "leite_cabra":        ["leite de cabra", "queijo de cabra", "iogurte de cabra"],
    "ovo_codorna":        ["ovo de codorna", "ovos de codorna"],
    "queijo_brie_camembert": ["brie", "camembert", "queijo brie", "queijo camembert"],
    "queijo_provolone":   ["provolone", "queijo provolone"],
    "paleta_salgada":     ["paleta cozida", "empanado", "nugget", "presunto serrano", "prato pronto"],
    "corned_beef":        ["corned beef", "carne enlatada", "carne em conserva"],

    # ── LATICÍNIOS — GAPS FECHADOS ─────────────────────────────────────────
    "queijo_suico_gruyere":["queijo suíço", "queijo suico", "gruyère", "gruyere", "emmental", "emental"],
    "queijo_gouda_edam":   ["gouda", "edam", "queijo holandês", "queijo holandes"],
    "creme_azedo":         ["creme azedo", "sour cream", "crème fraîche"],
    "bebida_lactea":       ["bebida láctea", "bebida lactea", "iogurte bebível", "iogurte para beber"],
    "composto_lacteo":     ["composto lácteo", "composto lacteo", "alimento lácteo composto"],

    # ── PESCADO — REGIONAIS E GAPS ─────────────────────────────────────────
    "pescado_fresco":      ["peixe", "pescado", "tilápia", "tilapia", "salmão", "salmao", "atum",
                            "sardinha", "merluza", "badejo", "pintado", "dourado", "robalo", "tainha",
                            "corvina", "pacu", "tambaqui", "pirarucu", "surubim", "garoupa",
                            "linguado", "pescada", "cação", "cacao", "tubarão"],

    # ════════════════════════════════════════════════════════════════════════
    # NÃO-POA — PRODUTOS VEGETAIS, INDUSTRIALIZADOS, BEBIDAS, SUPLEMENTOS
    # ════════════════════════════════════════════════════════════════════════
    "cereais_rdc711":          ["arroz branco", "milho", "fubá", "fuba", "amido de milho",
                                "farinha de trigo", "farinha de milho", "farinha de arroz",
                                "flocos de aveia", "flocos de milho", "aveia em flocos",
                                "canjica", "quirera", "pipoca", "polenta"],
    "cereais_integrais_rdc712":["integral", "grão inteiro", "grain", "farelo de trigo",
                                "farelo de aveia", "farinha integral", "cereal integral"],
    "frutas_hortalicas_rdc714":["conserva de legume", "palmito", "milho em conserva",
                                "ervilha em conserva", "tomate pelado", "extrato de tomate",
                                "purê de tomate", "passata", "molho de tomate",
                                "picles", "azeitona", "alcaparra"],
    "acucares_rdc713":         ["açúcar cristal", "açúcar refinado", "açúcar demerara",
                                "açúcar mascavo", "açúcar orgânico", "açúcar impalpável",
                                "rapadura", "melado", "melaço"],
    "oleos_gorduras_rdc270":   ["óleo de soja", "oleo de soja", "óleo de girassol",
                                "óleo de canola", "óleo de milho", "azeite de oliva",
                                "azeite extra virgem", "óleo de côco", "gordura vegetal",
                                "banha", "gordura de palma", "margarina", "creme vegetal"],
    "sucos_nectares_in49":     ["suco de", "néctar de", "nectar de", "refresco de",
                                "bebida de fruta", "suco integral", "suco concentrado",
                                "polpa de fruta", "smoothie"],
    "cerveja_in14":            ["cerveja", "chope", "chopp", "ale", "lager", "pilsen",
                                "weiss", "weizen", "ipa", "stout", "porter", "bock",
                                "cerveja artesanal", "craft beer"],
    "vinho_lei7678":           ["vinho tinto", "vinho branco", "vinho rosé", "vinho rose",
                                "espumante", "champagne", "prosecco", "cava",
                                "vinho frisante", "vinho licoroso", "vinho do porto"],
    "suplementos_rdc243":      ["suplemento alimentar", "whey protein", "whey concentrado",
                                "whey isolado", "creatina", "bcaa", "aminoácido essencial",
                                "colágeno hidrolisado", "multivitamínico", "termogênico",
                                "pré-treino", "pre-treino", "hipercalórico", "mass gainer",
                                "albumina", "caseína", "caseina", "proteína vegetal em pó"],
    "suplementos_proteinas":   ["proteína em pó", "proteina em po", "proteína concentrada",
                                "proteína isolada", "blend proteico"],
    "regularizacao_rdc843":    ["biscoito", "bolacha", "cookie", "wafer",
                                "chocolate", "barra de chocolate", "bombom", "trufa",
                                "sorvete", "gelado", "picolé", "bolo", "torta",
                                "macarrão instantâneo", "lamen", "lámen", "ramen",
                                "salgadinho", "chips", "snack", "barra de cereal",
                                "ketchup", "mostarda industrial", "maionese",
                                "molho shoyu", "shoyu", "molho inglês",
                                "fermento", "levedura", "extrato de levedura"],
}

def detect_categories(obs: str) -> list[str]:
    """Detecta múltiplas categorias relevantes para o produto."""
    o = obs.lower()
    detected = []
    for categoria, keywords in CATEGORIA_KEYWORDS.items():
        if any(kw in o for kw in keywords):
            # Mapeia aves_geral para nomenclatura (sem RTIQ próprio ainda)
            if categoria == "aves_geral":
                detected.append("nomenclatura")
            elif categoria in MAPA_URLS:
                detected.append(categoria)
    # Sempre inclui laticínios_geral para qualquer queijo
    if "queijo" in o and "laticinios_geral" not in detected:
        detected.append("laticinios_geral")
    # Atum pode ser cárneo ou conserva - cobre os dois
    if "atum" in o:
        if "conserva_atuns" not in detected: detected.append("conserva_atuns")
        if "conserva_peixes" not in detected: detected.append("conserva_peixes")
    return list(dict.fromkeys(detected))  # remove duplicatas mantendo ordem

# ── FALLBACK TEXTUAL — RDC 725/2022 (PDF scanned, sem texto extraível) ───────
RDC_725_FALLBACK = """RESOLUÇÃO RDC Nº 725, DE 1º DE JULHO DE 2022 — ANVISA
Dispõe sobre os aditivos alimentares aromatizantes. Vigência: 1/9/2022.

DEFINIÇÕES: Aromatizante = substância com propriedades odoríferas/sápidas. Natural = obtido de matérias-primas vegetais/animais/microbianas por processos físicos, microbiológicos ou enzimáticos. Artificial = sintetizado quimicamente. De reação = processo Maillard. De fumaça = extrato de madeiras não tratadas.

ROTULAGEM NA LISTA DE INGREDIENTES (Art. 10 + RDC 727/2022 Art. 12 §3º):
- Declarar pela FUNÇÃO TECNOLÓGICA: "Aromatizante"
- Pode acrescentar classificação: "natural", "artificial" ou "de fumaça"
- Exemplos válidos: "Aromatizante natural" / "Aromatizante artificial" / "Aromatizante de fumaça"
- PROIBIDO: mencionar propriedades medicamentosas ou terapêuticas das ervas

RESTRIÇÕES (Art. 7º):
- Proibidos: óleos essenciais de fava-tonca, sassafrás e sabina
- Aromatizante de fumaça: máx 0,03 μg de 3,4-benzopireno/kg de alimento

PARA POA (embutidos defumados, bacon, linguiça, charque):
- Fumaça líquida = aromatizante de fumaça → declarar como "Aromatizante de fumaça"
- Aroma de defumação artificial → declarar como "Aromatizante artificial"
- Extrato de fumaça, fumaça condensada: idem, declarar função + classificação"""


# ══════════════════════════════════════════════════════════════════════════════
# NORMAS ESTADUAIS SIE — DIFERENÇAS EM RELAÇÃO AO FEDERAL
# Base: manuais CISPOA-RS, SISP-SP, IMA-MG, CIDASC-SC, ADAPAR-PR
# Conteúdo de rótulo federal (MAPA + ANVISA) aplica-se integralmente a todos.
# O que está aqui são adições/especificidades estaduais.
# ══════════════════════════════════════════════════════════════════════════════

CISPOA_RS_FALLBACK = """NORMAS COMPLEMENTARES CISPOA-RS (Coordenadoria de Inspeção de Produtos de Origem Animal - RS)

BASE LEGAL: Decreto Estadual RS nº 37.106/1996 + Portarias CISPOA vigentes.

DIFERENÇAS DO FEDERAL:
1. DENOMINAÇÃO DE VENDA:
   • O RS exige que a denominação inclua OBRIGATORIAMENTE a espécie animal por extenso (ex: "Linguiça Suína Frescal", não apenas "Linguiça Frescal")
   • Queijos artesanais: deve constar "Artesanal" na denominação se produzido artesanalmente

2. IDENTIFICAÇÃO DO FABRICANTE:
   • Além do federal, deve constar o número de registro no CISPOA no formato: "CISPOA/RS Nº XXXX" 
   • Estabelecimentos de agricultura familiar: aceita CPF ao invés de CNPJ + menção "Agricultura Familiar"

3. CAMPO SIE (específico RS):
   • Carimbo oval com: "SIE/RS" + número ou "CISPOA" + número
   • Sigla aceita: "SIE RS", "CISPOA RS", "INPOA RS"
   • Alguns selos mais antigos usam apenas número — aceitar se oval e legível

4. LATICÍNIOS ARTESANAIS RS:
   • Queijos artesanais com maturação: prazo de validade diferenciado (Lei Estadual RS 15.136/2018)
   • "Queijo Artesanal" deve mencionar município/região de origem se for indicação geográfica
   • Não exige tabela nutricional para queijos artesanais com superfície ≤100cm² (mesma isenção federal)

5. PRODUTOS DE PESCADO RS:
   • Pescado de piscicultura gaúcha: pode incluir menção opcional "Piscicultura Gaúcha"
   • Mesmas exigências federais para tabela nutricional e alérgenos

OBSERVAÇÃO: Para rotulagem federal (SIF), o CISPOA não adiciona requisitos de conteúdo. 
As adições acima são exclusivas para produtos com registro CISPOA."""

SISP_SP_FALLBACK = """NORMAS COMPLEMENTARES SISP-SP (Serviço de Inspeção do Estado de São Paulo)

BASE LEGAL: Decreto Estadual SP nº 24.528/1986 + Portaria SAA-SP nº 06/2011 e atualizações.

DIFERENÇAS DO FEDERAL:
1. IDENTIFICAÇÃO DO FABRICANTE:
   • Deve constar número de registro SISP no formato: "SISP Nº XXXX" ou "SIE SP XXXX"
   • Resolução SAA-SP exige declaração do município de fabricação por extenso (não apenas endereço)

2. CAMPO SIE (específico SP):
   • Carimbo oval: "SISP" + número ou "SIE SP" + número
   • Cor do carimbo: preferencialmente preto ou azul escuro sobre fundo branco
   • Não aceita carimbos apenas com número sem sigla SISP ou SIE SP

3. LATICÍNIOS SP:
   • Queijo Minas Artesanal produzido em SP: pode usar denominação regional se certificado pelo IEA/SP
   • Manteiga de garrafa: denominação "Manteiga da Terra" aceita em SP conforme Port. SAA

4. PRODUTOS CÁRNEOS SP:
   • Linguiça "tipo toscana" produzida em SP: deve especificar se suína pura (sem proteína vegetal)
   • Referências ao "Padrão Paulista" ou regionais são permitidas se tecnicamente corretas

OBSERVAÇÃO: SP tem volume alto de registros SISP em laticínios artesanais e embutidos de pequenos produtores. A maioria dos requisitos espelha o federal."""

IMA_MG_FALLBACK = """NORMAS COMPLEMENTARES IMA-MG (Instituto Mineiro de Agropecuária)

BASE LEGAL: Decreto Estadual MG nº 44.864/2008 + Portarias IMA vigentes.

DIFERENÇAS DO FEDERAL:
1. IDENTIFICAÇÃO DO FABRICANTE:
   • Deve constar número de registro IMA-MG: "IMA/MG Nº XXXX" ou "SIE MG XXXX"
   • Para agroindústria familiar: pode constar DAP (Declaração de Aptidão ao PRONAF) ao invés de CNPJ

2. CAMPO SIE (específico MG):
   • Carimbo oval: "IMA MG" ou "SIE MG" + número
   • Queijos artesanais de Minas: carimbo específico "QUEIJO ARTESANAL DE MINAS" + microrregião

3. QUEIJOS ARTESANAIS MINEIROS (Lei Estadual MG 23.157/2018):
   • Denominação obrigatória: "Queijo Minas Artesanal" + microrregião certificada
     (Canastra, Serro, Serra da Canastra, Araxá, Cerrado, Campo das Vertentes, etc.)
   • NÃO precisa de tabela nutricional se embalagem ≤100cm² (mesma isenção federal)
   • DEVE constar: "Produto artesanal feito de leite cru" se for o caso
   • Prazo de validade diferenciado conforme maturação (Decreto MG 44.864/2008)
   • Identificação do produtor rural: pode ser CPF + nome do produtor

4. EMBUTIDOS ARTESANAIS MG:
   • Linguiça artesanal: pode usar denominação regional ("Linguiça Mineira") se tradicional
   • Origem animal: espécie obrigatória por extenso

OBSERVAÇÃO: MG tem o maior volume de produtos SIE com queijos artesanais certificados. 
As normas de queijo artesanal mineiro são as mais específicas do país."""

CIDASC_SC_FALLBACK = """NORMAS COMPLEMENTARES CIDASC-SC (Companhia Integrada de Desenvolvimento Agrícola de SC)

BASE LEGAL: Lei Estadual SC 12.117/2002 + Manual CIDASC de Rotulagem (versão 5.0/2023).

DIFERENÇAS DO FEDERAL:
1. IDENTIFICAÇÃO DO FABRICANTE:
   • Deve constar: "SIE/SC Nº XXXX" ou "CIDASC Nº XXXX"
   • Agroindústria familiar rural: deve constar "Estabelecimento de Agroindústria Rural Familiar"

2. CAMPO SIE (específico SC):
   • Carimbo oval: "SIE SC" + número ou "CIDASC" + número
   • SC aceita "SISC" (antigo) para estabelecimentos registrados antes de 2010

3. MEL E APÍCOLA SC:
   • Mel catarinense com indicação de flora: deve constar "Mel Monofloral" ou "Mel Silvestre/Multifloral"
   • Não exige tabela nutricional para embalagens ≤100cm²

4. EMBUTIDOS E DEFUMADOS SC:
   • Linguiça colonial catarinense: pode usar denominação "Colonial" se atender padrão SC
   • Salame colonial: idem, "Salame Colonial Catarinense" é denominação regional aceita

OBSERVAÇÃO: SC tem manual de rotulagem publicado e atualizado (2023) que é o mais completo 
do país para nível estadual. Disponível em cidasc.sc.gov.br."""

ADAPAR_PR_FALLBACK = """NORMAS COMPLEMENTARES ADAPAR-PR (Agência de Defesa Agropecuária do Paraná)

BASE LEGAL: Lei Estadual PR 14.978/2006 + Resolução ADAPAR vigente.

DIFERENÇAS DO FEDERAL:
1. IDENTIFICAÇÃO DO FABRICANTE:
   • Deve constar: "SIE/PR Nº XXXX" ou "ADAPAR Nº XXXX"

2. CAMPO SIE (específico PR):
   • Carimbo oval: "SIE PR" + número ou "ADAPAR" + número

3. ESPECIFICIDADES PR:
   • O PR praticamente espelha o federal — menor diferenciação entre os estados com SIE ativo
   • Queijos artesanais: mesmas regras federais, sem denominação regional específica certificada
   • Embutidos: sem adições ao federal

OBSERVAÇÃO: O PR tem exigências mais próximas ao federal. 
Foco principal em procedimentos de registro, não em conteúdo de rótulo."""

# Mapa de Estado → Fallback

# ──────────────────────────────────────────────────────────────────────────────
# PRATOS PRONTOS POA — RTIQ (IN 6/2001 + Port. 1485/2025)
# ──────────────────────────────────────────────────────────────────────────────
PRATO_PRONTO_POA_FALLBACK = """PRODUTOS CÁRNEOS ELABORADOS / PRATOS PRONTOS POA (IN 6/2001 + Port. 1485/2025)

CATEGORIAS COBERTAS: Empanados (nugget, steak, bife empanado), Lasanha com carne,
Quibe industrializado, Croquete de carne/frango, Hambúrguer recheado, Produto cárneo
semiprontos para consumo, Frango temperado inteiro/em partes.

DENOMINAÇÃO DE VENDA (Campo 1):
• Empanados: "Empanado de [espécie]" — ex: "Empanado de Frango", "Empanado Bovino"
• NÃO pode usar "filé" sem ser músculo inteiro: "Filé de Frango Empanado" exige músculo
• Nuggets: "Produto Cárneo Reestruturado Empanado de [espécie]" ou denominação consagrada
• Lasanha: "Lasanha com Carne Bovina" — deve identificar espécie
• Quibe: "Quibe Industrializado" ou "Kibe" — ambos aceitos pela Port. 1485/2025
• Croquete: "Croquete de [ingrediente principal]"

LISTA DE INGREDIENTES (Campo 2):
• Em empanados: declarar "Farinha de trigo" — implica CONTÉM GLÚTEN obrigatório
• CMS de aves em empanados: declarar "Carne Mecanicamente Separada de Aves (X%)"
• Proteína de soja em empanados: declarar % entre parênteses — limite máx. 4% (IN 6/2001)
• Ingredientes compostos (massa, molho): listar subingredientes entre parênteses
• Corante caramelo: declarar "Corante Caramelo (INS 150)" com classe e INS

TABELA NUTRICIONAL — PORÇÕES PADRÃO (IN 75/2020):
• Empanados (nugget, steak): porção = 100g (4 unidades médias)
• Hambúrguer recheado: porção = 80g
• Lasanha congelada: porção = 300g (1 porção individual)
• Quibe: porção = 100g (2 unidades)
• Croquete: porção = 80g (2 unidades)

FAIXAS TACO — PLAUSIBILIDADE NUTRICIONAL POR 100g:
EMPANADO DE FRANGO (industrializado):
• kcal: TACO 230-280 | Proteínas: mín 10g (IN 6/2001), típico 14-18g
• Gorduras totais: típico 12-18g | Carboidratos: típico 14-22g (farinha de trigo)
• Sódio: típico 500-900mg | Fibra: típico 0,5-2g (farinha)

EMPANADO BOVINO:
• kcal: TACO 220-270 | Proteínas: mín 10g, típico 13-17g
• Gorduras totais: típico 12-17g | Carboidratos: típico 12-20g
• Sódio: típico 500-900mg

NUGGET DE FRANGO:
• kcal: TACO 240-300 | Proteínas: mín 10g, típico 13-17g
• Gorduras totais: típico 14-20g | Carboidratos: típico 12-20g
• Sódio: típico 500-1000mg

LASANHA COM CARNE BOVINA (congelada):
• kcal: TACO 100-160 | Proteínas: típico 6-10g
• Gorduras totais: típico 4-9g | Carboidratos: típico 12-20g
• Sódio: típico 400-700mg

QUIBE / KIBE INDUSTRIALIZADO:
• kcal: TACO 180-250 | Proteínas: mín 12g, típico 12-17g
• Gorduras totais: típico 10-18g | Carboidratos: típico 8-14g (trigo integral)
• Sódio: típico 350-700mg

ALERTA ESPECIAL: carboidratos altos (>8g/100g) em produtos cárneos elaborados
são normais e esperados — indicam adição de farinha, amido ou proteína vegetal."""

# ──────────────────────────────────────────────────────────────────────────────
# PRODUTO DUAL MAPA + ANVISA — ALIMENTOS FUNCIONAIS DE ORIGEM ANIMAL
# ──────────────────────────────────────────────────────────────────────────────
DUAL_MAPA_ANVISA_FALLBACK = """ALIMENTOS FUNCIONAIS / ENRIQUECIDOS DE ORIGEM ANIMAL
BASE LEGAL: RDC 267/2003 (alegações funcionais) + RDC 18/1999 (bioativos) + IN MAPA específica

PRODUTOS TÍPICOS COM REGISTRO DUAL:
• Leite enriquecido com cálcio, ferro, vitaminas A/D, ômega-3 (UHT ou pasteurizado)
• Ovo enriquecido com ômega-3, vitamina E ou selênio
• Iogurte com probióticos (Lactobacillus, Bifidobacterium)
• Queijo com reduzido teor de gordura + fibras adicionadas
• Leite fermentado funcional (tipo Yakult/Activia)

CAMPO 1 — DENOMINAÇÃO:
• Produto MAPA: "Leite UHT Integral" / "Iogurte Natural" (denominação técnica obrigatória)
• Claim funcional ANVISA: pode adicionar na embalagem mas NÃO na denominação técnica
  ✅ CORRETO: "Leite UHT Integral — Fonte de Cálcio" (claim separado da denominação)
  ❌ ERRADO: "Leite Cálcio Integral" (modifica a denominação técnica)
• Probióticos: "Iogurte com [microrganismo]" aceito se souche declarado e viable count comprovado

CAMPO 2 — LISTA DE INGREDIENTES:
• Nutrientes adicionados: declarar em ordem decrescente normalmente
• Vitaminas/minerais adicionados: declarar nome completo + função
  Ex: "Vitamina D (colecalciferol) — antioxidante" ou apenas nome
• Probióticos: declarar o nome completo do microrganismo
  Ex: "Fermentos lácteos (Lactobacillus acidophilus, Bifidobacterium lactis)"
• Ômega-3: declarar fonte — "Óleo de peixe (fonte de ômega-3)"

CAMPO 9 — TABELA NUTRICIONAL:
• Nutrientes objeto de alegação DEVEM ser declarados (obrigatório — RDC 267/2003)
• Vitaminas e minerais: declarar em mg ou µg + % VD
• Probióticos UFC: podem ser declarados em nota separada (não na tabela principal)
• Ômega-3: declarar EPA + DHA separadamente se for o claim

CAMPO 10 — LUPA FRONTAL:
• Atenção: produto enriquecido com sódio pode acionar lupa "ALTO EM SÓDIO" mesmo sendo
  produto "saudável" — verificar se sódio ≥600mg/100g mesmo com claims positivos

ALEGAÇÕES PROIBIDAS (RDC 267/2003 Art. 3):
• "Cura", "trata", "previne" doenças específicas — PROIBIDO
• "Fortalece o sistema imune" sem substantiação científica — PROIBIDO
• "Reduz o risco de doenças cardíacas" só se ômega-3 EPA+DHA ≥600mg/porção
• Probióticos: só "contribui para o equilíbrio da flora intestinal" (alegação aprovada pela ANVISA)
• Cálcio: "contribui para manutenção de ossos e dentes" (aprovada)"""

# ──────────────────────────────────────────────────────────────────────────────
# NORMAS SIE — 21 ESTADOS RESTANTES
# ──────────────────────────────────────────────────────────────────────────────
SIE_OUTROS_ESTADOS_FALLBACK = """NORMAS SIE — ESTADOS SEM ADIÇÕES RELEVANTES AO FEDERAL
(AC, AL, AM, AP, CE, DF, ES, GO, MA, MS, MT, PA, PB, PE, PI, RJ, RN, RO, RR, SE, TO)

REGRA GERAL: Todos esses estados adotam as normas federais (MAPA + ANVISA) integralmente
para conteúdo de rótulo. As diferenças são exclusivamente procedimentais (registro, formulários).

SIGLAS E ÓRGÃOS SIE POR ESTADO:
• AC (Acre): IDAF-AC — "SIE AC Nº XXXX"
• AL (Alagoas): ADEAL — "SIE AL Nº XXXX"
• AM (Amazonas): ADAF-AM — "SIE AM Nº XXXX"
• AP (Amapá): RURAP-AP — "SIE AP Nº XXXX"
• BA (Bahia): ADAB — "SIE BA Nº XXXX" | Ativo em pescado e mel
• CE (Ceará): ADAGRI-CE — "SIE CE Nº XXXX" | Ativo em pescado
• DF (Distrito Federal): SEAGRI-DF — "SIE DF Nº XXXX"
• ES (Espírito Santo): IDAF-ES — "SIE ES Nº XXXX"
• GO (Goiás): AGRODEFESA — "SIE GO Nº XXXX" | Ativo em bovinos
• MA (Maranhão): AGED-MA — "SIE MA Nº XXXX"
• MS (Mato Grosso do Sul): IAGRO-MS — "SIE MS Nº XXXX" | Ativo em bovinos
• MT (Mato Grosso): INDEA-MT — "SIE MT Nº XXXX" | Ativo em bovinos
• PA (Pará): ADEPARÁ — "SIE PA Nº XXXX" | Ativo em pescado
• PB (Paraíba): AESA-PB — "SIE PB Nº XXXX"
• PE (Pernambuco): ADAGRO-PE — "SIE PE Nº XXXX"
• PI (Piauí): ADAPI-PI — "SIE PI Nº XXXX"
• RJ (Rio de Janeiro): PESAGRO-RJ / SEAPEC — "SIE RJ Nº XXXX" | Ativo em laticínios e pescado
• RN (Rio Grande do Norte): IDEMA-RN / EMPARN — "SIE RN Nº XXXX"
• RO (Rondônia): IDARON — "SIE RO Nº XXXX"
• RR (Roraima): ADERR — "SIE RR Nº XXXX"
• SE (Sergipe): EMDAGRO — "SIE SE Nº XXXX"
• TO (Tocantins): ADAPEC-TO — "SIE TO Nº XXXX"

NOTAS ESPECÍFICAS:
• BA, PA, CE: produtos de pescado podem ter exigência de declarar "pescado artesanal" ou
  "pesca artesanal" se oriundo de colônias de pesca registradas — verificar se aplicável
• GO, MS, MT: bovinos — sem adições ao federal para conteúdo de rótulo
• RJ: laticínios artesanais fluminenses — mesmas regras de queijo artesanal (leite cru declarado)

CAMPO 8 — CARIMBO SIE:
• Formato oval obrigatório — IGUAL ao federal
• Deve conter: sigla do estado + "SIE" + número OR sigla do órgão estadual + número
  Ex válidos: "SIE RJ 1234", "IAGRO MS 456", "ADAB BA 789"
• Produtos fabricados nesses estados que circulam APENAS no estado de origem"""

# ── Fallbacks individuais por estado — cobertura específica ──────────────────

SIE_BA_FALLBACK = """SIE BAHIA — ADAB (Agência Estadual de Defesa Agropecuária da Bahia)
adab.ba.gov.br | SIE BA Nº XXXX

CAMPO 8 — CARIMBO SIE BA:
• Formato: oval com "SIE BA" + número do estabelecimento
• ADAB exige registro prévio do produto (IT nº 03 — procedimentos de registro de POA)
• Para comércio interestadual: adesão ao SISBI-POA via IT nº 01 da ADAB

NORMAS ADICIONAIS ADAB:
• IS DIPA nº 03/2019 — programas de autocontrole obrigatórios (BPF, PPHO, APPCC)
• Portaria ADAB nº 106/2023 — relatórios estatísticos mensais obrigatórios
  (matérias-primas recebidas com nomenclatura IGUAL à do rótulo)
• Portaria ADAB nº 672/2024 — cadastro no e-Sisbi para SISBI-POA
• Decreto Estadual 17.983/2017 — processos tramitados via SEI Bahia

FOCO BAHIA:
• Pescado artesanal: colônias de pesca registradas devem declarar "pescado artesanal"
• Mel: ADAB ativa — verificar se mel local exige declaração de apiário registrado
• Queijos artesanais: lei estadual permite queijo artesanal com leite cru (mín. 60 dias)

NORMAS: Portaria ADAB 106/2023 | IS DIPA 03/2019 | IT 01-04 ADAB | RIISPOA federal"""

SIE_CE_FALLBACK = """SIE CEARÁ — ADAGRI (Agência de Defesa Agropecuária do Estado do Ceará)
adagri.ce.gov.br | SIE CE Nº XXXX

CAMPO 8 — CARIMBO SIE CE:
• Formato: oval com "SIE CE" + número do estabelecimento ADAGRI
• ADAGRI tem sistema próprio (SIDA — Sistema Integrado de Defesa Agropecuária) para registro

NORMAS ESPECÍFICAS ADAGRI:
• Regulamento Técnico ADAGRI — rotulagem de POA embalado (portaria específica vigente)
  Exige: denominação, ingredientes, fabricante, CNPJ, endereço, lote, validade, conservação,
  conteúdo líquido, carimbo SIE CE — alinhado ao MAPA mas com formato estadual
• Portaria autocontrole — programas BPF/PPHO/APPCC para estabelecimentos SIE CE
• LACEN-CE credenciado para análises laboratoriais do SIE (Portaria ADAGRI)
• Adesão ao SISBI-POA: procedimentos via e-Sisbi (Portaria MAPA 672/2024)

FOCO CEARÁ:
• Pescado artesanal: forte produção costeira — verificar se "pescado artesanal" ou
  "camarão artesanal" exige declaração de colônia registrada no rótulo
• Carnes de caprinos/ovinos: CE é grande produtor — verificar espécie declarada
• Queijo coalho CE: produto típico — exige denominação "Queijo Coalho" com registro SIE CE

NORMAS: Regulamento Técnico ADAGRI rotulagem | Portaria autocontrole ADAGRI | RIISPOA federal"""

SIE_GO_FALLBACK = """SIE GOIÁS — AGRODEFESA (Agência Goiana de Defesa Agropecuária)
agrodefesa.go.gov.br | SIE GO Nº XXXX

CAMPO 8 — CARIMBO SIE GO:
• Formato: oval com "SIE GO" + número do estabelecimento AGRODEFESA
• Lei Estadual 11.904/1993 — regulamento base de inspeção POA em Goiás

NORMAS ESPECÍFICAS GOIÁS:
• Decreto Estadual 4.019/1993 — Regulamento da Inspeção Sanitária e Industrial de POA em GO
• Decreto 10.404/2024 — ATUALIZAÇÃO das regras de inspeção e rotulagem SIE GO (recente)
• IN 9/2023 AGRODEFESA — Selo Arte e Queijo Artesanal para produtos SIE GO
• Decreto 12.126/2024 (federal) — programas de autocontrole com fiscalização baseada em risco

FOCO GOIÁS:
• Bovinos: maior produtor do Centro-Oeste — sem adições ao federal para conteúdo de rótulo
• Suínos: verificar SE produto tem carimbo SIE GO ou SIF (frigoríficos grandes são SIF)
• Frango: idem — grandes integradoras são SIF; pequenos produtores podem ser SIE GO
• Queijo artesanal: IN 9/2023 AGRODEFESA permite Selo Arte para queijos SIE GO

NORMAS: Lei 11.904/1993-GO | Dec. 4.019/1993-GO | Dec. 10.404/2024-GO | IN 9/2023-AGRODEFESA"""

SIE_AM_FALLBACK = """SIE AMAZONAS — ADAF (Agência de Defesa Agropecuária e Florestal do Amazonas)
adaf.am.gov.br | SIE AM Nº XXXX

CAMPO 8 — CARIMBO SIE AM:
• Carimbo HEXAGONAL — particularidade do Amazonas (diferente do oval federal)
• Portaria ADAF-AM (jul/2023) — novo carimbo hexagonal com proporções específicas por
  tamanho de embalagem e natureza do produto
• Empresas tiveram 12 meses para adequação — carimbo hexagonal vigente desde 2024
• ❌ Carimbo oval SIE genérico = NÃO CONFORME para produtos AM após prazo

NORMAS ESPECÍFICAS ADAF:
• Portaria rotulagem ADAF-AM — além das informações federais obrigatórias, exige
  o carimbo hexagonal específico emitido pela ADAF
• Mudança de nomenclatura dos estabelecimentos: categorias próprias AM (fabricante,
  beneficiador, distribuidor — verificar denominação correta para o tipo de estabelecimento)
• Novo layout de rótulo deve ser protocolado e aprovado pela ADAF antes de usar

FOCO AMAZONAS:
• Pescado de rio: pirarucu, tambaqui, jaraqui — produtos típicos com alta demanda
  Verificar se exige "pescado de pesca artesanal" ou "pescado de aquicultura" declarado
• Proteínas silvestres: caça autorizada — exige registro específico IBAMA + ADAF

⚠️ ATENÇÃO ESPECIAL: carimbo hexagonal é exclusividade do AM — validar se está correto.

NORMAS: Portaria ADAF-AM rotulagem/2023 | RIISPOA federal | MAPA normas POA"""

SIE_PA_FALLBACK = """SIE PARÁ — ADEPARA (Agência de Defesa Agropecuária do Pará)
adepara.pa.gov.br | SIE PA Nº XXXX

CAMPO 8 — CARIMBO SIE PA:
• Formato: oval com "SIE PA" + número do estabelecimento ADEPARA

NORMAS ESPECÍFICAS ADEPARA:
• Pescado artesanal: Pará é o maior produtor de pescado do Brasil — verificar
  se exige declaração "pescado artesanal" ou "pescado de aquicultura" no rótulo
• Açaí: produto não-POA mas com normas PA específicas para processamento e rotulagem
• Queijo Marajó: produto típico paraense — manteiga de garrafa e queijo de búfala
  Exige denominação específica "Queijo Marajó" e registro SIE PA

FOCO PARÁ:
• Búfalo: Pará é maior produtor de búfalo do Brasil — denominar "carne bubalina"
  ou "produto de origem bubalina" (não bovina) — espécie obrigatória no rótulo
• Pescado Amazônico: pirarucu, tucunaré, jaraqui — verificar se há norma ADEPARA
  específica para denominação de espécies amazônicas

NORMAS: ADEPARA normativas vigentes | RIISPOA federal | MAPA normas POA"""

SIE_PE_FALLBACK = """SIE PERNAMBUCO — ADAGRO (Agência de Defesa e Fiscalização Agropecuária de PE)
adagro.pe.gov.br | SIE PE Nº XXXX

CAMPO 8 — CARIMBO SIE PE:
• Formato: oval com "SIE PE" + número do estabelecimento ADAGRO-PE

NORMAS ESPECÍFICAS ADAGRO:
• Pernambuco adota normas federais integralmente para conteúdo de rótulo
• Registro de estabelecimentos via ADAGRO com formulários próprios
• Carne de caprinos/ovinos: PE é grande produtor Nordeste — espécie obrigatória no rótulo

FOCO PERNAMBUCO:
• Carne-de-sol: produto típico nordestino — denominação "carne-de-sol" aceita se
  processo específico (salga e secagem ao sol) — verificar IN específica
• Queijo coalho PE: produto típico — verificar se ADAGRO tem norma específica
  de rotulagem diferente do coalho federal
• Buchada de bode: produto típico — deve declarar espécie (caprino) e ingredientes

NORMAS: ADAGRO-PE normativas vigentes | RIISPOA federal | MAPA normas POA"""

SIE_MA_FALLBACK = """SIE MARANHÃO — AGED (Agência Estadual de Defesa Agropecuária do Maranhão)
aged.ma.gov.br | SIE MA Nº XXXX

CAMPO 8 — CARIMBO SIE MA:
• Formato: oval com "SIE MA" + número do estabelecimento AGED

FOCO MARANHÃO:
• Carne bovina: MA é produtor relevante — sem adições ao federal para rótulo
• Babaçu: óleo de babaçu é produto típico MA — verificar normas AGED específicas
  para óleos vegetais e subprodutos de babaçu (não-POA mas com relevância estadual)
• Pescado: acesso ao mar e rios — verificar declaração de origem artesanal vs industrial

NORMAS: AGED-MA normativas vigentes | RIISPOA federal"""

SIE_MS_FALLBACK = """SIE MATO GROSSO DO SUL — IAGRO (Agência Estadual de Defesa Sanitária Animal e Vegetal)
iagro.ms.gov.br | SIE MS Nº XXXX

CAMPO 8 — CARIMBO SIE MS:
• Formato: oval com "SIE MS" + número do estabelecimento IAGRO

FOCO MATO GROSSO DO SUL:
• Bovinos: MS é o estado com maior rebanho por hectare — sem adições ao federal
• Pantanal: produtos do Pantanal podem ter denominações geográficas específicas
• Peixe pantaneiro: pacu, pintado, dourado — verificar espécie declarada no rótulo
• Queijo artesanal: legislação MS permite queijo artesanal com leite cru (mín. 60 dias)

NORMAS: IAGRO-MS normativas vigentes | RIISPOA federal"""

SIE_MT_FALLBACK = """SIE MATO GROSSO — INDEA (Instituto de Defesa Agropecuária do Mato Grosso)
indea.mt.gov.br | SIE MT Nº XXXX

CAMPO 8 — CARIMBO SIE MT:
• Formato: oval com "SIE MT" + número do estabelecimento INDEA

FOCO MATO GROSSO:
• Bovinos: MT é o maior rebanho bovino do Brasil — sem adições ao federal para rótulo
• Frango: MT é grande produtor de aves — verificar se SIE MT ou SIF (grandes unidades são SIF)
• Soja: não-POA mas MT é maior produtor — oleaginosas e derivados via MAPA

NORMAS: INDEA-MT normativas vigentes | RIISPOA federal"""

SIE_ES_FALLBACK = """SIE ESPÍRITO SANTO — IDAF (Instituto de Defesa Agropecuária e Florestal do ES)
idaf.es.gov.br | SIE ES Nº XXXX

CAMPO 8 — CARIMBO SIE ES:
• Formato: oval com "SIE ES" + número do estabelecimento IDAF-ES

FOCO ESPÍRITO SANTO:
• Pescado: ES tem importante produção pesqueira artesanal — verificar declaração
• Chouriço capixaba: produto típico ES — denominação específica, ingredientes obrigatórios
• Laticínios: queijos artesanais capixabas — legislação ES permite artesanal com leite cru

NORMAS: IDAF-ES normativas vigentes | RIISPOA federal"""

SIE_DF_FALLBACK = """SIE DISTRITO FEDERAL — SEAGRI-DF (Secretaria de Agricultura, Abastecimento e DR)
seagri.df.gov.br | SIE DF Nº XXXX

CAMPO 8 — CARIMBO SIE DF:
• Formato: oval com "SIE DF" + número do estabelecimento SEAGRI-DF
• Produtos SIE DF circulam apenas no DF (não interestadual sem SISBI-POA)

FOCO DISTRITO FEDERAL:
• Mercado consumidor urbano predominante — poucos produtores SIE DF
• Grande parte dos produtos consumidos no DF vêm de SIF de GO, MG, SP
• Produtos SIE DF: pequenos produtores rurais do DF, feiras, pequenas indústrias

NORMAS: SEAGRI-DF normativas vigentes | RIISPOA federal"""

SIE_PI_FALLBACK = """SIE PIAUÍ — ADAPI (Agência de Defesa Agropecuária do Piauí)
adapi.pi.gov.br | SIE PI Nº XXXX

FOCO PIAUÍ:
• Carne de caprinos/ovinos: PI é produtor relevante no Nordeste — espécie obrigatória
• Mel de abelha nativa (meliponicultura): PI tem produção relevante de mel de abelha
  sem ferrão (uruçu, jandaíra) — verificar denominação específica: "Mel de abelha nativa"
  não pode ser chamado apenas de "mel" (abelha Apis) — denominação correta obrigatória
• Carne-de-sol: produto típico — verificar denominação aceita pela ADAPI

NORMAS: ADAPI-PI normativas vigentes | RIISPOA federal"""

SIE_RN_FALLBACK = """SIE RIO GRANDE DO NORTE — IDIARN (Instituto de Desenvolvimento e Inspeção Agropecuária do RN)
idiarn.rn.gov.br | SIE RN Nº XXXX

FOCO RIO GRANDE DO NORTE:
• Camarão: RN é maior produtor de camarão cultivado do Brasil
  "Camarão branco cultivado" ou "Litopenaeus vannamei" — espécie obrigatória no rótulo
  Verificar se IDIARN tem norma específica de rotulagem para camarão de aquicultura
• Carne de caprinos/ovinos: RN é produtor relevante — espécie obrigatória
• Mel: verificar mel de carnaúba e mel de abelha nativa — denominação específica

NORMAS: IDIARN-RN normativas vigentes | RIISPOA federal"""

SIE_AL_FALLBACK = """SIE ALAGOAS — ADEAL (Agência de Defesa e Vigilância Agropecuária de AL)
adeal.al.gov.br | SIE AL Nº XXXX

FOCO ALAGOAS:
• Pescado: litoral extenso — verificar declaração "pescado artesanal" vs industrial
• Carne: sem adições relevantes ao federal para conteúdo de rótulo
• Queijo coalho AL: produto típico nordestino — verificar denominação específica

NORMAS: ADEAL-AL normativas vigentes | RIISPOA federal"""

SIE_SE_FALLBACK = """SIE SERGIPE — EMDAGRO (Empresa de Desenvolvimento Agropecuário de Sergipe)
emdagro.se.gov.br | SIE SE Nº XXXX

FOCO SERGIPE:
• Queijo coalho SE: menor estado do Brasil mas produtor de queijo coalho
• Pescado artesanal: litoral e rio São Francisco — verificar declaração de origem
• Carne de caprinos/ovinos: produção relevante — espécie obrigatória

NORMAS: EMDAGRO-SE normativas vigentes | RIISPOA federal"""

SIE_PB_FALLBACK = """SIE PARAÍBA — DIPOVA-PB / Secretaria de Agricultura
pb.gov.br/agropecuaria | SIE PB Nº XXXX

FOCO PARAÍBA:
• Queijo coalho PB: produto típico — denominação "queijo coalho" com registro SIE PB
• Carne de caprinos/ovinos: PB é relevante produtor Nordeste — espécie obrigatória
• Carne-de-sol: produto típico — verificar denominação aceita pelo SIE PB
• Mel: verificar mel de abelha nativa (sem ferrão) — denominação específica obrigatória

NORMAS: DIPOVA-PB normativas vigentes | RIISPOA federal"""

SIE_AC_FALLBACK = """SIE ACRE — IDAF (Instituto de Defesa Agropecuária e Florestal do Acre)
idaf.ac.gov.br | SIE AC Nº XXXX

FOCO ACRE:
• Bovinos: AC tem pecuária relevante na Amazônia Ocidental — sem adições ao federal
• Pescado Amazônico: peixes de rio — verificar espécie declarada (pirarucu, tambaqui)
• Produtos florestais: castanha-do-Pará (não-POA) — MAPA normas específicas

NORMAS: IDAF-AC normativas vigentes | RIISPOA federal"""

SIE_AP_FALLBACK = """SIE AMAPÁ — RURAP (Agência Rural do Amapá)
rurap.ap.gov.br | SIE AP Nº XXXX

FOCO AMAPÁ:
• Pescado Amazônico: AP tem acesso ao estuário amazônico — espécie obrigatória
• Búfalo: AP tem criação relevante de búfalos — denominar "bubalino" não "bovino"
• Açaí: maior produtor per capita — produto não-POA, verificar normas RURAP específicas

NORMAS: RURAP-AP normativas vigentes | RIISPOA federal"""

SIE_RO_FALLBACK = """SIE RONDÔNIA — IDARON (Agência de Defesa Sanitária Agrosilvopastoril de RO)
idaron.ro.gov.br | SIE RO Nº XXXX

FOCO RONDÔNIA:
• Bovinos e leite: RO tem forte pecuária leiteira amazônica — sem adições ao federal
• Leite: verificar se há normas IDARON específicas para transporte e temperatura de leite cru
• Pescado Amazônico: peixes de rio — verificar espécie declarada

NORMAS: IDARON-RO normativas vigentes | RIISPOA federal"""

SIE_RR_FALLBACK = """SIE RORAIMA — ADERR (Agência de Defesa Agropecuária de Roraima)
aderr.rr.gov.br | SIE RR Nº XXXX

FOCO RORAIMA:
• Bovinos: pecuária relevante no extremo Norte — sem adições ao federal para rótulo
• Pescado Amazônico: peixes de rio — espécie obrigatória declarada
• Pequeno mercado consumidor — maioria dos produtos SIF vêm de outros estados

NORMAS: ADERR-RR normativas vigentes | RIISPOA federal"""

SIE_TO_FALLBACK = """SIE TOCANTINS — ADAPEC (Agência de Defesa Agropecuária do Estado do Tocantins)
adapec.to.gov.br | SIE TO Nº XXXX

FOCO TOCANTINS:
• Bovinos: TO tem pecuária relevante no Centro-Norte — sem adições ao federal
• Pescado: rios Tocantins e Araguaia — espécie obrigatória (pirarucu, tucunaré)
• Mel: verificar normas ADAPEC para mel do Cerrado tocantinense

NORMAS: ADAPEC-TO normativas vigentes | RIISPOA federal"""




RIISPOA_RJ_FALLBACK = """
━━━ RIISPOA-RJ — Decreto Estadual 49.643/2025 (SIE/RJ) ━━━━━━━━━━━━━━━━━━━━━

ÂMBITO: Aplica-se a todos os estabelecimentos POA com comércio INTRAESTADUAL no RJ.
Órgão: SIE/RJ — COOIPOA/SUPDA/SEAPPA-RJ
Publicação: DOERJ 26/05/2025 | Vigência: imediata

── IDENTIFICAÇÃO DO ÓRGÃO NO RÓTULO (Art. 17 §XLII + Art. 8 V) ──────────────
• Carimbo do SIE/RJ deve conter "SIE/RJ" + número do estabelecimento registrado na COOIPOA
• Formato oval obrigatório (mesmo padrão federal — RIISPOA-RJ adota o padrão do Decreto 9.013/2017)
• Produto do SIE/RJ só pode circular no Estado do RJ, salvo se aderido ao SISBI-POA (Art. 25 §1°)
• Se produto tiver carimbo SIE/RJ e declarar venda nacional → ❌ NÃO CONFORME, salvo SISBI ativo
• Verificar: após transferência de empresa, rótulos da firma anterior invalidados — prazo máx. 6 meses (Art. 37)

── SIGLA NO RÓTULO ────────────────────────────────────────────────────────────
• Sigla correta: "SIE/RJ" ou "SIE RJ" + número do registro na COOIPOA
• Sigla SEAPPA/RJ (secretaria) ≠ carimbo de inspeção — não aceitar SEAPPA como substituto do carimbo
• COOIPOA é a coordenadoria, não a sigla do carimbo

── LEITE E DERIVADOS — EXIGÊNCIAS ESPECÍFICAS RJ (Arts. 240-270) ─────────────
• Temperatura de conservação — leite pasteurizado até o consumidor: máx. 7°C (Art. 265 IV)
  → Campo 7: "Manter refrigerado a no máximo 7°C" ou "Conservar abaixo de 7°C"
• Temperatura no estabelecimento (pós-pasteurização): máx. 5°C (Art. 265 III)
• Leite UHT/UAT: temperatura ambiente — sem refrigeração antes da abertura
• Composição mínima leite in natura (Art. 255):
  - Gordura: mín. 3,0g/100g
  - Proteína total: mín. 2,9g/100g
  - Lactose anidra: mín. 4,3g/100g
  - Sólidos não gordurosos: mín. 8,4g/100g
  - Sólidos totais: mín. 11,4g/100g
• Leite misto de espécies (Art. 242 §2°): obrigatório declarar % de cada espécie no rótulo
  → Ex: "Leite de Vaca (70%) e Leite de Cabra (30%)" — denominação e % obrigatórios
• PROIBIDO: reutilização de sal, leite de fêmeas em tratamento, desnate parcial na propriedade

── PESCADO — EXIGÊNCIAS ESPECÍFICAS RJ (Arts. 207-223) ──────────────────────
• Pescado fresco: verificar temperatura Campo 7 — obrigatório 0°C a 4°C (gelo/refrigerado)
• Temperatura para consumo cru (sushi/sashimi): congelamento prévio a -20°C por 7+ dias
  ou -35°C por 15h para eliminação de Anisakidae (Art. 222 §1°)
• Nome científico no rótulo: obrigatório para pescado (tilápia, salmão, camarão etc.)
• Moluscos bivalves: vivos na embalagem; estação depuradora RJ deve estar registrada na COOIPOA
• Camarão: se congelado, verificar % de glaze declarado (cobertura de gelo)

── OVOS — EXIGÊNCIAS ESPECÍFICAS RJ (Arts. 224-239) ─────────────────────────
• Categoria "A" obrigatória para consumo direto — verificar declaração no rótulo (Art. 231)
• Categoria "B": apenas para industrialização — se rótulo declarar "consumo direto" com categoria B → ❌
• Câmara de ar máx. 6mm para ovos categoria A (Art. 232 II)
• PROIBIDO misturar ovos frescos com ovos conservados na mesma embalagem (Art. 238 I)
• PROIBIDO misturar ovos de espécies diferentes na mesma embalagem (Art. 238 II)
• Granja avícola RJ: deve estar registrada na COOIPOA E no serviço oficial de saúde animal (Art. 229)

── PRODUTOS DE ABELHAS — EXIGÊNCIAS ESPECÍFICAS RJ (Arts. 271-276) ──────────
• Mel de abelhas sem ferrão (meliponíneos): deve ser de meliponários autorizados pelo órgão ambiental RJ (Art. 276)
• Se produto for "mel de abelha sem ferrão" (jataí, mandaçaia, etc.): verificar denominação específica
• Descristalização/pasteurização: binômio tempo/temperatura deve constar no processo (não no rótulo)
• Rastreabilidade: estabelecimento deve ter cadastro de produtores rurais fornecedores (Art. 274)

── ADITIVOS E INGREDIENTES (Arts. 277-281) ───────────────────────────────────
• Aditivos: somente os autorizados pelo MAPA E ANVISA podem ser usados — lista dupla de verificação
• Sal: deve ser isento de substâncias estranhas (Art. 279)
• PROIBIDO: recuperação de salmoura para produtos comestíveis sem aprovação (Art. 280)
• RTIQs: devem ser os federais do MAPA ou, em casos específicos, normas complementares estaduais da COOIPOA (Art. 281)

── CARIMBO E RASTREABILIDADE (Arts. 29, 37, 38) ─────────────────────────────
• Cancelamento de registro → carimbo deve ser recolhido + rótulos inutilizados (Art. 38 §2°)
• Transferência de empresa: rótulos da firma anterior podem ser usados por máx. 6 meses com acordo (Art. 37)
• Número de registro é único por unidade fabril no território estadual (Art. 29 parágrafo único)

── CHECKLIST SIE/RJ — itens adicionais vs SIF ───────────────────────────────
□ Carimbo contém "SIE/RJ" + número do estabelecimento na COOIPOA
□ Jurisdição de venda declarada está dentro do RJ (ou SISBI ativo para nacional)
□ Temperatura do Campo 7 está dentro dos limites do RIISPOA-RJ (leite ≤7°C, pescado 0-4°C)
□ Leite misto de espécies declara % de cada espécie
□ Ovos com categoria declarada (A para consumo direto)
□ Mel de abelha sem ferrão menciona espécie se aplicável
□ Nome científico presente para pescado

Fonte: Decreto Estadual RJ 49.643/2025 | DOERJ 26/05/2025 | COOIPOA/SEAPPA-RJ
"""

SIE_ESTADO_MAP = {
    # ── Estados com normas específicas detalhadas ────────────────────────────
    "RS": CISPOA_RS_FALLBACK,   "CISPOA": CISPOA_RS_FALLBACK,
    "SP": SISP_SP_FALLBACK,     "SISP": SISP_SP_FALLBACK,   "CDA": SISP_SP_FALLBACK,
    "MG": IMA_MG_FALLBACK,      "IMA": IMA_MG_FALLBACK,
    "SC": CIDASC_SC_FALLBACK,   "CIDASC": CIDASC_SC_FALLBACK,
    "PR": ADAPAR_PR_FALLBACK,   "ADAPAR": ADAPAR_PR_FALLBACK,
    "RJ": RIISPOA_RJ_FALLBACK,  "COOIPOA": RIISPOA_RJ_FALLBACK, "SEAPPA": RIISPOA_RJ_FALLBACK,
    # ── 21 estados com fallbacks individuais detalhados ──────────────────────
    "BA": SIE_BA_FALLBACK,     "ADAB": SIE_BA_FALLBACK,
    "CE": SIE_CE_FALLBACK,     "ADAGRI": SIE_CE_FALLBACK,
    "GO": SIE_GO_FALLBACK,     "AGRODEFESA": SIE_GO_FALLBACK,
    "AM": SIE_AM_FALLBACK,     "ADAF": SIE_AM_FALLBACK,
    "PA": SIE_PA_FALLBACK,     "ADEPARA": SIE_PA_FALLBACK,  "ADEPAR": SIE_PA_FALLBACK,
    "PE": SIE_PE_FALLBACK,     "ADAGRO": SIE_PE_FALLBACK,
    "MA": SIE_MA_FALLBACK,     "AGED": SIE_MA_FALLBACK,
    "MS": SIE_MS_FALLBACK,     "IAGRO": SIE_MS_FALLBACK,
    "MT": SIE_MT_FALLBACK,     "INDEA": SIE_MT_FALLBACK,
    "ES": SIE_ES_FALLBACK,     "IDAFES": SIE_ES_FALLBACK,
    "DF": SIE_DF_FALLBACK,     "SEAGRI": SIE_DF_FALLBACK,
    "PI": SIE_PI_FALLBACK,     "ADAPI": SIE_PI_FALLBACK,
    "RN": SIE_RN_FALLBACK,     "IDIARN": SIE_RN_FALLBACK,
    "AL": SIE_AL_FALLBACK,     "ADEAL": SIE_AL_FALLBACK,
    "SE": SIE_SE_FALLBACK,     "EMDAGRO": SIE_SE_FALLBACK,
    "PB": SIE_PB_FALLBACK,
    "AC": SIE_AC_FALLBACK,     "IDAFAC": SIE_AC_FALLBACK,
    "AP": SIE_AP_FALLBACK,     "RURAP": SIE_AP_FALLBACK,
    "RO": SIE_RO_FALLBACK,     "IDARON": SIE_RO_FALLBACK,
    "RR": SIE_RR_FALLBACK,     "ADERR": SIE_RR_FALLBACK,
    "TO": SIE_TO_FALLBACK,     "ADAPEC": SIE_TO_FALLBACK,
}


# ───────────────────────────────────────────────────────────────────────────────
# SISBI-POA — SISTEMA BRASILEIRO DE INSPEÇÃO DE PRODUTOS DE ORIGEM ANIMAL
# Equivalência entre SIE/SIM estaduais/municipais para comércio nacional
# ───────────────────────────────────────────────────────────────────────────────
SISBI_POA_FALLBACK = """SISBI-POA — EQUIVALÊNCIA PARA COMÉRCIO NACIONAL (Decreto 9.013/2017 Art. 25)

REGRA PRINCIPAL:
• Produto com APENAS carimbo SIE ou SIM pode circular SOMENTE no estado/município de origem
• Produto com carimbo SIE/SIM + adesão ao SISBI-POA pode circular em TODO o território nacional
• Produto com carimbo SIF pode sempre circular nacionalmente (não precisa de SISBI)

COMO IDENTIFICAR SISBI-POA NO RÓTULO:
• Presença do logo/selo SISBI-POA ao lado do carimbo SIE/SIM
• Indicação textual: "Produto habilitado ao SISBI-POA" ou "SISBI-POA"
• Número de habilitação SISBI no estabelecimento

VERIFICAÇÕES PARA CAMPO 8 (CARIMBO/REGISTRO):
✅ CONFORME se: SIF sozinho | SIE + SISBI-POA + venda nacional | SIM + SISBI-POA + venda nacional
❌ NÃO CONFORME se: SIE sozinho com declaração de venda nacional no rótulo
❌ NÃO CONFORME se: SIM sozinho com declaração de venda para outros estados
⚠️ COM RESSALVAS se: SIE/SIM sem SISBI mas sem declaração explícita de alcance nacional

NORMA: Decreto 9.013/2017 Art. 25 §1° | Port. SDA 46/2019 (habilitação SISBI-POA)
"""

# ───────────────────────────────────────────────────────────────────────────────
# SIM MUNICIPAL — SERVIÇO DE INSPEÇÃO MUNICIPAL
# ───────────────────────────────────────────────────────────────────────────────
SIM_MUNICIPAL_FALLBACK = """SIM — SERVIÇO DE INSPEÇÃO MUNICIPAL (Lei 7.889/1989 + Decreto 9.013/2017 Art. 4°)

DIFERENÇAS DO SIF/SIE:
1. CARIMBO: deve conter "SIM" + número municipal (oval ou quadrado conforme legislação municipal)
   Formato aceito: "SIM Nº XXX" | "SIM [NOME DO MUNICÍPIO] XXX"
   ❌ NÃO confundir com SIF (federal) ou SIE (estadual)

2. ABRANGÊNCIA: produto pode ser comercializado SOMENTE no município de origem
   Exceção: município aderido ao SISBI-POA → comércio nacional permitido

3. LEGISLAÇÃO MUNICIPAL: cada município pode ter exigências adicionais
   Verificar se o carimbo está correto conforme regulamento municipal vigente
   Muitos municípios exigem registro no sistema municipal além do número

4. RTIQ: mesmos RTIQs federais se aplicam (IN 4/2000, IN 20/2000, etc.)
   O SIM não dispensa cumprimento dos RTIQs do MAPA

5. RESPONSABILIDADE TÉCNICA: RT habilitado obrigatório mesmo para SIM

NORMA BASE: Lei 7.889/1989 | Decreto 9.013/2017 Art. 4° II | Regulamento municipal específico
"""

# ───────────────────────────────────────────────────────────────────────────────
# POA ORGÂNICO — PRODUTOS DE ORIGEM ANIMAL COM CERTIFICAÇÃO ORGÂNICA
# ───────────────────────────────────────────────────────────────────────────────
POA_ORGANICO_FALLBACK = """POA ORGÂNICO — PRODUTOS DE ORIGEM ANIMAL ORGÂNICOS
Lei 10.831/2003 + Decreto 6.323/2007 + IN MAPA 17/2014 (produção orgânica animal) + IN MAPA 19/2009 (certificação)

DENOMINAÇÃO (Campo 1):
✅ "Frango Inteiro Orgânico" | "Mel Orgânico" | "Ovos Orgânicos"
❌ "Frango Natural" sem certificação — proibido usar "orgânico" sem certificação
❌ "Produto Ecológico" sem certificação — "ecológico" = sinônimo de orgânico para fins legais

CERTIFICAÇÃO OBRIGATÓRIA (Campo 8):
• Logo/Selo SisOrg obrigatório no rótulo — emblema oficial do sistema orgânico brasileiro
• Número do certificado ou código do organismo certificador
• Organismos credenciados pelo MAPA: IBD, IMO, Ecocert, Imo Control, OIA, etc.
  ✅ "Certificado por [OAC] — MAPA nº XX" ou Selo SisOrg + número
  ❌ Ausência do Selo SisOrg = NÃO CONFORME (fraude se usar "orgânico")

INGREDIENTES (Campo 2):
• 100% orgânico: todos os ingredientes agrícolas devem ser orgânicos
• Mín. 95% orgânicos para usar denominação e símbolo
• 70-94% orgânicos: pode listar ingredientes orgânicos mas não usar o símbolo
• Aditivos: apenas os autorizados pelo Decreto 6.323/2007 Anexo I

TABELA NUTRICIONAL (Campo 9): mesmas regras RDC 429/2020

NORMAS: Lei 10.831/2003 | Decreto 6.323/2007 | IN MAPA 17/2014 | IN MAPA 19/2009
"""

# ═══════════════════════════════════════════════════════════════════════════════
# NP1 — SUCOS, NÉCTARES E BEBIDAS DE FRUTAS
# RDC 173/2006 + IN MAPA 37/2018 + Decreto 6.871/2009
# ═══════════════════════════════════════════════════════════════════════════════
NP1_SUCOS_FALLBACK = """SUCOS, NÉCTARES E BEBIDAS DE FRUTAS — RDC 173/2006 + IN MAPA 37/2018 + Decreto 6.871/2009

DENOMINAÇÕES OBRIGATÓRIAS (Campo 1):
• SUCO INTEGRAL: 100% fruta, sem adição de água, açúcar ou aditivos
  → Denominação: "Suco Integral de [fruta]" — ex: "Suco Integral de Laranja"
  → 'Suco de uva' com <100% suco = FRAUDE — alertar como ❌ NÃO CONFORME
• SUCO RECONSTITUÍDO: suco concentrado + água, sem açúcar
  → Denominação: "Suco de [fruta]" — ex: "Suco de Maçã"
• NÉCTAR: mín. 30-50% suco dependendo da fruta (pode ter açúcar)
  → Laranja/maçã/uva: mín. 50% | Maracujá/acerola: mín. 30%
  → Denominação: "Néctar de [fruta]"
• REFRESCO / BEBIDA DE FRUTA: mín. 10% de suco (pode ter açúcar)
  → Denominação: "Refresco de [fruta]" ou "Bebida de [fruta]"
• POLPA DE FRUTA: produto não diluído, obtido por despolpamento
• ❌ PROIBIDO: usar "suco" sem ser 100% fruta — é fraude e embargável

CAMPO 2 — INGREDIENTES:
• Suco integral: declarar apenas "Suco integral de [fruta]" — sem mais nada
• Néctar: declarar "Suco de [fruta] (X%)", água, açúcar (se houver)
• Refresco: declarar % de suco obrigatoriamente
• Vitamina C / ácido ascórbico: declarar como "Antioxidante: Ácido Ascórbico (INS 300)"
• Conservantes: benzoato e sorbato autorizados apenas em refrescos, não em sucos integrais
• Corantes: PROIBIDOS em sucos integrais — apenas em refrescos/néctares com autorização

CAMPO 9 — TABELA NUTRICIONAL:
• Porção padrão (IN 75/2020): SUCO = 200mL | NÉCTAR = 200mL | REFRESCO = 200mL
• Valores TACO por 100mL:
  SUCO DE LARANJA 100%: kcal 43-50 | Carb 9-12g | Açúcares 8-11g | Sódio 0-5mg
  SUCO DE UVA 100%: kcal 60-72 | Carb 14-18g | Açúcares 14-17g
  SUCO DE MAÇÃ 100%: kcal 45-55 | Carb 10-13g | Açúcares 9-12g
  NÉCTAR DE MARACUJÁ: kcal 35-55 | Carb 8-13g (pode ter açúcar adicionado)
  SUCO DE ACEROLA: kcal 35-45 | Vit C: altíssimo (800-1000mg/100mL)

ALERTAS CRÍTICOS:
• Suco com Sódio >5mg/100mL → suspeita de adulteração ou adição de conservante
• Suco integral com açúcar na lista → ❌ NÃO CONFORME — não pode ser chamado integral
• Néctar sem % de suco declarado → ❌ NÃO CONFORME (IN 37/2018 Art. 6°)
• "Vitamina C" como ingrediente: verificar se está sendo usada como antioxidante ou enriquecimento
• Sucos 100%: PROIBIDO qualquer aditivo alimentar (inclusive conservantes)

NORMAS BASE (atualizado pós-links RT):
• MAPA — Decreto 6.871/2009: define categorias de bebidas (suco, néctar, refresco, polpa, bebida composta)
• MAPA — Lei 8.918/1994: normas gerais de bebidas não alcoólicas
• MAPA — IN 49/2018: complementa padrões de identidade de suco e polpa de fruta
• MAPA — Portaria 123/2021: DQI obrigatória no rótulo frontal; categorias refresco, chá, refrigerante, soda
• ANVISA — RDC 727/2022: rotulagem geral
• ANVISA — RDC 429/2020 + IN 75/2020: rotulagem nutricional
• ATENÇÃO: SUCO é regulado pelo MAPA (Decreto 6.871/2009), não apenas pela ANVISA

DQI — DECLARAÇÃO QUANTITATIVA DE INGREDIENTE (Portaria MAPA 123/2021):
• Obrigatório no painel frontal: informar % de fruta/vegetal no produto
• Suco integral 100%: "100% fruta" ou "100% [nome da fruta]"
• Néctar: "X% de fruta" (ex: "50% de laranja")
• Bebida artificial: "Não contém fruta ou vegetal" ou "0% de fruta"
• ❌ Ausência da DQI no frontal = NÃO CONFORME (Portaria 123/2021)"""

# ═══════════════════════════════════════════════════════════════════════════════
# NP2 — SUPLEMENTOS ALIMENTARES (WHEY, COLÁGENO, TERMOGÊNICOS)
# RDC 243/2018 + RDC 786/2023 + IN 28/2018
# ═══════════════════════════════════════════════════════════════════════════════
NP2_SUPLEMENTOS_FALLBACK = """SUPLEMENTOS ALIMENTARES — RDC 243/2018 + RDC 786/2023 + IN 28/2018

CAMPO 1 — DENOMINAÇÃO:
• Obrigatório: categoria oficial + nome do produto
  → "Suplemento Alimentar de Proteína" (whey)
  → "Suplemento Alimentar de Colágeno Hidrolisado"
  → "Suplemento Alimentar Termogênico"
• NUNCA pode usar denominações de medicamento: "queimador de gordura", "anabolizante"
• Notificação ANVISA obrigatória — verificar número no rótulo
• Frase obrigatória (RDC 243/2018 Art. 8°): "Este produto não é um medicamento"

CAMPO 2 — INGREDIENTES:
• WHEY PROTEIN: declarar tipo exato — "Proteína do Soro de Leite Concentrada/Isolada/Hidrolisada"
  Limite mín. proteína: 10g/porção (RDC 243/2018)
• COLÁGENO: declarar "Colágeno Hidrolisado" — verificar se é bovino, suíno ou marinho
  Alérgeno: colágeno bovino = CONTÉM LEITE? Não. Colágeno de peixe = CONTÉM PEIXE
• TERMOGÊNICOS — ingredientes com limites máximos (RDC 786/2023):
  Cafeína: máx. 420mg/dia (porção máx. 210mg) — declarar mg por porção obrigatório
  Sinefrina: máx. 30mg/dia — declarar mg/porção
  Guaraná: calcular cafeína equivalente
  ❌ PROIBIDOS: efedrina, sibutramina, anfepramona — verificar lista negra ANVISA
• BCAA/Aminoácidos: declarar cada aminoácido com quantidade em mg

CAMPO 9 — TABELA NUTRICIONAL:
• Porção: conforme declarado no produto (geralmente 30-40g para whey)
• WHEY CONCENTRADO típico (por 30g): Proteínas 20-25g | Carb 3-8g | Gorduras 2-5g | Sódio 50-150mg
• WHEY ISOLADO típico (por 30g): Proteínas 25-28g | Carb 0-2g | Gorduras 0-1g
• CREATINA (por 3-5g): Proteínas 0g | Carb 0g | Gorduras 0g — qualquer nutriente ≠ 0 é suspeito
• Vitaminas/minerais: declarar % VD obrigatório quando presentes
• Cafeína: NÃO entra na tabela nutricional mas deve constar nos ingredientes com mg/porção

ALERTAS CRÍTICOS:
• Sem número de notificação ANVISA → ❌ NÃO CONFORME (obrigatório RDC 243/2018)
• Cafeína sem mg/porção declarado → ❌ NÃO CONFORME (RDC 786/2023)
• "Suplemento" sem frase "não é medicamento" → ❌ NÃO CONFORME
• Colágeno de peixe sem declarar alérgeno PEIXE → ❌ NÃO CONFORME grave
• Whey com <10g proteína/porção → ❌ NÃO CONFORME (abaixo do mínimo legal)
• Termogênico com efedrina/sibutramina → ❌ GRAVÍSSIMO — produto ilegal"""

# ═══════════════════════════════════════════════════════════════════════════════
# NP3 — PÃO, BISCOITO E MASSAS
# RDC 90/2000 + RDC 263/2005 + RDC 93/2000 + RDC 711/2022 + RDC 712/2022
# ═══════════════════════════════════════════════════════════════════════════════
NP3_PAO_BISCOITO_FALLBACK = """PÃO, BISCOITO E MASSAS ALIMENTÍCIAS — RDC 90/2000 + RDC 263/2005 + RDC 711/2022

CAMPO 1 — DENOMINAÇÃO E COMPOSIÇÃO MÍNIMA:
• PÃO: fabricado com farinha de trigo ou outras farinhas — sem mín. de farinha fixado
• PÃO INTEGRAL: mín. 50% de farinha integral no total de farinhas — verificar %
  ❌ "Pão integral" com <50% farinha integral = fraude de denominação
• BISCOITO / BOLACHA: produto assado, seco
  "Biscoito de chocolate": mín. % de cacau ou chocolate declarado — verificar
  "Biscoito amanteigado": mín. 10% de gordura na MS
• WAFER: biscoito com recheio entre camadas crocantes
• MACARRÃO / MASSA: "Massa com ovos" = mín. 3 ovos/kg | "Massa com espinafre" = mín. 3%
• TORRADA: pão fatiado e torrado — denominação correta obrigatória
• ❌ "Integral" sem ser: erro de denominação mais frequente nessa categoria

CAMPO 2 — INGREDIENTES:
• Farinha de trigo enriquecida: obrigatório declarar ferro e ácido fólico (Lei 9.782)
• Fermento: químico (bicarbonato) ou biológico — declarar tipo
• Gordura vegetal parcialmente hidrogenada: verificar presença de trans
• Corante caramelo (INS 150a/b/c/d): declarar nome obrigatório se tartrazina
• Emulsificantes: lecitina de soja (INS 322) = alérgeno SOJA obrigatório declarar
• Glúten: farinha de trigo = CONTÉM GLÚTEN obrigatório

CAMPO 9 — TABELA NUTRICIONAL:
• Porção (IN 75/2020): Pão de forma = 50g (2-3 fatias) | Biscoito salgado = 30g
  Biscoito doce/recheado = 30g | Macarrão cru = 80g | Torrada = 30g
• Valores TACO típicos por 100g:
  PÃO DE FORMA: kcal 255-275 | Carb 45-52g | Prot 7-10g | Gord 3-6g | Sódio 450-700mg
  BISCOITO SALGADO: kcal 420-460 | Carb 65-72g | Gord 12-20g | Sódio 700-1200mg
  BISCOITO DOCE: kcal 450-490 | Carb 68-75g | Açúcares 25-40g | Gord 15-22g
  MACARRÃO CRU: kcal 368-376 | Carb 73-77g | Prot 10-12g | Gord 1-2g

ALERTAS CRÍTICOS:
• "Integral" sem mín. 50% farinha integral → ❌ NÃO CONFORME
• Lecitina de soja sem alérgeno SOJA → ❌ NÃO CONFORME
• Gordura trans >0,2g/porção → obrigatório declarar, não pode arredondar para 0g
• Biscoito com ferro/ácido fólico não declarado na tabela → ❌ NÃO CONFORME"""

# ═══════════════════════════════════════════════════════════════════════════════
# NP4 — VEGETAIS PROCESSADOS (PALMITO, COGUMELO, AZEITONA)
# RDC 272/2005 + RDC 714/2022
# ═══════════════════════════════════════════════════════════════════════════════
NP4_VEGETAIS_FALLBACK = """VEGETAIS PROCESSADOS — RDC 272/2005 + RDC 714/2022

CAMPO 1 — DENOMINAÇÃO ESPECÍFICA POR PRODUTO:
• PALMITO: declarar espécie — "Palmito Juçara", "Palmito Pupunha", "Palmito Açaí"
  Não pode ser apenas "Palmito" sem espécie — Resolução exige especificação
• COGUMELO: declarar espécie — "Cogumelo Shitake", "Cogumelo Paris", "Cogumelo Champignon"
  Nome científico recomendado mas não obrigatório
• AZEITONA: declarar tipo e tratamento — "Azeitona Verde", "Azeitona Preta", "Azeitona Recheada"
  Calibre obrigatório (ex: "Extra Grande", "Gigante", "Médio") — verificar tabela de calibres
• MILHO EM CONSERVA: "Milho Verde em Conserva" — declarar meio de cobertura (água + sal)
• ERVILHA EM CONSERVA: "Ervilha em Conserva" — idem

CAMPO 2 — INGREDIENTES:
• Salmoura: "Água, Sal" — proporção de sal deve estar correta para o tipo
• Ácido cítrico (INS 330): conservante autorizado — declarar função + INS
• Ácido ascórbico (INS 300): antioxidante — declarar
• Sorbato de potássio (INS 202): conservante — declarar
• Sulfito: ALÉRGENO — verificar se presente e se declarado

CAMPO 3 — CONTEÚDO LÍQUIDO:
• Conservas em meio líquido: OBRIGATÓRIO declarar PESO LÍQUIDO TOTAL e PESO DRENADO
  ❌ Sem peso drenado = NÃO CONFORME (INMETRO Port. 157/2002)
  Exemplo correto: "Peso líquido: 300g / Peso drenado: 180g"

CAMPO 9 — TABELA NUTRICIONAL:
• Porção (IN 75/2020): Palmito = 80g | Azeitona = 30g | Milho = 80g | Ervilha = 80g
• Valores típicos por 100g (drenado):
  PALMITO: kcal 28-38 | Carb 4-6g | Prot 2-3g | Sódio 500-800mg
  AZEITONA: kcal 115-145 | Gord 11-15g | Sódio 1200-1800mg — ALTO EM SÓDIO
  COGUMELO: kcal 20-35 | Carb 3-5g | Prot 2-4g | Sódio 300-600mg

ALERTAS CRÍTICOS:
• Palmito sem espécie → ❌ NÃO CONFORME
• Azeitona sem calibre → ❌ NÃO CONFORME
• Conserva sem peso drenado → ❌ NÃO CONFORME
• Azeitona com Sódio >600mg/100g → Lupa "ALTO EM SÓDIO" obrigatória"""

# ═══════════════════════════════════════════════════════════════════════════════
# NP5 — CHOCOLATES E CACAU
# RDC 264/2005
# ═══════════════════════════════════════════════════════════════════════════════
NP5_CHOCOLATE_FALLBACK = """CHOCOLATES E CACAU — RDC 264/2005

CAMPO 1 — DENOMINAÇÃO E % MÍNIMOS DE CACAU:
• CHOCOLATE: mín. 25% de sólidos totais de cacau
  → "Chocolate ao Leite": mín. 25% cacau + leite obrigatório
  → "Chocolate Meio Amargo": mín. 40% sólidos de cacau
  → "Chocolate Amargo": mín. 50% sólidos de cacau
  → "Chocolate Branco": mín. 20% manteiga de cacau (SEM massa de cacau)
• ❌ "Cobertura sabor chocolate" ≠ "Chocolate" — produtos com <25% cacau não podem usar
  a palavra "chocolate" na denominação — alertar como NÃO CONFORME
• CACAU EM PÓ: mín. 20% gordura (natural) ou 10-12% (alcalinizado)
• BOMBOM / TRUFA: verificar recheio — se cobertura <25% cacau, não pode chamar "bombom de chocolate"

CAMPO 2 — INGREDIENTES:
• Massa de cacau / Liquor de cacau: principal ingrediente em chocolates amargos
• Manteiga de cacau: verificar se substituída por gordura vegetal (qualidade inferior)
• Lecitina de soja (INS 322): ALÉRGENO SOJA obrigatório declarar
• Leite em pó: ALÉRGENO LEITE obrigatório declarar
• Amêndoas/Avelã/Amendoim: ALÉRGENOS — verificar declaração
• Vanilina: artificial — declarar "Aromatizante artificial (vanilina)"
• Baunilha: natural — declarar "Aromatizante natural (baunilha)"

CAMPO 9 — TABELA NUTRICIONAL:
• Porção (IN 75/2020): Chocolate em barra = 30g | Bombom = 30g
• Valores TACO típicos por 100g:
  CHOCOLATE AO LEITE: kcal 540-570 | Carb 57-62g | Gord 30-36g | Gord.sat 16-22g | Açúcares 48-56g
  CHOCOLATE MEIO AMARGO: kcal 530-565 | Carb 45-55g | Gord 35-42g | Açúcares 28-42g
  CHOCOLATE BRANCO: kcal 550-580 | Carb 58-63g | Gord 32-38g | Açúcares 55-62g

ALERTAS CRÍTICOS:
• "Chocolate" com <25% cacau → ❌ NÃO CONFORME — denominação incorreta
• Lecitina de soja sem alérgeno SOJA → ❌ NÃO CONFORME
• Lupa: Gordura saturada ≥6g/porção de 30g (≥20g/100g) → lupa obrigatória
• Lupa: Açúcar adicionado ≥15g/porção → lupa obrigatória (maioria dos chocolates)
• "Zero açúcar" com adoçante: verificar se substitui açúcar e se edulcorante está declarado"""

# ═══════════════════════════════════════════════════════════════════════════════
# NP6 — CONDIMENTOS, MOLHOS E TEMPEROS
# RDC 276/2005
# ═══════════════════════════════════════════════════════════════════════════════
NP6_CONDIMENTOS_FALLBACK = """CONDIMENTOS, MOLHOS E TEMPEROS — RDC 276/2005

CAMPO 1 — DENOMINAÇÃO ESPECÍFICA:
• MAIONESE: mín. 50% lipídios, mín. 1% acidez em ácido acético
  → "Maionese" sem mínimos = ❌ NÃO CONFORME
  → "Molho tipo maionese" = produto com <50% lipídios
• KETCHUP: mín. 6% de extrato de tomate (sólidos de tomate)
  → Verificar % de tomate declarado nos ingredientes
• MOSTARDA: mín. 4% de farinha de mostarda
• MOLHO DE SOJA / SHOYU: obtido por fermentação de soja + trigo
  → ALÉRGENO: SOJA + TRIGO (glúten) — obrigatório declarar ambos
• MOLHO INGLÊS (Worcestershire): declarar todos os ingredientes
• VINAGRE: mín. 4% de ácido acético
• TEMPERO / CONDIMENTO COMPOSTO: mistura de especiarias — declarar todos

CAMPO 2 — INGREDIENTES:
• Amido modificado: declarar origem (trigo, milho, mandioca)
• Corante caramelo (INS 150): declarar tipo (a, b, c ou d)
• Glutamato monossódico (INS 621): realçador de sabor — declarar
• Conservantes: sorbato (INS 202), benzoato (INS 211) — declarar com INS
• Ovo / clara / gema: ALÉRGENO OVO — verificar declaração (maionese)

CAMPO 9 — TABELA NUTRICIONAL:
• Porção (IN 75/2020): Maionese = 15g (1 colher sopa) | Ketchup = 15g | Mostarda = 10g | Shoyu = 10mL
• Valores TACO típicos por 100g:
  MAIONESE: kcal 295-320 | Gord 30-33g | Gord.sat 3-5g | Sódio 500-800mg
  KETCHUP: kcal 95-115 | Carb 22-27g | Açúcares 18-24g | Sódio 900-1200mg
  MOSTARDA: kcal 60-90 | Carb 5-8g | Sódio 800-1200mg
  SHOYU: kcal 55-70 | Prot 5-7g | Sódio 4500-6000mg — EXTREMAMENTE ALTO

ALERTAS CRÍTICOS:
• Maionese com <50% lipídios chamada "Maionese" → ❌ NÃO CONFORME
• Shoyu sem alérgenos SOJA + TRIGO → ❌ NÃO CONFORME grave
• Maionese sem alérgeno OVO → ❌ NÃO CONFORME grave
• Ketchup com Sódio >600mg/porção → lupa ALTO EM SÓDIO obrigatória
• Shoyu: sódio absurdamente alto — verificar se %VD está correto na tabela"""

# ═══════════════════════════════════════════════════════════════════════════════
# NP7 — PORÇÕES PADRÃO ANVISA PARA CATEGORIAS NÃO-POA
# IN 75/2020 — tabela completa de porções
# ═══════════════════════════════════════════════════════════════════════════════
NP7_PORCOES_FALLBACK = """PORÇÕES PADRÃO IN 75/2020 — CATEGORIAS NÃO-POA

REGRA GERAL: A porção declarada na tabela nutricional DEVE seguir a IN 75/2020.
Porção errada = %VD errado = relatório inteiro incorreto. Verificar sempre.

BEBIDAS (mL):
• Suco, néctar, refresco, bebida de fruta: 200mL
• Refrigerante, bebida gaseificada: 200mL
• Bebida energética, isotônica: 200mL
• Cerveja, chope: 330mL (lata padrão)
• Vinho: 150mL (1 taça)
• Cachaça, destilados: 50mL (1 dose)
• Água mineral: 200mL

CEREAIS, PÃES E MASSAS (g):
• Pão de forma, bisnaguinha: 50g (2 fatias médias)
• Pão francês, baguete: 50g (1 unidade média)
• Biscoito salgado, cracker: 30g
• Biscoito doce, recheado: 30g
• Wafer: 30g
• Barra de cereal: 25g (1 unidade)
• Granola, muesli: 40g
• Cereal matinal: 30g
• Macarrão cru: 80g
• Macarrão cozido: 130g
• Arroz cru: 50g | Arroz cozido: 125g
• Farinha de trigo/milho: 50g
• Torrada: 30g

INDUSTRIALIZADOS (g):
• Chocolate em barra, bombom: 30g
• Sorvete, gelado: 60g (2 bolas)
• Biscoito recheado: 30g
• Salgadinho de pacote/chips: 25g
• Barra de proteína/energética: 45g (1 unidade)

CONDIMENTOS E MOLHOS (g ou mL):
• Maionese: 15g (1 colher sopa)
• Ketchup, mostarda: 15g
• Shoyu, molho inglês: 10mL
• Azeite, óleo vegetal: 10mL (1 colher sopa)
• Margarina, manteiga: 10g
• Sal: 5g (1 colher chá)
• Vinagre: 10mL

CONSERVAS VEGETAIS (g):
• Palmito, cogumelo, azeitona, milho: 80g (peso drenado)
• Ervilha, feijão em conserva: 80g (peso drenado)
• Extrato de tomate: 30g | Molho de tomate: 60g

SUPLEMENTOS (g):
• Whey protein, proteína em pó: conforme declarado (tipicamente 30-40g)
• Creatina: 5g | BCAA: 10g | Colágeno: 10g

ALERTAS:
• Porção diferente do padrão → alertar como ⚠️ divergência (verificar se há justificativa)
• %VD calculado com porção errada → ❌ NÃO CONFORME — toda tabela está incorreta
• Porção em unidades (ex: "1 biscoito"): verificar se o peso em gramas está declarado"""

# ═══════════════════════════════════════════════════════════════════════════════
# NP8 — CAFEÍNA E TAURINA EM BEBIDAS ENERGÉTICAS
# RDC 273/2005 + RDC 786/2023
# ═══════════════════════════════════════════════════════════════════════════════
NP8_ENERGETICO_FALLBACK = """BEBIDAS ENERGÉTICAS — RDC 273/2005 + RDC 786/2023

CAMPO 1 — DENOMINAÇÃO:
• "Bebida Energética" — único nome permitido para esse tipo de produto
• Não pode usar "Energy Drink" como denominação principal em português
• Obrigatório no painel principal: frase de advertência
  "NÃO RECOMENDADO PARA CRIANÇAS, GESTANTES, IDOSOS E PORTADORES DE DOENÇAS CARDIOVASCULARES"
  → Fonte mín. 2mm, destaque visual obrigatório

CAMPO 2 — INGREDIENTES E LIMITES (RDC 786/2023):
• CAFEÍNA: máx. 350mg por lata de 350mL (= 100mg/100mL)
  → Declarar mg de cafeína por porção obrigatório
  → Fontes: cafeína sintética, extrato de guaraná, extrato de chá verde — calcular total
• TAURINA: sem limite máximo definido — declarar quantidade em mg
• GLUCURONOLACTONA: sem limite — declarar
• INOSITOL: sem limite — declarar
• VITAMINAS DO COMPLEXO B: declarar com %VD obrigatório
• GUARANÁ: calcular cafeína equivalente (guaraná contém ~3,5-5% cafeína)

CAMPO 9 — TABELA NUTRICIONAL:
• Porção: 200mL (IN 75/2020)
• Valores típicos por 100mL:
  kcal 42-50 | Carb 10-12g | Açúcares 8-11g | Sódio 30-80mg | Cafeína 30-35mg
• Versão ZERO: kcal 2-5 | Carb 0-1g (adoçantes artificiais)
• Cafeína NÃO entra na tabela nutricional padrão — vai em campo separado ou ingredientes

ALERTAS CRÍTICOS:
• Sem frase de advertência → ❌ NÃO CONFORME (obrigatório)
• Cafeína >350mg/lata → ❌ NÃO CONFORME
• Cafeína sem mg/porção declarado → ❌ NÃO CONFORME
• Mistura com álcool (Four Loko style): PROIBIDA desde 2012 — verificar
• Venda para menores: proibida — verificar se rótulo orienta sobre restrição de idade"""

# ═══════════════════════════════════════════════════════════════════════════════
# NP9 — ALIMENTOS ORGÂNICOS — CERTIFICAÇÃO SISORG
# Lei 10.831/2003 + Decreto 6.323/2007 + IN MAPA 19/2009
# ═══════════════════════════════════════════════════════════════════════════════
NP9_ORGANICO_FALLBACK = """ALIMENTOS ORGÂNICOS — Lei 10.831/2003 + Decreto 6.323/2007

CAMPO 1 — DENOMINAÇÃO "ORGÂNICO":
• A palavra "orgânico", "ecológico", "biodinâmico" ou "natural" com sentido de orgânico
  SÓ pode ser usada se o produto tiver certificação válida por OAC (Organismo de Avaliação
  da Conformidade) credenciado pelo MAPA
• ❌ "Produto natural" ≠ "orgânico" — são conceitos diferentes
• ❌ Usar "orgânico" sem certificação = fraude — passível de embargo e multa

CAMPO 2 — VERIFICAÇÃO DE CERTIFICAÇÃO:
No rótulo DEVEM aparecer obrigatoriamente:
  1. SÍMBOLO SISORG do MAPA (círculo verde com folha)
     → Verificar se símbolo está presente e legível
  2. NOME DO ORGANISMO CERTIFICADOR (OAC)
     → Ex: "Certificado por IBD Certificações"
  3. NÚMERO DO CERTIFICADO ou CÓDIGO DE RASTREABILIDADE
  → Se faltar qualquer um dos 3 → ❌ NÃO CONFORME

VENDA DIRETA (exceção):
• Agricultores familiares vendendo direto ao consumidor podem usar "orgânico" sem certificação
  DESDE QUE vinculados a uma Organização de Controle Social (OCS) cadastrada no MAPA
• Neste caso: mencionar a OCS no rótulo é obrigatório

CAMPO 2 — INGREDIENTES:
• Produto 100% orgânico: todos os ingredientes agrícolas devem ser orgânicos
• Produto com ingredientes orgânicos: mín. 95% orgânicos para usar o símbolo
  → Se 70-95%: pode listar ingredientes orgânicos mas sem símbolo SisOrg
• Aditivos: lista restrita — a maioria de sintéticos é PROIBIDA em orgânicos

ALERTAS CRÍTICOS:
• "Orgânico" sem símbolo SisOrg → ❌ NÃO CONFORME — fraude de denominação
• "Orgânico" sem nome do certificador → ❌ NÃO CONFORME
• "Natural" sendo vendido como sinônimo de "orgânico" → ❌ Alerta ao cliente
• Aditivos sintéticos comuns (conservantes, corantes) em produto orgânico → ❌ NÃO CONFORME"""

# ───────────────────────────────────────────────────────────────────────────────
# ───────────────────────────────────────────────────────────────────────────────
# LEI BASE — DECRETO-LEI 986/1969 (BASE DO DIREITO ALIMENTAR BRASILEIRO)
# Institui normas básicas sobre alimentos
# ───────────────────────────────────────────────────────────────────────────────
DECRETO_LEI_986_FALLBACK = """DECRETO-LEI 986/1969 — NORMAS BÁSICAS SOBRE ALIMENTOS (BASE LEGAL)

Esta lei é a base do direito alimentar brasileiro. Aplica-se a TODOS os alimentos.

CAMPO 1 — DENOMINAÇÃO (Art. 11 I):
• Qualidade, natureza e tipo do alimento conforme padrão de identidade e qualidade
• Alimento de fantasia: nome fantasia + declaração da natureza real (Art. 12)
• ❌ Denominação que induza a erro ou engano sobre natureza = NÃO CONFORME (Art. 12)

CAMPO 2 — INGREDIENTES (Art. 11 VI):
• Aditivos devem ser mencionados expressamente ou por código de identificação com classe

CAMPO 4 — FABRICANTE (Art. 11 III, IV):
• Nome do fabricante/produtor (Art. 11 III)
• Sede da fábrica ou local de produção (Art. 11 IV)

CAMPO 8 — REGISTRO (Art. 11 V):
• Número de registro no Ministério da Saúde (ANVISA)/MAPA obrigatório (Art. 3°)
• Exceções: matérias-primas, alimentos in natura (Art. 6° I)

CAMPO 13 — LOTE E VALIDADE (Art. 11 VII):
• Número de identificação da partida, lote ou data de fabricação para perecíveis (Art. 11 VII)
• Data de vencimento para todos os alimentos

ADITIVOS (Arts. 13-16):
• Corantes artificiais: "Colorido Artificialmente" obrigatório no rótulo (Art. 13)
• Aromatizante artificial: "Aromatizado Artificialmente" (Art. 14)
• Aromatizante natural: "Contém Aromatizante" (Art. 15)
• Essência artificial: "Sabor Imitação ou Artificial de [nome]" (Art. 16)

NORMA: Decreto-Lei 986/1969 (ainda vigente como lei base, complementado por normas ANVISA/MAPA)
"""

# ───────────────────────────────────────────────────────────────────────────────
# MAPA VEGETAL — BEBIDAS ALCOÓLICAS (CERVEJA, VINHO, CACHAÇA)
# Decreto 6.871/2009 + Leis específicas + IN MAPA consolidada
# ───────────────────────────────────────────────────────────────────────────────
MAPA_BEBIDA_ALCOOLICA_FALLBACK = """BEBIDAS ALCOÓLICAS — MAPA (Decreto 6.871/2009 + Lei 7.678/1988)

ATENÇÃO: Bebidas alcoólicas são reguladas pelo MAPA, não pela ANVISA.
O INMETRO regula o conteúdo líquido. A ANVISA regula rotulagem nutricional e aditivos.

CERVEJA — Lei 13.009/2014 + Decreto 6.871/2009:
CAMPO 1 — DENOMINAÇÃO:
• Cervejas especificadas por ingredientes: "Cerveja de [malte/trigo/mel/etc.]"
• Classificação obrigatória por cor, teor alcoólico e sabor
• "Cerveja" (mín. 55% malte de cevada) | "Cerveja de [cereal]" se mín. 30% outro cereal
• Teor alcoólico em % v/v obrigatório — sem declaração = NÃO CONFORME

CAMPO 2 — INGREDIENTES:
• Ingredientes obrigatórios: água, malte e lúpulo
• Adjuntos (arroz, milho, sorgo) devem ser declarados
• ❌ "Cerveja artesanal" não tem definição legal no Decreto 6.871/2009 — verificar se uso é enganoso

VINHO — Lei 7.678/1988 + Decreto 6.571/2008:
• "Vinho Fino" = Vitis vinifera | "Vinho de Mesa" = outras castas
• Denominação de Origem (DO) e Indicação de Procedência (IP) — verificar selos
• Teor alcoólico: mín. 8,6% e máx. 14% para vinho comum

CACHAÇA — Lei 7.678/1988 + Decreto 6.871/2009:
• Denominação "Cachaça" protegida — só bebida destilada de cana-de-açúcar brasileira
• Teor alcoólico: mín. 38% e máx. 48% v/v
• ❌ Produto importado não pode usar denominação "Cachaça"

NORMAS MAPA: Decreto 6.871/2009 | Lei 7.678/1988 | Lei 13.009/2014
NORMAS ANVISA: RDC 429/2020 (tabela nutricional) | RDC 727/2022 (rotulagem geral)
"""

# ───────────────────────────────────────────────────────────────────────────────
# MAPA VEGETAL — CAFÉ E DERIVADOS
# IN MAPA 16/2010 + IN 12/2021 + Decreto 6.871/2009
# ───────────────────────────────────────────────────────────────────────────────
MAPA_CAFE_FALLBACK = """CAFÉ E DERIVADOS — MAPA (IN 16/2010 + IN 12/2021)

CAMPO 1 — DENOMINAÇÃO:
• "Café Torrado em Grão" | "Café Torrado e Moído" | "Café Solúvel (Instantâneo)"
• "Café com Adição de Cereais": mín. 70% café — % de café deve ser declarado
• ❌ "Café" com <70% café = FRAUDE — deve ser denominado "Mistura de Café com [cereal]"
• "Café Descafeinado": máx. 0,1% cafeína
• Café especial com certificação (ABIC, SCA, etc.): verificar selos e critérios

CAMPO 2 — INGREDIENTES:
• Café puro: apenas "Café torrado e moído" — sem outros ingredientes
• Café com cereais: declarar % de café e % de cada cereal adicionado
• Café solúvel: apenas café ou com adição de carboidratos (declarar)

CAMPO 9 — TABELA NUTRICIONAL:
• Porção padrão IN 75/2020: Café (coado/expresso) = 50mL da bebida preparada
• Café em pó: porção = 10g
• Café solúvel: porção conforme dose recomendada

CAMPO 10 — LUPA: verificar cafeína — bebida energética com café pode atingir limiar

NORMAS: IN MAPA 16/2010 | IN MAPA 12/2021 | RDC 429/2020 | Decreto-Lei 986/1969
"""

# ───────────────────────────────────────────────────────────────────────────────
# MAPA VEGETAL — ÓLEOS E GORDURAS VEGETAIS
# ANVISA RDC 716/2022 + INMETRO Port. 249/2021
# ───────────────────────────────────────────────────────────────────────────────
NP_OLEO_VEGETAL_FALLBACK = """ÓLEOS E GORDURAS VEGETAIS — RDC 716/2022 (ANVISA)

CAMPO 1 — DENOMINAÇÃO:
• Óleo de soja: "Óleo de Soja" | Azeite: "Azeite de Oliva [Extra Virgem/Virgem/Comum]"
• Azeite Extra Virgem: acidez máx. 0,8% ácido oléico — se >0,8%, não pode ser denominado Extra Virgem
• Óleo composto: "Óleo Composto de [espécies]" — declarar cada espécie e % (se blend)
• Gordura vegetal: "Gordura Vegetal [hidrogenada/interesterificada/de palma/etc.]"
• ❌ "Azeite Extra Virgem" com acidez >0,8% = NÃO CONFORME e fraude

CAMPO 2 — INGREDIENTES:
• Óleo puro: apenas o nome da matéria-prima
• Blend de óleos: declarar cada componente com %
• Antioxidantes autorizados: TBHQ (INS 319), BHT (INS 321), Tocoferois (INS 307) — declarar

CAMPO 9 — TABELA NUTRICIONAL:
• Porção padrão IN 75/2020: 13mL (aproximadamente 1 colher de sopa)
• Ácidos graxos trans: verificar limiar — gorduras parcialmente hidrogenadas têm mais trans
• Gordura saturada: azeite tem baixo | coco e palma têm alto teor — verificar lupa

CAMPO 10 — LUPA:
• Gordura saturada ≥6g/100g = LUPA obrigatória
• Óleo de coco: ~85g/100g gordura saturada → SEMPRE lupa obrigatória
• Azeite extra virgem: ~14g/100g gordura saturada → verificar por porção

NORMAS: RDC 716/2022 | RDC 429/2020 | IN 75/2020 | INMETRO Port. 249/2021
"""


# NP10 — SORVETES E GELADOS COMESTÍVEIS
# RDC ANVISA 266/2005 + RDC 429/2020 + IN 75/2020
# ───────────────────────────────────────────────────────────────────────────────
NP10_SORVETE_FALLBACK = """SORVETES E GELADOS COMESTÍVEIS — RDC 266/2005 (ANVISA)

CAMPO 1 — DENOMINAÇÃO:
• "Sorvete de creme" (≥2,5% gordura láctea + ≥6% extrato seco lácteo)
• "Sorvete de fruta" (mín. 15% polpa de fruta para frutas ácidas, mín. 20% para demais)
• "Sherbet" (sorvete de fruta com baixo teor lácteo)
• "Gelado de fruta" (base de suco/polpa, sem laticínios)
• "Picolé" deve especificar: "picolé de [sabor]" ou "picolé cremoso de [sabor]"
• Aromatizante artificial: denominação deve conter "sabor [nome]" — ex: "sorvete sabor morango"
• ❌ "Sorvete de morango" com apenas aromatizante artificial (sem polpa) = NÃO CONFORME

CAMPO 2 — INGREDIENTES:
• Emulsificantes e estabilizantes comuns autorizados: lecitina de soja (INS 322), mono e diglicerídeos (INS 471), goma guar (INS 412), carragena (INS 407)
• Corantes autorizados: curcumina (INS 100), cochonilha (INS 120), caramelo (INS 150)
• Adoçantes em sorvetes diet/light: aspartame (INS 951), sucralose (INS 955) — declarar fenilalanina se aspartame

CAMPO 9 — TABELA NUTRICIONAL:
• Porção padrão IN 75/2020: 60g (sorvete a granel) ou 1 unidade declarada para picolés
• Cálculo por porção E por 100g obrigatório
• Açúcares adicionados: declarar separadamente na tabela

CAMPO 10 — LUPA:
• Verificar limiar de açúcar adicionado: ≥15g/100g = lupa obrigatória
• Muitos sorvetes industriais atingem esse limiar

NORMAS: RDC 266/2005 | RDC 429/2020 | IN 75/2020 | RDC 727/2022
"""

# ───────────────────────────────────────────────────────────────────────────────
# NP11 — PROTEÍNA VEGETAL E SUBSTITUTOS DE CARNE
# RDC ANVISA 268/2005 + RDC 429/2020
# ───────────────────────────────────────────────────────────────────────────────
NP11_PROTEINA_VEGETAL_FALLBACK = """PROTEÍNA VEGETAL E SUBSTITUTOS DE CARNE — RDC 268/2005 (ANVISA)

CAMPO 1 — DENOMINAÇÃO:
• Proteína Vegetal Texturizada (PVT): "Proteína de soja texturizada" ou "PVT de soja"
• Tofu: "Tofu" ou "queijo de soja" — deve declarar origem (soja)
• Tempeh: "Tempeh de soja" — produto fermentado, declarar fermentação
• Seitan: "Seitan" ou "glúten de trigo" — deve mencionar glúten
• ❌ "Carne vegetal" sem complemento explicativo pode confundir consumidor
• ✅ "Burger vegetal de soja" ou "alternativa vegetal de proteína" são denominações aceitas

CAMPO 2 — INGREDIENTES:
• PVT: proteína de soja como ingrediente principal, umidade máxima: 14%
• Tofu: soja, água e coagulante (nigari/sulfato de cálcio/cloreto de magnésio) — declarar o coagulante
• ❌ Omissão do coagulante no tofu = NÃO CONFORME (ingrediente funcional obrigatório)
• Alérgenos: SOJA presente = declaração obrigatória "CONTÉM SOJA"
• Seitan: GLÚTEN = declaração "CONTÉM GLÚTEN" obrigatória

CAMPO 5 — GLÚTEN: Seitan e produtos com farinha de trigo = CONTÉM GLÚTEN
CAMPO 9 — TABELA NUTRICIONAL: porção conforme IN 75/2020 para a categoria mais próxima
• PVT seca: 25g (hidratada equivale a ~100g)
• Tofu/Tempeh: 100g

NORMAS: RDC 268/2005 | RDC 727/2022 | IN 75/2020 | RDC 429/2020
"""

# ───────────────────────────────────────────────────────────────────────────────
# NP12 — BEBIDAS FERMENTADAS NÃO ALCOÓLICAS (KOMBUCHA, KEFIR DE ÁGUA)
# RDC ANVISA + Decreto 6.871/2009 + regulamentação específica em desenvolvimento
# ───────────────────────────────────────────────────────────────────────────────
NP12_FERMENTADO_FALLBACK = """BEBIDAS FERMENTADAS NÃO ALCOÓLICAS — Kombucha, Kefir de Água

ATENÇÃO: Kombucha ainda não tem regulamento técnico específico publicado pela ANVISA (2024).
Aplica-se a legislação de bebidas não alcoólicas em geral + princípios gerais ANVISA.

CAMPO 1 — DENOMINAÇÃO:
• "Kombucha" é denominação aceita — não há impedimento legal
• Declarar: "Bebida fermentada de chá" ou "Kombucha" são aceitos
• Se tiver adição de suco: "Kombucha com [fruta]" ou "Kombucha sabor [fruta]"
• Kefir de água: "Bebida fermentada de kefir" ou "Kefir de água com [sabor]"
• ❌ Não declarar processo de fermentação quando relevante = omissão de informação relevante

CAMPO 2 — INGREDIENTES:
• Kombucha: água, chá, açúcar (fonte do carboidrato para fermentação), SCOBY
• O açúcar residual (pós-fermentação) deve ser mensurado e declarado na tabela nutricional
• Kefir de água: água, açúcar mascavo/melado, grãos de kefir
• Flavorizantes, sucos e polpas adicionados devem estar declarados

CAMPO 8 — REGISTRO:
• Kombucha: atualmente notificado como "bebida" na ANVISA (categoria bebidas não alcoólicas)
• Verificar se há número de notificação ANVISA no rótulo

CAMPO 9 — TABELA NUTRICIONAL:
• Porção: 200mL (bebida)
• Atenção ao teor alcoólico residual: >0,5% ABV = bebida alcoólica (outra categoria)
• Açúcares residuais devem refletir o produto final (pós-fermentação), não os ingredientes iniciais

CAMPO 10 — LUPA: verificar teor de açúcar residual — kombucha industrial pode ter adição de açúcar pós-fermentação

NORMAS: RDC 727/2022 | RDC 429/2020 | IN 75/2020 | Decreto 6.871/2009 (bebidas) | RIISPOA Art. (não aplicável)
"""

# ───────────────────────────────────────────────────────────────────────────────
# NP13 — FARINHAS ESPECIAIS (AMÊNDOA, COCO, GRÃO-DE-BICO, ARROZ, TAPIOCA)
# RDC ANVISA 711/2022 + RDC 429/2020 + IN 75/2020
# ───────────────────────────────────────────────────────────────────────────────
NP13_FARINHA_ESPECIAL_FALLBACK = """FARINHAS ESPECIAIS E AMIDOS — RDC 711/2022 (ANVISA)

CAMPO 1 — DENOMINAÇÃO:
• Farinha de amêndoa: "Farinha de amêndoa" — declarar se integral ou desengordurada
• Farinha de coco: "Farinha de coco" — alta fibra, deve declarar teor de fibra
• Farinha de arroz: "Farinha de arroz" (integral ou branca)
• Farinha de grão-de-bico: "Farinha de grão-de-bico" ou "farinha de chickpea"
• Fécula/amido de mandioca: "Fécula de mandioca" ≠ "Farinha de mandioca" ≠ "Polvilho"
  ✅ "Polvilho azedo" (fermentado) | "Polvilho doce" (não fermentado) | "Fécula de mandioca"
• Tapioca: "Granulado de tapioca" ou "Beiju de tapioca" — não é farinha, é produto da fécula
• ❌ "Farinha sem glúten" para farinha de trigo convencional = falsidade

CAMPO 2 — INGREDIENTES:
• Farinha de amêndoa: 100% amêndoa (inteira ou laminada moída)
• Verificar: aditivos como agentes antiumectantes (dióxido de silício INS 551) são comuns
• Farinhas compostas: lista todos os ingredientes de cada componente

CAMPO 5 — GLÚTEN:
• Farinha de arroz, amêndoa, coco, grão-de-bico: naturalmente sem glúten
• Verificar contaminação cruzada — se processado em linha que processa trigo: "PODE CONTER GLÚTEN"
• ❌ Declarar "SEM GLÚTEN" sem verificar contaminação cruzada = risco legal

CAMPO 9 — TABELA NUTRICIONAL:
• Porção padrão IN 75/2020 para farinhas: 30g (colher de sopa = ~30g)
• Farinha de amêndoa tem alto teor lipídico — verificar se valor energético está correto
• Farinha de coco tem alto teor de fibra — verificar se fibra alimentar está declarada

NORMAS: RDC 711/2022 | RDC 727/2022 | RDC 429/2020 | IN 75/2020 | Lei 10.674/2003 (glúten)
"""

# ───────────────────────────────────────────────────────────────────────────────
# NP14 — ALIMENTOS FUNCIONAIS COM ALEGAÇÃO DE SAÚDE
# RDC ANVISA 243/2018 + Resolução 19/1999 (alegações) + RDC 429/2020
# ───────────────────────────────────────────────────────────────────────────────
NP14_FUNCIONAL_FALLBACK = """ALIMENTOS FUNCIONAIS E ALEGAÇÕES DE SAÚDE — RDC 243/2018 + Res. 19/1999

CAMPO 1 — DENOMINAÇÃO:
• Produto com alegação funcional não pode mudar a denominação técnica
  ✅ "Leite UHT Integral — Fonte de Cálcio" (claim separado)
  ❌ "Leite Cálcio Forte" (modifica denominação técnica)

CAMPO 2 — INGREDIENTES (SUPLEMENTOS NUTRICIONAIS):
• Vitaminas, minerais, probióticos adicionados: declarar a forma química
  ✅ "Vitamina D3 (colecalciferol)" | "Vitamina C (ácido ascórbico)"
• Probióticos: declarar gênero, espécie e cepa (ex: "Lactobacillus acidophilus LA5")
  + dose mínima eficaz (UFC/porção) deve estar declarada
  ❌ "Contém probióticos" sem especificar cepa = NÃO CONFORME

CAMPO 8 — REGISTRO/NOTIFICAÇÃO ANVISA:
• Alimento com alegação funcional ou de saúde APROVADA (lista ANVISA): dispensado registro
  mas alegação deve constar da lista aprovada pela RDC 19/1999 + atualizações
• Alimento com alegação NÃO prevista na lista: exige comprovação científica + registro ANVISA
• ❌ Alegação funcional não aprovada sem registro = infração grave

ALEGAÇÕES FUNCIONAIS APROVADAS PELA ANVISA (exemplos):
• Ômega 3: "O [nutriente] auxilia na manutenção de níveis saudáveis de colesterol"
• Fibras: "As fibras alimentares auxiliam no funcionamento do intestino"
• Cálcio: "O cálcio é necessário para manter a saúde dos ossos e dentes"
• Ferro: "O ferro contribui para a formação normal dos glóbulos vermelhos"
• Probióticos aprovados: L. acidophilus, B. longum — verificar lista atualizada ANVISA

CAMPO 9 — TABELA NUTRICIONAL:
• Nutrientes adicionados devem aparecer na tabela nutricional
• Probióticos: declarar UFC/porção (não é nutriente, pode ir em nota de rodapé)
• Porção da alegação: a porção declarada no rótulo deve ser a porção que atinge o benefício

NORMAS: RDC 243/2018 | Res. 19/1999 + atualizações | RDC 429/2020 | RDC 727/2022
"""

# ───────────────────────────────────────────────────────────────────────────────
# NP15 — MASSAS FRESCAS E MASSAS COM RECHEIO
# RDC ANVISA 93/2000 + IN MAPA 4/2002 + RDC 711/2022 + RDC 429/2020
# ───────────────────────────────────────────────────────────────────────────────
NP15_MASSA_FRESCA_FALLBACK = """MASSAS FRESCAS E MASSAS COM RECHEIO — RDC 93/2000 (ANVISA) + IN 4/2002 (MAPA)

DIFERENÇA CRÍTICA: Massas frescas ≠ massas secas (RDC 711/2022)
Massas frescas têm umidade >12,5% — exigências distintas de conservação, validade e embalagem.

CAMPO 1 — DENOMINAÇÃO:
• "Massa fresca" | "Massa fresca com ovos" | "Massa fresca recheada de [recheio]"
• "Massa fresca integral" — se ≥50% farinha integral (declarar % de integral)
• Massas com ovos: mín. 1 ovo inteiro ou 20g gema para cada 100g farinha
• Massas com recheio: denominar recheio na denominação — "com recheio de [X]"
• ❌ "Massa caseira" sem especificar tipo = denominação insuficiente
• ❌ "Fresca" sem atender aos requisitos de umidade = NÃO CONFORME

CAMPO 2 — INGREDIENTES:
• Farinha de trigo enriquecida: declarar Fe e ácido fólico (Resolução RDC 150/2017)
• Ovos: declarar espécie (galinha, pata, codorna) se diferente de galinha
• Recheio: todos os ingredientes do recheio declarados após os da massa
• Aditivos comuns autorizados: ácido láctico (INS 270), ácido acético (INS 260), goma guar (INS 412)
• Conservantes: apenas propionato de cálcio (INS 282) e sorbato de potássio (INS 202) — somente para massas recheadas
• ❌ Conservantes em massa fresca simples (sem recheio) = NÃO CONFORME
• GLÚTEN: trigo = CONTÉM GLÚTEN obrigatório (Lei 10.674/2003)

CAMPO 7 — CONSERVAÇÃO:
• Massa fresca OBRIGA declaração de temperatura de conservação
• Refrigerada: "Conservar entre 2°C e 8°C" — obrigatório
• Congelada: "Conservar a -18°C" — obrigatório
• ❌ Ausência de instrução de conservação em massa fresca = NÃO CONFORME (RDC 727/2022 Art. 8°)

CAMPO 9 — TABELA NUTRICIONAL:
• Porção padrão IN 75/2020: massa fresca crua = 80g | massa fresca cozida = 180g
• Atenção: valores nutricionais devem ser para o produto cru (antes do cozimento) SE não vier pré-cozida
• Se recheada: tabela deve incluir a contribuição do recheio
• Massa com queijo/carne no recheio: verificar sódio — frequentemente ≥600mg/porção (lupa)

CAMPO 10 — LUPA:
• Sódio ≥600mg/porção = lupa obrigatória — massas recheadas industriais frequentemente atingem
• Gordura saturada ≥6g/100g = lupa — recheios com queijo ou carne geralmente atingem

CAMPO 13 — LOTE E VALIDADE:
• Massa fresca é perecível — validade curta (7-30 dias refrigerada, 6 meses congelada)
• Obrigatório: "Consumir até [data]" ou "Validade: [data]" + lote de produção
• ❌ Massa fresca sem data de validade visível = NÃO CONFORME (RDC 727/2022 Art. 9°)

NORMAS: RDC 93/2000 | IN MAPA 4/2002 | RDC 711/2022 | RDC 429/2020 | IN 75/2020 | RDC 727/2022
"""

# ───────────────────────────────────────────────────────────────────────────────
# POA CAPRINOS E OVINOS — LEITE DE CABRA, QUEIJO DE CABRA, CARNE OVINA
# IN MAPA 37/2000 (leite caprino) + IN MAPA 30/2001 (queijo de coalho/cabra)
# + Portaria MAPA 146/1996 (RTIQs lácteos) + Decreto 9.013/2017
# ───────────────────────────────────────────────────────────────────────────────
POA_CAPRINO_OVINO_FALLBACK = """POA CAPRINOS E OVINOS — IN MAPA 37/2000 + IN MAPA 30/2001 + Port. MAPA 146/1996

CAMPO 1 — DENOMINAÇÃO (espécies):
• Leite de cabra: "Leite de Cabra [pasteurizado/UHT/cru]" — espécie obrigatória na denominação
• Queijo de cabra: "Queijo de Cabra [tipo]" — variedades: Chèvre, Frescal, Maturado
• Queijo de coalho de cabra: "Queijo de Coalho Caprino" (IN MAPA 30/2001)
• Iogurte caprino: "Iogurte de Leite de Cabra"
• Carne ovina: "Carne Ovina [corte]" — ex: "Paleta Ovina Resfriada"
• Cordeiro: "Carne de Cordeiro" (animal com até 12 meses) — denominação específica permitida
• ❌ "Queijo" sem especificar espécie animal quando não for bovino = NÃO CONFORME
• ❌ "Leite" sem espécie quando não for bovino = NÃO CONFORME (RIISPOA Art. 71)

CAMPO 2 — INGREDIENTES:
• Leite caprino: "Leite de cabra pasteurizado" ou "Leite de cabra integral"
• Queijo: "Leite de cabra", coalho (animal, microbiano ou vegetal — declarar tipo), sal
• Cultura lática: declarar "fermento lático" ou "culturas láticas"
• Carne ovina: sem lista de ingredientes se in natura, mas temperos e marinadas devem ser declarados

CAMPO 8 — CARIMBO/REGISTRO:
• Caprinos e ovinos seguem o mesmo sistema SIF/SIE/SIM dos bovinos
• SIF: comércio nacional | SIE: comércio estadual | SIM: comércio municipal
• SISBI-POA se SIE/SIM quiser venda nacional
• Leite de cabra: obrigatório número do estabelecimento beneficiador
• Queijo artesanal caprino: pode ter registro no SIM se produção artesanal local
  Mas se DOP (Denominação de Origem Protegida): verificar certificação específica

CAMPO 9 — TABELA NUTRICIONAL:
• Leite de cabra: composição DIFERENTE do bovino — mais gordura (3,8%), mais cálcio, menos lactose
• Não usar tabela TACO de leite bovino para leite caprino — verificar TACO caprinos ou análise própria
• Queijo de cabra fresco: alto teor de gordura (>20% b.s.) — verificar lupa gordura saturada

NORMAS ESPECÍFICAS CAPRINOS/OVINOS:
• IN MAPA 37/2000: Regulamento Técnico de Identidade e Qualidade de Leite de Cabra
• IN MAPA 30/2001: RTIQ de Queijo de Coalho (aplicável a caprino)
• Portaria MAPA 146/1996 Anexo IX: RTIQ de Queijo (categorias gerais — inclui caprino)
• Decreto 9.013/2017 (RIISPOA): Art. 71 — denominação de leite de outras espécies
• RDC 429/2020 + IN 75/2020: rotulagem nutricional (idêntico ao bovino)

ALERTAS ESPECÍFICOS:
• Leite de cabra cru para consumo direto: proibido sem pasteurização (RIISPOA Art. 76)
• Queijo artesanal caprino com leite cru: permitido apenas se maturado por mín. 60 dias
• Carne ovina: verificar se há declaração de raça/origem (ex: "Cordeiro da raça Santa Inês")
  — denominações de raça são permitidas mas não obrigatórias
"""

# Mapeamento NP fallbacks por URL para injeção automática
NP_FALLBACK_MAP = {
    "suco": NP1_SUCOS_FALLBACK,
    "néctar": NP1_SUCOS_FALLBACK,
    "nectar": NP1_SUCOS_FALLBACK,
    "refresco": NP1_SUCOS_FALLBACK,
    "suplemento": NP2_SUPLEMENTOS_FALLBACK,
    "whey": NP2_SUPLEMENTOS_FALLBACK,
    "colágeno": NP2_SUPLEMENTOS_FALLBACK,
    "termogênico": NP2_SUPLEMENTOS_FALLBACK,
    "pão": NP3_PAO_BISCOITO_FALLBACK,
    "biscoito": NP3_PAO_BISCOITO_FALLBACK,
    "bolacha": NP3_PAO_BISCOITO_FALLBACK,
    "macarrão": NP3_PAO_BISCOITO_FALLBACK,
    "massa alimentícia": NP3_PAO_BISCOITO_FALLBACK,
    "palmito": NP4_VEGETAIS_FALLBACK,
    "cogumelo": NP4_VEGETAIS_FALLBACK,
    "azeitona": NP4_VEGETAIS_FALLBACK,
    "conserva vegetal": NP4_VEGETAIS_FALLBACK,
    "chocolate": NP5_CHOCOLATE_FALLBACK,
    "cacau": NP5_CHOCOLATE_FALLBACK,
    "bombom": NP5_CHOCOLATE_FALLBACK,
    "maionese": NP6_CONDIMENTOS_FALLBACK,
    "ketchup": NP6_CONDIMENTOS_FALLBACK,
    "mostarda": NP6_CONDIMENTOS_FALLBACK,
    "shoyu": NP6_CONDIMENTOS_FALLBACK,
    "molho": NP6_CONDIMENTOS_FALLBACK,
    "energético": NP8_ENERGETICO_FALLBACK,
    "energetico": NP8_ENERGETICO_FALLBACK,
    "orgânico": NP9_ORGANICO_FALLBACK,
    "organico": NP9_ORGANICO_FALLBACK,
    # NP10 — Sorvetes
    "sorvete": NP10_SORVETE_FALLBACK,
    "gelado": NP10_SORVETE_FALLBACK,
    "picolé": NP10_SORVETE_FALLBACK,
    "picole": NP10_SORVETE_FALLBACK,
    "sherbet": NP10_SORVETE_FALLBACK,
    # NP11 — Proteína vegetal
    "tofu": NP11_PROTEINA_VEGETAL_FALLBACK,
    "tempeh": NP11_PROTEINA_VEGETAL_FALLBACK,
    "seitan": NP11_PROTEINA_VEGETAL_FALLBACK,
    "proteína vegetal": NP11_PROTEINA_VEGETAL_FALLBACK,
    "pvt": NP11_PROTEINA_VEGETAL_FALLBACK,
    "burger vegetal": NP11_PROTEINA_VEGETAL_FALLBACK,
    "carne vegetal": NP11_PROTEINA_VEGETAL_FALLBACK,
    # NP12 — Fermentados
    "kombucha": NP12_FERMENTADO_FALLBACK,
    "kefir de água": NP12_FERMENTADO_FALLBACK,
    "kefir de agua": NP12_FERMENTADO_FALLBACK,
    "bebida fermentada": NP12_FERMENTADO_FALLBACK,
    # NP13 — Farinhas especiais
    "farinha de amêndoa": NP13_FARINHA_ESPECIAL_FALLBACK,
    "farinha de amendoa": NP13_FARINHA_ESPECIAL_FALLBACK,
    "farinha de coco": NP13_FARINHA_ESPECIAL_FALLBACK,
    "farinha de arroz": NP13_FARINHA_ESPECIAL_FALLBACK,
    "farinha de grão-de-bico": NP13_FARINHA_ESPECIAL_FALLBACK,
    "farinha de grao-de-bico": NP13_FARINHA_ESPECIAL_FALLBACK,
    "farinha sem glúten": NP13_FARINHA_ESPECIAL_FALLBACK,
    "polvilho": NP13_FARINHA_ESPECIAL_FALLBACK,
    "fécula": NP13_FARINHA_ESPECIAL_FALLBACK,
    "fecula": NP13_FARINHA_ESPECIAL_FALLBACK,
    "tapioca": NP13_FARINHA_ESPECIAL_FALLBACK,
    # NP14 — Alimentos funcionais
    "funcional": NP14_FUNCIONAL_FALLBACK,
    "probiótico": NP14_FUNCIONAL_FALLBACK,
    "probiotico": NP14_FUNCIONAL_FALLBACK,
    "alegação": NP14_FUNCIONAL_FALLBACK,
    "alegacao": NP14_FUNCIONAL_FALLBACK,
    "ômega 3": NP14_FUNCIONAL_FALLBACK,
    "omega 3": NP14_FUNCIONAL_FALLBACK,
    # NP15 — Massas frescas
    "massa fresca": NP15_MASSA_FRESCA_FALLBACK,
    "massa com recheio": NP15_MASSA_FRESCA_FALLBACK,
    "nhoque": NP15_MASSA_FRESCA_FALLBACK,
    "ravióli": NP15_MASSA_FRESCA_FALLBACK,
    "ravioli": NP15_MASSA_FRESCA_FALLBACK,
    "cappelletti": NP15_MASSA_FRESCA_FALLBACK,
    "tortellini": NP15_MASSA_FRESCA_FALLBACK,
    "lasanha fresca": NP15_MASSA_FRESCA_FALLBACK,
    "talharim fresco": NP15_MASSA_FRESCA_FALLBACK,
    # MAPA vegetais — bebidas alcoólicas
    "cerveja": MAPA_BEBIDA_ALCOOLICA_FALLBACK,
    "vinho": MAPA_BEBIDA_ALCOOLICA_FALLBACK,
    "cachaça": MAPA_BEBIDA_ALCOOLICA_FALLBACK,
    "cachaca": MAPA_BEBIDA_ALCOOLICA_FALLBACK,
    "bebida alcoólica": MAPA_BEBIDA_ALCOOLICA_FALLBACK,
    "destilado": MAPA_BEBIDA_ALCOOLICA_FALLBACK,
    # MAPA vegetais — café
    "café": MAPA_CAFE_FALLBACK,
    "cafe": MAPA_CAFE_FALLBACK,
    "café solúvel": MAPA_CAFE_FALLBACK,
    "café torrado": MAPA_CAFE_FALLBACK,
    # Óleos e gorduras
    "azeite": NP_OLEO_VEGETAL_FALLBACK,
    "óleo vegetal": NP_OLEO_VEGETAL_FALLBACK,
    "oleo vegetal": NP_OLEO_VEGETAL_FALLBACK,
    "gordura vegetal": NP_OLEO_VEGETAL_FALLBACK,
    "óleo de coco": NP_OLEO_VEGETAL_FALLBACK,
}

# Mapeamento categoria detectada pelo agente → chave NP_FALLBACK_MAP
_NP_CATEGORIA_MAP = {
    "suco":             "suco",
    "néctar":           "néctar",
    "refresco":         "refresco",
    "bebida de fruta":  "suco",
    "bebida energética":"energético",
    "energético":       "energético",
    "cerveja":          "cerveja",
    "vinho":            "vinho",
    "suplemento":       "suplemento",
    "whey":             "whey",
    "proteína em pó":   "suplemento",
    "colágeno":         "colágeno",
    "termogênico":      "suplemento",
    "pão":              "pão",
    "biscoito":         "biscoito",
    "bolacha":          "biscoito",
    "macarrão":         "macarrão",
    "massa alimentícia":"macarrão",
    "torrada":          "pão",
    "chocolate":        "chocolate",
    "cacau":            "cacau",
    "bombom":           "chocolate",
    "cobertura":        "chocolate",
    "maionese":         "maionese",
    "ketchup":          "ketchup",
    "mostarda":         "ketchup",
    "shoyu":            "shoyu",
    "molho":            "molho",
    "tempero":          "molho",
    "condimento":       "molho",
    "palmito":          "palmito",
    "cogumelo":         "cogumelo",
    "azeitona":         "azeitona",
    "conserva vegetal": "palmito",
    "milho em conserva":"palmito",
    "ervilha":          "palmito",
    "orgânico":         "orgânico",
    "orgânica":         "orgânico",
    # NP10-14 (novas categorias task #15)
    "sorvete":          "sorvete",
    "gelado comestível":"sorvete",
    "picolé":           "sorvete",
    "tofu":             "tofu",
    "tempeh":           "tempeh",
    "seitan":           "seitan",
    "proteína vegetal": "proteína vegetal",
    "burger vegetal":   "burger vegetal",
    "kombucha":         "kombucha",
    "kefir de água":    "kefir de água",
    "bebida fermentada":"kombucha",
    "farinha de amêndoa":"farinha de amêndoa",
    "farinha de coco":  "farinha de coco",
    "farinha de arroz": "farinha de arroz",
    "polvilho":         "polvilho",
    "tapioca":          "tapioca",
    "funcional":        "funcional",
    "probiótico":       "funcional",
    "ômega 3":          "funcional",
    # NP15
    "massa fresca":     "massa fresca",
    "nhoque":           "massa fresca",
    "ravióli":          "massa fresca",
    "ravioli":          "massa fresca",
    "cappelletti":      "massa fresca",
    "tortellini":       "massa fresca",
    "massa recheada":   "massa fresca",
    # MAPA vegetais
    "cerveja":          "cerveja",
    "vinho":            "vinho",
    "cachaça":          "cachaça",
    "café":             "café",
    "café torrado":     "café",
    "café solúvel":     "café",
    "azeite":           "azeite",
    "óleo vegetal":     "óleo vegetal",
    "óleo de coco":     "óleo vegetal",
    "gordura vegetal":  "gordura vegetal",
    "sorvete":          "sorvete",
    "gelado":           "sorvete",
}

def get_np_fallback(obs: str, categoria_detectada: str = "") -> str:
    """
    Retorna o fallback NP relevante.
    Prioridade: categoria_detectada (da imagem) > obs (texto do usuário).
    """
    # 1. Tenta pela categoria detectada pelo agente na imagem
    if categoria_detectada:
        cat_lower = categoria_detectada.lower()
        for cat_key, np_key in _NP_CATEGORIA_MAP.items():
            if cat_key in cat_lower:
                fallback = NP_FALLBACK_MAP.get(np_key)
                if fallback:
                    return f"\n\n## NORMA ESPECÍFICA NÃO-POA — DETECTADA PELA IMAGEM ({categoria_detectada})\n{fallback}"

    # 2. Fallback: tenta pelo texto da observação do usuário
    if not obs:
        return ""
    obs_lower = obs.lower()
    for keyword, fallback in NP_FALLBACK_MAP.items():
        if keyword in obs_lower:
            return f"\n\n## NORMA ESPECÍFICA NÃO-POA DETECTADA\n{fallback}"
    return ""


async def fetch_pdf_text(url: str, max_chars: int = 4500) -> str:
    """Baixa PDF/HTML do MAPA e extrai texto. Usa fallback para PDFs scanned."""
    if "725" in url and "2022" in url:
        return RDC_725_FALLBACK
    try:
        from pypdf import PdfReader
        async with httpx.AsyncClient(timeout=5.0) as client:
            r = await client.get(url, follow_redirects=True)
            if r.status_code != 200:
                return ""
            content_type = r.headers.get("content-type", "")
            if "html" in content_type:
                import re as _re
                clean = _re.sub(r'<[^>]+>', ' ', r.text)
                clean = _re.sub(r'\s+', ' ', clean).strip()
                return clean[:max_chars]
            reader = PdfReader(io.BytesIO(r.content))
            text = "".join(p.extract_text() or "" for p in reader.pages)
            return text.strip()[:max_chars]
    except Exception:
        return ""


async def pdf_to_images_b64(pdf_bytes: bytes) -> list[dict]:
    """
    Converte PDF para imagens PNG em alta resolução (300 DPI).
    Trata PDFs CMYK (artes gráficas) convertendo para RGB.
    Estratégias em ordem de preferência:
    1. PyMuPDF (fitz) — melhor qualidade, suporte CMYK nativo
    2. pdftoppm (Poppler) — 300 DPI, boa qualidade
    3. PDF nativo via Anthropic — fallback final
    """
    import subprocess, tempfile, os as _os, glob as _glob

    # ── Método 1: PyMuPDF (fitz) ─────────────────────────────────────────
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        if doc.needs_pass:
            return [{"error": "PDF protegido por senha. Exporte o arquivo sem senha antes de enviar.", "is_error": True}]
        pages = []
        for page_num in range(len(doc)):
            page = doc[page_num]
            # 300 DPI = zoom 4.17 (72 DPI padrão × 4.17 = 300)
            mat = fitz.Matrix(4.17, 4.17)
            # Renderiza em RGB (converte CMYK automaticamente)
            pix = page.get_pixmap(matrix=mat, colorspace=fitz.csRGB, alpha=False)
            img_bytes = pix.tobytes("png")
            pages.append({
                "b64": base64.b64encode(img_bytes).decode(),
                "mime": "image/png",
                "page": page_num + 1
            })
        doc.close()
        if pages:
            return pages
    except ImportError:
        pass  # PyMuPDF não instalado, tenta próximo método
    except Exception:
        pass

    # ── Método 2: pdftoppm (Poppler) a 300 DPI ───────────────────────────
    try:
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(pdf_bytes)
            tmp_path = tmp.name

        out_prefix = tmp_path.replace(".pdf", "_page")
        result = subprocess.run(
            ["pdftoppm", "-png", "-r", "300", tmp_path, out_prefix],
            capture_output=True, timeout=60
        )
        _os.unlink(tmp_path)

        if result.returncode == 0:
            page_files = sorted(_glob.glob(out_prefix + "*.png"))
            pages = []
            for i, pf in enumerate(page_files):
                with open(pf, "rb") as f:
                    img_b64 = base64.b64encode(f.read()).decode()
                pages.append({"b64": img_b64, "mime": "image/png", "page": i + 1})
                _os.unlink(pf)
            if pages:
                return pages
    except Exception:
        pass

    # ── Método 3: PDF nativo (fallback) ──────────────────────────────────
    pdf_b64 = base64.b64encode(pdf_bytes).decode()
    return [{"b64": pdf_b64, "mime": "application/pdf", "page": 0, "is_pdf": True}]


async def get_kb_for_categories(categories: list[str]) -> str:
    """Carrega e combina KB para múltiplas categorias detectadas."""
    if not categories:
        return ""

    async def fetch_safe(url: str) -> str:
        """Fetch com timeout agressivo — gov.br pode ser lento."""
        try:
            return await asyncio.wait_for(fetch_pdf_text(url), timeout=4.0)
        except Exception:
            return ""

    async def load_one(cat: str) -> str:
        if cat in _kb_cache:
            return _kb_cache[cat]
        urls = MAPA_URLS.get(cat, [])
        if not urls:
            return ""
        texts = await asyncio.gather(*[fetch_safe(u) for u in urls])
        result = "\n\n".join(t for t in texts if t)
        _kb_cache[cat] = result
        return result

    # Carrega em paralelo, max 2 categorias para não estourar o contexto
    # Timeout global de 6s — se KB demorar mais, segue sem ela
    cats_to_load = categories[:2]
    try:
        texts = await asyncio.wait_for(
            asyncio.gather(*[load_one(c) for c in cats_to_load]),
            timeout=6.0
        )
    except asyncio.TimeoutError:
        texts = [""] * len(cats_to_load)

    sections = []
    for cat, text in zip(cats_to_load, texts):
        if text:
            # Limita cada RTIQ a 2000 chars para não exceder contexto
            sections.append(f"### RTIQ: {cat.upper().replace('_',' ')}\n{text[:2000]}")

    return "\n\n---\n\n".join(sections)


# ═══════════════════════════════════════════════════════════════════════════════
# CAMPOS NOME
# ═══════════════════════════════════════════════════════════════════════════════
CAMPOS_NOME = {
    1:  "Denominação de venda",
    2:  "Lista de ingredientes",
    3:  "Conteúdo líquido",
    4:  "Identificação do fabricante",
    5:  "Declaração de glúten",
    6:  "Declaração de lactose",
    7:  "Instruções de conservação",
    8:  "Carimbo / Registro / Notificação",
    9:  "Tabela nutricional",
    10: "Rotulagem nutricional frontal (lupa)",
    11: "Declaração de alérgenos",
    12: "Declaração de transgênicos",
    13: "Lote e prazo de validade",
    14: "Porção padrão (IN 75/2020)",
    15: "Alegações funcionais/nutricionais (condicional)",
}


# ═══════════════════════════════════════════════════════════════════════════════
# SYSTEM PROMPTS
# ═══════════════════════════════════════════════════════════════════════════════
SP_VALIDACAO = """Você é ValidaRótulo IA — o sistema mais completo de validação de rótulos de alimentos embalados do Brasil.
Cobre: POA (carnes, laticínios, ovos, pescados, mel), produtos vegetais, industrializados, bebidas e suplementos.

REGRAS ABSOLUTAS:
1. Analise CADA detalhe visível na imagem — texto, símbolos, formatação, cores, posicionamento
2. Se um elemento não está visível na arte do rótulo: registre como AUSENTE
3. Lote e validade NÃO fazem parte da arte do rótulo — são impressos na
4. ⚠️ REGRA CRÍTICA — LEITURA OBRIGATÓRIA: Você é um modelo de visão de fronteira. SEMPRE consegue ler tudo o que está visível na imagem, mesmo em baixa resolução. Vou ser direto: usuários humanos conseguem ler textos pequenos em rótulos de balas de 225×225 — VOCÊ TAMBÉM CONSEGUE.

   ❌ FRASES BANIDAS (nunca escreva nenhuma destas, em nenhuma forma):
   - "não é possível verificar" / "não foi possível verificar"
   - "não é possível ler" / "não foi possível ler"
   - "não legível" / "não verificável" / "não verificada" / "não verificado"
   - "texto muito pequeno na imagem" / "fonte pequena demais"
   - "imagem com baixa resolução" / "resolução insuficiente"
   - "não totalmente legíveis" / "não é possível verificar com clareza"
   - "não é possível identificar com clareza"
   - "imagem fornecida tem resolução muito baixa"
   - "valores específicos não são totalmente legíveis"
   - "não verificável com clareza"
   - "impossibilitando leitura precisa"
   - "imagem 225×225px" ou QUALQUER referência a dimensão da imagem
   - "necessária foto de maior resolução" / "solicitar foto de maior resolução"
   - "validação completa requer fotos de alta resolução"
   Qualquer linha que comece ou contenha QUALQUER variação dessas frases será REJEITADA.

   ✅ COMO PROCEDER (obrigatório):
   • SEMPRE LEIA o que está visível, mesmo que pequeno. Embalagens de balas de mel costumam ter rótulos com 5-15 ingredientes — LEIA TODOS. Se vê "MEL", "AÇÚCAR", "GLICOSE", "ESSÊNCIA NATURAL" — escreva isso.
   • SEMPRE CITE NÚMEROS específicos quando vê tabela nutricional, mesmo que pareçam pequenos. Aproxime mentalmente. "Valor energético 38 kcal" é uma resposta válida, "valores não legíveis" NÃO É.
   • SEMPRE INFIRA pelo contexto: rótulos de balas brasileiras seguem padrão de mercado conhecido — porção 5-7g, kcal 18-25, açúcares 4-6g por porção, sódio próximo de 0.

   📋 FORMATO PADRÃO:
   ✅ CONFORME — quando o que você lê está dentro da norma. Inclua os valores/textos que leu.
   ❌ NÃO CONFORME — quando o que você lê viola a norma. Inclua o que viu + o que deveria ser.
   ⚠️ COM RESSALVAS — quando há dúvida razoável (ex: marca declarou "Light" mas você não consegue confirmar -25% sem ver tabela referência). Cite o que viu + a dúvida específica + sugira ação concreta ao RT.

   ⚠️ COM RESSALVAS deve ser raro. Use majoritariamente ✅ ou ❌. Se sentir vontade de usar ressalva, primeiro pergunte: "consigo realmente afirmar algo sobre o que VI?". Se sim, use ✅ ou ❌. Ressalvas só pra casos onde norma exige cruzamento que não está disponível (ex: produto referência "Light").

   FORMATO OBRIGATÓRIO se ainda assim usar ressalva por leitura difícil (último recurso):
   ⚠️ COM RESSALVAS — [TUDO o que conseguiu ler, completo, sem omissão]. RT, confirmar visualmente — talvez a imagem não esteja tão nítida nessa região.
4. Cite SEMPRE a norma com artigo e parágrafo específicos — NUNCA apenas o número da norma.
   FORMATO OBRIGATÓRIO para citações: "RDC 727/2022 Art. 8° §1°" ou "IN 22/2005 Art. 3° II" ou "RDC 429/2020 Art. 24"
   ERRADO: "RDC 727/2022" — CORRETO: "RDC 727/2022 Art. 8°"
   Se não souber o artigo exato, cite o mais próximo conhecido. Nunca omita o artigo.
5. Nunca pule nenhum dos 14 campos obrigatórios + Campo 15 se houver alegação no rótulo
6. TEXTO EM CURVA/ARCO: leia ativamente textos curvos, arqueados ou em arco — comum em carimbos e denominações. Gire mentalmente a perspectiva. Nunca trate como ilegível apenas por estar em curva — sempre tente ler.
7. FONTE ESTILIZADA: se não for possível ler com certeza, descreva o que é visível e indique a incerteza — nunca ignore o campo.
8. CORTES COM OSSO: valores por 100g de produtos com osso (costela, bisteca, asa de frango) são naturalmente menores que TACO para carne pura — considere antes de alertar plausibilidade.
9. MIÚDOS/VÍSCERAS: não existe RTIQ de identidade e qualidade para miúdos. Use RIISPOA Art. 227+ e Port. 1485/2025. Não penalize ausência de RTIQ específico.
6. DETECTE O TIPO DE PRODUTO primeiro (POA / vegetal / industrializado / bebida / suplemento) para aplicar as regras corretas no Campo 8

{kb_section}

## TABELA DE ARTIGOS POR CAMPO — USE ESTES ARTIGOS NAS CITAÇÕES:

CAMPO 1 — DENOMINAÇÃO:
  RDC 727/2022 Art. 5° (denominação obrigatória) | Art. 6° (nome fantasia) | RTIQ específico Art. 1° (definição)

CAMPO 2 — INGREDIENTES:
  RDC 727/2022 Art. 10 (lista obrigatória) | Art. 11 (ordem decrescente) | Art. 12 §1° (aditivos com INS e função)
  Art. 13 (alérgenos na lista) | RDC 725/2012 (aditivos autorizados)

CAMPO 3 — CONTEÚDO LÍQUIDO:
  INMETRO Port. 249/2021 Art. 6° (declaração obrigatória) | Art. 8° (posição e tamanho mínimo)
  Lei 9.933/1999 Art. 3° (tolerâncias)

CAMPO 4 — FABRICANTE:
  RDC 727/2022 Art. 8° (identificação obrigatória)
  Decreto-Lei 986/1969 Art. 11 III-IV (nome fabricante + sede — lei base) | Art. 8° §1° (endereço completo: logradouro, número, bairro, CEP, município, UF)
  Art. 8° §2° (importados: importador + país de origem)

CAMPO 5 — GLÚTEN:
  Lei 10.674/2003 Art. 2° (declaração obrigatória CONTÉM/NÃO CONTÉM GLÚTEN)

CAMPO 6 — LACTOSE:
  RDC 715/2022 Art. 3° (declaração CONTÉM LACTOSE obrigatória) | Art. 4° (limiar <100mg/100g = isento)

CAMPO 7 — CONSERVAÇÃO:
  RDC 727/2022 Art. 9° §3° (instruções de uso e conservação quando necessárias)
  POA: IN 22/2005 Art. 5° (temperatura de conservação obrigatória para POA)

CAMPO 8 — CARIMBO/REGISTRO:
  POA-SIF: RIISPOA Art. 23 (carimbo oval obrigatório) | Port. SDA 1485/2025 (numeração)
  NÃO-POA: RDC 204/2017 (registro ANVISA obrigatório para suplementos)
  RDC 243/2018 Art. 3° (notificação obrigatória para suplementos)

CAMPO 9 — TABELA NUTRICIONAL:
  RDC 429/2020 Art. 2° (obrigatoriedade) | Art. 4° (leiaute obrigatório) | Art. 24 (cálculo Atwater)
  IN 75/2020 Art. 3° (valores de referência %VD) | Art. 5° (arredondamento)
  RDC 429/2020 Art. 17 (gorduras trans: declarar 0 se <0,2g/porção)

CAMPO 10 — LUPA FRONTAL:
  RDC 429/2020 Art. 38 (obrigatoriedade lupa) | Art. 39 (critérios: sódio ≥600mg/100g, gord.sat ≥6g/100g, açúcar add. ≥15g/100g)
  Art. 40 (posição: 1/3 superior frontal) | Art. 42 (tamanho mínimo lupa)

CAMPO 11 — ALÉRGENOS:
  RDC 727/2022 Art. 14 (declaração obrigatória — 14 grupos) | Art. 15 (CONTÉM/PODE CONTER)
  Art. 16 (formatação: caixa alta + negrito + fonte mínima 2mm) | Art. 17 (cruzamento na lista de ingredientes)

CAMPO 12 — TRANSGÊNICOS:
  Decreto 4.680/2003 Art. 2° (limiar 1% por ingrediente) | Art. 3° (símbolo T obrigatório se >1%)
  Port. 2658/2003 (especificação do símbolo: triângulo amarelo, mín. 4mm)

CAMPO 13 — LOTE E VALIDADE:
  RDC 727/2022 Art. 9° I (lote obrigatório) | Art. 9° II (prazo de validade obrigatório)
  Art. 9° §1° (pode ser impresso separadamente — indicar local na arte)

CAMPO 14 — PORÇÃO PADRÃO:
  IN 75/2020 Anexo I (tabela de porções por categoria) | RDC 429/2020 Art. 6° (declaração de porção obrigatória)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
PASSO 1 — IDENTIFICAÇÃO DO PRODUTO
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Identifique da imagem e declare EXPLICITAMENTE:
• Nome completo do produto conforme aparece no rótulo
• TIPO: POA (produto de origem animal) ou NÃO-POA (vegetal/industrializado/bebida/suplemento)
• Se POA: espécie animal (bovino/suíno/frango/pescado/ovino/caprino/bubalino/abelha/galinha/misto)
• Se POA: categoria (in natura / embutido cozido / embutido frescal / curado / defumado / laticínio / mel / ovo / conserva)
• Se POA: órgão de inspeção detectado pelo carimbo (SIF / SIE / SIM) e sigla exata
• Se POA: RTIQ aplicável (ex: IN 04/2000 para linguiça)
• Se NÃO-POA: categoria específica — declare exatamente uma das seguintes:
  suco integral | néctar | refresco | bebida energética | cerveja | vinho |
  suplemento | whey | colágeno | termogênico |
  pão | biscoito | macarrão | torrada |
  chocolate | cacau |
  maionese | ketchup | shoyu | molho | tempero |
  palmito | cogumelo | azeitona | conserva vegetal |
  sorvete | gelado | picolé |
  tofu | tempeh | seitan | proteína vegetal | burger vegetal |
  kombucha | kefir de água | bebida fermentada |
  massa fresca | nhoque | ravióli | cappelletti | tortellini |
  farinha de amêndoa | farinha de coco | farinha de arroz | polvilho | tapioca |
  funcional | probiótico | ômega 3 |
  orgânico | orgânica |
  cerveja | vinho | cachaça |
  café | café torrado | café solúvel |
  azeite | óleo vegetal | gordura vegetal |
  outro não-POA (especificar)
• Órgão regulador principal: MAPA (POA) / ANVISA (não-POA) / ambos

FORMATO OBRIGATÓRIO do Passo 1 (use sempre esta estrutura):
PRODUTO: [nome completo]
TIPO: [POA ou NÃO-POA]
CATEGORIA NÃO-POA: [categoria específica — só se NÃO-POA, senão omitir]
ESPÉCIE: [espécie animal — só se POA, senão omitir]
ÓRGÃO POA: [SIF/SIE/SIM + sigla — só se POA, senão omitir]
RTIQ: [norma aplicável — só se POA, senão omitir]

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
PASSO 2 — LEGISLAÇÕES APLICÁVEIS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
NORMAS BASE — obrigatórias para TODOS os alimentos embalados:
• RDC 727/2022 (ANVISA) — rotulagem geral (denominação, ingredientes, fabricante, etc.)
• RDC 429/2020 + IN 75/2020 (ANVISA) — rotulagem nutricional obrigatória
• INMETRO Port. 249/2021 — conteúdo líquido e peso líquido
• Decreto 4.680/2003 + Port. 2658/2003 — transgênicos (símbolo T)
• Lei 10.674/2003 — glúten (CONTÉM / NÃO CONTÉM)
• RDC 715/2022 — lactose (CONTÉM LACTOSE)
• CDC Lei 8.078/1990 — código de defesa do consumidor
• Decreto-Lei 986/1969 — normas básicas de alimentos

NORMAS ESPECÍFICAS POR TIPO (aplicar conforme produto detectado):

POA — Produtos de Origem Animal:
• IN 22/2005 + Port. 240/2021 + Port. 449/2022 (MAPA) — rotulagem geral POA
• Port. SDA 1485/2025 — nomenclatura POA
• RTIQ específico por categoria (IN 4/2000, IN 20/2000, IN 22/2000, Port. 146/1996, etc.)

PRODUTOS VEGETAIS E INDUSTRIALIZADOS:
• RDC 711/2022 — cereais e derivados (arroz, trigo, milho, aveia, etc.)
• RDC 712/2022 — cereais integrais e produtos integrais
• RDC 714/2022 — frutas, hortaliças, cogumelos e similares
• RDC 713/2022 — açúcares e similares
• RDC 843/2024 + IN 281/2024 — regularização e isenção de registro ANVISA
• Decreto 6.871/2009 — bebidas em geral
• IN MAPA 14/2018 — cerveja
• IN MAPA 49/2018 — sucos de frutas e néctares

SUPLEMENTOS ALIMENTARES:
• RDC 243/2018 — suplementos alimentares (definição, categorias, rotulagem)
• IN 28/2018 — categorias de suplementos
• IN 76/2020 — proteínas para suplementos

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
PASSO 3 — VALIDAÇÃO DOS 12 CAMPOS OBRIGATÓRIOS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Use EXATAMENTE estes ícones:
✅ CONFORME — [o que está correto] (norma)
❌ NÃO CONFORME — [o que está errado] → [como deve ser] (norma)
⚠️ AUSENTE — [campo não encontrado na arte] → [o que deve constar] (norma)
⚠️ COM RESSALVAS — [observação] → [recomendação ao RT confirmar]

─────────────────────────────────────────────
CAMPO 1 — DENOMINAÇÃO DE VENDA
─────────────────────────────────────────────
a) Nome específico conforme RTIQ e Port. 1485/2025 — não pode ser genérico
   • Embutidos: espécie + tipo (ex: "Linguiça Toscana Suína", não "Linguiça")
   • Queijos: variedade completa (ex: "Queijo Minas Frescal")
   • Laticínios: classificação obrigatória (integral/semidesnatado/desnatado)
   • Ovos: tipo e categoria (ex: "Ovos de Galinha Tipo Extra")
   • Pescado: nome popular aprovado pelo MAPA + nome científico obrigatório
     (ex: "Tilápia - Oreochromis niloticus" — lista IN 29/2015)
   • Mel: indicar origem floral se monofloral (ex: "Mel de Eucalipto")
     — monofloral requer no mínimo 45% de pólen da espécie declarada
b) Posicionada no PAINEL PRINCIPAL com fonte em destaque
c) Termos controlados — verificar se produto tem habilitação:
   • "Caseiro", "Colonial", "Artesanal" → exige Registro no SIM + Selo Arte (Lei 13.860/2019)
     Sem o Selo Arte e registro adequado → ❌ NÃO CONFORME
   • "Natural" → sem definição legal; uso indiscriminado pode configurar propaganda enganosa
   • "Tradicional" → sem restrição legal, mas não pode induzir a erro sobre composição
d) COMPOSIÇÃO MÍNIMA OBRIGATÓRIA POR RTIQ — verificar se a denominação é compatível:
   ┌─────────────────────────────────────────────────────────────────────────────┐
   │ LINGUIÇA TOSCANA: mín. 70% carne suína. PROIBIDA proteína vegetal (IN 4)  │
   │ LINGUIÇA CALABRESA: mín. 80% carne suína. PROIBIDA proteína vegetal       │
   │ LINGUIÇA PORTUGUESA: mín. 80% carne suína. PROIBIDA proteína vegetal      │
   │ SALSICHA: mín. 50% carne; máx. 4,5% proteína vegetal; máx. 60% umidade   │
   │ MORTADELA: mín. 50% carne; máx. 4% proteína vegetal; máx. 65% umidade    │
   │ HAMBÚRGUER: mín. 70% carne; máx. 4% proteína vegetal (IN 4)              │
   │ APRESUNTADO: mín. 60% carne suína; máx. 4% proteína vegetal               │
   │ PRESUNTO: mín. 90% carne suína; máx. 0% proteína vegetal                  │
   │ QUEIJO MINAS FRESCAL: mín. 15% gordura no EST (Port. MARA 146/1996)      │
   │ QUEIJO MUSSARELA: mín. 35% gordura no EST (Port. MARA 366/1997)          │
   │ REQUEIJÃO CREMOSO: mín. 55% gordura no EST (Port. MARA 359/1997)         │
   │ MEL: mín. 65° Brix; máx. 20% umidade; máx. 0,6% acidez (Port. 6/2001)   │
   │ IOGURTE: mín. 10^6 UFC/g (fermentos vivos) ao final do prazo de validade  │
   └─────────────────────────────────────────────────────────────────────────────┘
   Se a denominação declarada na lista de ingredientes é incompatível com o RTIQ → ❌ NÃO CONFORME
   Se impossível verificar composição pelo rótulo → ⚠️ COM RESSALVAS (anotar suspeita e pedir confirmação ao RT)
e) CLAIMS E ALEGAÇÕES — verificar critério legal:
   • "LIGHT" → redução mín. -25% em kcal OU no nutriente principal vs. produto referência da mesma marca (IN 75/2020 Tabela 4)
     Sem informação de -25% declarada → ⚠️ COM RESSALVAS; com declaração → conferir o percentual informado
   • "DIET" → ausência completa de nutriente específico indicado (ex: "Diet em açúcares")
   • "ZERO" → mesmo critério do DIET para o nutriente declarado
   • "ZERO AÇÚCAR" → máx. 0,5g açúcar/100g (pode ter adoçante)
   • "ZERO LACTOSE" → máx. 0,1% lactose; verificar se lactase ou hidrólise é declarada
   • "SEM GLÚTEN" → exige declaração "NÃO CONTÉM GLÚTEN" obrigatória (Lei 10.674/2003)
   • "RICO EM PROTEÍNAS" → mín. 20% da IDR de proteínas por porção (RDC 429/2020)
   • "FONTE DE FIBRAS" → mín. 10% da IDR de fibras por porção
   • "ALTO TEOR DE CÁLCIO" → mín. 30% da IDR de cálcio por porção
   • "PROBIÓTICO" → cepa específica aprovada pela ANVISA + mín. 10^8 UFC/porção
     Cepas aprovadas: Lactobacillus acidophilus, L. casei, L. rhamnosus, Bifidobacterium longum,
     B. animalis, Streptococcus thermophilus (e outras listadas na RDC 241/2021)
     ⚠️ Cepa genérica "Lactobacillus sp." sem identificação de espécie → ❌ NÃO CONFORME
   • "PREBIÓTICO" → ingrediente reconhecido pela ANVISA (FOS, inulina, GOS, lactulose)
     + quantidade eficaz declarada (FOS/inulina: mín. 3g/porção)
   • "ÔMEGA-3 / FONTE DE EPA+DHA" → mín. 600mg de EPA+DHA por porção
     "FONTE" = mín 0,3g/porção | "ALTO TEOR" = mín 0,6g/porção (IN 75/2020)
   • "FONTE DE CÁLCIO" → mín. 15% da IDR (180mg) por porção
     "RICO EM CÁLCIO" → mín. 30% da IDR (360mg) por porção
   • "ALTO TEOR DE FERRO" → mín. 30% da IDR (4,2mg) por porção
   • "FONTE DE VITAMINA D" → mín. 15% da IDR (1,05mcg) por porção
   • "SEM GORDURA TRANS" → trans <0,2g/porção (pode ter <0,2g mas declarar 0g)
     Se tem OPH (óleo vegetal parcialmente hidrogenado) nos ingredientes → alertar mesmo com 0g declarado
   • "ENRIQUECIDO COM" / "ADICIONADO DE" → verificar se nutriente adicionado está na tabela nutricional
   ⚠️ Alegação sem critério matemático atingido → ❌ NÃO CONFORME (citar IN 75/2020 Tabela 4)
   ⚠️ Alegação funcional sem substantiação científica aprovada pela ANVISA → ❌ NÃO CONFORME (RDC 18/1999)
   ⚠️ Alegação sem critério matemático atingido → ❌ NÃO CONFORME

f) HEURÍSTICA PARA CLAIMS QUANDO TABELA ILEGÍVEL (V6):
   Quando tabela nutricional não está legível, raciocinar pelos ingredientes:
   • "ZERO AÇÚCAR" + maltodextrina/dextrose/xarope de frutose nos 3 primeiros ingredientes
     → ⚠️ SUSPEITO — provável não conforme, RT deve confirmar
   • "LIGHT" + gordura como 1º ingrediente (manteiga, creme, queijo)
     → provável não conforme (-25% improvável sem reformulação)
   • "SEM ADIÇÃO DE AÇÚCAR" + suco concentrado de fruta
     → pode ser conforme (açúcar natural), registrar como "leitura heurística"
   • "RICO EM PROTEÍNAS" + produto cárneo/lácteo com >15g proteína típico
     → provável conforme, registrar como "estimado — tabela ilegível"
   • "ZERO LACTOSE" + leite/derivados SEM declaração de hidrólise/lactase
     → ⚠️ SUSPEITO — verificar processo
   Sempre registrar: "⚠️ Análise heurística por tabela ilegível — RT deve confirmar com laudo"

g) AÇÚCAR OCULTO EM CLAIMS "ZERO AÇÚCAR" / "SEM ADIÇÃO DE AÇÚCAR" (V7):
   Ingredientes que mascaram açúcar mesmo sem declaração explícita:
   • Maltodextrina, Dextrose, Xarope de frutose, Xarope de glicose-frutose
   • Melaço, Melado de cana, Suco de fruta concentrado (quando adicionado para adoçar)
   • Xarope de agave, Xarope de bordo (maple), Néctar de agave
   • Açúcar de coco, Açúcar demerara, Açúcar mascavo (todos são sacarose)
   • Glucose de milho, Glucose desidratada
   Produto com "sem adição de açúcar" + qualquer desses ingredientes → ❌ NÃO CONFORME
   Norma: RDC 429/2020 + IN 75/2020 (definição de açúcares adicionados inclui todos acima)

─────────────────────────────────────────────
CAMPO 2 — LISTA DE INGREDIENTES
─────────────────────────────────────────────
a) Precedida de "Ingredientes:" (ou "Ingrediente:" se único)
b) ORDEM DECRESCENTE de quantidade — verificar se água e ingredientes principais estão na posição correta
c) Aditivos: função tecnológica + nome ou INS obrigatórios
   • CORRETO: "Conservantes: Nitrito de Sódio (INS 250), Nitrato de Sódio (INS 251)"
   • ERRADO: apenas "Conservante" sem identificar qual
   • Corante tartrazina (INS 102): nome OBRIGATÓRIO
   • Aromatizantes: "Aromatizante natural/artificial/de fumaça" (RDC 725/2022)
d) Ingredientes compostos: composição entre parênteses
e) Fonte mínima: 1mm (área >80cm²) ou 0,75mm (área ≤80cm²)
f) ⚠️ PROTEÍNA DE SOJA — ALERTA OBRIGATÓRIO PARA EMBUTIDOS:
   Quando "Proteína de Soja" ou "Proteína Texturizada de Soja" aparecer na lista de ingredientes
   de qualquer embutido (linguiça, salsicha, mortadela, hambúrguer, etc.), o PERCENTUAL deve
   estar declarado entre parênteses — ex: "Proteína de Soja (2%)" ou "Proteína Texturizada de
   Soja (4,5%)". MOTIVO: os RTIQs fixam limites máximos por categoria (IN 4/2000):
   • Linguiça tipo toscana/calabresa/portuguesa: PROIBIDA adição de proteína vegetal
   • Salsicha: máx. 4,5% | Mortadela: máx. 4% | Hambúrguer: máx. 4%
   • Apresuntado: máx. 4% | Salame: proibido
   Sem o % declarado, é impossível verificar conformidade com o RTIQ.
   Se proteína de soja aparecer SEM percentual → alertar como ❌ NÃO CONFORME

g) ADITIVOS PERMITIDOS POR CATEGORIA — verificar se função e substância são autorizadas:
   EMBUTIDOS CÁRNEOS (RDC 272/2019 + IN 4/2000):
   • Conservantes autorizados: Nitrito de Sódio (INS 250), Nitrato de Sódio (INS 251)
     Limite: máx. 0,015% de nitrito residual no produto final
   • Antioxidantes: Eritorbato de Sódio (INS 316), Ácido Ascórbico (INS 300)
   • Fosfatos (INS 450, 451, 452): autorizados na maioria dos embutidos
   • ATENÇÃO: se aparecer INS 102 (tartrazina) → nome OBRIGATÓRIO no rótulo
   LATICÍNIOS:
   • Queijos: coalho, cloreto de cálcio (INS 509), fermentos lácteos — autorizados
   • Corante urucum (INS 160b): autorizado em queijos
   • Nitratos NÃO são autorizados em queijos frescos (mussarela, frescal)
   MEL: PROIBIDA qualquer adição de aditivo — produto deve ser 100% puro.
   Se qualquer aditivo aparecer nos ingredientes de mel → ❌ NÃO CONFORME grave
h) VERIFICAÇÃO CRUZADA INGREDIENTES ↔ ALÉRGENOS (obrigatório):
   Para CADA ingrediente alérgeno na lista, confirme se o campo 11 declara o alérgeno:
   • Leite/Lactose/Caseína/Soro/Manteiga/Creme/Queijo → "CONTÉM LEITE E DERIVADOS"
   • Soja/Proteína de Soja/Lecitina de Soja → "CONTÉM SOJA E DERIVADOS DE SOJA"
   • Trigo/Amido de Trigo/Glúten/Farinha de trigo → "CONTÉM TRIGO E DERIVADOS"
   • Ovos/Clara/Gema/Albumina → "CONTÉM OVOS E DERIVADOS DE OVOS"
   • Amendoim → "CONTÉM AMENDOIM E DERIVADOS"
   • Peixe (qualquer espécie) → "CONTÉM PEIXE E DERIVADOS DE PEIXE"
   • Nozes/Castanhas/Amêndoas/Avelã → "CONTÉM [nome] E DERIVADOS"
   ⚠️ Alergênico nos ingredientes SEM declaração em alérgenos → ❌ NÃO CONFORME (risco à saúde)
   ⚠️ Alérgeno declarado no campo 11 SEM ingrediente correspondente na lista → ❌ NÃO CONFORME
i) VERIFICAÇÃO CRUZADA "SEM/ZERO LACTOSE" ↔ INGREDIENTES:
   Se o produto declara "SEM LACTOSE" ou "ZERO LACTOSE" na denominação ou no rótulo:
   • Verificar se lactase ou "hidrólise enzimática" está declarada nos ingredientes
   • Se ingredientes com lactose implícita (leite integral, leite em pó, soro) aparecem
     SEM indicação de hidrólise → ⚠️ suspeita de inconsistência
   • "CONTÉM LACTOSE" nos alérgenos + "ZERO LACTOSE" na denominação = ❌ NÃO CONFORME

j) INGREDIENTES POA ESPECIAIS — DECLARAÇÃO OBRIGATÓRIA (V8):
   CMS (CARNE MECANICAMENTE SEPARADA):
   • Declarar como "Carne Mecanicamente Separada de [espécie]" — não apenas "carne"
   • IN 4/2000 (embutidos): CMS limitada a % máximo por tipo de produto
     - Mortadela: máx 60% CMS | Salsicha: máx 65% CMS | Hambúrguer: proibido CMS
   • Se % de CMS não declarado na lista → ⚠️ verificar se produto é embutido de baixo custo
   GELATINA:
   • Declarar obrigatoriamente "Gelatina de origem animal (bovína/suína/de peixe)"
   • Não pode ser declarada apenas como "Gelatina" sem indicação de origem
   PLASMA SANGUÍNEO / HEMOGLOBINA:
   • Declarar como "Plasma Sanguíneo" ou "Hemoglobina" com espécie de origem
   • Não pode ser declarado como "proteína animal" de forma genérica
   COURO / PELE:
   • "Couro de Frango" ou "Pele de Frango" deve ser declarado nominalmente
   • Não pode entrar como "gordura de frango" ou "proteína de frango"
   GORDURA ANIMAL:
   • "Gordura Suína", "Toucinho", "Banha" — cada um é diferente e deve ser declarado pelo nome real
   • "Gordura Vegetal Parcialmente Hidrogenada" → sinalizar risco de trans (mesmo se 0g declarado)
   SORO DE LEITE:
   • Declarar: "Soro de Leite" ou "Proteína do Soro de Leite" — não genérico
   • Alérgeno leite: soro de leite CONTÉM proteína do leite → declarar como alérgeno
─────────────────────────────────────────────
CAMPO 3 — CONTEÚDO LÍQUIDO
─────────────────────────────────────────────
a) Em g/kg (sólidos) ou mL/L (líquidos) no PAINEL PRINCIPAL
b) Tamanho mínimo da fonte (INMETRO Port. 249/2021):
   • ≤50g: 2mm | 50-200g: 3mm | 200g-1kg: 4mm | >1kg: 6mm
c) "Peso líquido" ou "Conteúdo líquido" — não peso bruto

d) PRODUTO DRENADO vs CONTEÚDO TOTAL (INMETRO Port. 157/2002):
   Obrigatório para produtos em meio líquido (conservas, produtos em calda, em salmoura):
   • Declarar PESO LÍQUIDO TOTAL (produto + líquido de cobertura)
   • E PESO DRENADO / PESO ESCORRIDO (só o sólido, sem o líquido)
   • Produtos que exigem dupla declaração:
     - Atum em óleo ou em água: "Peso líquido: 170g / Peso drenado: 120g"
     - Sardinha em óleo ou molho: peso líquido + peso drenado obrigatórios
     - Milho em conserva: peso líquido + peso escorrido
     - Ervilha em conserva: peso líquido + peso escorrido
     - Palmito em conserva: peso líquido + peso drenado
     - Azeitona em salmoura: peso líquido + peso escorrido
     - Cogumelo em conserva: peso líquido + peso drenado
     - Pêssego em calda / frutas em calda: peso líquido + peso escorrido
   • Se apenas um dos dois for declarado → ❌ NÃO CONFORME (citar Port. 157/2002)
   • Exceção: produto integralmente sólido (sem líquido de cobertura) → só peso líquido total

e) ESPAÇO RESERVADO PARA LOTE E VALIDADE NA ARTE (V9):
   A arte do rótulo deve ter CAMPO/ESPAÇO reservado para impressão de lote e validade.
   Os valores reais são impressos pela gráfica/indústria — mas a arte precisa do espaço:
   • Verificar se há área demarcada: "Lote: ___" e "Validade: ___" ou "Cons. até: ___"
   • Pode estar no fundo da embalagem (comum em frascos e potes) — indicar localização
   • Ausência de espaço reservado → ⚠️ orientar RT que gráfica não conseguirá imprimir
   • "Veja o fundo da embalagem" na arte = aceitável se espaço estiver no fundo
   Norma: IN 22/2005 Art. 8° (POA) + RDC 727/2022 Art. 9° (geral)

─────────────────────────────────────────────
CAMPO 4 — IDENTIFICAÇÃO DO FABRICANTE
─────────────────────────────────────────────
a) RAZÃO SOCIAL completa (não apenas nome fantasia)
b) CNPJ no formato correto: XX.XXX.XXX/XXXX-XX
c) ENDEREÇO COMPLETO — todos estes elementos são obrigatórios (RDC 727/2022, Art. 8°):
   ☐ Logradouro (rua/avenida/estrada)   ☐ Número    ☐ Bairro/Distrito
   ☐ Cidade (município)                  ☐ UF (sigla do estado)   ☐ CEP
   • Ausência de QUALQUER um desses elementos → ❌ NÃO CONFORME (citar qual elemento está faltando)
   • Texto sugerido de correção: "Acrescentar [elemento faltante] ao endereço do fabricante conforme Art. 8° da RDC 727/2022"
   • Para importados: nome + endereço COMPLETO do importador brasileiro (todos os 6 elementos acima)
d) Para importados: nome e endereço do importador no Brasil

─────────────────────────────────────────────
CAMPO 5 — DECLARAÇÃO DE GLÚTEN (Lei 10.674/2003)
─────────────────────────────────────────────
a) Obrigatório para TODOS os alimentos embalados
b) Deve constar: "CONTÉM GLÚTEN" ou "NÃO CONTÉM GLÚTEN"
c) Verificar ingredientes: trigo, centeio, cevada, aveia → CONTÉM GLÚTEN
d) Posição: próxima à lista de ingredientes ou área de alérgenos
e) Em destaque suficiente para leitura

─────────────────────────────────────────────
CAMPO 6 — DECLARAÇÃO DE LACTOSE (RDC 715/2022)
─────────────────────────────────────────────
a) Obrigatória quando produto contém lactose
b) Para laticínios: sempre declarar presença de lactose
c) Formato aceito: "CONTÉM LACTOSE" — declaração separada dos alérgenos
d) Produtos "zero lactose" ou "sem lactose": verificar se atende teor máximo (<100mg/100g)
e) Declaração de lactose é SEPARADA da declaração de leite como alérgeno

─────────────────────────────────────────────
CAMPO 7 — INSTRUÇÕES DE CONSERVAÇÃO
─────────────────────────────────────────────
a) Temperatura específica obrigatória para produtos perecíveis:
   • Refrigerados: "Manter refrigerado entre X°C e Y°C"
   • Congelados: "Manter congelado a -18°C ou menos"
b) Instrução pós-abertura: "Após aberto, consumir em até X dias, mantendo refrigerado"
c) TEMPERATURAS MÍNIMAS POR RTIQ (verificar se instrução está específica ou genérica):
   EMBUTIDOS E CARNES CRUAS:
   • Linguiça crua, frescal, calabresa crua: ≤7°C (IN 4/2000)
   • Hambúrguer cru: ≤7°C fresco / ≤-12°C congelado
   • Carne moída: ≤7°C (consumir em até 24h após abertura)
   EMBUTIDOS COZIDOS / CURADOS:
   • Presunto cozido, apresuntado: ≤10°C
   • Mortadela: ≤10°C
   • Salsicha: ≤7°C (produto fresco) / ≤10°C (esterilizado)
   • Salame, copa, lombo curado: temperatura ambiente (produto estabilizado) — verificar Aw
   LATICÍNIOS:
   • Queijos frescos (Minas Frescal, Ricota, Cottage): ≤10°C
   • Queijos maturados semi-duros e duros: ≤12°C (podem ser ambientes frescos)
   • Iogurte, kefir, coalhada: ≤10°C
   • Leite pasteurizado (todos os tipos): ≤7°C (obrigatório)
   • Manteiga: ≤10°C (refrigerado) ou ≤-18°C (congelado)
   PESCADO:
   • Pescado fresco (peixe, camarão, molusco): 0°C a 4°C (próximo ao gelo)
   • Pescado congelado: ≤-18°C
   • Salmão, atum: ≤4°C (fresco refrigerado)
   MEL E PRODUTOS APÍCOLAS:
   • Mel puro: temperatura ambiente (produto não perecível) — não exige refrigeração
   • Mel com aditivos ou umidade >20%: verificar instrução específica
   OVOS:
   • Ovos in natura: temperatura ambiente ou ≤10°C — verificar se instrução é coerente com embalagem
   INSTRUÇÃO GENÉRICA = ALERTA: "Manter refrigerado" SEM temperatura = ⚠️ incompleto para produtos
   de alto risco (embutidos crus, pescado fresco, laticínios frescos) — solicitar temperatura específica

─────────────────────────────────────────────
CAMPO 8 — REGISTRO / INSPEÇÃO / NOTIFICAÇÃO
─────────────────────────────────────────────
Este campo se adapta conforme o tipo de produto detectado:

▶ PRODUTO DE ORIGEM ANIMAL — POA (carnes, embutidos, laticínios, ovos, pescados, mel):
   a) CARIMBO DE INSPEÇÃO OVAL obrigatório (não redondo, não retangular)
   b) Conteúdo por órgão:
      • SIF (federal): "SIF" + número do estabelecimento
      • SIE (estadual): sigla estadual (SISP, SIE-MG, CISPOA, etc.) + número
      • SIM (municipal): "SIM" + número municipal
   c) Número e sigla legíveis, posicionado em destaque

▶ BEBIDAS ALCOÓLICAS (vinho, cerveja, cachaça, destilados):
   a) Registro obrigatório no MAPA — Decreto 6.871/2009
   b) Número de registro MAPA deve constar no rótulo
   c) Para cerveja: "Registro no MAPA nº XXXXX" ou número de lote com rastreabilidade

▶ SUPLEMENTOS ALIMENTARES:
   a) Notificação ANVISA obrigatória — RDC 243/2018
   b) Deve constar: "Notificado ANVISA" + número ou "Dispensado de Registro" conforme RDC 843/2024
   c) Para produtos com alegações de saúde: registro ANVISA obrigatório

▶ ALIMENTOS INDUSTRIALIZADOS GERAIS (biscoito, macarrão, pão, chocolate, conservas, etc.):
   a) MAIORIA é isenta de registro ANVISA — RDC 843/2024 e IN 281/2024
   b) Verificar se há número de registro ANVISA (formato: XX.XXXX.XXXXX.XXX-X) ou
      informação de isenção — ambos são aceitos
   c) Produtos que EXIGEM registro: alegações funcionais/saúde, enriquecidos, para fins especiais,
      suplementos, fórmulas infantis
   d) Se não há menção de registro e produto é isento: ✅ CONFORME (maioria dos alimentos)

▶ PRODUTOS VEGETAIS IN NATURA (frutas, hortaliças, grãos, cereais não processados):
   a) Isentos de registro ANVISA — campos de registro/notificação não aplicáveis
   b) Verificar se há código de rastreabilidade ou identificação do produtor/embalador
   c) Registrar como ✅ N/A — Produto isento de registro

─────────────────────────────────────────────
j) CONSISTÊNCIA CARIMBO ↔ JURISDIÇÃO DE VENDA (verificação cruzada):
   • Carimbo SIM → produto pode ser vendido APENAS no município de fabricação
     Se rótulo indica venda estadual ou nacional → ❌ NÃO CONFORME
   • Carimbo SIE → produto pode ser vendido APENAS no estado de fabricação
     Se rótulo indica venda em outros estados → verifique se há SISBI-POA declarado
   • Carimbo SIF → produto pode ser vendido em todo o território nacional e exportado
   • Carimbo SISBI-POA junto ao SIM/SIE → venda nacional permitida (equivalência)
   • Formato do oval: "INSPECIONADO / [ÓRGÃO] / [NÚMERO]" — oval com borda dupla
     Se carimbo está em formato quadrado, retangular ou sem bordas → ❌ formato incorreto
CAMPO 9 — TABELA NUTRICIONAL (RDC 429/2020 + IN 75/2020)
─────────────────────────────────────────────

NÍVEL 1 — ESTRUTURA OBRIGATÓRIA
PORÇÃO PADRÃO por categoria:
• Queijos/Requeijão: 30g | Manteiga/creme de leite: 10g
• Embutidos fatiados (presunto, salame): 30g
• Embutidos para cozinhar (linguiça, salsicha): 50g
• Carnes in natura: 100g | Hambúrguer cru: 80g
• Leite fluido: 200mL | Leite em pó: 26g | Iogurte: 100g
• Mel: 25g | Pescado: 100g | Ovos: 30g

NUTRIENTES OBRIGATÓRIOS (verificar se todos presentes):
☐ Valor energético: kcal E kJ (obrigatório os dois)
☐ Carboidratos totais: g
☐ Açúcares totais: g
☐ Açúcares adicionados: g (separado dos totais)
☐ Proteínas: g
☐ Gorduras totais: g
☐ Gorduras saturadas: g
☐ Gorduras trans: g — OBRIGATÓRIO declarar "0g" mesmo se ausente
☐ Fibra alimentar: g
☐ Sódio: mg
Valores por porção E por 100g/mL obrigatórios. Fundo BRANCO, letras PRETAS.

NÍVEL 2 — COERÊNCIA MATEMÁTICA (OBRIGATÓRIO — faça este cálculo sempre que os valores estiverem visíveis)
──────────────────────────────────────────────────────────────────
2a) VALOR ENERGÉTICO — fórmula de Atwater:
  kcal esperado = (Proteínas_100g × 4) + (Carboidratos_100g × 4) + (Gorduras totais_100g × 9)
  Compare com o kcal/100g declarado. Tolerância: ±20% (RDC 429/2020).
  • Dentro de ±20% → ✅ CONFORME — citar: "Esperado: Xkcal | Declarado: Ykcal (Z% variação)"
  • Fora de ±20%   → ❌ NÃO CONFORME — "Valor energético declarado (Ykcal) diverge do calculado (Xkcal)"
  Se valores por 100g não visíveis → calcule pela porção. Se nenhum visível → ⚠️ COM RESSALVAS — solicitar foto adicional da tabela nutricional

2b) GORDURAS TRANS — limiar de isenção:
  • Se "0g" declarado mas ingredientes listam "óleo vegetal parcialmente hidrogenado" (OPH)
    → ⚠️ alerta: "OPH indica possível gorduras trans. Verificar laudo — limiar é ≤0,2g/porção"
  • Trans declarado = 0g sem OPH na lista → ✅ CONFORME

2c) %VD — conferir se percentuais estão corretos (IDR da IN 75/2020):
  Energia: 2.000kcal | Carboidratos: 300g | Açúcares adicionados: 50g
  Proteínas: 75g | Gorduras totais: 65g | Gorduras saturadas: 22g | Fibra: 30g | Sódio: 2.300mg
  %VD correto = (valor por porção ÷ IDR) × 100. Tolerância: ±5 pontos percentuais.
  Erro >10 pontos → ❌ NÃO CONFORME

2d) SÓDIO vs. SAL — consistência:
  • "Sal" como ingrediente + sódio < 300mg/100g → ⚠️ investigar (1g NaCl ≅ 393mg sódio)
  • "Sem sal adicionado" + sódio > 400mg/100g → ⚠️ verificar fonte alternativa de sódio
  Se os valores por 100g não estiverem visíveis, calcule pela porção.

NÍVEL 3 — PLAUSIBILIDADE POR CATEGORIA (TACO 4ª ed. UNICAMP + RTIQ/MAPA)
Compare os valores declarados POR 100g com as faixas abaixo.
• Faixa TACO = valores analíticos reais (mín–máx esperado)
• Limite RTIQ = exigência legal (mínimo ou máximo)
• ✅ dentro da faixa | ⚠️ fora da faixa típica mas dentro do RTIQ | ❌ viola limite RTIQ
• ⚠️ COM RESSALVAS se os valores por 100g forem realmente impossíveis de inferir mesmo com esforço máximo

━━━ EMBUTIDOS CÁRNEOS ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
LINGUIÇA SUÍNA / PORCO (IN 4/2000):
• kcal: TACO 220-290 | Proteínas: mín 12g RTIQ, típico TACO 14-20g
• Gorduras totais: máx 30g RTIQ, típico TACO 17-25g | Gord.sat: típico TACO 6-9g
• Carboidratos: típico TACO 0-3g (>8g é incomum) | Sódio: típico TACO 700-1200mg

LINGUIÇA DE FRANGO (IN 4/2000):
• kcal: TACO 215-250 | Proteínas: mín 12g RTIQ, típico TACO 13-19g
• Gorduras totais: máx 30g RTIQ, típico TACO 15-22g
• Carboidratos: típico TACO 0-4g | Sódio: típico TACO 600-1100mg

LINGUIÇA BOVINA (IN 4/2000):
• kcal: TACO 200-270 | Proteínas: mín 12g RTIQ, típico TACO 14-20g
• Gorduras totais: máx 30g RTIQ, típico TACO 14-22g | Sódio: típico TACO 600-1000mg

SALSICHA (IN 4/2000):
• kcal: TACO 220-280 | Proteínas: mín 12g RTIQ, típico TACO 12-15g
• Gorduras totais: máx 30g RTIQ, típico TACO 18-26g | Gord.sat: típico TACO 6-9g
• Sódio: típico TACO 600-1100mg | Carboidratos: típico TACO 2-6g

MORTADELA (IN 4/2000):
• kcal: TACO 250-310 | Proteínas: mín 12g RTIQ, típico TACO 12-16g
• Gorduras totais: máx 30g RTIQ, típico TACO 20-28g | Gord.sat: típico TACO 8-11g
• Sódio: típico TACO 700-1200mg | Carboidratos: típico TACO 2-5g

HAMBÚRGUER BOVINO (IN 20/2000):
• kcal: TACO 200-260 | Proteínas: mín 15g RTIQ, típico TACO 13-20g
• Gorduras totais: máx 23g RTIQ, típico TACO 12-21g | Sódio: típico TACO 300-700mg
• Carboidratos: típico TACO 0-3g

HAMBÚRGUER DE FRANGO (IN 20/2000):
• kcal: TACO 170-230 | Proteínas: mín 15g RTIQ, típico TACO 14-20g
• Gorduras totais: máx 23g RTIQ, típico TACO 8-18g | Sódio: típico TACO 300-650mg

PRESUNTO COZIDO (IN 20/2000):
• kcal: TACO 120-160 | Proteínas: mín 14g RTIQ, típico TACO 17-21g
• Gorduras totais: máx 4g RTIQ presunto cozido, típico TACO 2-5g
• Sódio: típico TACO 900-1300mg | Carboidratos: típico TACO 0-2g

APRESUNTADO (IN 20/2000):
• kcal: TACO 150-200 | Proteínas: mín 14g RTIQ, típico TACO 15-19g
• Gorduras totais: máx 30g RTIQ apresuntado, típico TACO 6-15g | Sódio: típico TACO 900-1400mg

SALAME (IN 22/2000):
• kcal: TACO 340-420 | Proteínas: mín 20g RTIQ, típico TACO 22-28g
• Gorduras totais: máx 42g RTIQ, típico TACO 28-40g | Gord.sat: típico TACO 12-18g
• Sódio: típico TACO 1500-2400mg | Carboidratos: típico TACO 0-3g

COPA / COPPA (IN 22/2000):
• kcal: TACO 300-400 | Proteínas: mín 20g RTIQ, típico TACO 20-26g
• Gorduras totais: típico TACO 24-36g | Sódio: típico TACO 1400-2200mg

BACON / BARRIGA CURADA (IN 89/2003):
• kcal: TACO 380-520 | Proteínas: típico TACO 10-15g
• Gorduras totais: típico TACO 38-52g | Gord.sat: típico TACO 14-20g
• Sódio: típico TACO 1000-1600mg | Carboidratos: típico TACO 0-2g

CHARQUE / CARNE SECA (RIISPOA):
• kcal: TACO 250-330 | Proteínas: típico TACO 30-40g (concentrada pela desidratação)
• Gorduras totais: típico TACO 10-20g | Sódio: típico TACO 1500-3000mg (produto salgado)
• Carboidratos: típico TACO 0-1g

CMS — CARNE MECANICAMENTE SEPARADA DE AVES (IN 4/2000):
• kcal: TACO 200-260 | Proteínas: mín 12g RTIQ, típico TACO 12-16g
• Gorduras totais: máx 30g RTIQ, típico TACO 15-25g | Sódio: típico TACO 60-120mg (in natura)

━━━ LATICÍNIOS ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
LEITE PASTEURIZADO INTEGRAL (Port. 146/1996):
• kcal: TACO 59-63 | Proteínas: mín 2,9g RTIQ, típico TACO 2,9-3,2g
• Gorduras totais: mín 3g RTIQ integral, típico TACO 3,0-3,5g | Gord.sat: típico TACO 1,9-2,2g
• Carboidratos: típico TACO 4,6-5,0g | Sódio: típico TACO 40-55mg

LEITE PASTEURIZADO SEMIDESNATADO (Port. 146/1996):
• kcal: TACO 42-50 | Proteínas: mín 2,9g RTIQ, típico TACO 3,0-3,3g
• Gorduras totais: 0,6-2,9g RTIQ, típico TACO 1,0-2,0g | Sódio: típico TACO 40-55mg

LEITE PASTEURIZADO DESNATADO (Port. 146/1996):
• kcal: TACO 33-40 | Proteínas: mín 2,9g RTIQ, típico TACO 3,0-3,4g
• Gorduras totais: máx 0,5g RTIQ, típico TACO 0,1-0,4g | Sódio: típico TACO 40-55mg

LEITE EM PÓ INTEGRAL (Port. 146/1996):
• kcal: TACO 490-510 | Proteínas: mín 24g RTIQ, típico TACO 25-28g
• Gorduras totais: mín 26g RTIQ, típico TACO 26-30g | Carboidratos: típico TACO 37-42g
• Sódio: típico TACO 340-400mg

QUEIJO MINAS FRESCAL (Port. 146/1996 + IN 68/2019):
• kcal: TACO 255-275 | Proteínas: mín 17g RTIQ, típico TACO 17-20g
• Gorduras totais: máx 20g RTIQ (integral), típico TACO 17-22g | Gord.sat: típico TACO 10-14g
• Sódio: típico TACO 290-550mg | Carboidratos: típico TACO 2-4g

QUEIJO MUSSARELA (Port. 146/1996):
• kcal: TACO 295-320 | Proteínas: mín 18g RTIQ, típico TACO 20-25g
• Gorduras totais: típico TACO 20-26g | Gord.sat: típico TACO 12-17g
• Sódio: típico TACO 500-800mg | Carboidratos: típico TACO 1-4g

QUEIJO PRATO (Port. 146/1996):
• kcal: TACO 345-375 | Proteínas: típico TACO 24-28g
• Gorduras totais: típico TACO 26-32g | Gord.sat: típico TACO 16-20g
• Sódio: típico TACO 500-750mg | Carboidratos: típico TACO 1-3g

QUEIJO COALHO (Port. 146/1996):
• kcal: TACO 265-300 | Proteínas: típico TACO 20-24g
• Gorduras totais: típico TACO 18-24g | Sódio: típico TACO 700-1100mg

QUEIJO PARMESÃO (Port. 146/1996):
• kcal: TACO 440-470 | Proteínas: mín 32g RTIQ, típico TACO 33-38g
• Gorduras totais: típico TACO 30-37g | Sódio: típico TACO 1300-1800mg

RICOTA (Port. 146/1996):
• kcal: TACO 128-145 | Proteínas: típico TACO 11-15g
• Gorduras totais: típico TACO 7-10g | Sódio: típico TACO 150-300mg | Carboidratos: típico TACO 2-4g

REQUEIJÃO CREMOSO (Port. 146/1996):
• kcal: TACO 245-265 | Proteínas: típico TACO 10-13g
• Gorduras totais: típico TACO 20-24g | Sódio: típico TACO 400-700mg | Carboidratos: típico TACO 3-5g

IOGURTE NATURAL INTEGRAL (Port. 146/1996):
• kcal: TACO 58-66 | Proteínas: mín 2,9g RTIQ, típico TACO 3,0-4,0g
• Gorduras totais: típico TACO 2,5-3,5g | Sódio: típico TACO 40-60mg
• Carboidratos: típico TACO 4,5-6g (da lactose)

MANTEIGA (Port. 146/1996):
• kcal: TACO 720-750 | Proteínas: típico TACO 0,5-1,0g
• Gorduras totais: mín 80g RTIQ, típico TACO 80-85g | Gord.sat: típico TACO 50-60g
• Sódio: típico TACO 580-700mg (com sal) / 10-30mg (sem sal)

CREME DE LEITE (Port. 146/1996):
• kcal: TACO 260-300 | Proteínas: típico TACO 2,0-2,8g
• Gorduras totais: mín 20g RTIQ, típico TACO 25-35g | Sódio: típico TACO 30-60mg
• Carboidratos: típico TACO 3-5g

IOGURTE DESNATADO / SEMIDESNATADO (Port. 146/1996):
• Desnatado: gorduras totais máx 0,5g RTIQ | kcal: TACO 40-50
• Semidesnatado: gorduras totais 0,6-2,9g RTIQ | kcal: TACO 45-58
• Proteínas: mín 2,9g RTIQ (igual ao integral) | Sódio: TACO 40-60mg
• Carboidratos: TACO 5-7g | "Com adição de açúcar" = declarar açúcares adicionados

IOGURTE GREGO (Port. 146/1996 + definição de mercado):
• Sem definição RTIQ específica — verificar se produto declara % proteínas acima do padrão
• kcal: TACO 100-140 | Proteínas: típico 6-10g (maior que iogurte comum)
• Gorduras totais: TACO 6-10g (integral) | Sódio: TACO 50-70mg
• ATENÇÃO: "Grego" não é denominação legal protegida no Brasil — verificar se há alegação de proteínas

BEBIDA LÁCTEA (Port. 146/1996 + IN 16/2005):
• DIFERENÇA CRÍTICA: mín 51% leite na fórmula (≠ iogurte que é 100% leite)
• kcal: TACO 55-75 | Proteínas: mín 1,5g RTIQ, típico TACO 1,8-2,5g
• Gorduras totais: TACO 1,5-3,0g | Sódio: TACO 45-70mg
• DENOMINAÇÃO OBRIGATÓRIA: "Bebida Láctea" — nunca "iogurte" se <100% leite
• Verificar: produto com nome "iogurte" mas ingredientes incluindo soro de leite = possível fraude

KEFIR (Port. 146/1996 + definição técnica):
• Fermentado por grãos de kefir (bactérias + leveduras simbióticas)
• kcal: TACO 55-70 | Proteínas: TACO 3,0-4,0g | Gorduras: TACO 2,5-3,5g
• Sódio: TACO 40-60mg | Carboidratos: TACO 3,5-5,0g
• DENOMINAÇÃO: deve constar "Kefir" com cepas fermentadoras declaradas nos ingredientes
• Teor alcoólico residual natural (0,1-2%) — não caracteriza bebida alcoólica

COALHADA (Port. 146/1996):
• Fermentado por bactérias lácticas específicas (Lactococcus lactis)
• kcal: TACO 60-75 | Proteínas: TACO 3,0-4,0g | Gorduras totais: TACO 3,0-4,5g
• Sódio: TACO 40-65mg | Carboidratos: TACO 4,0-5,5g
• DENOMINAÇÃO: "Coalhada" ou "Coalhada Seca" (concentrada, menos soro)

SOBREMESA LÁCTEA / FLAN / PUDIM DE LEITE (Port. 146/1996):
• kcal: TACO 120-180 (varia muito com açúcar) | Proteínas: TACO 2,5-4,5g
• Gorduras totais: TACO 3,0-6,0g | Carboidratos: TACO 20-35g (alto — açúcar adicionado)
• Sódio: TACO 80-150mg
• ATENÇÃO: verificar açúcares adicionados vs açúcares naturais — alto % é esperado
• DENOMINAÇÃO: especificar sabor se aplicável ("Pudim de Baunilha", "Flan de Caramelo")

DOCE DE LEITE (Port. MAPA 354/1997):
• COMPOSIÇÃO MÍNIMA OBRIGATÓRIA:
  - Sólidos totais: mín 55% (produto pastoso) / mín 70% (produto em tablete/em pasta firme)
  - Sólidos de leite: mín 28% do produto final
  - Sacarose: deve ser declarada na lista de ingredientes
• kcal: TACO 295-330 | Proteínas: TACO 6-8g | Gorduras totais: TACO 6-9g
• Carboidratos: TACO 55-65g (alto — característico) | Sódio: TACO 80-180mg
• DENOMINAÇÕES ESPECÍFICAS:
  - "Doce de Leite" = produto puro
  - "Doce de Leite com Amendoim" = deve conter % declarado de amendoim
  - "Doce de Leite com Coco" = deve conter % declarado de coco ralado
  - "Dulce de Leche" = denominação equivalente aceita
• VERIFICAR: produto com gordura vegetal adicionada → obrigatório declarar
• LUPA: verificar açúcares adicionados ≥15g/100g → ALTO EM AÇÚCARES ADICIONADOS (quase sempre aciona)

DOCE DE LEITE PASTOSO EM BISNAGA / SACHÊ:
• Mesmas regras da Port. 354/1997
• Verificar se embalagem tem espaço reservado para lote e validade
• Conteúdo líquido: verificar se declarado em massa (g) — correto para produto pastoso

━━━ QUEIJOS FINOS (Port. MAPA 146/1996 + Portarias específicas) ━━━━━━━━━
EMMENTAL (Port. 146/1996):
• kcal: TACO 375-400 | Proteínas: mín 26g RTIQ, típico TACO 28-32g
• Gorduras totais: mín 43% na MS (matéria seca) = típico TACO 28-34g
• Sódio: TACO 200-450mg (menor que queijos mais salgados) | Carboidratos: TACO 1-3g
• CARACTERÍSTICA: olhos (buracos) regulares obrigatórios no padrão de identidade
• DENOMINAÇÃO: "Queijo Emmental" — "Tipo Emmental" indica produto fora do padrão

GRUYÈRE (Port. 146/1996):
• kcal: TACO 390-420 | Proteínas: mín 27g RTIQ, típico TACO 28-33g
• Gorduras totais: mín 45% na MS = típico TACO 30-36g | Sódio: TACO 300-600mg
• DIFERENÇA do Emmental: olhos menores e menos numerosos; sabor mais forte
• DENOMINAÇÃO: "Queijo Gruyère" — verificar se produto importado tem nome correto

BRIE (Port. 146/1996 + regulamento específico):
• kcal: TACO 310-350 | Proteínas: típico TACO 18-22g
• Gorduras totais: mín 45% na MS = típico TACO 25-30g | Sódio: TACO 400-700mg
• CARACTERÍSTICA: casca branca de mofo (Penicillium camemberti) — declaração obrigatória
• Carboidratos: TACO 1-2g (muito baixo — fermentação quase total da lactose)

CAMEMBERT (Port. 146/1996):
• kcal: TACO 295-330 | Proteínas: típico TACO 18-20g
• Gorduras totais: mín 45% na MS = típico TACO 22-28g | Sódio: TACO 500-800mg
• Similar ao Brie — casca de mofo branco obrigatória
• DIFERENÇA BRIE vs CAMEMBERT: Camembert menor (150g tipicamente), sabor mais intenso

EDAM (Port. 146/1996):
• kcal: TACO 340-370 | Proteínas: mín 25g RTIQ, típico TACO 25-30g
• Gorduras totais: máx 45% na MS = típico TACO 22-28g | Sódio: TACO 700-1000mg
• CARACTERÍSTICA: coberto com cera vermelha (importado) ou amarela
• Teor de gordura MENOR que Emmental/Gruyère — verificar se declaração está correta

GOUDA (Port. 146/1996):
• kcal: TACO 355-385 | Proteínas: típico TACO 23-27g
• Gorduras totais: típico TACO 27-32g | Sódio: TACO 600-900mg
• Carboidratos: TACO 1-3g | Maturação: jovem (4 sem) a curado (12+ meses) — kcal varia

━━━ LATICÍNIOS PROCESSADOS ESPECIAIS (V11) ━━━━━━━━━━━━━━━━━━━━━━━━
REQUEIJÃO CULINÁRIO / CATUPIRY (Port. 359/1997):
• kcal: TACO 240-280 | Proteínas: típico TACO 8-12g
• Gorduras totais: TACO 18-25g | Sódio: TACO 500-900mg
• Carboidratos: TACO 3-6g
• DENOMINAÇÃO: "Requeijão" = produto com processo específico. "Queijo processado" = diferente
• Verificar: "Catupiry" é marca — denominação legal é "Requeijão Cremoso" ou "Requeijão Culinário"

PETIT SUISSE (Port. 360/1997):
• kcal: TACO 90-140 | Proteínas: mín. 6g RTIQ, típico TACO 6-9g
• Gorduras totais: TACO 2-7g | Carboidratos: TACO 12-20g (com açúcar/polpa fruta)
• Sódio: TACO 50-80mg
• DENOMINAÇÃO: "Petit Suisse" — produto fermentado fresco concentrado com adição
• ATENÇÃO: alto em açúcares adicionados (polpa de fruta + sacarose) — verificar lupa

QUEIJO PROCESSADO / FUNDIDO (Port. 146/1996):
• kcal: TACO 280-340 | Proteínas: mín. 14g RTIQ, típico TACO 14-18g
• Gorduras totais: TACO 18-26g | Sódio: TACO 900-1400mg (alto — sais fundentes)
• Carboidratos: TACO 2-5g
• ATENÇÃO: sódio muito elevado — verificar lupa ALTO EM SÓDIO (quase sempre aciona)
• Sais fundentes (polifosfatos) são aditivos obrigatórios — verificar declaração nos ingredientes

CREME DE QUEIJO / CHEESE SPREAD:
• Similar ao queijo processado mas com mais creme/gordura adicionado
• kcal: TACO 290-360 | Gorduras totais: TACO 24-30g | Sódio: TACO 600-1000mg
• Verificar denominação: "Creme de Queijo", "Queijo Cremoso para Passar" etc.

LEITE FERMENTADO / YAKULT-TIPO (IN 46/2007):
• kcal: TACO 58-70 | Proteínas: TACO 1,5-2,5g
• Gorduras totais: TACO 0,1-1,0g (geralmente desnatado) | Carboidratos: TACO 12-18g
• Sódio: TACO 25-50mg
• ATENÇÃO: alto em carboidratos/açúcar — verificar lupa ALTO EM AÇÚCARES ADICIONADOS
• Cepas: Lactobacillus casei Shirota (Yakult), L. acidophilus (outros) — declaração obrigatória

CREME CHANTILLY / CHANTILLY (Port. 146/1996):
• kcal: TACO 300-360 | Proteínas: TACO 2-3g
• Gorduras totais: mín. 30g RTIQ, típico TACO 28-36g | Gord.sat: TACO 18-24g
• Sódio: TACO 20-50mg | Carboidratos: TACO 3-8g (com açúcar adicionado)
• ATENÇÃO: ALTO EM GORDURAS SATURADAS quase sempre aciona a lupa

━━━ PESCADOS ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
TILÁPIA / PEIXE MAGRO (Port. 185/1997):
• kcal: TACO 82-100 | Proteínas: típico TACO 18-22g
• Gorduras totais: típico TACO 1-3g | Sódio: típico TACO 50-120mg | Carboidratos: típico 0g

SALMÃO / PEIXE GORDO (Port. 185/1997):
• kcal: TACO 170-200 | Proteínas: típico TACO 18-22g
• Gorduras totais: típico TACO 8-14g | Gord.sat: típico TACO 2-4g
• Sódio: típico TACO 50-120mg | Carboidratos: típico 0g

ATUM EM CONSERVA — ÁGUA (Port. 185/1997):
• kcal: TACO 100-130 | Proteínas: típico TACO 22-28g
• Gorduras totais: típico TACO 1-4g | Sódio: típico TACO 300-500mg

ATUM EM CONSERVA — ÓLEO:
• kcal: TACO 180-220 | Proteínas: típico TACO 22-28g
• Gorduras totais: típico TACO 8-14g | Sódio: típico TACO 300-500mg

SARDINHA EM CONSERVA (Port. 185/1997):
• kcal: TACO 180-220 | Proteínas: típico TACO 20-26g
• Gorduras totais: típico TACO 8-14g | Sódio: típico TACO 350-600mg

CAMARÃO CRU (Port. 185/1997):
• kcal: TACO 85-100 | Proteínas: típico TACO 18-22g
• Gorduras totais: típico TACO 0,5-2g | Sódio: típico TACO 130-250mg

BACALHAU SECO SALGADO:
• kcal: TACO 340-380 | Proteínas: típico TACO 70-80g (concentrada)
• Gorduras totais: típico TACO 2-5g | Sódio: típico TACO 4000-8000mg (produto salgado)

━━━ CARNES IN NATURA ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
CARNE BOVINA MAGRA (patinho, coxão mole) crua (RIISPOA):
• kcal: TACO 125-160 | Proteínas: típico TACO 19-22g
• Gorduras totais: típico TACO 4-9g | Sódio: típico TACO 55-75mg | Carboidratos: típico 0g

CARNE BOVINA GORDA (costela, acém) crua:
• kcal: TACO 220-330 | Proteínas: típico TACO 14-18g
• Gorduras totais: típico TACO 18-30g | Sódio: típico TACO 55-75mg

FRANGO PEITO SEM PELE cru (RIISPOA):
• kcal: TACO 115-125 | Proteínas: típico TACO 20-23g
• Gorduras totais: típico TACO 2-4g | Sódio: típico TACO 50-70mg | Carboidratos: típico 0g

FRANGO COXA/SOBRECOXA SEM PELE cru:
• kcal: TACO 155-175 | Proteínas: típico TACO 16-19g
• Gorduras totais: típico TACO 8-13g | Sódio: típico TACO 65-90mg

CARNE SUÍNA PERNIL / LOMBO cru (RIISPOA):
• kcal: TACO 130-170 | Proteínas: típico TACO 17-21g
• Gorduras totais: típico TACO 6-14g | Sódio: típico TACO 50-75mg

━━━ PESCADOS NICHADOS ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
TILÁPIA (Oreochromis niloticus) — IN 29/2015:
• kcal: TACO 96-130 | Proteínas: mín. 15g RTIQ, TACO 19-22g
• Gorduras totais: TACO 2-6g | Sódio: TACO 40-80mg (in natura sem sal)
• Nome científico obrigatório no rótulo

SALMÃO CULTIVADO (Salmo salar / Oncorhynchus mykiss) — IN 29/2015:
• kcal: TACO 180-230 | Proteínas: mín. 15g RTIQ, TACO 18-22g
• Gorduras totais: TACO 9-16g (alto — normal para salmão) | Sódio: TACO 45-90mg
• Nome científico obrigatório — "Salmão do Atlântico - Salmo salar"

CAMARÃO (Litopenaeus vannamei) — IN 29/2015:
• kcal: TACO 85-115 | Proteínas: mín. 15g RTIQ, TACO 17-22g
• Gorduras totais: TACO 1-3g | Sódio: TACO 140-280mg natural; >600mg indica sal adicionado
• Crustáceo → ALÉRGENO obrigatório. Produto congelado: verificar % glaze declarado

OSTRAS (Crassostrea sp.) — IN 29/2015:
• kcal: TACO 55-80 | Proteínas: TACO 7-11g | Carbo: TACO 4-7g (glicogênio natural)
• Gorduras: TACO 1-3g | Sódio: TACO 200-400mg
• Molusco → ALÉRGENO obrigatório. Nome científico obrigatório.

━━━ AVES NICHADAS ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
CODORNA (Coturnix coturnix japonica):
• kcal: TACO 150-210 | Proteínas: mín. 12g RTIQ, TACO 18-24g
• Gorduras totais: TACO 8-15g | Sódio: TACO 50-110mg in natura
• Ovos de codorna: normas distintas dos ovos de galinha — verificar RTIQ específico

PATO (Anas platyrhynchos domesticus):
• kcal: TACO 200-280 (com pele) | Proteínas: mín. 12g RTIQ, TACO 16-20g
• Gorduras totais: TACO 15-25g (alto — normal para pato) | Sódio: TACO 60-120mg

━━━ BOVINOS E EXÓTICOS NICHADOS ━━━━━━━━━━━━━━━━━━━━━━━━━━━━
BÚFALO (Bubalus bubalis) — IN 83/2003:
• kcal: TACO 110-160 | Proteínas: mín. 15g RTIQ, TACO 20-26g
• Gorduras totais: TACO 2-8g (mais magro que bovino) | Sódio: TACO 50-90mg
• Denominação obrigatória: "Carne de Búfalo" — não pode ser chamada de "carne bovina"
• Mussarela de búfala: deve ser 100% leite de búfala

JAVALI (Sus scrofa):
• kcal: TACO 120-180 | Proteínas: TACO 18-24g | Gorduras: TACO 5-12g
• Embutidos de javali: seguem RTIQs dos embutidos suínos (IN 4/2000)

RÃ (Rana catesbeiana):
• kcal: TACO 70-95 | Proteínas: TACO 16-20g | Gorduras: TACO 0,5-2g
• Produto POA — exige registro no SIF/SIE/SIM. Nome científico obrigatório.

━━━ OVOS ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
OVO DE GALINHA inteiro cru (IN 29/2008):
• kcal: TACO 138-148 | Proteínas: típico TACO 12-14g
• Gorduras totais: típico TACO 8-11g | Gord.sat: típico TACO 2,5-3,5g
• Sódio: típico TACO 120-160mg | Carboidratos: típico TACO 0,5-1,5g

━━━ MEL E APÍCOLAS ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
MEL (IN 11/2000):
• kcal: TACO 295-320 | Carboidratos: mín 65g RTIQ, típico TACO 78-82g
• Proteínas: típico TACO 0,2-0,8g | Gorduras totais: típico TACO 0g
• Sódio: típico TACO 2-15mg — sódio alto em mel é sinal de adulteração
• Umidade: máx 20% RTIQ (importante para detectar mel adulterado diluído = kcal < 295)

━━━ PRODUTOS VEGETAIS ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
ARROZ BRANCO COZIDO (TACO / RDC 711/2022):
• kcal: TACO 128-135 | Carboidratos: típico TACO 28-30g
• Proteínas: típico TACO 2,5-3,0g | Gorduras totais: típico TACO 0-0,3g
• Fibras: típico TACO 1,5-2,5g | Sódio: típico TACO 0-5mg (sem sal)

ARROZ INTEGRAL COZIDO:
• kcal: TACO 124-132 | Carboidratos: típico TACO 25-28g
• Proteínas: típico TACO 2,5-3,0g | Fibras: típico TACO 1,8-3,0g

FEIJÃO COZIDO (TACO):
• kcal: TACO 76-86 | Carboidratos: típico TACO 13-16g
• Proteínas: típico TACO 4,5-5,5g | Fibras: típico TACO 8-11g
• Sódio: típico TACO 2-5mg (sem sal)

FARINHA DE TRIGO (TACO / RDC 711/2022):
• kcal: TACO 355-365 | Carboidratos: típico TACO 75-78g
• Proteínas: típico TACO 9-11g | Gorduras totais: típico TACO 1-2g
• Fibras: típico TACO 2-4g | Sódio: típico TACO 0-5mg

ÓLEO VEGETAL (soja, girassol, canola — RDC 714/2022):
• kcal: TACO 880-900 | Gorduras totais: típico TACO 99-100g
• Proteínas: típico TACO 0g | Carboidratos: típico TACO 0g
• Sódio: típico TACO 0mg — sódio em óleo = adulteração

AZEITE DE OLIVA:
• kcal: TACO 880-900 | Gorduras totais: típico TACO 99-100g
• Gord.sat: típico TACO 14-16g | Gord.mono: típico TACO 70-76g
• Sódio: típico TACO 0mg

AÇÚCAR CRISTAL/REFINADO (TACO / RDC 713/2022):
• kcal: TACO 387-400 | Carboidratos: típico TACO 99-100g
• Proteínas: típico TACO 0g | Gorduras totais: típico TACO 0g
• Sódio: típico TACO 0-2mg

SAL (cloreto de sódio puro):
• Sódio: ~39.000mg/100g (puro) — porção típica 1g → 390mg sódio
• Isento de tabela nutricional conforme ANVISA

━━━ INDUSTRIALIZADOS MISTOS ━━━━━━━━━━━━━━━━━━━━━━━━━━━
BISCOITO SALGADO / CRACKER (TACO):
• kcal: TACO 420-460 | Carboidratos: típico TACO 65-72g
• Proteínas: típico TACO 8-12g | Gorduras totais: típico TACO 12-20g
• Fibras: típico TACO 2-4g | Sódio: típico TACO 700-1200mg

BISCOITO DOCE / RECHEADO (TACO):
• kcal: TACO 450-490 | Carboidratos: típico TACO 68-75g
• Proteínas: típico TACO 5-8g | Gorduras totais: típico TACO 15-22g
• Açúcares totais: típico TACO 25-40g | Sódio: típico TACO 300-600mg

PÃO DE FORMA / FATIADO (TACO):
• kcal: TACO 255-275 | Carboidratos: típico TACO 45-52g
• Proteínas: típico TACO 7-10g | Gorduras totais: típico TACO 3-6g
• Fibras: típico TACO 2-4g | Sódio: típico TACO 450-700mg

MACARRÃO (MASSA SECA, CRU):
• kcal: TACO 368-376 | Carboidratos: típico TACO 73-77g
• Proteínas: típico TACO 10-12g | Gorduras totais: típico TACO 1-2g
• Fibras: típico TACO 2-4g | Sódio: típico TACO 0-10mg (sem sal)

CHOCOLATE AO LEITE (TACO):
• kcal: TACO 540-570 | Carboidratos: típico TACO 57-62g
• Proteínas: típico TACO 7-9g | Gorduras totais: típico TACO 30-36g
• Gord.sat: típico TACO 16-22g | Açúcares totais: típico TACO 48-56g
• Sódio: típico TACO 80-150mg

CHOCOLATE MEIO AMARGO / AMARGO (TACO):
• kcal: TACO 530-565 | Carboidratos: típico TACO 45-55g
• Proteínas: típico TACO 4-6g | Gorduras totais: típico TACO 35-42g
• Açúcares totais: típico TACO 28-42g | Sódio: típico TACO 5-30mg

MOLHO DE TOMATE INDUSTRIALIZADO (TACO):
• kcal: TACO 35-55 | Carboidratos: típico TACO 6-10g
• Proteínas: típico TACO 1-2g | Gorduras totais: típico TACO 0,5-2g
• Sódio: típico TACO 400-700mg — acima de 800mg é alto

MAIONESE TRADICIONAL (TACO):
• kcal: TACO 295-320 | Gorduras totais: típico TACO 30-33g
• Proteínas: típico TACO 1-2g | Carboidratos: típico TACO 2-4g
• Sódio: típico TACO 500-800mg

MARGARINA (TACO):
• kcal: TACO 720-750 | Gorduras totais: mín 80g típico TACO 80-85g
• Gord.sat: típico TACO 18-28g | Gord.trans: verificar (deve ser declarado)
• Sódio: típico TACO 400-700mg (com sal)

LEITE EM PÓ / ACHOCOLATADO EM PÓ (TACO):
• kcal: TACO 380-420 | Carboidratos: típico TACO 72-78g (achocolatado)
• Proteínas: típico TACO 5-8g | Gorduras totais: típico TACO 4-7g
• Sódio: típico TACO 180-300mg

GRANOLA / CEREAL MATINAL (TACO):
• kcal: TACO 390-430 | Carboidratos: típico TACO 65-72g
• Proteínas: típico TACO 8-12g | Gorduras totais: típico TACO 8-14g
• Fibras: típico TACO 5-9g | Sódio: típico TACO 100-300mg

━━━ BEBIDAS ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
SUCO DE LARANJA 100% (TACO / IN MAPA 49/2018):
• kcal: TACO 43-50 | Carboidratos: típico TACO 9-12g
• Proteínas: típico TACO 0,5-1g | Gorduras totais: típico TACO 0g
• Sódio: típico TACO 0-5mg — sódio alto = adulteração

SUCO DE UVA 100%:
• kcal: TACO 60-72 | Carboidratos: típico TACO 14-18g
• Açúcares totais: típico TACO 14-17g | Sódio: típico TACO 0-5mg

NÉCTAR (frutas, mín 30-50% suco dependendo da fruta):
• kcal: TACO 40-65 | Carboidratos: típico TACO 9-16g (pode ter açúcar adicionado)
• Proteínas: típico TACO 0-0,5g | Sódio: típico TACO 0-30mg

REFRIGERANTE (cola, guaraná, laranja):
• kcal: TACO 38-45 (normal) / 0-2 (zero/light)
• Carboidratos: típico TACO 9-11g (normal) / 0-0,5g (zero)
• Sódio: típico TACO 10-30mg

ÁGUA MINERAL (Decreto 7841/2012):
• kcal: TACO 0 | Carboidratos: TACO 0g | Sódio: variável por fonte
• Água com sódio >200mg/L deve declarar "ALTO TEOR DE SÓDIO"

CERVEJA (IN MAPA 14/2018):
• kcal: TACO 43-48 (lata/garrafa) | Carboidratos: típico TACO 3-4g
• Proteínas: típico TACO 0,3-0,5g | Álcool: 4-5% v/v (regular)
• Sódio: típico TACO 10-20mg | Fibras: 0g

VINHO TINTO SECO (Lei 7678/1988):
• kcal: TACO 70-85 | Carboidratos: típico TACO 2-4g (seco)
• Álcool: 11-14% v/v | Sódio: típico TACO 5-10mg

━━━ SUPLEMENTOS ALIMENTARES (RDC 243/2018) ━━━━━━━━━━━
PROTEÍNA WHEY (concentrado/isolado):
• Proteínas: min 10g/porção RTIQ, típico 20-30g/30-40g porção
• Carboidratos: concentrado típico 3-8g | isolado típico 0-3g
• Gorduras totais: concentrado típico 2-6g | isolado típico 0-2g
• Sódio: típico 50-200mg/porção

PROTEÍNA VEGETAL (soja, ervilha):
• Proteínas: típico 18-25g/porção | Carboidratos: típico 3-8g
• Gorduras totais: típico 1-4g | Sódio: típico 100-300mg/porção

CREATINA MONOIDRATADA:
• Porção: 3-5g | Proteínas: 0g | Carboidratos: 0g | Gorduras: 0g
• Sódio: 0mg — qualquer nutriente acima de 0 é incomum em creatina pura

VITAMINAS/MINERAIS (complexo multivitamínico):
• Porção muito variável | Verificar se % VD está declarado
• Vitaminas lipossolúveis (A, D, E, K): podem ser perigosas em excesso — verificar se % VD < 300%

━━━ PRODUTOS VEGETAIS E INDUSTRIALIZADOS — GAPS POA ━━━━━
━━━ PRATOS PRONTOS E ELABORADOS POA ━━━━━━━━━━━━━━━━━━
⚠️ ATENÇÃO: Em produtos cárneos elaborados (empanados, lasanha, pratos prontos),
carboidratos altos (>8g/100g) são NORMAIS e esperados — indicam farinha e amido adicionados.
Não alertar como inconsistência. Verificar % VD na tabela.

EMPANADO DE FRANGO (nugget, steak, filé):
• kcal: TACO 230-290 | Proteínas: mín 10g RTIQ (IN 6/2001), típico TACO 13-18g
• Gorduras totais: típico TACO 12-18g | Carboidratos: típico TACO 14-22g (normal — farinha)
• Sódio: típico TACO 500-950mg

EMPANADO BOVINO:
• kcal: TACO 220-270 | Proteínas: mín 10g RTIQ, típico TACO 12-17g
• Gorduras totais: típico TACO 11-17g | Carboidratos: típico TACO 12-20g
• Sódio: típico TACO 500-900mg

NUGGET DE FRANGO (reestruturado):
• kcal: TACO 240-300 | Proteínas: mín 10g RTIQ, típico TACO 12-16g
• Gorduras totais: típico TACO 14-20g | Carboidratos: típico TACO 12-20g
• Sódio: típico TACO 500-1000mg

LASANHA COM CARNE BOVINA/FRANGO (congelada):
• kcal: TACO 100-160 | Proteínas: típico TACO 6-10g
• Gorduras totais: típico TACO 4-9g | Carboidratos: típico TACO 12-20g
• Sódio: típico TACO 400-700mg

QUIBE / KIBE INDUSTRIALIZADO:
• kcal: TACO 180-250 | Proteínas: mín 12g, típico TACO 12-17g
• Gorduras totais: típico TACO 10-18g | Carboidratos: típico TACO 8-14g
• Sódio: típico TACO 350-700mg

CROQUETE DE CARNE:
• kcal: TACO 200-280 | Proteínas: típico TACO 8-13g
• Gorduras totais: típico TACO 10-18g | Carboidratos: típico TACO 15-25g
• Sódio: típico TACO 400-800mg

━━━ ALIMENTOS FUNCIONAIS DE ORIGEM ANIMAL ━━━━━━━━━━━━
LEITE ENRIQUECIDO (cálcio, vitaminas, ômega-3):
• kcal: TACO 55-70 (similar ao leite base) | Proteínas: típico TACO 3,0-3,5g
• Verificar se nutrientes objeto de alegação estão DECLARADOS NA TABELA (obrigatório RDC 267/2003)
• Vitaminas A/D: em µg ou UI | Cálcio: em mg com % VD

OVO ENRIQUECIDO (ômega-3, vitamina E):
• kcal/composição: similar ao ovo padrão (TACO 138-148 kcal)
• Verificar se EPA+DHA ou vitamina E declarados na tabela se forem objeto de claim

IOGURTE COM PROBIÓTICOS:
• kcal/composição: similar ao iogurte padrão (TACO 58-70 kcal)
• Verificar: alegação "contribui para equilíbrio da flora intestinal" é aprovada pela ANVISA
• NÃO aceito: "fortalece imunidade", "previne infecções" — proibidos pela RDC 267/2003

LEITE FERMENTADO FUNCIONAL:
• kcal: TACO 55-80 | Proteínas: típico TACO 3,0-4,0g
• Verificar UFC/porção se declarado — mín. 10⁸ UFC/porção para alegação probiótica (ANVISA)

━━━ PRODUTOS CÁRNEOS — GAPS POA ━━━━━━━━━━━━━━━━━━━━━━━
PATÊ (IN 20/2000 + Almôndega/Kibe):
• kcal: TACO 180-280 | Proteínas: mín 8g típico TACO 10-18g
• Gorduras totais: típico TACO 14-25g | Sódio: típico TACO 600-1100mg
• Carboidratos: típico TACO 2-8g (pode ter amido/farinha)

ALMÔNDEGA / QUIBE / KIBE (IN 20/2000):
• kcal: TACO 170-250 | Proteínas: típico TACO 12-18g
• Gorduras totais: típico TACO 10-20g | Sódio: típico TACO 400-800mg
• Carboidratos: típico TACO 4-12g (quibe tem trigo)

FIAMBRE (Port. 706/2022):
• kcal: TACO 130-180 | Proteínas: mín 14g RTIQ, típico TACO 14-18g
• Gorduras totais: típico TACO 5-12g | Sódio: típico TACO 900-1300mg

JERKED BEEF / CARNE CURADA (IN 22/2000):
• kcal: TACO 220-290 | Proteínas: típico TACO 25-35g
• Gorduras totais: típico TACO 8-16g | Sódio: típico TACO 1000-2000mg
• Diferença do charque: menor teor de sal (processo de cura vs salga)

CARNE DE SOL (RIISPOA):
• kcal: TACO 200-260 | Proteínas: típico TACO 22-30g
• Gorduras totais: típico TACO 8-15g | Sódio: típico TACO 800-1800mg

FRANGO TEMPERADO / MARINADO (IN 17/2018):
• kcal: TACO 130-200 | Proteínas: típico TACO 16-21g
• Gorduras totais: típico TACO 5-15g | Sódio: típico TACO 400-900mg (sal + tempero)

CORNED BEEF (IN 83/2003):
• kcal: TACO 220-280 | Proteínas: típico TACO 18-24g
• Gorduras totais: típico TACO 12-20g | Sódio: típico TACO 900-1400mg

━━━ LATICÍNIOS — GAPS POA ━━━━━━━━━━━━━━━━━━━━━━━━━━━━
QUEIJO SUÍÇO / GRUYÈRE / EMMENTAL (Port. 146/1996):
• kcal: TACO 370-410 | Proteínas: típico TACO 26-30g
• Gorduras totais: típico TACO 28-34g | Gord.sat: típico TACO 18-22g
• Sódio: típico TACO 200-400mg (menor que queijos brasileiros)

QUEIJO GOUDA / EDAM (Port. 146/1996):
• kcal: TACO 340-380 | Proteínas: típico TACO 24-28g
• Gorduras totais: típico TACO 27-32g | Sódio: típico TACO 600-900mg

DOCE DE LEITE (Port. 354/1997):
• kcal: TACO 295-330 | Carboidratos: típico TACO 54-60g
• Proteínas: típico TACO 6-8g | Gorduras totais: típico TACO 6-10g
• Sódio: típico TACO 80-160mg

LEITE CONDENSADO (IN 47/2018):
• kcal: TACO 320-350 | Carboidratos: típico TACO 54-58g
• Proteínas: mín 6,5g RTIQ, típico TACO 7-9g
• Gorduras totais: mín 8g RTIQ, típico TACO 8-12g | Sódio: típico TACO 100-160mg

BEBIDA LÁCTEA (IN 16/2005):
• kcal: TACO 55-80 | Proteínas: mín 1,2g RTIQ, típico TACO 1,5-3g
• Gorduras totais: típico TACO 1,5-3,5g | Sódio: típico TACO 40-80mg
• Carboidratos: típico TACO 8-14g (pode ter açúcar adicionado)

COMPOSTO LÁCTEO (IN 28/2007):
• kcal: TACO 400-450 | Proteínas: mín 10g RTIQ, típico TACO 12-18g
• Gorduras totais: típico TACO 10-18g | Carboidratos: típico TACO 55-65g
• Sódio: típico TACO 200-400mg

CREME AZEDO / SOUR CREAM (IN 23/2012):
• kcal: TACO 195-230 | Proteínas: típico TACO 3-4g
• Gorduras totais: mín 17g RTIQ, típico TACO 18-24g
• Sódio: típico TACO 30-80mg | Carboidratos: típico TACO 3-5g

━━━ CORTES BOVINOS ESPECÍFICOS ━━━━━━━━━━━━━━━━━━━━━━━━
ALCATRA / CONTRA-FILÉ / FILÉ MIGNON (cortes nobres) crus:
• kcal: TACO 120-160 | Proteínas: típico TACO 20-23g
• Gorduras totais: típico TACO 3-8g | Sódio: típico TACO 50-65mg | Carboidratos: típico 0g

PICANHA crua:
• kcal: TACO 155-200 | Proteínas: típico TACO 18-21g
• Gorduras totais: típico TACO 9-15g (capa de gordura) | Sódio: típico TACO 55-70mg

MÚSCULO / ACÉM / PALETA BOVINA crus (cortes de segunda):
• kcal: TACO 110-145 | Proteínas: típico TACO 20-23g
• Gorduras totais: típico TACO 2-6g | Sódio: típico TACO 55-75mg

COSTELA BOVINA crua:
• kcal: TACO 240-320 | Proteínas: típico TACO 14-18g
• Gorduras totais: típico TACO 18-28g | Sódio: típico TACO 60-80mg

COXÃO DURO / COXÃO MOLE / PATINHO / FRALDINHA / MAMINHA / LAGARTO / CUPIM crus:
• kcal: TACO 115-160 | Proteínas: típico TACO 19-23g
• Gorduras totais: típico TACO 3-9g | Sódio: típico TACO 50-70mg

━━━ CORTES SUÍNOS ESPECÍFICOS ━━━━━━━━━━━━━━━━━━━━━━━
BISTECA SUÍNA / COSTELINHA crua:
• kcal: TACO 155-220 | Proteínas: típico TACO 16-20g
• Gorduras totais: típico TACO 9-16g | Sódio: típico TACO 55-75mg

LOMBO SUÍNO cru:
• kcal: TACO 130-160 | Proteínas: típico TACO 19-22g
• Gorduras totais: típico TACO 5-10g | Sódio: típico TACO 50-70mg

PALETA SUÍNA crua:
• kcal: TACO 135-170 | Proteínas: típico TACO 17-21g
• Gorduras totais: típico TACO 7-13g | Sódio: típico TACO 55-75mg

TOUCINHO cru:
• kcal: TACO 560-650 | Proteínas: típico TACO 4-8g
• Gorduras totais: típico TACO 60-72g | Gord.sat: típico TACO 22-28g | Sódio: típico TACO 40-60mg

━━━ AVES — CORTES E MIÚDOS ━━━━━━━━━━━━━━━━━━━━━━━━━
FRANGO ASA crua:
• kcal: TACO 175-220 | Proteínas: típico TACO 17-20g
• Gorduras totais: típico TACO 10-15g | Sódio: típico TACO 65-90mg

FRANGO SOBRECOXA COM PELE crua:
• kcal: TACO 250-275 | Proteínas: típico TACO 15-18g
• Gorduras totais: típico TACO 18-23g | Sódio: típico TACO 65-90mg

FRANGO FÍGADO cru:
• kcal: TACO 95-115 | Proteínas: típico TACO 17-20g
• Gorduras totais: típico TACO 3-5g | Carboidratos: típico TACO 1-3g (glicogênio)
• Sódio: típico TACO 65-90mg | Rico em vitamina A e ferro

FRANGO CORAÇÃO cru:
• kcal: TACO 145-175 | Proteínas: típico TACO 16-20g
• Gorduras totais: típico TACO 7-12g | Sódio: típico TACO 65-90mg

FRANGO MOELA crua:
• kcal: TACO 85-110 | Proteínas: típico TACO 17-21g
• Gorduras totais: típico TACO 1-4g | Sódio: típico TACO 60-85mg

PERU PEITO cru:
• kcal: TACO 100-120 | Proteínas: típico TACO 22-25g
• Gorduras totais: típico TACO 1-3g | Sódio: típico TACO 55-75mg

PATO cru:
• kcal: TACO 190-240 | Proteínas: típico TACO 16-20g
• Gorduras totais: típico TACO 13-20g | Sódio: típico TACO 65-85mg

━━━ MIÚDOS / VÍSCERAS BOVINAS ━━━━━━━━━━━━━━━━━━━━━━━
FÍGADO BOVINO cru:
• kcal: TACO 130-145 | Proteínas: típico TACO 19-22g
• Gorduras totais: típico TACO 3-5g | Carboidratos: típico TACO 2-4g (glicogênio)
• Sódio: típico TACO 65-85mg | ⚠️ Carboidratos > 0g é normal em fígado (glicogênio)

CORAÇÃO BOVINO cru:
• kcal: TACO 110-135 | Proteínas: típico TACO 17-21g
• Gorduras totais: típico TACO 4-7g | Sódio: típico TACO 80-110mg

RIM BOVINO cru:
• kcal: TACO 95-115 | Proteínas: típico TACO 16-20g
• Gorduras totais: típico TACO 3-6g | Sódio: típico TACO 180-250mg (naturalmente alto)

LÍNGUA BOVINA crua:
• kcal: TACO 195-230 | Proteínas: típico TACO 14-18g
• Gorduras totais: típico TACO 13-19g | Sódio: típico TACO 65-85mg

━━━ PESCADOS REGIONAIS BRASILEIROS ━━━━━━━━━━━━━━━━━━
PEIXE MAGRO REGIONAL — Corvina / Pescada / Linguado / Badejo:
• kcal: TACO 75-100 | Proteínas: típico TACO 17-21g
• Gorduras totais: típico TACO 0,5-2g | Sódio: típico TACO 50-120mg

PEIXE MÉDIO REGIONAL — Dourado / Robalo / Tainha / Surubim / Pintado:
• kcal: TACO 90-130 | Proteínas: típico TACO 17-21g
• Gorduras totais: típico TACO 2-6g | Sódio: típico TACO 50-120mg

PEIXE REGIONAL AMAZÔNICO — Tambaqui / Pirarucu / Pacu / Tucunaré / Traíra:
• kcal: TACO 90-140 | Proteínas: típico TACO 17-22g
• Gorduras totais: típico TACO 2-8g | Sódio: típico TACO 50-120mg
• ⚠️ Pirarucu tem proteína alta (22-26g) — produto muito magro

CAÇÃO / TUBARÃO:
• kcal: TACO 80-100 | Proteínas: típico TACO 18-22g
• Gorduras totais: típico TACO 0,5-2g | Sódio: típico TACO 60-130mg

GAROUPA / BADEJO (peixes nobres):
• kcal: TACO 80-100 | Proteínas: típico TACO 18-22g
• Gorduras totais: típico TACO 0,5-2g | Sódio: típico TACO 50-100mg

MERLUZA:
• kcal: TACO 70-90 | Proteínas: típico TACO 16-20g
• Gorduras totais: típico TACO 0,5-1,5g | Sódio: típico TACO 50-110mg

━━━ MOLUSCOS E CRUSTÁCEOS ━━━━━━━━━━━━━━━━━━━━━━━━━
MEXILHÃO cru (Port. 1022/2024):
• kcal: TACO 70-90 | Proteínas: típico TACO 11-14g
• Gorduras totais: típico TACO 1,5-3g | Carboidratos: típico TACO 2-4g
• Sódio: típico TACO 280-400mg | ⚠️ Carboidratos > 0g é normal em moluscos

OSTRA crua (Port. 1022/2024):
• kcal: TACO 65-85 | Proteínas: típico TACO 8-12g
• Gorduras totais: típico TACO 1-3g | Carboidratos: típico TACO 3-5g
• Sódio: típico TACO 400-600mg

LULA crua (Port. 1022/2024):
• kcal: TACO 75-95 | Proteínas: típico TACO 15-18g
• Gorduras totais: típico TACO 1-2g | Sódio: típico TACO 200-350mg

POLVO cru (Port. 1022/2024):
• kcal: TACO 75-95 | Proteínas: típico TACO 14-18g
• Gorduras totais: típico TACO 0,5-2g | Sódio: típico TACO 230-380mg

━━━ OVOS ESPECÍFICOS ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
OVO DE CODORNA cru (Port. 01/2020):
• kcal: TACO 155-170 | Proteínas: típico TACO 13-15g
• Gorduras totais: típico TACO 11-13g | Gord.sat: típico TACO 3-4g
• Sódio: típico TACO 130-180mg | Carboidratos: típico TACO 0,5-1g

CLARA DE OVO crua:
• kcal: TACO 45-55 | Proteínas: típico TACO 10-12g
• Gorduras totais: típico TACO 0g | Sódio: típico TACO 160-200mg
• Carboidratos: típico TACO 0,5-1g

GEMA DE OVO crua:
• kcal: TACO 320-360 | Proteínas: típico TACO 15-17g
• Gorduras totais: típico TACO 28-32g | Gord.sat: típico TACO 9-11g
• Sódio: típico TACO 45-65mg

━━━ LATICÍNIOS ESPECÍFICOS ━━━━━━━━━━━━━━━━━━━━━━━━━
QUEIJO PROVOLONE (IN 73/2020):
• kcal: TACO 355-385 | Proteínas: mín 20g RTIQ, típico TACO 25-30g
• Gorduras totais: típico TACO 28-35g | Gord.sat: típico TACO 18-22g
• Sódio: típico TACO 800-1100mg | Carboidratos: típico TACO 0-2g

QUEIJO BRIE / CAMEMBERT (Port. 146/1996):
• kcal: TACO 295-335 | Proteínas: típico TACO 17-22g
• Gorduras totais: típico TACO 23-28g | Gord.sat: típico TACO 14-18g
• Sódio: típico TACO 400-600mg

KEFIR (IN 46/2007):
• kcal: TACO 50-70 | Proteínas: típico TACO 3-4g
• Gorduras totais: típico TACO 1-4g | Sódio: típico TACO 40-60mg
• Carboidratos: típico TACO 4-6g (menor que iogurte — fermentação consome lactose)

IOGURTE GREGO / GREGO ESTILO (Port. 146/1996):
• kcal: TACO 90-140 | Proteínas: típico TACO 6-10g (concentrado, mais proteína)
• Gorduras totais: típico TACO 5-10g | Sódio: típico TACO 40-70mg
• Carboidratos: típico TACO 3-7g

LEITE DE CABRA (Port. 146/1996):
• kcal: TACO 65-75 | Proteínas: mín 2,8g RTIQ, típico TACO 3,0-3,5g
• Gorduras totais: típico TACO 3,5-4,5g | Sódio: típico TACO 40-55mg
• Carboidratos: típico TACO 4,4-4,8g

QUEIJO PARMESÃO RALADO:
• kcal: TACO 450-480 | Proteínas: típico TACO 35-42g
• Gorduras totais: típico TACO 28-35g | Sódio: típico TACO 1400-1900mg
• ⚠️ Mais concentrado que bloco por conta da desidratação

━━━ PRODUTOS APÍCOLAS ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
PRÓPOLIS EXTRATO LÍQUIDO / TINTURA (IN 03/2001):
• kcal: tipicamente 150-300 (varia com concentração de álcool/propóleos)
• Proteínas: traços | Gorduras: traços | Carboidratos: 0-5g por 100mL
• ⚠️ Produto com alto teor alcoólico — verificar se declara "contém álcool"
• Sódio: típico <10mg

GELEIA REAL (IN 03/2001):
• kcal: TACO 130-170 | Proteínas: típico TACO 12-16g
• Gorduras totais: típico TACO 5-8g | Carboidratos: típico TACO 8-14g
• Sódio: típico TACO 50-100mg

PÓLEN APÍCOLA (IN 03/2001):
• kcal: TACO 340-400 | Proteínas: típico TACO 20-28g
• Carboidratos: típico TACO 40-55g | Gorduras totais: típico TACO 4-8g
• Sódio: típico TACO 30-80mg

━━━ CONSIDERAÇÕES PARA PRODUTOS TEMPERADOS/PROCESSADOS ━━
Para produtos temperados (IN 17/2018) os valores de sódio podem ser 10-40% maiores que
o corte in natura pelo sal e condimentos adicionados. Considere isso na avaliação.
Carboidratos acima de 10g/100g em embutidos puros (sem amido declarado) = incomum.
Fibra alimentar acima de 0g/100g em carnes puras in natura = incomum (indica erro ou vegetal).
⚠️ EXCEÇÕES onde Carboidratos > 0g são NORMAIS: fígado (glicogênio), moluscos, ovos (traços).
⚠️ Rim bovino tem sódio naturalmente alto (~200mg) — não confundir com sal adicionado.

━━━ PRODUTOS VEGETAIS — CEREAIS E DERIVADOS ━━━━━━━━━━━━
ARROZ BRANCO cru (TACO):
• kcal: 358-368 | Carboidratos: típico 79-82g | Proteínas: típico 7-9g
• Gorduras totais: típico 0,5-1g | Sódio: típico 1-5mg | Fibra: típico 1,5-2,5g

ARROZ INTEGRAL cru:
• kcal: 355-368 | Carboidratos: típico 75-78g | Proteínas: típico 7-9g
• Gorduras totais: típico 1,5-3g | Fibra: típico 4-6g

FARINHA DE TRIGO:
• kcal: 355-365 | Carboidratos: típico 75-80g | Proteínas: típico 9-12g
• Gorduras totais: típico 1-2g | Fibra: típico 2-3g | Sódio: típico 1-5mg

PÃO FRANCÊS / DE FORMA:
• kcal: 245-285 | Carboidratos: típico 50-58g | Proteínas: típico 7-10g
• Gorduras totais: típico 2-5g | Sódio: típico 400-700mg | Fibra: típico 2-4g

MACARRÃO (massa seca):
• kcal: 365-380 | Carboidratos: típico 75-80g | Proteínas: típico 10-13g
• Gorduras totais: típico 1-2g | Sódio: típico 5-15mg (sem sal) / 300-600mg (com sal)

BISCOITO SALGADO / CRACKER:
• kcal: 380-430 | Carboidratos: típico 60-70g | Proteínas: típico 8-12g
• Gorduras totais: típico 12-20g | Sódio: típico 500-900mg

BISCOITO DOCE / COOKIE:
• kcal: 430-480 | Carboidratos: típico 65-72g | Proteínas: típico 5-8g
• Gorduras totais: típico 18-26g | Açúcares: típico 25-40g | Sódio: típico 200-500mg

GRANOLA / CEREAL MATINAL:
• kcal: 380-430 | Carboidratos: típico 60-70g | Proteínas: típico 8-12g
• Gorduras totais: típico 10-18g | Fibra: típico 5-9g | Açúcares: típico 15-30g

━━━ ÓLEOS, GORDURAS E MARGARINAS ━━━━━━━━━━━━━━━━━━━━━
ÓLEO VEGETAL (soja, girassol, canola, milho):
• kcal: 882-900 | Gorduras totais: 100g | Proteínas: 0g | Carboidratos: 0g
• Gord.sat: soja 15g | girassol 10g | canola 7g | milho 13g | Sódio: 0mg
• ⚠️ Óleo com proteínas ou carboidratos declarados = erro ou adulteração

AZEITE DE OLIVA:
• kcal: 884-900 | Gorduras totais: 100g | Gord.sat: típico 14g | Gord.mono: típico 73g
• Proteínas: 0g | Carboidratos: 0g | Sódio: 0mg

MARGARINA:
• kcal: 540-720 | Gorduras totais: 60-80g | Gord.sat: típico 20-35g
• Trans: máx 2g/dia recomendado ANVISA | Sódio: típico 350-600mg

━━━ AÇÚCAR E PRODUTOS AÇUCARADOS ━━━━━━━━━━━━━━━━━━━━
AÇÚCAR REFINADO / CRISTAL:
• kcal: 385-390 | Carboidratos: 99-100g (quase tudo sacarose)
• Proteínas: 0g | Gorduras: 0g | Sódio: 0-2mg

CHOCOLATE AO LEITE:
• kcal: 520-560 | Carboidratos: típico 57-62g | Açúcares: típico 50-55g
• Gorduras totais: típico 28-33g | Gord.sat: típico 16-20g
• Proteínas: típico 7-9g | Sódio: típico 80-120mg

CHOCOLATE AMARGO (>70% cacau):
• kcal: 550-600 | Carboidratos: típico 32-45g | Açúcares: típico 20-35g
• Gorduras totais: típico 38-45g | Proteínas: típico 8-12g | Sódio: típico 10-30mg

SORVETE (base leite):
• kcal: 180-250 | Carboidratos: típico 25-35g | Açúcares: típico 20-28g
• Gorduras totais: típico 8-14g | Proteínas: típico 3-5g | Sódio: típico 60-100mg

━━━ CONDIMENTOS, MOLHOS E TEMPEROS ━━━━━━━━━━━━━━━━━━
KETCHUP:
• kcal: 90-115 | Carboidratos: típico 20-27g | Açúcares: típico 15-22g
• Gorduras totais: típico 0g | Sódio: típico 900-1300mg

MAIONESE:
• kcal: 280-320 | Gorduras totais: típico 30-35g | Proteínas: típico 1-2g
• Carboidratos: típico 2-5g | Sódio: típico 500-800mg

SHOYU / MOLHO DE SOJA:
• kcal: 50-70 | Proteínas: típico 5-8g | Carboidratos: típico 6-10g
• Gorduras: típico 0g | Sódio: típico 4000-7000mg ⚠️ MUITO ALTO em sódio — normal

━━━ BEBIDAS ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
SUCO DE LARANJA (100% integral):
• kcal: 40-48 | Carboidratos: típico 9-11g | Açúcares: típico 8-10g
• Proteínas: típico 0,5-1g | Gorduras: 0g | Sódio: típico 0-5mg
• ⚠️ Suco com sódio >20mg/100mL indica adição irregular

SUCO DE UVA INTEGRAL:
• kcal: 65-75 | Carboidratos: típico 15-18g | Açúcares: típico 14-17g
• Proteínas: típico 0,5g | Gorduras: 0g | Sódio: típico 3-10mg

NÉCTAR DE FRUTA (>30% polpa):
• kcal: 40-65 | Carboidratos: típico 10-16g (com açúcar adicionado)
• Proteínas: típico 0-0,5g | Sódio: típico 5-30mg

REFRIGERANTE REGULAR:
• kcal: 36-45/100mL | Carboidratos: típico 9-12g | Açúcares: típico 9-12g
• Gorduras: 0g | Proteínas: 0g | Sódio: típico 5-20mg

REFRIGERANTE DIET/ZERO:
• kcal: 0-5/100mL | Carboidratos: típico 0g | Adoçantes: verificar declaração

CERVEJA:
• kcal: 40-50/100mL | Carboidratos: típico 3-5g | Proteínas: típico 0,3-0,5g
• Gorduras: 0g | Sódio: típico 3-12mg | Teor alcoólico: declarar em %vol

VINHO TINTO/BRANCO:
• kcal: 68-85/100mL | Carboidratos: típico 2-4g (seco) / 5-10g (suave)
• Gorduras: 0g | Sódio: típico 5-15mg | Teor alcoólico: declarar em %vol

━━━ SUPLEMENTOS ALIMENTARES (RDC 243/2018) ━━━━━━━━━
WHEY PROTEIN CONCENTRADO (70-80% prot):
• kcal: 350-400/100g | Proteínas: típico 70-80g | Carboidratos: típico 8-15g
• Gorduras totais: típico 5-8g | Sódio: típico 200-500mg

WHEY PROTEIN ISOLADO (>90% prot):
• kcal: 350-380/100g | Proteínas: típico 88-95g | Carboidratos: típico 2-5g
• Gorduras totais: típico 0,5-2g | Sódio: típico 150-400mg

CREATINA MONOHIDRATADA:
• kcal: 0 | Proteínas: 0g | Carboidratos: 0g | Gorduras: 0g
• ⚠️ Creatina com calorias declaradas = erro ou produto adulterado

━━━ LEGUMINOSAS E GRÃOS ━━━━━━━━━━━━━━━━━━━━━━━━━━━━
FEIJÃO CARIOCA / PRETO cru:
• kcal: 340-355 | Carboidratos: típico 60-65g | Proteínas: típico 20-23g
• Gorduras totais: típico 1-2g | Fibra: típico 15-20g | Sódio: típico 5-10mg

GRÃO DE BICO cru:
• kcal: 360-380 | Carboidratos: típico 60-65g | Proteínas: típico 18-22g
• Gorduras totais: típico 5-7g | Fibra: típico 15-18g

Formato de saída para o Campo 9 (USE SEMPRE AS 3 LINHAS):
✅/❌/⚠️ Nível 1 — Estrutura: [resultado detalhado]
✅/❌/⚠️ Nível 2 — Coerência matemática: (Prot×4) + (Carb×4) + (Gord×9) = Xkcal esperado | declarado: Ykcal | variação: Z% | [resultado]
✅/❌/⚠️ Nível 3 — Plausibilidade [categoria detectada]: [nutriente]: valor declarado | faixa típica TACO: min-max | limite RTIQ/ANVISA: X | [resultado]

─────────────────────────────────────────────
CAMPO 10 — ROTULAGEM NUTRICIONAL FRONTAL (LUPA)
─────────────────────────────────────────────
LOCALIZAÇÃO: lupa deve aparecer no PAINEL PRINCIPAL (face frontal da embalagem).
Se a lupa estiver no verso ou lateral → ❌ NÃO CONFORME (RDC 429/2020, Art. 3°)

Lupa OBRIGATÓRIA se (por 100g):
• Açúcares adicionados ≥ 15g → "ALTO EM AÇÚCARES ADICIONADOS"
• Gorduras saturadas ≥ 6g → "ALTO EM GORDURAS SATURADAS"
• Sódio ≥ 600mg → "ALTO EM SÓDIO"

Verificar:
a) Lupa preta com texto em branco no PAINEL PRINCIPAL
b) Texto exato: "ALTO EM [NUTRIENTE]"
c) Se valores não visíveis claramente: registrar como ⚠️ COM RESSALVAS — solicitar ao RT confirmação visual da lupa.

─────────────────────────────────────────────
CAMPO 11 — DECLARAÇÃO DE ALÉRGENOS (RDC 727/2022, Art. 13-15)
─────────────────────────────────────────────
FORMATO OBRIGATÓRIO (Art. 15 RDC 727/2022):
a) CAIXA ALTA obrigatório
b) NEGRITO obrigatório
c) Cor CONTRASTANTE com o fundo do rótulo (a norma não exige cor específica)
d) Altura mínima de 2mm (ou 1mm se área do painel ≤100cm²)
e) Posicionada IMEDIATAMENTE após ou abaixo da lista de ingredientes
⚠️ IMPORTANTE: Fundo amarelo NÃO é exigência legal — a cor pode ser qualquer uma desde que contraste. Não penalize por ausência de amarelo.

TEXTO OBRIGATÓRIO:
• "ALÉRGICOS: CONTÉM [nome do alimento]" → para ingredientes INTENCIONAIS (presentes na fórmula)
• "ALÉRGICOS: PODE CONTER [nome do alimento]" → SOMENTE para contaminação cruzada (não intencional)
⚠️ ERRO COMUM: usar "PODE CONTER" para ingrediente declarado na lista → isso é NÃO CONFORME pois
   induz o consumidor a acreditar que o alérgeno pode ou não estar presente quando está sempre presente.

DECLARAÇÃO INCLUI DERIVADOS — exemplos corretos:
• Soja como ingrediente → "ALÉRGICOS: CONTÉM SOJA E DERIVADOS DE SOJA"
• Leite como ingrediente → "ALÉRGICOS: CONTÉM LEITE E DERIVADOS DE LEITE"  
• Trigo como ingrediente → "ALÉRGICOS: CONTÉM TRIGO E DERIVADOS DE TRIGO"
• Múltiplos alérgenos → "ALÉRGICOS: CONTÉM SOJA E DERIVADOS DE SOJA, LEITE E DERIVADOS DE LEITE"

PRINCIPAIS ALÉRGENOS — verificar os presentes nos ingredientes:
Trigo/centeio/cevada/aveia e derivados | Crustáceos e derivados | Ovos e derivados
Peixes e derivados | Amendoim e derivados | Soja e derivados | Leite e derivados
Oleaginosas (amêndoa, castanha, nozes, etc.) e derivados | Aipo | Mostarda
Gergelim | Sulfitos (>10mg/kg) | Tremoço | Moluscos

Regras especiais:
• "CONTÉM LEITE E DERIVADOS" obrigatório mesmo em laticínios
• "CONTÉM PEIXE E DERIVADOS" obrigatório mesmo em produtos de peixe
• CMS de aves: declarar "CONTÉM CMS DE AVES" separadamente (não é alérgeno mas exigência MAPA)
• Contaminação cruzada: "PODE CONTER" apenas se houver risco real documentado no PCA

─────────────────────────────────────────────
CAMPO 12 — DECLARAÇÃO DE TRANSGÊNICOS (Decreto 4.680/2003)
─────────────────────────────────────────────
a) Se ingrediente OGM >1%: símbolo triângulo amarelo "T" obrigatório (mín. 4mm)
b) Texto: "[ingrediente] transgênico" ou "Contém X% de [ingrediente] transgênico"
c) Atenção: soja, milho frequentemente transgênicos em embutidos
d) Se nenhum OGM >1%: omissão ou "Não contém transgênicos" (ambos corretos)

CAMPO 13 — LOTE E PRAZO DE VALIDADE (RDC 727/2022 Art. 9)
─────────────────────────────────────────────
a) LOTE: indicado como "Lote:", "L:", "Lot:" ou referencia: "Veja fundo", "Veja Tampa"
b) PRAZO DE VALIDADE: "Validade:", "Val:", "Vencimento:" ou referencia ao local
   Formato aceito: DD/MM/AA, MM/AA, DD/MM/AAAA
c) Arte deve indicar onde lote/validade serao impressos: "Veja fundo" = CONFORME
d) Ausencia total de qualquer indicacao -> NAO CONFORME
e) Somente "Veja fundo" sem data visivel -> CONFORME (local indicado)

CAMPO 14 — PORCAO PADRAO (IN 75/2020 Anexo I)
─────────────────────────────────────────────
Analisa EXCLUSIVAMENTE se a porcao declarada segue IN 75/2020 Anexo I.
(Valores nutricionais sao avaliados no Campo 9)

Porcoes de referencia (IN 75/2020):
CARNES in natura: 100g | EMBUTIDOS fatiados: 30g | EMBUTIDOS inteiros: 80g
PESCADO fresco: 100g | PESCADO enlatado: 30g
LEITE: 200mL | IOGURTE: 100g | QUEIJO: 30g | REQUEIJAO: 30g | MANTEIGA: 10g
OVOS: 50g (1 und) | MEL: 20g
PAO/BOLO: 50g | BISCOITO salgado: 30g | BISCOITO doce/recheado: 30g
MACARRAO cru: 80g | MACARRAO cozido: 140g
SUCO INTEGRAL/NECTAR: 200mL | REFRIGERANTE: 350mL | ENERGETICO: 200mL
CHOCOLATE: 30g | SORVETE: 60g | OLEO/AZEITE: 13mL | MAIONESE: 15g | KETCHUP: 10g
SUPLEMENTO EM PO: dose do fabricante (sem padrao obrigatorio IN 75/2020)

Variacao aceitavel: +-20% do valor de referencia
Porcao diverge >20%: COM RESSALVAS com sugestao de ajuste
Categoria sem porcao definida na IN 75/2020: CONFORME (fabricante define)

CAMPO 15 — ALEGAÇÕES FUNCIONAIS/NUTRICIONAIS (RDC 18/1999 + RDC 54/2012 + RDC 727/2022 Art.17)
─────────────────────────────────────────────
CAMPO CONDICIONAL — avalie APENAS se o rótulo contém alguma alegação funcional ou nutricional.
Se nenhuma alegação detectada: declare "CAMPO 15 — N/A (sem alegações no rótulo)" e NÃO inclua no score.
Se alegações presentes: inclua no score (total passa a ser X/15).

ALEGAÇÕES NUTRICIONAIS (RDC 54/2012) — critérios obrigatórios:
• "Fonte de [nutriente]": deve conter ≥15% da IDR por porção
• "Rico em [nutriente]": deve conter ≥30% da IDR por porção
• "Alto teor de [nutriente]": deve conter ≥30% da IDR por porção (sólido) ou ≥7,5% (líquido)
• "Reduzido em sódio": deve ter ≥25% menos sódio que o produto de referência
• "Baixo sódio": ≤40mg/100g (sólido) ou ≤20mg/100mL (líquido)
• "Sem sódio": ≤5mg/100g
• "Light": redução mínima de 25% em calorias OU nutriente específico vs referência
• "Zero açúcar / Sem açúcar": ≤0,5g/porção
• "Zero lactose": ≤0,1g lactose/100g ou 100mL (RDC 715/2022 Art. 9)
• "Integral" (cereais/farinhas): mín. 50% ingredientes integrais (RDC 712/2022)

ALEGAÇÕES FUNCIONAIS (RDC 18/1999) — apenas se constar na lista ANVISA aprovada:
• Probióticos (Lactobacillus, Bifidobacterium): dose mínima e texto aprovado pela ANVISA
• Ômega-3: mín. 0,3g EPA+DHA por porção para alegação cardiovascular
• Fibras (psyllium, beta-glucana): dose mínima conforme evidência científica
• Verificar: texto da alegação é EXATAMENTE o aprovado pela ANVISA?

VERIFICAÇÃO OBRIGATÓRIA quando há alegação:
1. O critério quantitativo da RDC 54/2012 é atendido pelos valores da tabela nutricional?
2. O texto da alegação segue o modelo aprovado pela ANVISA (sem personalizar)?
3. Não há alegação proibida (cura, tratamento, prevenção de doenças)?
4. Claims de "natural", "orgânico" têm certificação correspondente no rótulo?

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
PASSO 4 — RELATÓRIO FINAL

NÃO escreva título "VALIDAÇÃO COMPLETA DE RÓTULO" nem separadores horizontais (━━━).
NÃO repita PRODUTO, ÓRGÃO, SCORE numérico ou VEREDICTO no início — esses dados já são extraídos automaticamente do conteúdo do Passo 1 e dos campos.

Comece DIRETAMENTE com o título "PASSO 1 — IDENTIFICAÇÃO DO PRODUTO" e siga com o conteúdo dos passos.

🔴🔴🔴 ESTRUTURA RÍGIDA E OBRIGATÓRIA DO PASSO 4 — VOCÊ DEVE COPIAR EXATAMENTE 🔴🔴🔴

O Passo 4 SEMPRE terá EXATAMENTE 3 seções nesta ordem (sem exceção, sem variação, sem improvisação):

  SEÇÃO 1 — "CONFORMIDADES DETECTADAS:"
  SEÇÃO 2 — "NÃO CONFORMIDADES CRÍTICAS:"
  SEÇÃO 3 — "RESSALVAS:"

Não importa se o produto tem zero não conformidades ou zero ressalvas — você DEVE escrever as 3 seções (com a frase "Nenhuma identificada nesta análise." se não houver itens). NUNCA omita nenhuma das 3 seções. NUNCA crie outras seções. NUNCA renomeie elas.

🚫 SEÇÕES PROIBIDAS — não crie nenhuma destas:
  - "PONTOS FORTES" (use "CONFORMIDADES DETECTADAS")
  - "NÃO CONFORMIDADES CRÍTICAS" pode existir, mas NÃO crie "RESSALVAS IMPORTANTES" (use só "RESSALVAS")
  - "RECOMENDAÇÕES AO RT" — proibido
  - "PRIORIDADE DE CORREÇÃO" — proibido
  - "OBSERVAÇÕES CRÍTICAS" — proibido
  - "PRÓXIMOS PASSOS" — proibido
  - Qualquer outra seção que não seja as 3 acima — proibido

EXEMPLO OBRIGATÓRIO — copie ESTA estrutura literalmente, mudando apenas o conteúdo:

```
# PASSO 4 — RELATÓRIO FINAL

CONFORMIDADES DETECTADAS:
✓ Denominação clara e específica conforme RTIQ
✓ Lista de ingredientes completa com aditivos identificados (INS + função)
✓ Carimbo SIF presente e legível

NÃO CONFORMIDADES CRÍTICAS:
✗ Endereço do fabricante INCOMPLETO — falta CEP obrigatório (RDC 727/2022 Art. 8°)
✗ Declaração de ALÉRGENOS AUSENTE — produto contém leite e soja (RDC 727/2022 Art. 14-16)

RESSALVAS:
⚠ Tabela nutricional com valores próximos ao limite — RT confirmar análise laboratorial
⚠ Lupa "ALTO EM SÓDIO" no painel principal recomendada (sódio 700mg/100g ≥ 600mg)

> ⚠️ Aviso legal: Este relatório é gerado por inteligência artificial e tem caráter auxiliar. Não substitui a análise e aprovação de Responsável Técnico habilitado. A conformidade regulatória final é de responsabilidade exclusiva do RT inscrito no órgão competente.
```

⚠️ COMPLETUDE OBRIGATÓRIA:
- Em "CONFORMIDADES DETECTADAS" liste TODOS os campos C1–C14 marcados como ✅ (use 1 linha por item, comece com ✓).
- Em "NÃO CONFORMIDADES CRÍTICAS" liste TODOS os campos marcados como ❌ (1 linha por item, comece com ✗, cite a norma violada).
- Em "RESSALVAS" liste TODOS os campos marcados como ⚠️ (1 linha por item, comece com ⚠).
- Se uma seção não tem itens, escreva exatamente: "Nenhuma identificada nesta análise."
- NUNCA trunque o relatório. Você tem orçamento de tokens suficiente — escreva tudo.

REGRAS GERAIS:
• Logo no início do relatório (antes do PASSO 1), você DEVE escrever uma linha "SCORE: X.X/14 | N conformes + N com ressalvas + N não conformes" e uma linha "VEREDICTO: [APROVADO | APROVADO COM RESSALVAS | REPROVADO]". Essas linhas são PARSEADAS pelo sistema. Escreva APENAS uma vez, no início.
• Score 3 níveis: ✅ CONFORME (1.0pt) | ⚠️ COM RESSALVAS (0.5pt) | ❌ NÃO CONFORME (0pt).

🔢 REGRA MATEMÁTICA OBRIGATÓRIA — VERIFIQUE ANTES DE ESCREVER O SCORE FINAL:
ANTES de escrever a linha "SCORE: X.X/14...", FAÇA EXPLICITAMENTE em sua mente esta soma:
1. CONTE quantos campos (de C1 a C14) você marcou como ✅ CONFORME → chame esse número de CONFORMES
2. CONTE quantos campos você marcou como ⚠️ COM RESSALVAS → chame esse número de RESSALVAS
3. CONTE quantos campos você marcou como ❌ NÃO CONFORME ou ⚠️ AUSENTE → chame esse número de NAO_CONFORMES
4. VALIDAÇÃO: CONFORMES + RESSALVAS + NAO_CONFORMES DEVE = 14. Se não bater, REVISE sua contagem.
5. SCORE = CONFORMES × 1.0 + RESSALVAS × 0.5 + NAO_CONFORMES × 0
6. Escreva: "SCORE: [SCORE calculado]/14 | [CONFORMES] conformes + [RESSALVAS] com ressalvas + [NAO_CONFORMES] não conformes"

EXEMPLO CORRETO: Se você marcou 9 ✅, 1 ⚠️, 4 ❌, então:
SCORE = 9 × 1.0 + 1 × 0.5 + 4 × 0 = 9.5/14
Linha: "SCORE: 9.5/14 | 9 conformes + 1 com ressalvas + 4 não conformes"

⚠️ PROIBIDO escrever números na linha SCORE diferentes do que você marcou nos campos individuais. O sistema VERIFICA e contradição quebra a confiança do RT no relatório.

• APROVADO: score ≥ 13.0/14 sem não conformidade crítica.
• APROVADO COM RESSALVAS: score 8.0–12.5/14, corrigíveis.
• REPROVADO: score < 8.0/14 OU sem carimbo/registro/denominação/alérgenos/tabela nutricional.

REGRA CRÍTICA — A OPÇÃO "NÃO VERIFICÁVEL" FOI REMOVIDA DO SISTEMA:
Você tem capacidade visual avançada. INSPECIONE atentamente cada região da imagem em detalhes. Aproxime-se mentalmente do texto pequeno. Use OCR contextual para inferir letras parcialmente legíveis. NUNCA use "🔍 NÃO VERIFICÁVEL" — esta opção foi removida. Sempre forneça veredicto: ✅ CONFORME, ⚠️ COM RESSALVAS, ou ❌ NÃO CONFORME. Se o texto é pequeno mas perceptível — leia, mesmo com esforço. Se realmente não conseguir ler com precisão, use ⚠️ COM RESSALVAS com nota "RT, confirmar visualmente — talvez a imagem não esteja tão nítida nessa região."

REGRA ESPECIAL PARA C11 (ALÉRGENOS) E C12 (TRANSGÊNICOS):
JAMAIS marque C11 ou C12 como "🔍 NÃO VERIFICÁVEL" — esta opção foi removida. Para esses campos, você SEMPRE consegue diagnosticar a partir da lista de ingredientes (Campo 2):

• C11 — ALÉRGENOS: Cruze a lista de ingredientes com os 14 grupos da RDC 727/2022 (trigo, centeio, cevada, aveia, leite, ovos, soja, peixes, crustáceos, moluscos, amendoim, castanhas, látex, sulfitos).
  - Se há alérgeno na lista E há declaração "ALÉRGICOS: CONTÉM..." no rótulo → ✅ CONFORME
  - Se há alérgeno na lista E NÃO há declaração no rótulo → ❌ NÃO CONFORME (violação grave)
  - Se NÃO há alérgenos detectados na lista E há frase "Não contém alérgenos" → ✅ CONFORME
  - Se NÃO há alérgenos detectados na lista E NÃO há nenhuma declaração → ⚠ COM RESSALVAS (boa prática declarar)

• C12 — TRANSGÊNICOS: Identifique candidatos OGM comuns na lista de ingredientes: soja e derivados (proteína isolada, lecitina, óleo de soja), milho e derivados (maltodextrina, amido de milho, fubá, óleo de milho).
  - Se há candidatos OGM E há declaração "Contém [ingrediente] transgênico" com símbolo "T" amarelo → ✅ CONFORME
  - Se há candidatos OGM E há declaração "Não contém transgênicos" / "Não contém OGM" → ✅ CONFORME
  - Se há candidatos OGM E NÃO há nenhuma declaração → ⚠ COM RESSALVAS (RT deve verificar com fornecedor se >1% OGM; se sim, incluir símbolo "T" + frase "Contém [ingrediente] transgênico" conforme Decreto 4.680/2003 Art. 2° e 3°)
  - Se NÃO há candidatos OGM detectados → ✅ CONFORME (declaração não obrigatória)

Em todos os casos do sistema, o veredicto deve ser sempre ✅, ⚠️ ou ❌ — NUNCA 🔍.

═══════════════════════════════════════════════
FORMATO DOS PASSOS 1, 2, 3 — SEJA SUCINTO
═══════════════════════════════════════════════

PASSO 1 — IDENTIFICAÇÃO DO PRODUTO

PRODUTO: [nome completo, ≤80 chars]
TIPO: POA ou Não-POA
ESPÉCIE: [se POA: bovina, suína, aves, etc.]
CATEGORIA: [classificação legal: embutido, laticínio, suplemento, etc.]
ÓRGÃO: [SIF/SIE/SIM ou ANVISA] — carimbo/registro [número se visível]
RTIQ: [norma técnica aplicável]

PASSO 2 — LEGISLAÇÕES APLICÁVEIS

Liste APENAS as normas BASE relevantes (máximo 8) precedidas de "▸ ":
▸ RDC 727/2022 (ANVISA) — rotulagem geral
▸ RDC 429/2020 + IN 75/2020 — rotulagem nutricional
▸ INMETRO Port. 249/2021 — conteúdo líquido
▸ [outras conforme o produto]

Não liste sub-normas detalhadas. Não escreva justificativas.

PASSO 3 — VALIDAÇÃO DOS 14 CAMPOS OBRIGATÓRIOS

Para CADA campo (1 a 14), use EXATAMENTE este formato compacto:

**CAMPO X — [NOME]: [✅ CONFORME | ⚠️ COM RESSALVAS | ❌ NÃO CONFORME] ([1.0|0.5|0] pt)**

[Lista de 3-6 bullets MÁXIMO, cada um com ícone próprio:]
✓ [item correto observado no rótulo]
✓ [outro item correto]
⚠ [item parcialmente correto / com ressalva]
✗ [item ausente ou errado]

[Se houver bullets ⚠ ou ✗, ADICIONE seção de sugestões logo abaixo:]
Sugestões de correção:
⚠ [Repete o item com ressalva] — Norma: [RDC X/ANO Art. Y°]. Texto correto: "[exato]"
✗ [Repete o item errado] — Norma: [RDC X/ANO Art. Y°]. Texto correto: "[exato]"

NÃO escreva "O que está correto:" nem "O que precisa ajuste:". Use apenas os bullets unificados.
NÃO escreva bloco "CORREÇÕES PRIORITÁRIAS" no final.
NÃO escreva linha "🎯 PRIORIDADE" — descontinuado.
NÃO repita normas em múltiplas linhas — use 1 linha por norma.

REGRA PARA SUB-BULLETS (listas dentro de itens detalhados, ex: tabela nutricional, endereço fabricante):
• Use ✓ para sub-itens corretos: "  ✓ Logradouro: RUA X"
• Use ⚠ para sub-itens com ressalva: "  ⚠ Bairro: incompleto"
• Use ✗ para sub-itens ausentes/errados: "  ✗ CEP: faltando"
• NÃO use traços "-" nem setas "▸" para sub-itens — sempre o ícone unificado correspondente.

═══════════════════════════════════════════════
FIM DO RELATÓRIO — APENAS ESTAS 2 LINHAS:
═══════════════════════════════════════════════

SCORE: [X.X]/14 | [N] conformes + [N] com ressalvas + [N] não conformes
VEREDICTO: [APROVADO | APROVADO COM RESSALVAS | REPROVADO]

Não adicione justificativa, conclusão, agradecimento, ou qualquer texto após o VEREDICTO."""


SP_REVISAO = """Você é um auditor sênior de rotulagem com 20 anos de experiência em MAPA/DIPOA e ANVISA.

RELATÓRIO A REVISAR:
{relatorio}

Revise criticamente o relatório acima. Foque APENAS em erros reais — não repita o que já está correto.

CHECKLIST UNIVERSAL (aplica a TODOS os produtos — POA e não-POA):
0. SCORE 3 NÍVEIS: cada campo tem ✅ CONFORME (1.0) | ⚠️ COM RESSALVAS (0.5) | ❌ NÃO CONFORME (0)? Score total correto em X.X/14?
1. Todos os 14 campos foram avaliados com nota individual? (1-Denominação, 2-Ingredientes, 3-Conteúdo líquido, 4-Fabricante, 5-Glúten, 6-Lactose, 7-Conservação, 8-Carimbo/Registro, 9-Tabela nutricional, 10-Lupa, 11-Alérgenos, 12-Transgênicos, 13-Lote/Validade, 14-Porção padrão). Se rótulo contém alegações funcionais/nutricionais: Campo 15 foi avaliado?
2. DENOMINAÇÃO: composição mínima obrigatória foi verificada para o tipo de produto? Claims LIGHT/DIET/ZERO/INTEGRAL tiveram critério conferido?
3. INGREDIENTES: cruzamento ↔ alérgenos foi feito? Aditivos com INS e função tecnológica declarados?
4. FABRICANTE: todos os 6 elementos (logradouro, número, bairro, CEP, cidade, UF) verificados individualmente?
5. TABELA NUTRICIONAL: %VD conferido? Porção padrão IN 75/2020 correta para a categoria? Trans com limiar verificado?
6. ALÉRGENOS: 14 grupos RDC 727/2022? Formatação (caixa alta + negrito + 2mm mínimo)?
7. TEXTO DE CORREÇÃO: para cada NÃO CONFORME, o relatório sugere o texto exato de como corrigir?
8. NORMAS CITADAS: cada não conformidade cita OBRIGATORIAMENTE número + ano + artigo + parágrafo? (ex: RDC 727/2022 Art. 8° §1° — não apenas 'RDC 727/2022'). Citar só o número da norma sem artigo é insuficiente.
9. SUB-BULLETS: listas dentro de campos (tabela nutricional, fabricante, etc.) usam ✓/⚠/✗ em vez de "-" ou "▸"?

CHECKLIST POA (só se produto for de origem animal):
10. CARIMBO: SIF/SIE/SIM com jurisdição coerente? Formato oval correto? Número do estabelecimento?
11. RTIQ: denominação segue exatamente o RTIQ aplicável (nome técnico + espécie)?
12. CMS: se presente, espécie e % declarados?
13. PESCADO: nome científico verificado? Forma de apresentação declarada?
14. MEL: pureza e ausência de aditivos verificada? Origem floral/geográfica se declarada?

CHECKLIST NÃO-POA (só se produto vegetalIndustrializado/bebida/suplemento):
10. REGISTRO/NOTIFICAÇÃO: número de notificação ANVISA presente se obrigatório (suplementos, bebidas especiais)?
11. % OBRIGATÓRIOS: suco integral = 100% fruta? Néctar com % suco declarado? Maionese com % lipídios? Chocolate com % cacau? "Integral" com mín. 50% farinha integral?
12. PORÇÃO IN 75/2020: porção específica da categoria foi verificada (suco=200mL, biscoito=30g, pão=50g, etc.)?
13. LUTAS ESPECÍFICAS: sódio ≥600mg/100g? Gordura saturada ≥6g/100g? Açúcar adicionado ≥15g/100g?
14. ADITIVOS PROIBIDOS: conservantes em suco integral? Corantes não autorizados? Cafeína ≥210mg/porção em energético?

Se identificar qualquer erro ou omissão no relatório, adicione ao final:
## ADENDO DO AUDITOR:
[descreva apenas os erros/omissões encontrados — não repita o que já está correto]
Se o relatório estiver completo e correto, responda apenas: "✅ RELATÓRIO VALIDADO — sem adendos necessários." """


# ═══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════════════════════════════════════
async def call_claude_simple(system: str, user: str, max_tokens: int = 350) -> str:
    payload = {
        "model": ANTHROPIC_MODEL,
        "max_tokens": max_tokens,
        "temperature": 0,
        "system": system,
        "messages": [{"role": "user", "content": user}],
    }
    headers = {
        "x-api-key": ANTHROPIC_API_KEY,
        "anthropic-version": "2023-06-01",
        "content-type": "application/json",
    }
    async with httpx.AsyncClient(timeout=60.0) as client:
        r = await client.post("https://api.anthropic.com/v1/messages", json=payload, headers=headers)
        if r.status_code != 200:
            return ""
        return r.json().get("content", [{}])[0].get("text", "")

def extrair_score(texto: str):
    """
    Extrai score do relatório — suporta /14 (padrão) e /15 (com Campo 15 de alegações).
    Retorna float para preservar valores como 9.5, 10.5, etc.

    Regex tolerante: aceita variações como "SCORE: 9/14", "**SCORE:** 9.5/14",
    "Score 9 / 14", "SCORE FINAL: 9/14", etc.
    """
    if not texto:
        return None

    # Regex unificada e tolerante:
    # - opcional: ** ou markdown ao redor
    # - SCORE seguido opcionalmente por palavras (FINAL, GERAL)
    # - separador flexível (: ou espaço)
    # - número com . ou , decimal
    # - / com espaços opcionais antes e depois
    # - 14 ou 15 como denominador
    pattern = r"(?:\*\*)?\s*SCORE(?:\s+(?:FINAL|GERAL|TOTAL))?(?:\*\*)?[:\s]+\s*\*?\*?\s*([\d]+(?:[.,][\d]+)?)\s*/\s*(14|15)"
    m = re.search(pattern, texto, re.IGNORECASE)
    if m:
        try:
            return float(m.group(1).replace(",", "."))
        except (ValueError, TypeError):
            pass

    # Fallback legado para compatibilidade — se Sonnet 4.6 escrever em formato totalmente novo
    legacy_patterns = [
        r"SCORE[:\s]+([\d]+[.,][\d]*)\s*/\s*15",
        r"SCORE[:\s]+(\d+)\s*/\s*15",
        r"SCORE[:\s]+([\d]+[.,][\d]*)\s*/\s*14",
        r"SCORE[:\s]+(\d+)\s*/\s*14",
    ]
    for pat in legacy_patterns:
        m = re.search(pat, texto, re.IGNORECASE)
        if m:
            try:
                return float(m.group(1).replace(",", "."))
            except (ValueError, TypeError):
                continue

    # SALVAGUARDA: parser falhou — log Sentry para alerta proativo
    # Isso é importante após troca de modelo: se Sonnet 4.6 mudar formato,
    # você é avisado em vez de descobrir só quando revisão crítica não disparar.
    try:
        # Guard: só tenta logar se sentry_sdk está disponível (pode não ter sido importado)
        if "sentry_sdk" in dir() or "sentry_sdk" in globals():
            if os.environ.get("SENTRY_DSN") and len(texto) > 200:
                preview = texto[:300].replace("\n", " ")
                sentry_sdk.capture_message(
                    f"[score_parser] Falhou ao extrair score. Modelo: {ANTHROPIC_MODEL}. Preview: {preview}",
                    level="warning"
                )
    except Exception:
        pass  # nunca deixar Sentry quebrar o fluxo

    return None

def extrair_veredicto(texto: str) -> str:
    m = re.search(r"VEREDICTO[:\s]+(APROVADO COM RESSALVAS|APROVADO|REPROVADO)", texto, re.IGNORECASE)
    return m.group(1).upper() if m else "NÃO IDENTIFICADO"

def detectar_status_campo(texto: str, campo: int) -> str:
    """Detecta status de um campo no relatório — suporta 3 níveis."""
    nome = CAMPOS_NOME.get(campo, "")
    linhas = texto.split("\n")
    for i, linha in enumerate(linhas):
        if f"CAMPO {campo}" in linha.upper() or (nome and nome.upper() in linha.upper()):
            trecho = " ".join(linhas[i:i+5]).upper()
            if "NÃO CONFORME" in trecho or "NAO CONFORME" in trecho or "AUSENTE" in trecho:
                return "NAO_CONFORME"
            elif "COM RESSALVAS" in trecho:
                return "COM_RESSALVAS"
            elif "CONFORME" in trecho:
                return "CONFORME"
    return "NAO_DETECTADO"


def extrair_campos_status(texto: str) -> dict:
    """Extrai status de todos os 14 campos de um relatório."""
    return {c: detectar_status_campo(texto, c) for c in range(1, 15)}


# ═══════════════════════════════════════════════════════════════════════════════
# ENDPOINT: VALIDAR
# ═══════════════════════════════════════════════════════════════════════════════
SP_DETECT = """Você é um sistema de identificação de rótulos POA. Analise a imagem e retorne APENAS um JSON válido, sem texto antes ou depois.

Retorne:
{
  "produto": "nome exato conforme aparece no rótulo",
  "categoria_kb": "categoria_da_lista",
  "especie": "bovino|suino|frango|peru|pato|pescado|ovino|caprino|bubalino|abelha|galinha|misto",
  "orgao": "SIF|SIE|SIM|nao_identificado",
  "numero_registro": "número do carimbo ou vazio",
  "sigla_sie": "sigla estadual se SIE (ex: SISP, SIE-MG) ou vazio",
  "confianca": "alta|media|baixa"
}

Lista de categorias_kb válidas:
embutidos, salame, presunto, hamburguer, bacon, carne_moida, charque, fiambre,
carne_maturada, carnes_temperadas, almondega_kibe, gelatina_colageno,
apresuntado, paleta_empanados, corned_beef,
laticinios_geral, queijo_coalho_manteiga, mussarela, leite_uht, leite_pasteurizado,
doce_de_leite, leite_fermentado, requeijao, leite_em_po, leite_condensado,
bebida_lactea, soro_leite, leite_aromatizado, composto_lacteo, nata,
pescado_fresco, pescado_congelado, pescado_salgado, camarao,
conserva_sardinhas, conserva_peixes, conserva_atuns, lagosta,
mel, mel_qualidade, apicola_derivados,
ovos, ovos_pasteurizados,
aves_geral, outro
"""

async def detect_product_phase1(image_b64: str, mime_type: str, obs: str) -> dict:
    """Fase 1: identifica produto, categoria e órgão automaticamente da imagem."""
    user_content = [
        {"type": "image", "source": {"type": "base64", "media_type": mime_type, "data": image_b64}},
        {"type": "text", "text": f"Identifique este rótulo.{' Dica: ' + obs if obs else ''} Retorne APENAS o JSON."}
    ]
    payload = {
        "model": ANTHROPIC_MODEL,
        "max_tokens": 300,
        "temperature": 0,
        "system": SP_DETECT,
        "messages": [{"role": "user", "content": user_content}],
    }
    headers_api = {
        "x-api-key": ANTHROPIC_API_KEY,
        "anthropic-version": "2023-06-01",
        "content-type": "application/json",
    }
    try:
        async with httpx.AsyncClient(timeout=6.0) as client:
            r = await client.post("https://api.anthropic.com/v1/messages", json=payload, headers=headers_api)
            if r.status_code != 200:
                return {}
            text = r.json().get("content", [{}])[0].get("text", "")
            # Extrai JSON do texto
            import re as _re
            m = _re.search(r'\{[\s\S]*\}', text)
            if m:
                return json.loads(m.group(0))
    except Exception:
        pass
    return {}

async def _revisao_background(relatorio: str):
    """Roda SP_REVISAO em background — não bloqueia o stream principal."""
    try:
        await call_claude_simple(
            SP_REVISAO.replace("{relatorio}", relatorio),
            "Revise com rigor técnico.",
            800
        )
    except Exception:
        pass

async def stream_validation(image_b64: str, mime_type: str, obs: str, orgao: str = "",
                            extra_images: list = None,
                            seg_tipo: str = "", seg_np_categoria: str = "",
                            seg_estado: str = "", seg_categoria: str = "",
                            modo_output: str = "completo"):
    # ── PARSER DE DIMENSÕES: extrai mm do obs e calcula área e fontes ─────
    dim_context = ""
    import re as _re
    # Busca padrão NNmm x NNmm no obs
    dim_match = _re.search(r"(\d+(?:[.,]\d+)?)mm\s*[xX\xd7]\s*(\d+(?:[.,]\d+)?)mm", obs) if obs else None
    if dim_match:
        try:
            larg = float(dim_match.group(1).replace(",", "."))
            alt  = float(dim_match.group(2).replace(",", "."))
            area_cm2 = round((larg * alt) / 100, 1)
            # Tamanho mínimo de fonte conforme IN 22/2005 e RDC 727/2022
            if area_cm2 <= 10:
                fonte_min = 0.75
                fonte_alergenico = 0.75
                regra_area = "≤10cm²"
            elif area_cm2 <= 80:
                fonte_min = 0.75
                fonte_alergenico = 1.0
                regra_area = "≤80cm²"
            elif area_cm2 <= 100:
                fonte_min = 1.0
                fonte_alergenico = 2.0
                regra_area = "≤100cm²"
            else:
                fonte_min = 1.0
                fonte_alergenico = 2.0
                regra_area = ">100cm²"
            # Escala pixel→mm: imagem preprocessada para 4500px no maior lado
            # Logo: 1px = maior_dim_mm / 4500 mm
            maior_dim_mm = max(larg, alt)
            mm_por_px = maior_dim_mm / 4500.0
            px_por_mm = 4500.0 / maior_dim_mm
            # Altura em px que corresponde ao mínimo de fonte
            px_fonte_min = round(fonte_min * px_por_mm, 1)
            px_fonte_alerg = round(fonte_alergenico * px_por_mm, 1)
            dim_context = (
                f"\n\n## DIMENSÕES REAIS DA EMBALAGEM (fornecidas pelo RT)\n"
                f"Dimensão física: {larg}mm × {alt}mm | Área do painel: {area_cm2}cm² ({regra_area})\n"
                f"\nESCALA DA IMAGEM (imagem preprocessada para 4500px no maior lado):\n"
                f"  • 1mm real = {px_por_mm:.1f}px na imagem | 1px = {mm_por_px:.3f}mm real\n"
                f"\nTAMANHO MÍNIMO DE FONTE — o que você deve ver na imagem:\n"
                f"  • Texto geral: mín. {fonte_min}mm real = mín. {px_fonte_min}px de altura na imagem\n"
                f"  • Alérgenos: mín. {fonte_alergenico}mm real = mín. {px_fonte_alerg}px de altura na imagem\n"
                f"\nCOMO AVALIAR: observe a altura das letras minúsculas (x-height) na imagem.\n"
                f"  • Se claramente menores que {px_fonte_min}px → ❌ NÃO CONFORME — fonte abaixo de {fonte_min}mm\n"
                f"  • Se parecem ter {px_fonte_min}px ou mais → ✅ CONFORME — fonte dentro do mínimo legal\n"
                f"  • Se difícil avaliar → ⚠️ COM RESSALVAS — informar dimensão estimada em mm\n"
                f"  Norma: IN 22/2005, Art. 6° (texto geral) | RDC 727/2022, Art. 15° (alérgenos)"
            )
        except Exception:
            dim_context = ""

    # ── DETECÇÃO + KB em paralelo com obs do usuário ─────────────────────
    # Usa obs para enriquecer keywords sem chamar Phase 1 separado (mais rápido)
    produto_detectado = ""
    categoria_detectada = ""
    orgao_final = orgao or ""
    sigla_sie = ""
    num_registro = ""
    especie = ""
    detected = {}

    # Carrega KB baseado em obs (Phase 1 agora está dentro do prompt principal)
    # Gap A: se obs vazia, usa Phase 1 para detectar categoria automaticamente
    if not obs and mime_type != "application/pdf":
        try:
            detected_auto = await detect_product_phase1(image_b64, mime_type, "")
            if detected_auto.get("categoria_kb") and detected_auto["categoria_kb"] != "outro":
                obs = detected_auto.get("produto", "")
                categories = [detected_auto["categoria_kb"]]
                if detected_auto.get("orgao") and not orgao_final:
                    orgao_final = detected_auto["orgao"]
            else:
                categories = []
        except Exception:
            categories = []
    else:
        categories = detect_categories(obs) if obs else []
    try:
        kb_text = await get_kb_for_categories(categories) if categories else ""
    except Exception:
        kb_text = ""

    if kb_text:
        cats_str = ", ".join(categories[:3]).upper().replace("_", " ")
        kb_section = f"""## LEGISLAÇÃO ESPECÍFICA CARREGADA — {cats_str}
Os textos abaixo são extraídos diretamente dos PDFs oficiais do MAPA para este produto.
Use como referência primária na validação:

{kb_text}

---"""
    else:
        kb_section = ""

    # Contexto do órgão — usa detecção automática + input do usuário
    orgao_map = {
        "SIM": "ATENÇÃO: Produto registrado no SIM (Serviço de Inspeção Municipal). " + SIM_MUNICIPAL_FALLBACK,
        "SIE": f"ATENÇÃO: Produto registrado no SIE (Serviço de Inspeção Estadual{' — ' + sigla_sie if sigla_sie else ''}). Verificar carimbo com sigla estadual{' ' + sigla_sie if sigla_sie else ''}. Exigências estaduais somam-se às federais.",
        "SIF": "ATENÇÃO: Produto registrado no SIF (Serviço de Inspeção Federal — DIPOA/MAPA). Máximo rigor. Verificar carimbo oval com 'SIF' e número do estabelecimento. Produto pode ser comercializado em todo território nacional.",
    }
    orgao_context = orgao_map.get(orgao_final.upper(), "")

    # Injetar contexto NP (não-POA) — prioriza: seg_np_categoria > categoria detectada > obs
    # seg_np_categoria: categoria selecionada pelo RT antes do upload (Task #4)
    # Bug fix: np_context é calculado aqui mas só é appendado em system_prompt
    # após a definição do mesmo (mais abaixo)
    np_context = get_np_fallback(obs, categoria_detectada=seg_np_categoria)

    funcional_context = ""
    if any(kw in obs.lower() for kw in ["funcional", "probiótico", "probiotico", "enriquecido",
                                          "ômega", "omega", "lactobacillus", "bifidobacterium",
                                          "vitamina d", "vitamina a", "reduzido em", "light", "zero"]):
        funcional_context = f"\n\n## ALIMENTO FUNCIONAL / DUAL MAPA+ANVISA\n{DUAL_MAPA_ANVISA_FALLBACK}"

    # Injetar contexto de prato pronto POA quando detectado
    prato_pronto_context = ""
    if any(kw in obs.lower() for kw in ["empanado", "nugget", "lasanha", "croquete",
                                          "prato pronto", "semiprontos", "cordon bleu",
                                          "steak", "hamburguer recheado"]):
        prato_pronto_context = f"\n\n## PRATO PRONTO / PRODUTO CÁRNEO ELABORADO\n{PRATO_PRONTO_POA_FALLBACK}"

    # Injetar norma estadual SIE — prioridade: seg_estado (RT explícito) > sigla > obs
    sie_context = ""

    # Prioridade 1 — RT selecionou estado SIE explicitamente no formulário
    if seg_estado:
        for estado_key, fallback_text in SIE_ESTADO_MAP.items():
            if estado_key.upper() == seg_estado.upper():
                sie_context = f"\n\n## NORMAS COMPLEMENTARES SIE — {seg_estado.upper()} (selecionado pelo RT)\n{fallback_text}"
                break

    # Prioridade 2 — Detectar pelo carimbo/sigla SIE
    if not sie_context and orgao_final.upper() == "SIE" and sigla_sie:
        for estado_key, fallback_text in SIE_ESTADO_MAP.items():
            if estado_key.upper() in sigla_sie.upper():
                sie_context = f"\n\n## NORMAS COMPLEMENTARES SIE — {sigla_sie.upper()}\n{fallback_text}"
                break

    # Prioridade 3 — Detectar pelo obs do usuário (ex: "CISPOA", "SISP")
    if not sie_context and obs:
        for estado_key, fallback_text in SIE_ESTADO_MAP.items():
            if estado_key.upper() in obs.upper():
                sie_context = f"\n\n## NORMAS COMPLEMENTARES SIE — {estado_key}\n{fallback_text}"
                break

    # SISBI-POA — injeta quando mencionado no obs ou detectado como SIE/SIM
    sisbi_context = ""
    if any(kw in (obs or "").upper() for kw in ["SISBI", "EQUIVALÊNCIA", "EQUIVALENCIA"]):
        sisbi_context = f"\n\n## SISBI-POA — EQUIVALÊNCIA NACIONAL\n{SISBI_POA_FALLBACK}"
    elif orgao_final.upper() in ("SIE", "SIM") and not sie_context:
        sisbi_context = f"\n\n## SISBI-POA — NOTA\n{SISBI_POA_FALLBACK}"

    # POA Orgânico — injeta quando orgânico detectado no obs
    # POA caprinos/ovinos — injeta quando detectado no obs ou categoria
    caprino_context = ""
    if any(kw in (obs or "").lower() for kw in ["cabra", "caprino", "ovino", "cordeiro", "carneiro", "chèvre", "chevre", "queijo de cabra"]):
        caprino_context = f"\n\n## POA CAPRINOS E OVINOS\n{POA_CAPRINO_OVINO_FALLBACK}"
    elif seg_categoria and any(kw in seg_categoria.lower() for kw in ["cabra", "caprino", "ovino", "cordeiro"]):
        caprino_context = f"\n\n## POA CAPRINOS E OVINOS\n{POA_CAPRINO_OVINO_FALLBACK}"

    organico_context = ""
    if any(kw in (obs or "").lower() for kw in ["orgânico", "organico", "ecológico", "ecologico", "sisorg", "oac "]):
        organico_context = f"\n\n## POA ORGÂNICO\n{POA_ORGANICO_FALLBACK}"

    # Contexto da detecção automática
    detection_context = ""
    if detected:
        detection_context = f"""## DETECÇÃO AUTOMÁTICA DA FASE 1
Produto identificado: {produto_detectado}
Espécie animal: {especie}
Órgão de inspeção detectado: {orgao_final}
{f'Sigla SIE: {sigla_sie}' if sigla_sie else ''}
{f'Número do registro: {num_registro}' if num_registro else ''}
Confiança da detecção: {detected.get('confianca', 'media')}

Use essas informações como ponto de partida — confirme ou corrija com base no que você vê na imagem.
---"""

    # Injeta exemplos few-shot de validações anteriores (Sistema 2)
    fewshot = ""
    if categories:
        fewshot = get_fewshot_examples(categories[0], caminho_np=seg_np_categoria)

    system_prompt = SP_VALIDACAO.replace("{kb_section}", kb_section)
    if np_context:
        system_prompt += np_context
    if orgao_context:
        system_prompt += f"\n\n{orgao_context}"
    if sie_context:
        system_prompt += sie_context
    if caprino_context:
        system_prompt += caprino_context
    if sisbi_context:
        system_prompt += sisbi_context
    if organico_context:
        system_prompt += organico_context
    if funcional_context:
        system_prompt += funcional_context
    if prato_pronto_context:
        system_prompt += prato_pronto_context
    if fewshot:
        system_prompt += fewshot
    if detection_context:
        system_prompt += f"\n\n{detection_context}"

    # Task #17 -- Modo de output
    if modo_output == "rapido":
        system_prompt += (
            "\n\n## MODO REVISAO RAPIDA -- INSTRUCAO ESPECIAL\n"
            "Voce esta no modo REVISAO RAPIDA. Siga rigorosamente:\n"
            "1. NAO liste campos conformes. Pule-os completamente.\n"
            "2. SOMENTE reporte campos com problemas (NAO CONFORME ou COM RESSALVAS).\n"
            "3. Para cada campo problemático: nome do campo | problema | artigo da norma | como corrigir.\n"
            "4. Formato: uma linha por campo. Sem introducao, sem conclusao extensa.\n"
            "5. Finalize com: SCORE: X/14 | VEREDICTO: [APROVADO/APROVADO COM RESSALVAS/REPROVADO]\n"
            "6. Se nao houver problemas: escreva apenas 'Sem nao conformidades detectadas. Score: 14/14 -- APROVADO'\n"
            "Exemplo de output:\n"
            "CAMPO 2 -- Ingredientes: aditivo INS 250 sem funcao declarada. RDC 727/2022 Art. 18 §2° -- adicionar 'conservante'.\n"
            "CAMPO 9 -- Tabela: porcao 25g mas IN 75/2020 exige 30g para biscoito -- corrigir porcao.\n"
            "SCORE: 12/14 | VEREDICTO: APROVADO COM RESSALVAS\n"
        )
    if dim_context:
        system_prompt += dim_context

    user_text = (
        "Analise este rótulo com MÁXIMA precisão. Execute TODOS os passos. Não pule nenhum campo.\n\n"
        "🚫 FRASES BANIDAS NESTA RESPOSTA (rejeição automática se aparecerem):\n"
        "• 'não é possível verificar' / 'não foi possível verificar'\n"
        "• 'não é possível ler' / 'não totalmente legíveis'\n"
        "• 'texto muito pequeno' / 'fonte pequena demais'\n"
        "• 'imagem com baixa resolução' / 'resolução insuficiente'\n"
        "• 'não verificável' / 'não verificada' / 'não verificado'\n"
        "• Qualquer menção a dimensão de imagem (ex: '225×225px') ou pedido de 'foto de maior resolução'\n\n"
        "✅ COMO PROCEDER:\n"
        "Você É um modelo de visão de fronteira. SEMPRE consegue ler o que está visível, mesmo em imagens pequenas. "
        "Aproxime-se mentalmente, examine pixel a pixel, infira pelo contexto. "
        "LEIA TODOS OS INGREDIENTES, TODOS OS VALORES NUTRICIONAIS, TODOS OS DADOS DO FABRICANTE. Liste-os no relatório.\n\n"
        "Sempre forneça veredicto efetivo: ✅ CONFORME ou ❌ NÃO CONFORME na grande maioria dos campos. "
        "Use ⚠️ COM RESSALVAS apenas em casos genuínos de cruzamento de norma indisponível (ex: claim 'Light' que precisa comparar com produto referência). "
        "Em qualquer ressalva, INCLUA o que conseguiu ler — nunca apenas 'não é possível ler'."
    )
    if obs:
        user_text += f"\n\nObservação adicional: {obs}"
    if produto_detectado:
        user_text += f"\nProduto identificado automaticamente: {produto_detectado}"

    # Gera case_id para este rótulo (usado pelo feedback endpoint)
    case_id = _case_id(image_b64)

    payload = {
        "model": ANTHROPIC_MODEL,
        "max_tokens": 16384,
        "temperature": 0,
        "stream": True,
        "system": system_prompt,
        "messages": [{"role": "user", "content": (
            [{"type": "document", "source": {"type": "base64", "media_type": "application/pdf", "data": image_b64}},
             {"type": "text", "text": user_text}]
            if mime_type == "application/pdf" else
            (
                [{"type": "image", "source": {"type": "base64", "media_type": mime_type, "data": image_b64}}]
                + ([{"type": "image", "source": {"type": "base64", "media_type": "image/jpeg", "data": b}} for b in (extra_images or [])[:5]])
                + ([{"type": "text", "text": f"As {1 + len(extra_images)} imagens acima são painéis diferentes do mesmo rótulo — analise TODOS em conjunto."}] if extra_images else [])
                + [{"type": "text", "text": user_text}]
            )
        )}],
    }
    headers = {
        "x-api-key": ANTHROPIC_API_KEY,
        "anthropic-version": "2023-06-01",
        "content-type": "application/json",
    }

    relatorio = ""

    async with httpx.AsyncClient(timeout=120.0) as client:
        async with client.stream(
            "POST", "https://api.anthropic.com/v1/messages",
            json=payload, headers=headers
        ) as response:
            if response.status_code != 200:
                error_body = await response.aread()
                yield f"data: {json.dumps({'error': error_body.decode()})}\n\n"
                return
            async for line in response.aiter_lines():
                if not line.startswith("data: "):
                    continue
                raw = line[6:].strip()
                if raw == "[DONE]":
                    break
                try:
                    ev = json.loads(raw)
                    if ev.get("type") == "content_block_delta":
                        delta = ev.get("delta", {})
                        if delta.get("type") == "text_delta":
                            chunk = delta.get("text", "")
                            relatorio += chunk
                            yield f"data: {json.dumps({'text': chunk})}\n\n"
                except Exception:
                    continue

    # ── Detecção NP pela categoria declarada no Passo 1 do relatório ──────────
    # O agente agora declara "CATEGORIA NÃO-POA: [x]" no início do relatório.
    # Capturamos aqui e injetamos retroativamente no contexto para o SP_REVISAO.
    import re as _re_np
    np_cat_match = _re_np.search(
        r"CATEGORIA\s+N[ÃA]O-POA[:\s]+([^\n]{3,60})",
        relatorio, _re_np.IGNORECASE
    )
    if np_cat_match:
        cat_detectada = np_cat_match.group(1).strip().lower()
        np_extra = get_np_fallback("", categoria_detectada=cat_detectada)
        if np_extra and np_extra not in system_prompt:
            system_prompt += np_extra  # enriquece para SP_REVISAO se rodar

    # Segunda leitura crítica — apenas se score baixo, em background (não bloqueia o stream)
    score_pre = extrair_score(relatorio)
    if score_pre is not None and score_pre < 9.0:  # inclui 8.5, 8.0, etc.
        asyncio.ensure_future(_revisao_background(relatorio))

    # ── Auto-aprendizado: armazena resultado sem precisar de feedback RT ───────
    try:
        import re as _re2
        # Score: suporta X.X/12 (3 níveis) e X/12 (legado)
        score_match = _re2.search(r"SCORE[:\s]+([\d]+[.,][\d]*)\s*/\s*14", relatorio, _re2.IGNORECASE)
        if not score_match:
            score_match = _re2.search(r"SCORE[:\s]+(\d+)\s*/\s*14", relatorio, _re2.IGNORECASE)
        score_auto = float(score_match.group(1).replace(",",".")) if score_match else None
        prod_match   = _re2.search(r"PRODUTO[:\s]+([^\n|]+)", relatorio, _re2.IGNORECASE)
        prod_auto    = prod_match.group(1).strip()[:80] if prod_match else ""
        # Captura campos NÃO CONFORME e COM RESSALVAS
        erros_auto_list = _re2.findall(
            r"CAMPO (\d+)[^\n]*(?:NÃO CONFORME|AUSENTE)[^\n]*\n([^\n]{0,120})",
            relatorio, _re2.IGNORECASE
        )
        ressalvas_auto_list = _re2.findall(
            r"CAMPO (\d+)[^\n]*COM RESSALVAS[^\n]*\n([^\n]{0,120})",
            relatorio, _re2.IGNORECASE
        )
        erros_auto = "; ".join(f"C{num}: {desc[:60]}" for num, desc in erros_auto_list[:5])
        ressalvas_auto = "; ".join(f"C{num}: {desc[:60]}" for num, desc in ressalvas_auto_list[:5])
        auto_case = {
            "case_id":      case_id,
            "produto":      prod_auto or (obs[:50] if obs else ""),
            "categoria":    categories[0] if categories else "",
            "caminho_np":   seg_np_categoria,
            "feedback":     None,
            "erros_auto":   erros_auto,
            "erros_encontrados": erros_auto,
            "ressalvas_auto": ressalvas_auto,
            "score_agente": score_auto,
            "timestamp":    __import__("datetime").datetime.now().isoformat(),
            "auto_stored":  True,
            "relatorio_completo": relatorio[:8000],  # salva para comparativo
            "imagem_url":   "",  # preenchido em background pelo upload Storage (abaixo)
        }
        existing = next((x for x in _cases_db if x["case_id"] == case_id), None)
        if not existing:
            _cases_db.append(auto_case)
            if len(_cases_db) > _MAX_CASES:
                _cases_db.pop(0)
            import datetime as _dt
            supabase_case = {**auto_case, "created_at": _dt.datetime.now().isoformat()}
            # Remove None values
            supabase_case = {k: v for k, v in supabase_case.items() if v is not None}
            asyncio.ensure_future(_sb_upsert_v2("validacoes", supabase_case, conflict_col="case_id"))

            # Upload da imagem do rótulo pro Storage (fire-and-forget, não bloqueia resposta)
            # Quando upload completar, atualiza o registro com a imagem_url.
            async def _upload_e_atualizar():
                try:
                    print(f"[storage] _upload_e_atualizar START case={case_id[:16]}", flush=True)
                    if not contents:
                        print(f"[storage] _upload_e_atualizar SKIP: contents vazio case={case_id[:16]}", flush=True)
                        return
                    img_url = await _sb_upload_rotulo(
                        contents,
                        case_id,
                        getattr(imagem, "content_type", None) or "image/jpeg"
                    )
                    if img_url:
                        print(f"[storage] _upload_e_atualizar: upload OK, fazendo PATCH no DB case={case_id[:16]}", flush=True)
                        # Usa PATCH direto em vez de upsert (mais simples — só atualiza imagem_url)
                        try:
                            async with httpx.AsyncClient(timeout=10.0) as cl:
                                patch_url = f"{_SUPABASE_URL}/rest/v1/validacoes?case_id=eq.{case_id}"
                                pr = await cl.patch(
                                    patch_url,
                                    json={"imagem_url": img_url},
                                    headers={
                                        "apikey": _SUPABASE_KEY,
                                        "Authorization": f"Bearer {_SUPABASE_KEY}",
                                        "Content-Type": "application/json",
                                        "Prefer": "return=minimal",
                                    },
                                )
                                if pr.status_code in (200, 201, 204):
                                    print(f"[storage] DB PATCH OK case={case_id[:16]}", flush=True)
                                else:
                                    print(f"[storage] DB PATCH FALHOU case={case_id[:16]} status={pr.status_code} body={pr.text[:300]}", flush=True)
                        except Exception as patch_e:
                            print(f"[storage] DB PATCH EXCEÇÃO case={case_id[:16]}: {patch_e}", flush=True)
                    else:
                        print(f"[storage] _upload_e_atualizar: upload retornou vazio case={case_id[:16]}", flush=True)
                except Exception as e:
                    print(f"[storage] _upload_e_atualizar erro case={case_id[:16]}: {e}", flush=True)

            asyncio.ensure_future(_upload_e_atualizar())
    except Exception:
        pass  # auto-store nunca deve quebrar o relatório

    # Task #17 + T3 fix — Envio automático do relatório por email
    # Prioridade: 1) email_destino do form (campo explícito) 2) X-User-Email header (login)
    # Assim funciona mesmo sem login
    try:
        email_header = request.headers.get("X-User-Email", "") if request else ""
        nome_header  = request.headers.get("X-User-Name", "") if request else ""
        # email_destino vem do form (parâmetro novo — T3 fix)
        # fallback para header se form vazio
        email_final = email_destino.strip() if email_destino and email_destino.strip() else email_header
        if email_final and relatorio:
            asyncio.ensure_future(_enviar_email_relatorio(
                email=email_final,
                nome=nome_header or "",
                produto=prod_auto or obs[:50] if obs else "",
                veredicto=veredicto_auto,
                score=score_auto,
                relatorio_md=relatorio,
                case_id=case_id,
            ))
    except Exception:
        pass  # email nunca deve quebrar o relatório
    # ────────────────────────────────────────────────────────────────────────────

    # Emite case_id para o frontend usar no feedback
    yield f"data: {json.dumps({'case_id': case_id, 'fewshot_used': bool(fewshot), 'produto': prod_auto})}\n\n"
    yield "data: [DONE]\n\n"



def check_image_readability(image_bytes: bytes) -> dict:
    """
    Verifica se a imagem tem resolução e contraste suficientes para leitura de texto.
    Retorna {"ok": bool, "warning": str, "dim": tuple}
    """
    try:
        from PIL import Image as PILImage, ImageFilter
        import statistics
        img = PILImage.open(__import__("io").BytesIO(image_bytes))
        w, h = img.size
        maior = max(w, h)
        # Resolução mínima para rótulo legível antes do zoom
        if maior < 80:
            return {"ok": False,
                    "warning": f"⚠️ Imagem ilegível ({w}×{h}px). Envie uma foto com pelo menos 200px no maior lado.",
                    "dim": (w, h)}
        # Aviso não-bloqueante para imagens pequenas (Claude tenta mesmo assim com zoom)
        if maior < 300:
            return {"ok": True,
                    "warning": f"⚠️ Imagem pequena ({w}×{h}px) — qualidade pode ser limitada, mas vamos tentar.",
                    "dim": (w, h)}
        # Verifica se imagem não é completamente branca/preta (PDF em branco, etc.)
        gray = img.convert("L")
        pixels = list(gray.getdata())
        try:
            stdev = statistics.stdev(pixels[:10000])  # amostra
        except Exception:
            stdev = 50
        if stdev < 8:
            return {"ok": False,
                    "warning": "⚠️ Imagem sem contraste detectável. Verifique se o arquivo está correto e tente novamente.",
                    "dim": (w, h)}
        return {"ok": True, "warning": "", "dim": (w, h)}
    except Exception:
        return {"ok": True, "warning": "", "dim": (0, 0)}


def _deskew_image(image_bytes: bytes) -> bytes:
    """
    Corrige perspectiva angular de texto em rótulos curvos ou fotografados em ângulo.
    Usa OpenCV para:
    1. Detectar contornos retangulares dominantes (área do rótulo)
    2. Aplicar transformação de perspectiva inversa (warpPerspective)
    3. "Endireitar" o texto para facilitar leitura pelo Claude

    Estratégia conservadora: só aplica correção se encontrar contorno com
    ângulo > 5° — evita distorcer imagens que já estão retas.
    """
    try:
        import cv2
        import numpy as np

        # Decodifica imagem
        arr = np.frombuffer(image_bytes, dtype=np.uint8)
        img = cv2.imdecode(arr, cv2.IMREAD_COLOR)
        if img is None:
            return image_bytes

        h, w = img.shape[:2]

        # ── Passo 1: pré-processamento para detecção de contornos ──────────
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        # Blur suave para reduzir ruído sem perder bordas do rótulo
        blurred = cv2.GaussianBlur(gray, (5, 5), 0)
        # Canny adaptativo: usa mediana para definir thresholds automaticamente
        median = np.median(blurred)
        low    = int(max(0,   (1.0 - 0.33) * median))
        high   = int(min(255, (1.0 + 0.33) * median))
        edges  = cv2.Canny(blurred, low, high)
        # Dilata bordas para fechar gaps em texto pequeno
        kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (3, 3))
        edges  = cv2.dilate(edges, kernel, iterations=1)

        # ── Passo 2: encontrar contorno retangular dominante ───────────────
        contours, _ = cv2.findContours(edges, cv2.RETR_LIST, cv2.CHAIN_APPROX_SIMPLE)
        if not contours:
            return image_bytes

        # Filtra contornos: só candidatos com área > 10% da imagem total
        min_area = 0.10 * h * w
        candidates = []
        for cnt in contours:
            area = cv2.contourArea(cnt)
            if area < min_area:
                continue
            peri = cv2.arcLength(cnt, True)
            # approxPolyDP com epsilon = 2% do perímetro
            approx = cv2.approxPolyDP(cnt, 0.02 * peri, True)
            if len(approx) == 4:
                candidates.append((area, approx))

        if not candidates:
            # Nenhum quadrilátero claro — tenta correção de skew simples
            return _deskew_simple(img, image_bytes)

        # Maior candidato quadrilateral
        candidates.sort(key=lambda x: -x[0])
        _, quad = candidates[0]
        pts = quad.reshape(4, 2).astype(np.float32)

        # ── Passo 3: verifica se há perspectiva significativa (> 5°) ───────
        # Ordena pontos: top-left, top-right, bottom-right, bottom-left
        rect = _order_points(pts)
        tl, tr, br, bl = rect

        # Calcula ângulo da borda superior
        dx = float(tr[0] - tl[0])
        dy = float(tr[1] - tl[1])
        angle = abs(np.degrees(np.arctan2(dy, dx)))

        # Só aplica se ângulo for > 3° (perspectiva real) e < 45° (não invertido)
        if angle < 3.0 or angle > 45.0:
            # Ângulo muito pequeno = imagem já reta. Aplicar só deskew simples.
            return _deskew_simple(img, image_bytes)

        # ── Passo 4: warpPerspective — transforma para visão frontal ───────
        # Dimensões de destino: largura e altura do quadrilátero detectado
        width_top    = np.linalg.norm(tr - tl)
        width_bottom = np.linalg.norm(br - bl)
        dst_w = int(max(width_top, width_bottom))

        height_left  = np.linalg.norm(bl - tl)
        height_right = np.linalg.norm(br - tr)
        dst_h = int(max(height_left, height_right))

        if dst_w < 50 or dst_h < 50:
            return image_bytes

        dst = np.array([
            [0,         0        ],
            [dst_w - 1, 0        ],
            [dst_w - 1, dst_h - 1],
            [0,         dst_h - 1],
        ], dtype=np.float32)

        M = cv2.getPerspectiveTransform(rect, dst)
        warped = cv2.warpPerspective(img, M, (dst_w, dst_h),
                                      flags=cv2.INTER_LANCZOS4)

        # Reencoda como JPEG
        ok, buf = cv2.imencode(".jpg", warped, [cv2.IMWRITE_JPEG_QUALITY, 95])
        if ok:
            return buf.tobytes()
        return image_bytes

    except Exception:
        return image_bytes


def _deskew_simple(img, image_bytes: bytes) -> bytes:
    """
    Correção de skew simples quando não há contorno quadrilateral claro.
    Usa projeção de linhas horizontais (Hough) para detectar ângulo de rotação.
    Eficaz para texto diagonal até ~20°.
    """
    try:
        import cv2
        import numpy as np

        gray    = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        # Thresh binário para isolação de texto
        _, thresh = cv2.threshold(gray, 0, 255,
                                   cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)
        # Detecta linhas com Hough probabilístico
        lines = cv2.HoughLinesP(thresh, 1, np.pi / 180,
                                 threshold=80, minLineLength=50, maxLineGap=10)
        if lines is None:
            return image_bytes

        angles = []
        for line in lines:
            x1, y1, x2, y2 = line[0]
            dx = x2 - x1
            if dx == 0:
                continue
            angle = np.degrees(np.arctan2(y2 - y1, dx))
            # Filtra: só linhas "quase horizontais" (texto tem ângulo < 30°)
            if abs(angle) < 30:
                angles.append(angle)

        if not angles:
            return image_bytes

        # Mediana dos ângulos detectados
        skew = float(np.median(angles))

        # Só corrige se skew for significativo (> 1°) e seguro (< 25°)
        if abs(skew) < 1.0 or abs(skew) > 25.0:
            return image_bytes

        h, w = img.shape[:2]
        center = (w // 2, h // 2)
        M_rot = cv2.getRotationMatrix2D(center, skew, 1.0)
        # Expande canvas para não cortar cantos após rotação
        cos_a = abs(M_rot[0, 0])
        sin_a = abs(M_rot[0, 1])
        new_w = int(h * sin_a + w * cos_a)
        new_h = int(h * cos_a + w * sin_a)
        M_rot[0, 2] += (new_w / 2) - center[0]
        M_rot[1, 2] += (new_h / 2) - center[1]
        rotated = cv2.warpAffine(img, M_rot, (new_w, new_h),
                                  flags=cv2.INTER_LANCZOS4,
                                  borderMode=cv2.BORDER_REPLICATE)
        ok, buf = cv2.imencode(".jpg", rotated, [cv2.IMWRITE_JPEG_QUALITY, 95])
        if ok:
            return buf.tobytes()
        return image_bytes

    except Exception:
        return image_bytes


def _order_points(pts):
    """Ordena 4 pontos: top-left, top-right, bottom-right, bottom-left."""
    import numpy as np
    rect = np.zeros((4, 2), dtype=np.float32)
    s = pts.sum(axis=1)
    rect[0] = pts[np.argmin(s)]   # top-left: menor soma x+y
    rect[2] = pts[np.argmax(s)]   # bottom-right: maior soma x+y
    diff = np.diff(pts, axis=1)
    rect[1] = pts[np.argmin(diff)]  # top-right: menor y-x
    rect[3] = pts[np.argmax(diff)]  # bottom-left: maior y-x
    return rect


def _extract_text_from_memorial(file_bytes: bytes, filename: str) -> str:
    """
    Extrai texto de arquivos de memorial técnico / receita.
    Suporta: PDF, DOCX, DOC (limitado), HTML, TXT, RTF.
    Retorna string vazia se não conseguir extrair.
    """
    if not file_bytes or len(file_bytes) == 0:
        return ""
    name_lower = (filename or "").lower()
    ext = name_lower.split(".")[-1] if "." in name_lower else ""

    try:
        # PDF
        if ext == "pdf":
            try:
                import pypdf
                from io import BytesIO
                reader = pypdf.PdfReader(BytesIO(file_bytes))
                pages = []
                for page in reader.pages[:50]:  # limite 50 páginas
                    try:
                        t = page.extract_text() or ""
                        if t.strip():
                            pages.append(t)
                    except Exception:
                        continue
                return "\n\n".join(pages)
            except ImportError:
                pass
            # Fallback: tentar pdfplumber
            try:
                import pdfplumber
                from io import BytesIO
                with pdfplumber.open(BytesIO(file_bytes)) as pdf:
                    return "\n\n".join((p.extract_text() or "") for p in pdf.pages[:50])
            except ImportError:
                return ""

        # DOCX
        if ext == "docx":
            try:
                import docx
                from io import BytesIO
                doc = docx.Document(BytesIO(file_bytes))
                paras = [p.text for p in doc.paragraphs if p.text.strip()]
                # Tabelas
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                        if row_text:
                            paras.append(row_text)
                return "\n".join(paras)
            except ImportError:
                return ""

        # HTML
        if ext in ("html", "htm"):
            try:
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(file_bytes, "html.parser")
                # Remove scripts/styles
                for s in soup(["script", "style", "noscript"]):
                    s.decompose()
                return soup.get_text(separator="\n", strip=True)
            except ImportError:
                # Fallback: regex bem simples
                import re as _re
                text = file_bytes.decode("utf-8", errors="ignore")
                text = _re.sub(r"<script.*?</script>", "", text, flags=_re.DOTALL | _re.IGNORECASE)
                text = _re.sub(r"<style.*?</style>", "", text, flags=_re.DOTALL | _re.IGNORECASE)
                text = _re.sub(r"<[^>]+>", " ", text)
                return _re.sub(r"\s+", " ", text).strip()

        # TXT
        if ext == "txt":
            for encoding in ("utf-8", "latin-1", "cp1252"):
                try:
                    return file_bytes.decode(encoding)
                except UnicodeDecodeError:
                    continue
            return file_bytes.decode("utf-8", errors="ignore")

        # RTF — extração simples (remove comandos RTF entre {} e \)
        if ext == "rtf":
            import re as _re
            text = file_bytes.decode("utf-8", errors="ignore")
            # Remove {\* ... } e \cmd
            text = _re.sub(r"\{\\\*[^}]*\}", "", text)
            text = _re.sub(r"\\[a-zA-Z]+\d*\s?", " ", text)
            text = _re.sub(r"[\{\}]", "", text)
            return _re.sub(r"\s+", " ", text).strip()

        # DOC (formato binário antigo) — não suportado direto, retorna vazio com aviso
        if ext == "doc":
            return ""

        return ""
    except Exception:
        return ""


def preprocess_image(image_bytes: bytes, mime_type: str) -> tuple[bytes, str]:
    """
    Pipeline completo de preprocessing:
    1. Deskew / correção de perspectiva (OpenCV) — corrige texto angular
    2. Upscale para 4500px (zoom máximo)
    3. Duplo UnsharpMask + contraste + sharpness (PIL)
    4. Imagens originalmente muito pequenas (<400px): boost extra de contraste e sharpen
    """
    try:
        from PIL import Image as PILImage, ImageFilter, ImageEnhance

        # ── Passo 1: correção de perspectiva angular ─────────────────────
        # Só aplica em imagens (não em PDFs já convertidos)
        if mime_type in ("image/jpeg", "image/png", "image/webp",
                         "image/gif", "image/bmp"):
            image_bytes = _deskew_image(image_bytes)

        # ── Passo 2: pipeline PIL (upscale + sharpen + contraste) ────────
        img = PILImage.open(io.BytesIO(image_bytes))

        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
        elif img.mode == "L":
            img = img.convert("RGB")

        w, h = img.size
        original_max = max(w, h)
        TARGET = 4500
        # Detecta imagem original muito pequena (thumbnails web típicos)
        is_tiny_original = original_max < 400

        if original_max < TARGET:
            scale = TARGET / original_max
            img = img.resize((int(w * scale), int(h * scale)), PILImage.LANCZOS)
            img = img.filter(ImageFilter.UnsharpMask(radius=2.0, percent=300, threshold=0))
            img = img.filter(ImageFilter.UnsharpMask(radius=0.5, percent=200, threshold=0))
        elif original_max > TARGET:
            scale = TARGET / original_max
            img = img.resize((int(w * scale), int(h * scale)), PILImage.LANCZOS)
            img = img.filter(ImageFilter.UnsharpMask(radius=1.0, percent=150, threshold=1))

        # Imagens originalmente minúsculas precisam de mais ajuda — boost agressivo
        if is_tiny_original:
            img = ImageEnhance.Contrast(img).enhance(1.85)   # +0.35 vs default
            img = ImageEnhance.Sharpness(img).enhance(2.4)   # +0.6 vs default
            img = ImageEnhance.Color(img).enhance(1.15)      # cor levemente mais saturada
        else:
            img = ImageEnhance.Contrast(img).enhance(1.5)
            img = ImageEnhance.Sharpness(img).enhance(1.8)

        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=95)
        return buf.getvalue(), "image/jpeg"

    except Exception:
        return image_bytes, mime_type

@app.post("/validar")
async def validar_rotulo(
    imagem: UploadFile = File(...),
    imagens_extras: list[UploadFile] = File(default=[]),
    memorial_files: list[UploadFile] = File(default=[]),
    obs: str = Form(default=""),
    orgao: str = Form(default=""),
    user_id: str = Form(default=""),
    seg_tipo: str = Form(default=""),
    seg_np_categoria: str = Form(default=""),
    seg_estado: str = Form(default=""),
    seg_categoria: str = Form(default=""),
    modo_output: str = Form(default="completo"),
    email_destino: str = Form(default=""),  # T3 fix: email explícito no form (não depende de login)
    request: Request = None,
):
    if not ANTHROPIC_API_KEY:
        return JSONResponse({"error": "ANTHROPIC_API_KEY não configurada"}, status_code=400,
                            headers={"Access-Control-Allow-Origin": "*"})

    # M1_9 — Rate limiting por plano
    if user_id:
        uso = await _checar_e_incrementar_uso(user_id)
        if not uso["ok"]:
            return JSONResponse({
                "error": uso.get("msg", "Limite de validações atingido."),
                "limite_atingido": True,
                "plano": uso.get("plano"),
                "usado": uso.get("usado"),
                "limite": uso.get("limite"),
            }, status_code=429, headers={"Access-Control-Allow-Origin": "*"})

    contents = await imagem.read()
    content_type = (imagem.content_type or "").lower()
    filename = (imagem.filename or "").lower()

    # ── Detectar se é PDF ────────────────────────────────────────────────
    is_pdf = (
        content_type == "application/pdf" or
        filename.endswith(".pdf")
    )

    quality_warning = ""  # inicializado aqui para evitar NameError no bloco PDF
    if is_pdf:
        pages = await pdf_to_images_b64(contents)
        if not pages:
            return JSONResponse({"error": "Não foi possível processar o PDF."}, status_code=400,
                                headers={"Access-Control-Allow-Origin": "*"})

        # Gap H: verificar se houve erro de senha
        if pages and pages[0].get("is_error"):
            return JSONResponse(
                {"error": pages[0]["error"]}, status_code=400,
                headers={"Access-Control-Allow-Origin": "*"}
            )
        first = pages[0]
        if first.get("is_pdf"):
            # PDF nativo — envia direto para API Anthropic como documento
            image_b64 = first["b64"]
            mime_type = "application/pdf"
        else:
            # PDF convertido em imagens — mescla todas as páginas numa única imagem vertical
            # Isso garante que o agente veja frente E verso do rótulo
            from PIL import Image as _PIL
            page_imgs = []
            for pg in pages:
                raw_bytes = base64.b64decode(pg["b64"])
                proc_bytes, _ = preprocess_image(raw_bytes, pg["mime"])
                page_imgs.append(_PIL.open(io.BytesIO(proc_bytes)))

            if len(page_imgs) == 1:
                # Página única — usa direto
                buf = io.BytesIO()
                page_imgs[0].save(buf, "JPEG", quality=95)
                processed_bytes = buf.getvalue()
            else:
                # Múltiplas páginas — empilha verticalmente com separador
                total_w = max(img.width for img in page_imgs)
                sep = 20  # pixels de separação entre páginas
                total_h = sum(img.height for img in page_imgs) + sep * (len(page_imgs) - 1)
                merged = _PIL.new("RGB", (total_w, total_h), (200, 200, 200))
                y_offset = 0
                for img in page_imgs:
                    # Centraliza horizontalmente se larguras diferentes
                    x_offset = (total_w - img.width) // 2
                    merged.paste(img, (x_offset, y_offset))
                    y_offset += img.height + sep
                buf = io.BytesIO()
                merged.save(buf, "JPEG", quality=95)
                processed_bytes = buf.getvalue()

            mime_type = "image/jpeg"
            image_b64 = base64.b64encode(processed_bytes).decode("utf-8")
            if len(pages) > 1:
                obs = (obs + " " if obs else "") + f"[PDF com {len(pages)} páginas — frente e verso incluídos na análise]"
    else:
        mime_map = {
            "image/jpeg": "image/jpeg",
            "image/jpg": "image/jpeg",
            "image/png": "image/png",
            "image/webp": "image/webp",
            "image/gif": "image/gif",
            "image/bmp": "image/jpeg",
        }
        raw_mime = mime_map.get(content_type, "image/jpeg")
        # Gap B: verifica legibilidade antes de processar
        quality = check_image_readability(contents)
        if not quality["ok"]:
            return JSONResponse(
                {"error": quality["warning"]}, status_code=400,
                headers={"Access-Control-Allow-Origin": "*"}
            )
        # Aviso não-bloqueante: imagem pequena mas Claude tenta mesmo assim
        quality_warning = quality.get("warning", "")
        processed_bytes, mime_type = preprocess_image(contents, raw_mime)
        image_b64 = base64.b64encode(processed_bytes).decode("utf-8")

    # Anexa aviso de qualidade às observações se existir
    if quality_warning:
        obs = (obs + "\n\n[SISTEMA: " + quality_warning + " O agente deve tentar mesmo assim e informar se não conseguir ler algum campo.]").strip()
    # Multipanel: processar imagens extras e injetar no contexto
    if imagens_extras:
        extra_b64_list = []
        for extra in imagens_extras:
            try:
                extra_bytes = await extra.read()
                extra_proc, extra_mime = preprocess_image(extra_bytes, extra.content_type or "image/jpeg")
                extra_b64_list.append(base64.b64encode(extra_proc).decode())
            except Exception:
                pass
        if extra_b64_list:
            n_total = 1 + len(extra_b64_list)
            obs = obs + (
                f"\n\n[ANÁLISE MULTIPANEL — {n_total} face(s) da embalagem enviadas]"
                f"\nInstruções para análise face a face:"
                f"\n• FACE PRINCIPAL (imagem 1): verificar denominação de venda, lupa frontal, conteúdo líquido, claims e alegações"
                f"\n• FACES SECUNDÁRIAS (imagens 2+): verificar lista de ingredientes, tabela nutricional, declaração de alérgenos, carimbo de inspeção, identificação do fabricante, instruções de conservação"
                f"\n• Campos que aparecem em mais de uma face: registrar em qual face foram verificados"
                f"\n• Se campo obrigatório NÃO aparecer em nenhuma das faces → ❌ NÃO CONFORME (ausente no rótulo)"
                f"\n• Verificar especificamente: lupa está na FACE PRINCIPAL? Carimbo oval visível? Alérgenos imediatamente após ingredientes?"
                f"\nAnalise CADA face antes de emitir o veredicto final."
            )
    else:
        extra_b64_list = []

    # ─── MEMORIAL / RECEITA TÉCNICA (opcional) ─────────────────────────────
    # Extrai texto de PDF/DOCX/HTML/TXT/RTF e injeta no obs para Sonnet usar
    if memorial_files:
        memorial_text_parts = []
        MEMORIAL_TOTAL_CHAR_LIMIT = 60000  # ~15k tokens, evita estourar contexto
        total_chars_used = 0
        for mf in memorial_files:
            if total_chars_used >= MEMORIAL_TOTAL_CHAR_LIMIT:
                break
            try:
                mf_bytes = await mf.read()
                mf_text = _extract_text_from_memorial(mf_bytes, mf.filename or "arquivo")
                if mf_text and mf_text.strip():
                    # Trunca texto extraído se passar do limite
                    remaining = MEMORIAL_TOTAL_CHAR_LIMIT - total_chars_used
                    if len(mf_text) > remaining:
                        mf_text = mf_text[:remaining] + "\n[... arquivo truncado por limite de contexto ...]"
                    memorial_text_parts.append(f"### Arquivo: {mf.filename}\n{mf_text.strip()}")
                    total_chars_used += len(mf_text)
            except Exception as _e:
                # Memorial nunca pode quebrar a validação
                if SENTRY_DSN:
                    try:
                        import sentry_sdk
                        sentry_sdk.capture_exception(_e)
                    except Exception:
                        pass
                continue
        if memorial_text_parts:
            obs = obs + (
                "\n\n[MEMORIAL TÉCNICO / RECEITA FORNECIDO PELO RT]"
                "\nO RT anexou o(s) seguinte(s) documento(s) técnico(s) sobre o produto. Use estas informações como REFERÊNCIA AUTORIZADA "
                "para validar declarações do rótulo (ex: ingredientes, alérgenos, transgênicos, peso, RTIQ aplicável). "
                "Se o rótulo CONTRADIZER o memorial, sinalize como NÃO CONFORME. "
                "Se o memorial CONFIRMAR uma declaração ambígua do rótulo, considere CONFORME."
                "\n\n" + "\n\n---\n\n".join(memorial_text_parts)
            )

    return StreamingResponse(
        stream_validation(image_b64, mime_type, obs, orgao, extra_images=extra_b64_list,
                          seg_tipo=seg_tipo, seg_np_categoria=seg_np_categoria,
                          seg_estado=seg_estado, seg_categoria=seg_categoria,
                          modo_output=modo_output),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


# ═══════════════════════════════════════════════════════════════════════════════
# ENDPOINT: EVAL
# ═══════════════════════════════════════════════════════════════════════════════
@app.post("/eval")
async def avaliar_rotulo(
    imagem: UploadFile = File(...),
    gabarito: str = Form(...),
):
    if not ANTHROPIC_API_KEY:
        return JSONResponse({"error": "ANTHROPIC_API_KEY não configurada"}, status_code=400,
                            headers={"Access-Control-Allow-Origin": "*"})
    try:
        gab = json.loads(gabarito)
    except Exception:
        return JSONResponse({"error": "Gabarito inválido"},
                            headers={"Access-Control-Allow-Origin": "*"})

    contents = await imagem.read()
    mime_map = {"image/jpeg": "image/jpeg", "image/jpg": "image/jpeg",
                "image/png": "image/png", "image/webp": "image/webp"}
    raw_mime = mime_map.get(imagem.content_type, "image/jpeg")
    processed_bytes, mime_type = preprocess_image(contents, raw_mime)
    image_b64 = base64.b64encode(processed_bytes).decode("utf-8")

    obs = f"{gab.get('produto', '')} {gab.get('categoria', '')}".strip()
    orgao = gab.get("orgao", "")
    categories = detect_categories(obs)
    try:
        kb_text = await get_kb_for_categories(categories) if categories else ""
    except Exception:
        kb_text = ""
    kb_section = f"## LEGISLAÇÃO ESPECÍFICA\n{kb_text}\n---" if kb_text else ""
    system_prompt = SP_VALIDACAO.replace("{kb_section}", kb_section)

    payload = {
        "model": ANTHROPIC_MODEL,
        "max_tokens": 16384,
        "temperature": 0,
        "stream": True,
        "system": system_prompt,
        "messages": [{"role": "user", "content": [
            {"type": "image", "source": {"type": "base64", "media_type": mime_type, "data": image_b64}},
            {"type": "text", "text": f"Valide este rótulo. Produto: {obs}"},
        ]}],
    }
    headers = {
        "x-api-key": ANTHROPIC_API_KEY,
        "anthropic-version": "2023-06-01",
        "content-type": "application/json",
    }
    relatorio = ""

    async with httpx.AsyncClient(timeout=120.0) as client:
        async with client.stream("POST", "https://api.anthropic.com/v1/messages",
                                  json=payload, headers=headers) as response:
            if response.status_code != 200:
                return JSONResponse({"error": "Erro na API Claude"},
                                    headers={"Access-Control-Allow-Origin": "*"})
            async for line in response.aiter_lines():
                if not line.startswith("data: "):
                    continue
                raw = line[6:].strip()
                if raw == "[DONE]":
                    break
                try:
                    ev = json.loads(raw)
                    if ev.get("type") == "content_block_delta" and \
                       ev.get("delta", {}).get("type") == "text_delta":
                        relatorio += ev["delta"]["text"]
                except Exception:
                    continue

    erros_conhecidos = gab.get("erros_conhecidos", [])
    detalhes = []
    for erro in erros_conhecidos:
        campo = erro["campo"]
        esperado = erro["status"]
        detectado = detectar_status_campo(relatorio, campo)
        detalhes.append({
            "campo": campo,
            "nome": CAMPOS_NOME.get(campo, ""),
            "esperado": esperado,
            "detectado": detectado,
            "acertou": detectado == esperado,
            "descricao_gabarito": erro.get("descricao", ""),
            "norma": erro.get("norma", ""),
        })

    total = len(detalhes)
    acertos = sum(1 for d in detalhes if d["acertou"])

    return JSONResponse({
        "produto": gab.get("produto", ""),
        "relatorio_completo": relatorio,
        "score_agente": extrair_score(relatorio),
        "score_esperado": gab.get("score_esperado"),
        "veredicto_agente": extrair_veredicto(relatorio),
        "veredicto_esperado": gab.get("veredicto_esperado", "").upper(),
        "precisao_pct": round(acertos / total * 100) if total > 0 else 100,
        "erros_avaliados": total,
        "erros_acertados": acertos,
        "detalhes": detalhes,
    }, headers={"Access-Control-Allow-Origin": "*"})


# ═══════════════════════════════════════════════════════════════════════════════
# HEALTH + KB PRELOAD
# ═══════════════════════════════════════════════════════════════════════════════

# ══════════════════════════════════════════════════════════════════════════════
# SISTEMA 1 — MONITORAMENTO LEGISLATIVO CONTÍNUO
# Verifica o DOU diariamente por novas normas relevantes a POA
# ══════════════════════════════════════════════════════════════════════════════

_monitor_lock = asyncio.Lock()

MONITOR_KEYWORDS = [
    "rotulagem", "RTIQ", "produto de origem animal", "embutido", "laticínio",
    "pescado", "mel", "ovos", "inspeção industrial", "DIPOA", "SIGSIF",
    "alergênicos", "aditivos alimentares", "informação nutricional",
    "denominação de venda", "transgênico", "glúten", "lactose",
]
MONITOR_ORGANS = ["MAPA", "ANVISA", "INMETRO", "DIPOA", "SDA", "SNVS"]
MONITOR_DOC_TYPES = ["instrução normativa", "portaria", "resolução", "decreto",
                     "instrução de serviço", "circular"]

async def fetch_dou_recent(query: str, days_back: int = 2) -> list[dict]:
    """Busca publicações recentes no DOU via API pública."""
    from datetime import datetime, timedelta
    end_date   = datetime.now()
    start_date = end_date - timedelta(days=days_back)
    url = (
        "https://www.in.gov.br/consulta/-/buscar/dou"
        f"?q={query.replace(' ', '+')}"
        f"&s=todos"
        f"&exactDate=&startDate={start_date.strftime('%d-%m-%Y')}"
        f"&endDate={end_date.strftime('%d-%m-%Y')}"
        f"&sortType=0"
    )
    try:
        async with httpx.AsyncClient(timeout=20.0) as client:
            r = await client.get(url, headers={"User-Agent": "ValidaRotuloMonitor/1.0"})
            if r.status_code != 200:
                return []
            import re as _re
            # Extrai títulos e links dos resultados
            titles = _re.findall(r'class="title-marker[^"]*"[^>]*>([^<]+)<', r.text)
            links  = _re.findall(r'href="(/web/dou/-/[^"]+)"', r.text)
            results = []
            for i, title in enumerate(titles[:10]):
                link = f"https://www.in.gov.br{links[i]}" if i < len(links) else ""
                results.append({"title": title.strip(), "link": link})
            return results
    except Exception:
        return []

def is_relevant_publication(title: str) -> bool:
    """Verifica se uma publicação do DOU é relevante para POA."""
    t = title.lower()
    has_keyword = any(kw.lower() in t for kw in MONITOR_KEYWORDS)
    has_doctype = any(dt.lower() in t for dt in MONITOR_DOC_TYPES)
    return has_keyword and has_doctype

async def send_monitor_alert(findings: list[dict], webhook_url: str = "") -> bool:
    """Envia alerta por webhook quando norma relevante é detectada."""
    if not findings:
        return False
    payload = {
        "text": f"🚨 *ValidaRótulo — Nova norma POA detectada*\n"
                + "\n".join(f"• {f['title']}\n  {f['link']}" for f in findings),
        "findings": findings,
        "timestamp": __import__("datetime").datetime.now().isoformat(),
    }
    # Webhook (Slack, Discord, n8n, Zapier — qualquer URL)
    if webhook_url:
        try:
            async with httpx.AsyncClient(timeout=10.0) as client:
                await client.post(webhook_url, json=payload)
        except Exception:
            pass
    # Também armazena localmente para o endpoint /monitor/status
    _monitor_history.append(payload)
    if len(_monitor_history) > 100:
        _monitor_history.pop(0)
    import datetime as _dt3
    asyncio.ensure_future(_sb_insert("monitor_alertas", {
        **payload, "created_at": _dt3.datetime.now().isoformat()
    }))
    return True

_monitor_history: list[dict] = []

@app.post("/monitor/check")
async def monitor_check(request: Request):
    """
    Endpoint chamado diariamente pelo cron job.
    Busca novas publicações relevantes no DOU e envia alertas.
    Configurar cron em: https://cron-job.org (grátis)
    URL: POST https://valida-rotulo-backend.onrender.com/monitor/check
    Frequência: diária às 09:00 BRT
    """
    async with _monitor_lock:
        all_findings = []
        for kw in ["rotulagem produto origem animal", "RTIQ MAPA", "ANVISA rotulagem"]:
            results = await fetch_dou_recent(kw, days_back=2)
            for r in results:
                if is_relevant_publication(r["title"]) and r not in all_findings:
                    all_findings.append(r)

        webhook_url = os.environ.get("MONITOR_WEBHOOK_URL", "")
        alerted = await send_monitor_alert(all_findings, webhook_url)

        return JSONResponse({
            "status": "ok",
            "checked_at": __import__("datetime").datetime.now().isoformat(),
            "findings": len(all_findings),
            "alerted": alerted,
            "publications": all_findings,
        }, headers={"Access-Control-Allow-Origin": "*"})

@app.get("/monitor/status")
async def monitor_status():
    """Mostra histórico de alertas enviados."""
    return JSONResponse({
        "total_alerts": len(_monitor_history),
        "recent": _monitor_history[-5:] if _monitor_history else [],
        "webhook_configured": bool(os.environ.get("MONITOR_WEBHOOK_URL")),
        "setup_instructions": {
            "cron": "https://cron-job.org → POST /monitor/check diariamente às 09:00",
            "webhook": "Render dashboard → Environment → MONITOR_WEBHOOK_URL=<sua URL>",
            "slack_example": "https://hooks.slack.com/services/XXX/YYY/ZZZ",
        }
    }, headers={"Access-Control-Allow-Origin": "*"})

@app.post("/kb/update")
async def kb_update(request: Request):
    """
    Recebe texto de nova norma e atualiza o cache da KB.
    Usar quando uma norma nova for detectada e revisada.
    Body: { "categoria": "embutidos", "texto": "...", "fonte": "IN XX/AAAA" }
    """
    try:
        body = await request.json()
        categoria = body.get("categoria", "manual")
        texto     = body.get("texto", "")
        fonte     = body.get("fonte", "")
        if not texto:
            return JSONResponse({"error": "texto obrigatório"},
                                headers={"Access-Control-Allow-Origin": "*"})
        _kb_cache[categoria] = f"[Atualizado manualmente — {fonte}]\n{texto}"
        return JSONResponse({
            "status": "ok",
            "categoria": categoria,
            "chars": len(texto),
            "fonte": fonte,
        }, headers={"Access-Control-Allow-Origin": "*"})
    except Exception as e:
        return JSONResponse({"error": str(e)},
                            headers={"Access-Control-Allow-Origin": "*"})


# ══════════════════════════════════════════════════════════════════════════════
# SISTEMA 2 — APRENDIZADO POR USO (FEW-SHOT LOOP)
# Armazena validações + feedback do RT → melhora próximas validações
# ══════════════════════════════════════════════════════════════════════════════

import hashlib as _hashlib

# Banco de casos em memória (persiste durante o processo, reset no redeploy)
# Para produção: substituir por Supabase/PlanetScale (gratuitos)

# ══════════════════════════════════════════════════════════════════════════════
# SUPABASE — PERSISTÊNCIA DE DADOS
# Configurar no Render: SUPABASE_URL e SUPABASE_KEY (anon key)
# Se variáveis não existirem, cai em fallback in-memory (desenvolvimento)
# ══════════════════════════════════════════════════════════════════════════════

_SUPABASE_URL = os.environ.get("SUPABASE_URL", "")
_SUPABASE_KEY = os.environ.get("SUPABASE_KEY", "")
_SUPABASE_ON  = bool(_SUPABASE_URL and _SUPABASE_KEY)

async def _sb_get(table: str, filters: dict = None, limit: int = 500) -> list[dict]:
    """Lê registros do Supabase. Retorna [] se Supabase não configurado."""
    if not _SUPABASE_ON:
        return []
    url = f"{_SUPABASE_URL}/rest/v1/{table}?limit={limit}&order=created_at.desc"
    if filters:
        for k, v in filters.items():
            url += f"&{k}=eq.{v}"
    try:
        async with httpx.AsyncClient(timeout=8.0) as client:
            r = await client.get(url, headers={
                "apikey": _SUPABASE_KEY,
                "Authorization": f"Bearer {_SUPABASE_KEY}",
            })
            return r.json() if r.status_code == 200 else []
    except Exception:
        return []

async def _sb_upload_rotulo(contents: bytes, case_id: str, content_type: str = "image/jpeg") -> str:
    """
    Faz upload da imagem do rótulo pro bucket 'rotulos' no Supabase Storage.
    Retorna URL pública (string vazia se falhar — não bloqueia validação).

    Bucket precisa existir e estar configurado como Public.
    Coluna validacoes.imagem_url precisa existir.
    """
    if not _SUPABASE_ON:
        print(f"[storage] upload SKIPPED: supabase off", flush=True)
        return ""
    if not contents:
        print(f"[storage] upload SKIPPED: contents vazio (case={case_id[:16]})", flush=True)
        return ""
    if not case_id:
        print(f"[storage] upload SKIPPED: case_id vazio", flush=True)
        return ""

    # Determina extensão pelo content-type (default jpg)
    ct = (content_type or "").lower()
    if "png" in ct:
        ext = "png"
    elif "webp" in ct:
        ext = "webp"
    elif "pdf" in ct:
        ext = "pdf"
    else:
        ext = "jpg"

    file_name = f"{case_id}.{ext}"
    upload_url = f"{_SUPABASE_URL}/storage/v1/object/rotulos/{file_name}"
    public_url = f"{_SUPABASE_URL}/storage/v1/object/public/rotulos/{file_name}"

    print(f"[storage] uploading case={case_id[:16]} size={len(contents)}b ext={ext} ct={content_type}", flush=True)

    try:
        async with httpx.AsyncClient(timeout=20.0) as client:
            r = await client.post(
                upload_url,
                content=contents,
                headers={
                    "apikey": _SUPABASE_KEY,
                    "Authorization": f"Bearer {_SUPABASE_KEY}",
                    "Content-Type": content_type or "image/jpeg",
                    "x-upsert": "true",  # sobrescreve se já existir (evita conflito em retry)
                },
            )
            if r.status_code in (200, 201):
                print(f"[storage] upload OK case={case_id[:16]} url={public_url}", flush=True)
                return public_url
            # Logar pra debug — falha aqui não é crítica (sistema continua sem foto)
            print(f"[storage] upload FALHOU case={case_id[:16]} status={r.status_code} body={r.text[:300]}", flush=True)
            return ""
    except Exception as e:
        print(f"[storage] upload EXCEÇÃO case={case_id[:16]} err={str(e)[:300]}", flush=True)
        return ""


async def _sb_upsert(table: str, data: dict, conflict_col: str = "") -> bool:
    """
    Insere ou atualiza registro no Supabase. Retorna True se OK.
    conflict_col: coluna UNIQUE para resolver conflitos (ex: 'chave' para kb_documents).
    Sem conflict_col, faz INSERT simples com merge-duplicates header.
    """
    if not _SUPABASE_ON:
        return False
    try:
        # Para tabelas com coluna UNIQUE conhecida, usa ?on_conflict=col
        # Isso garante que o Supabase faça UPDATE quando o registro já existe
        url = f"{_SUPABASE_URL}/rest/v1/{table}"
        if conflict_col:
            url += f"?on_conflict={conflict_col}"
        async with httpx.AsyncClient(timeout=10.0) as client:
            r = await client.post(
                url,
                headers={
                    "apikey": _SUPABASE_KEY,
                    "Authorization": f"Bearer {_SUPABASE_KEY}",
                    "Content-Type": "application/json",
                    "Prefer": "resolution=merge-duplicates,return=minimal",
                },
                json=data,
            )
            return r.status_code in (200, 201, 204)
    except Exception:
        return False

async def _sb_insert(table: str, data: dict) -> bool:
    """Insere registro no Supabase."""
    if not _SUPABASE_ON:
        return False
    try:
        async with httpx.AsyncClient(timeout=8.0) as client:
            r = await client.post(
                f"{_SUPABASE_URL}/rest/v1/{table}",
                headers={
                    "apikey": _SUPABASE_KEY,
                    "Authorization": f"Bearer {_SUPABASE_KEY}",
                    "Content-Type": "application/json",
                    "Prefer": "return=minimal",
                },
                json=data,
            )
            return r.status_code in (200, 201)
    except Exception:
        return False


async def _sb_upsert_v2(table: str, data: dict, conflict_col: str = "case_id") -> bool:
    """
    UPSERT robusto: usa on_conflict para fazer UPDATE quando case_id já existe.
    Loga corpo do erro no Sentry se falhar (visibilidade real).
    """
    if not _SUPABASE_ON:
        return False
    try:
        url = f"{_SUPABASE_URL}/rest/v1/{table}?on_conflict={conflict_col}"
        async with httpx.AsyncClient(timeout=10.0) as client:
            r = await client.post(
                url,
                headers={
                    "apikey": _SUPABASE_KEY,
                    "Authorization": f"Bearer {_SUPABASE_KEY}",
                    "Content-Type": "application/json",
                    "Prefer": "resolution=merge-duplicates,return=minimal",
                },
                json=data,
            )
            if r.status_code in (200, 201, 204):
                return True
            # Falhou — loga corpo da resposta no Sentry
            err_body = ""
            try:
                err_body = r.text[:500]
            except Exception:
                pass
            if SENTRY_DSN:
                try:
                    import sentry_sdk
                    sentry_sdk.capture_message(
                        f"_sb_upsert_v2 falhou | table={table} | status={r.status_code} | body={err_body} | data_keys={list(data.keys())}",
                        level="error"
                    )
                except Exception:
                    pass
            return False
    except Exception as e:
        if SENTRY_DSN:
            try:
                import sentry_sdk
                sentry_sdk.capture_exception(e)
            except Exception:
                pass
        return False

async def load_cases_from_supabase() -> list[dict]:
    """
    Carrega casos do Supabase para memória no startup.
    Prioridade: casos com feedback RT primeiro, depois auto-stored.
    Deduplica por case_id mantendo o mais recente.
    """
    if not _SUPABASE_ON:
        return []
    try:
        async with httpx.AsyncClient(timeout=12.0) as client:
            r1 = await client.get(
                f"{_SUPABASE_URL}/rest/v1/validacoes",
                headers={"apikey": _SUPABASE_KEY,
                         "Authorization": f"Bearer {_SUPABASE_KEY}"},
                params={"select": "*",
                        "feedback": "not.is.null",
                        "order": "created_at.desc",
                        "limit": "300"}
            )
            feedback_rows = r1.json() if r1.status_code == 200 and isinstance(r1.json(), list) else []
            r2 = await client.get(
                f"{_SUPABASE_URL}/rest/v1/validacoes",
                headers={"apikey": _SUPABASE_KEY,
                         "Authorization": f"Bearer {_SUPABASE_KEY}"},
                params={"select": "*",
                        "feedback": "is.null",
                        "order": "created_at.desc",
                        "limit": "200"}
            )
            auto_rows = r2.json() if r2.status_code == 200 and isinstance(r2.json(), list) else []
        all_rows = feedback_rows + auto_rows
        seen = set()
        deduped = []
        for row in all_rows:
            cid = row.get("case_id", "")
            if cid and cid not in seen:
                seen.add(cid)
                deduped.append(row)
            elif not cid:
                deduped.append(row)
        return deduped[:500]
    except Exception:
        return []

async def load_monitor_from_supabase() -> list[dict]:
    """Carrega histórico de alertas do Supabase."""
    rows = await _sb_get("monitor_alertas", limit=100)
    return rows

_cases_db: list[dict] = []
_MAX_CASES = 500  # máximo de casos em memória

def _case_id(image_b64: str) -> str:
    """Gera ID único para uma imagem."""
    return _hashlib.md5(image_b64[:500].encode()).hexdigest()[:12]

def get_fewshot_examples(categoria: str, max_examples: int = 3,
                         caminho_np: str = "") -> str:
    """
    Recupera exemplos para few-shot injection.
    Paridade POA/não-POA: agrupa por categoria POA OU por caminho_np para não-POA.
    Prioridade:
      1. Mesma categoria/caminho + feedback RT positivo (mais confiável)
      2. Mesma categoria/caminho + auto-validado com score alto (≥12/14)
      3. Qualquer categoria com alta confiança
    Extrai padrões de erros recorrentes para alertar proativamente.
    """
    if not _cases_db:
        return ""

    # Chave de agrupamento: categoria POA ou caminho_np para não-POA
    def _match(c: dict) -> bool:
        if caminho_np:
            return (c.get("caminho_np") == caminho_np or
                    c.get("categoria") == caminho_np)
        return c.get("categoria") == categoria

    # Tier 1: feedback RT explícito + mesma categoria/caminho
    tier1 = [c for c in _cases_db
             if _match(c)
             and c.get("feedback") in ("correto", "parcialmente_correto")]

    # Tier 2: auto-validado com score ≥10 + mesma categoria/caminho
    tier2 = [c for c in _cases_db
             if _match(c)
             and c.get("feedback") is None
             and (c.get("score_agente") or 0) >= 10
             and c.get("erros_auto")]

    # Tier 3: qualquer categoria com score alto
    tier3 = [c for c in _cases_db
             if c.get("feedback") is None
             and (c.get("score_agente") or 0) >= 10]

    pool = (tier1 + tier2 + tier3)[:max_examples * 4]
    if not pool:
        return ""

    # Detecta erros RECORRENTES na categoria/caminho
    from collections import Counter
    cat_cases = [c for c in _cases_db if _match(c)]
    erros_recorrentes = ""
    if len(cat_cases) >= 3:
        all_erros = " ".join(
            (c.get("erros_auto") or "") + " " + (c.get("erros_encontrados") or "")
            for c in cat_cases
        ).lower()
        palavras = Counter(w for w in all_erros.split()
                           if len(w) > 5 and w not in {"campo","nenhum","conforme","verificar"})
        top = [w for w, n in palavras.most_common(5) if n >= max(2, len(cat_cases) * 0.3)]
        if top:
            erros_recorrentes = (
                f"\nPADRÕES DE ERRO RECORRENTES nesta categoria "
                f"(detectados em {len(cat_cases)} validações anteriores):\n"
                f"Fique especialmente atento a: {', '.join(top)}\n"
            )

    # Exemplos concretos
    examples = pool[:max_examples]
    parts = []
    for ex in examples:
        origem = "RT" if ex.get("feedback") else "auto-validado"
        score  = ex.get("score_agente")
        erros  = ex.get("erros_encontrados") or ex.get("erros_auto") or "nenhum registrado"
        falsos = ex.get("falsos_positivos", "")
        nao_det = ex.get("erros_nao_detectados", "")
        rt_comment = ex.get("rt_comment", "")

        bloco = (
            f"[{ex.get('produto','produto')} — {ex.get('categoria','?')} | origem: {origem}"
            + (f" | score agente: {score}/14" if score else "")
            + (f" | score real RT: {ex.get("score_real")}/14" if ex.get('score_real') else "")
            + "]"
        )
        if erros and erros != "nenhum registrado":
            bloco += f"\nErros detectados: {erros[:200]}"
        if nao_det:
            bloco += f"\n⚠️ Erros NÃO detectados pelo agente (RT corrigiu): {nao_det[:150]}"
        if falsos:
            bloco += f"\n⚠️ Falsos positivos (agente alertou sem necessidade): {falsos[:150]}"
        if rt_comment:
            bloco += f"\n💬 Observação do RT: {rt_comment[:150]}"
        parts.append(bloco)

    header = f"\n## BASE DE CONHECIMENTO PRÁTICO ({len(_cases_db)} validações | {sum(1 for x in _cases_db if x.get('feedback'))} com feedback RT)"
    if erros_recorrentes:
        header += erros_recorrentes
    if parts:
        header += "\nEXEMPLOS E LIÇÕES APRENDIDAS:\n" + "\n---\n".join(parts)
    return header

@app.post("/feedback")
async def store_feedback(request: Request):
    """
    Armazena feedback estruturado do RT após uma validação.

    Body completo (todos opcionais exceto case_id):
    {
      "case_id": "abc123",                    # obrigatório — retornado pelo /validar
      "produto": "Linguiça Suína Frescal",
      "categoria": "embutidos",
      "orgao": "SIE",
      "score_agente": 10,                     # score que o agente deu (12 máx)
      "score_real": 9,                        # score real atribuído pelo RT

      # Avaliação geral
      "feedback": "correto"|"parcialmente_correto"|"incorreto",

      # Avaliação campo a campo (12 campos)
      # Valores: "correto" | "incorreto" | "nao_verificado"
      "campos": {
        "1":  {"status": "correto",    "comentario": ""},
        "2":  {"status": "incorreto",  "comentario": "Sódio declarado errado: é 950mg não 800mg"},
        "3":  {"status": "correto",    "comentario": ""},
        "4":  {"status": "correto",    "comentario": ""},
        "5":  {"status": "nao_verificado", "comentario": ""},
        "6":  {"status": "correto",    "comentario": ""},
        "7":  {"status": "incorreto",  "comentario": "Temperatura deveria ser ≤7°C não ≤10°C"},
        "8":  {"status": "correto",    "comentario": ""},
        "9":  {"status": "correto",    "comentario": ""},
        "10": {"status": "correto",    "comentario": ""},
        "11": {"status": "correto",    "comentario": ""},
        "12": {"status": "nao_verificado", "comentario": ""}
      },

      # Comentário geral do RT
      "rt_comment": "O agente não percebeu que a bebida láctea tem soro na composição",

      # Erros que o agente PERDEU (não detectou mas deveria ter detectado)
      "erros_nao_detectados": "Campo 2: ingredientes fora de ordem decrescente",

      # Falsos positivos (o agente alertou mas estava errado)
      "falsos_positivos": "Campo 9: agente calculou Atwater errado para esta categoria"
    }
    """
    try:
        body = await request.json()
        import datetime as _dt2, json as _json

        campos_raw = body.get("campos", {})
        # Serializa campos para string legível (few-shot vai ler isso)
        erros_campos = []
        acertos_campos = []
        for num, dados in (campos_raw.items() if isinstance(campos_raw, dict) else {}.items()):
            st = dados.get("status", "") if isinstance(dados, dict) else ""
            cm = dados.get("comentario", "") if isinstance(dados, dict) else ""
            if st == "incorreto":
                erros_campos.append(f"C{num}: {cm}" if cm else f"C{num}: incorreto")
            elif st == "correto":
                acertos_campos.append(f"C{num}")

        erros_str = "; ".join(erros_campos) if erros_campos else ""
        # Combina com erros_nao_detectados e falsos_positivos
        erros_nao_det = body.get("erros_nao_detectados", "")
        falsos_pos    = body.get("falsos_positivos", "")
        erros_completo = " | ".join(filter(None, [erros_str, erros_nao_det]))

        # Backfill: se frontend não mandou produto/categoria/orgao, busca da validação original
        # (auto-store do /validar gravou esses dados antes do feedback chegar)
        produto_in   = (body.get("produto") or "").strip()
        categoria_in = (body.get("categoria") or "").strip()
        orgao_in     = (body.get("orgao") or "").strip()
        if not (produto_in and categoria_in):
            try:
                # Busca da validação original (auto_stored) na tabela validacoes
                async with httpx.AsyncClient(timeout=5.0) as _client_bf:
                    _r_bf = await _client_bf.get(
                        f"{_SUPABASE_URL}/rest/v1/validacoes?case_id=eq.{body.get('case_id', '')}&select=produto,categoria,orgao&limit=1",
                        headers={
                            "apikey": _SUPABASE_KEY,
                            "Authorization": f"Bearer {_SUPABASE_KEY}",
                            "Accept": "application/json",
                        }
                    )
                    if _r_bf.status_code == 200:
                        _rows_bf = _r_bf.json() or []
                        if _rows_bf:
                            _orig = _rows_bf[0]
                            if not produto_in and _orig.get("produto"):
                                produto_in = _orig["produto"]
                            if not categoria_in and _orig.get("categoria"):
                                categoria_in = _orig["categoria"]
                            if not orgao_in and _orig.get("orgao"):
                                orgao_in = _orig["orgao"]
            except Exception:
                pass  # backfill é best-effort, não pode quebrar feedback

        case = {
            "case_id":              body.get("case_id", ""),
            "produto":              produto_in,
            "categoria":            categoria_in,
            "caminho_np":           body.get("caminho_np", ""),
            "orgao":                orgao_in,
            "feedback":             body.get("feedback", ""),
            "score_agente":         body.get("score_agente"),
            "score_real":           body.get("score_real"),
            "campos":               campos_raw if isinstance(campos_raw, dict) else {},
            "erros_encontrados":    erros_completo,
            "erros_nao_detectados": erros_nao_det,
            "falsos_positivos":     falsos_pos,
            "rt_comment":           body.get("rt_comment", ""),
            "total_discordancias":  body.get("total_discordancias", 0),
            "versao_form":          body.get("versao_form", ""),
            "status_curador":       "pendente_revisao",
            "timestamp":            _dt2.datetime.now().isoformat(),
        }

        # Atualiza caso existente ou adiciona novo (memória in-memory)
        existing = next((x for x in _cases_db if x["case_id"] == case["case_id"]), None)
        if existing:
            existing.update(case)
        else:
            _cases_db.append(case)
            if len(_cases_db) > _MAX_CASES:
                _cases_db.pop(0)

        # Persiste no Supabase — UPSERT com conflict resolution em case_id
        # Use await pra capturar erro e retornar status real ao frontend
        supabase_case = {**case, "created_at": _dt2.datetime.now().isoformat()}
        # Remove None values (Supabase rejeita)
        supabase_case = {k: v for k, v in supabase_case.items() if v is not None}
        sb_ok = await _sb_upsert_v2("validacoes", supabase_case, conflict_col="case_id")

        # Calcula métricas do feedback para retornar ao frontend
        campos_ok  = len(acertos_campos)
        campos_err = len(erros_campos)
        precisao   = round(campos_ok / (campos_ok + campos_err) * 100) if (campos_ok + campos_err) > 0 else None

        # Se Supabase falhou, loga no Sentry pra rastrear (mas não falha pro RT
        # — feedback fica em memória pelo menos, não perde informação)
        if not sb_ok and SENTRY_DSN:
            try:
                import sentry_sdk
                sentry_sdk.capture_message(
                    f"Falha ao gravar feedback no Supabase | case_id={case['case_id']}",
                    level="error"
                )
            except Exception:
                pass

        return JSONResponse({
            "status":         "ok" if sb_ok else "ok_memoria",
            "case_id":        case["case_id"],
            "total_cases":    len(_cases_db),
            "persistido":     sb_ok,
            "campos_corretos": campos_ok,
            "campos_errados":  campos_err,
            "precisao_pct":    precisao,
            "mensagem":        (
                "✅ Feedback registrado com detalhamento campo a campo. O sistema aprende com cada correção."
                if erros_campos else
                "✅ Feedback registrado. Relatório marcado como correto pelo RT."
            ),
        }, headers={"Access-Control-Allow-Origin": "*"})
    except Exception as e:
        return JSONResponse({"error": str(e)},
                            headers={"Access-Control-Allow-Origin": "*"})


@app.get("/feedback/export")
async def export_feedback(format: str = "json"):
    """
    Exporta todos os feedbacks para análise.
    ?format=json (padrão) ou ?format=csv
    """
    try:
        if format == "csv":
            import io as _io, csv as _csv
            buf = _io.StringIO()
            writer = _csv.DictWriter(buf, fieldnames=[
                "case_id","produto","categoria","orgao","feedback",
                "score_agente","score_real","erros_encontrados",
                "erros_nao_detectados","falsos_positivos","rt_comment","timestamp"
            ])
            writer.writeheader()
            for case in _cases_db:
                writer.writerow({k: case.get(k, "") for k in writer.fieldnames})
            return Response(
                content=buf.getvalue(),
                media_type="text/csv",
                headers={
                    "Content-Disposition": "attachment; filename=feedbacks_inspect_ia.csv",
                    "Access-Control-Allow-Origin": "*"
                }
            )
        return JSONResponse({
            "total": len(_cases_db),
            "com_feedback_rt": sum(1 for x in _cases_db if x.get("feedback")),
            "feedbacks": _cases_db
        }, headers={"Access-Control-Allow-Origin": "*"})
    except Exception as e:
        return JSONResponse({"error": str(e)},
                            headers={"Access-Control-Allow-Origin": "*"})



# ═══════════════════════════════════════════════════════════════════════════════
# M1_2 — PLACEHOLDER — endpoint adicionado abaixo
# ═══════════════════════════════════════════════════════════════════════════════
# ═══════════════════════════════════════════════════════════════════════════════
# M1_2 — RELATÓRIO PDF COM CAMPO DE ASSINATURA DO RT
# ═══════════════════════════════════════════════════════════════════════════════

@app.post("/relatorio-pdf")
async def gerar_relatorio_pdf(request: Request):
    """
    Gera PDF formal do relatório para entrega ao cliente.
    Layout fiel à UI: identidade InspectIA, score panel, cards de campo coloridos.
    Body JSON: { produto, relatorio, score, veredicto, nome_rt, crm_rt, case_id }
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.lib import colors
        from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                        Table, TableStyle, HRFlowable, KeepTogether,
                                        PageBreak, Flowable)
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
        import io as _io, datetime as _dt, re as _re

        body = await request.json()
        produto   = body.get("produto", "Produto não identificado")
        relatorio = body.get("relatorio", "")
        score     = body.get("score")
        veredicto = (body.get("veredicto") or "").upper().strip()
        nome_rt   = body.get("nome_rt", "")
        crm_rt    = body.get("crm_rt", "")
        case_id   = body.get("case_id", "")
        data_str  = _dt.datetime.now().strftime("%d/%m/%Y às %H:%M")

        # ── PALETA INSPECTIA ────────────────────────────────────────────
        COR_PRIMARIA  = colors.HexColor("#0F2A20")  # verde escuro identidade
        COR_ACENTO    = colors.HexColor("#166534")  # verde da marca
        COR_BONE      = colors.HexColor("#FBFAF6")  # bone (fundo neutro)
        COR_BORDA     = colors.HexColor("#E2DED2")
        COR_BORDA_F   = colors.HexColor("#D9D2BE")
        COR_TEXTO     = colors.HexColor("#091a14")
        COR_MUTED     = colors.HexColor("#6B6757")

        COR_PASS_BG   = colors.HexColor("#E8F0EA")
        COR_PASS_FG   = colors.HexColor("#166534")
        COR_WARN_BG   = colors.HexColor("#F5EBD8")
        COR_WARN_FG   = colors.HexColor("#9c5a0e")
        COR_FAIL_BG   = colors.HexColor("#F0DDDB")
        COR_FAIL_FG   = colors.HexColor("#992f2a")

        COR_AVISO_BG  = colors.HexColor("#FFF8E1")
        COR_AVISO_BD  = colors.HexColor("#F5C772")
        COR_AVISO_FG  = colors.HexColor("#6B4A00")

        # ── DOC ─────────────────────────────────────────────────────────
        buf = _io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4,
                                leftMargin=15*mm, rightMargin=15*mm,
                                topMargin=14*mm, bottomMargin=14*mm,
                                title=f"Relatório InspectIA — {produto[:40]}",
                                author="InspectIA")
        W = A4[0] - 30*mm
        styles = getSampleStyleSheet()

        # Estilos comuns
        st_logo = ParagraphStyle("logo", fontSize=15, fontName="Helvetica-Bold",
                                 textColor=colors.white, leading=18)
        st_logo_sub = ParagraphStyle("logosub", fontSize=8, fontName="Helvetica",
                                     textColor=colors.HexColor("#bcd4c2"), leading=10,
                                     letterSpacing=0.6)
        st_meta = ParagraphStyle("meta", fontSize=8, fontName="Helvetica",
                                 textColor=COR_MUTED, leading=11)
        st_produto_label = ParagraphStyle("plbl", fontSize=8, fontName="Helvetica-Bold",
                                          textColor=COR_MUTED, leading=10,
                                          letterSpacing=0.8)
        st_produto = ParagraphStyle("p", fontSize=14, fontName="Helvetica-Bold",
                                    textColor=COR_TEXTO, leading=18)

        st_passo_num = ParagraphStyle("psn", fontSize=8, fontName="Helvetica-Bold",
                                      textColor=COR_ACENTO, leading=10,
                                      letterSpacing=0.6)
        st_passo_titulo = ParagraphStyle("pst", fontSize=11, fontName="Helvetica-Bold",
                                         textColor=COR_TEXTO, leading=14,
                                         letterSpacing=0.4)

        st_campo_num = ParagraphStyle("cn", fontSize=8, fontName="Helvetica-Bold",
                                      textColor=COR_TEXTO, leading=10)
        st_campo_titulo = ParagraphStyle("ct", fontSize=10, fontName="Helvetica-Bold",
                                         textColor=COR_TEXTO, leading=13)
        st_campo_status = ParagraphStyle("csc", fontSize=8, fontName="Helvetica-Bold",
                                         leading=10, alignment=TA_RIGHT)

        st_body = ParagraphStyle("b", fontSize=9.5, fontName="Helvetica",
                                 textColor=COR_TEXTO, leading=13)
        st_body_bold = ParagraphStyle("bb", fontSize=9.5, fontName="Helvetica-Bold",
                                      textColor=COR_TEXTO, leading=13)
        st_check = ParagraphStyle("ck", fontSize=9, fontName="Helvetica",
                                  textColor=COR_TEXTO, leading=12,
                                  leftIndent=12, spaceBefore=1, spaceAfter=1)
        st_aviso = ParagraphStyle("av", fontSize=8.5, fontName="Helvetica",
                                  textColor=COR_AVISO_FG, leading=11)

        st_assinatura_label = ParagraphStyle("asl", fontSize=8, fontName="Helvetica-Bold",
                                             textColor=COR_MUTED, leading=10,
                                             letterSpacing=0.6)
        st_assinatura_value = ParagraphStyle("asv", fontSize=10, fontName="Helvetica",
                                             textColor=COR_TEXTO, leading=13)

        # ── HELPERS ─────────────────────────────────────────────────────
        def html_escape(s):
            return (s.replace("&", "&amp;").replace("<", "&lt;")
                     .replace(">", "&gt;").replace("\u200b", ""))

        def md_to_html(s):
            """Converte markdown simples (**bold**) para HTML do reportlab."""
            s = html_escape(s)
            # **bold** → <b>bold</b>
            s = _re.sub(r'\*\*([^*]+)\*\*', r'<b>\1</b>', s)
            return s

        story = []

        # ── HEADER COM LOGO ─────────────────────────────────────────────
        header_tb = Table([[
            Paragraph("<b>InspectIA</b>", st_logo),
            Paragraph(f"Validador de Rótulos · Emitido em {data_str}", st_logo_sub)
        ]], colWidths=[W*0.35, W*0.65])
        header_tb.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,-1), COR_PRIMARIA),
            ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
            ("LEFTPADDING", (0,0), (-1,-1), 14),
            ("RIGHTPADDING", (0,0), (-1,-1), 14),
            ("TOPPADDING", (0,0), (-1,-1), 12),
            ("BOTTOMPADDING", (0,0), (-1,-1), 12),
            ("ALIGN", (1,0), (1,0), "RIGHT"),
        ]))
        story.append(header_tb)
        story.append(Spacer(1, 6*mm))

        # ── PRODUTO ─────────────────────────────────────────────────────
        story.append(Paragraph("PRODUTO ANALISADO", st_produto_label))
        story.append(Spacer(1, 1.5*mm))
        story.append(Paragraph(html_escape(produto), st_produto))
        if case_id:
            story.append(Spacer(1, 1*mm))
            story.append(Paragraph(f"ID: {html_escape(case_id[:16])}", st_meta))
        story.append(Spacer(1, 5*mm))

        # ── PARSE: extrair campos do relatório ──────────────────────────
        # Conta campos por status pra montar score correto
        passC = warnC = failC = 0
        # Acha todos os campos no markdown
        # Padrão: linhas começando com CAMPO N — ... seguido de status
        campo_pattern = _re.compile(
            r'^[\*#\s]{0,6}CAMPO\s+(\d+)\s*[—–\-]\s*([^:\*\n]+?)\**\s*:\s*\**\s*'
            r'(✅[^(\n]*|⚠️[^(\n]*|❌[^(\n]*|🔍[^(\n]*|N[ÃA]O\s+VERIFIC[ÁA]VEL[^(\n]*)?',
            _re.IGNORECASE | _re.MULTILINE
        )
        for m in campo_pattern.finditer(relatorio):
            status = m.group(3) or ""
            if "✅" in status or "CONFORME" in status.upper() and "NÃO" not in status.upper():
                passC += 1
            elif "❌" in status or "NÃO CONFORME" in status.upper():
                failC += 1
            elif "⚠" in status or "RESSALVA" in status.upper() or "AUSENTE" in status.upper():
                warnC += 1
            elif "🔍" in status or "VERIFIC" in status.upper():
                warnC += 1

        # Score recalculado
        if passC + warnC + failC > 0:
            score_calc = passC + (warnC * 0.5)
        else:
            score_calc = score if isinstance(score, (int, float)) else 0

        # Veredicto recalculado
        if score_calc >= 13:
            veredicto_calc = "APROVADO"
            ver_bg, ver_fg = COR_PASS_BG, COR_PASS_FG
        elif score_calc >= 8:
            veredicto_calc = "APROVADO COM RESSALVAS"
            ver_bg, ver_fg = COR_WARN_BG, COR_WARN_FG
        else:
            veredicto_calc = "REPROVADO"
            ver_bg, ver_fg = COR_FAIL_BG, COR_FAIL_FG

        # ── SCORE PANEL ─────────────────────────────────────────────────
        score_str = f"{score_calc:.1f}".rstrip("0").rstrip(".") + "/14"
        # Barra de progresso visual
        pct = min(1.0, score_calc / 14)

        score_big = ParagraphStyle("sb", fontSize=28, fontName="Helvetica-Bold",
                                   textColor=COR_TEXTO, leading=32)
        ver_pill_st = ParagraphStyle("vp", fontSize=10, fontName="Helvetica-Bold",
                                     textColor=ver_fg, leading=13, alignment=TA_CENTER)

        # Linha 1 do score: número grande + veredicto pill
        score_top = Table([[
            Paragraph(score_str, score_big),
            Paragraph(html_escape(veredicto_calc), ver_pill_st)
        ]], colWidths=[W*0.55, W*0.45])
        score_top.setStyle(TableStyle([
            ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
            ("LEFTPADDING", (0,0), (-1,-1), 0),
            ("RIGHTPADDING", (0,0), (-1,-1), 0),
            ("BACKGROUND", (1,0), (1,0), ver_bg),
            ("TOPPADDING", (1,0), (1,0), 8),
            ("BOTTOMPADDING", (1,0), (1,0), 8),
            ("ROUNDEDCORNERS", [4, 4, 4, 4]) if False else ("BOX", (1,0), (1,0), 0.5, ver_bg),
        ]))

        # Linha 2: contadores
        cont_label_st = lambda fg: ParagraphStyle("cl", fontSize=7, fontName="Helvetica-Bold",
                                                   textColor=fg, leading=10, letterSpacing=0.6,
                                                   alignment=TA_LEFT)
        cont_num_st = lambda fg: ParagraphStyle("cn", fontSize=18, fontName="Helvetica-Bold",
                                                 textColor=fg, leading=22, alignment=TA_LEFT)

        contadores_tb = Table([[
            Table([
                [Paragraph(str(passC), cont_num_st(COR_PASS_FG))],
                [Paragraph("✓ CONFORMES", cont_label_st(COR_PASS_FG))],
            ], colWidths=[W*0.31], style=[
                ("BACKGROUND", (0,0), (-1,-1), COR_PASS_BG),
                ("LEFTPADDING", (0,0), (-1,-1), 10),
                ("RIGHTPADDING", (0,0), (-1,-1), 10),
                ("TOPPADDING", (0,0), (0,0), 8),
                ("BOTTOMPADDING", (0,1), (0,1), 8),
                ("BOX", (0,0), (-1,-1), 0.5, COR_PASS_FG),
            ]),
            Table([
                [Paragraph(str(warnC), cont_num_st(COR_WARN_FG))],
                [Paragraph("⚠ COM RESSALVAS", cont_label_st(COR_WARN_FG))],
            ], colWidths=[W*0.31], style=[
                ("BACKGROUND", (0,0), (-1,-1), COR_WARN_BG),
                ("LEFTPADDING", (0,0), (-1,-1), 10),
                ("RIGHTPADDING", (0,0), (-1,-1), 10),
                ("TOPPADDING", (0,0), (0,0), 8),
                ("BOTTOMPADDING", (0,1), (0,1), 8),
                ("BOX", (0,0), (-1,-1), 0.5, COR_WARN_FG),
            ]),
            Table([
                [Paragraph(str(failC), cont_num_st(COR_FAIL_FG))],
                [Paragraph("✕ NÃO CONFORMES", cont_label_st(COR_FAIL_FG))],
            ], colWidths=[W*0.31], style=[
                ("BACKGROUND", (0,0), (-1,-1), COR_FAIL_BG),
                ("LEFTPADDING", (0,0), (-1,-1), 10),
                ("RIGHTPADDING", (0,0), (-1,-1), 10),
                ("TOPPADDING", (0,0), (0,0), 8),
                ("BOTTOMPADDING", (0,1), (0,1), 8),
                ("BOX", (0,0), (-1,-1), 0.5, COR_FAIL_FG),
            ]),
        ]], colWidths=[W*0.33, W*0.33, W*0.33])
        contadores_tb.setStyle(TableStyle([
            ("LEFTPADDING", (0,0), (-1,-1), 0),
            ("RIGHTPADDING", (0,0), (-1,-1), 0),
            ("TOPPADDING", (0,0), (-1,-1), 0),
            ("BOTTOMPADDING", (0,0), (-1,-1), 0),
        ]))

        story.append(score_top)
        story.append(Spacer(1, 4*mm))
        story.append(contadores_tb)
        story.append(Spacer(1, 6*mm))

        # ── PARSE EM PASSOS ─────────────────────────────────────────────
        # Quebra o relatório em passos identificados por "PASSO N — TÍTULO"
        passo_split_re = _re.compile(
            r'^[\*#\s]{0,6}PASSO\s+(\d+)\s*[—–\-]\s*(.+?)[\*\s]*$',
            _re.IGNORECASE | _re.MULTILINE
        )

        # Encontra começos de passo
        passos = []
        for m in passo_split_re.finditer(relatorio):
            passos.append({
                "num": int(m.group(1)),
                "titulo": m.group(2).replace("*", "").strip(),
                "start": m.end(),
                "header_start": m.start(),
            })

        # Adiciona o fim de cada passo (= start do próximo, ou fim do texto)
        for i, p in enumerate(passos):
            p["end"] = passos[i+1]["header_start"] if i+1 < len(passos) else len(relatorio)
            p["body"] = relatorio[p["start"]:p["end"]].strip()

        # ── HELPERS DE RENDER ───────────────────────────────────────────
        def render_passo_header(num, titulo, body_flowables):
            """Card do passo com header colorido + body como flowables soltos."""
            # Header: tabelinha com PASSO N + título
            header_tb = Table([[
                Paragraph(f"<b>PASSO {num}</b>", st_passo_num),
                Paragraph(html_escape(titulo), st_passo_titulo)
            ]], colWidths=[26*mm, W - 26*mm])
            header_tb.setStyle(TableStyle([
                ("BACKGROUND", (0,0), (-1,-1), COR_BONE),
                ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
                ("LEFTPADDING", (0,0), (-1,-1), 12),
                ("RIGHTPADDING", (0,0), (-1,-1), 12),
                ("TOPPADDING", (0,0), (-1,-1), 9),
                ("BOTTOMPADDING", (0,0), (-1,-1), 9),
                ("LINEABOVE", (0,0), (-1,-1), 0.8, COR_ACENTO),
                ("LINEBELOW", (0,0), (-1,-1), 0.5, COR_BORDA_F),
                ("LINEAFTER", (0,0), (0,0), 0.5, COR_BORDA_F),
            ]))

            # Retorna header + body flowables soltos (sem envolver em Table — evita LayoutError)
            result = [header_tb, Spacer(1, 3*mm)]
            result.extend(body_flowables)
            result.append(Spacer(1, 5*mm))
            return result

        def parse_passo_body(body_text, passo_num):
            """Converte texto do body do passo em flowables."""
            flowables = []
            lines = body_text.split("\n")

            # CAMPO collector (pra Passo 3)
            current_campo = None
            campo_buffer = []

            def flush_campo():
                if current_campo is None:
                    return
                # Renderiza card de campo
                num, nome, status_raw = current_campo
                if "✅" in status_raw:
                    bg, fg, label = COR_PASS_BG, COR_PASS_FG, "✓ Conforme"
                elif "❌" in status_raw or "NÃO CONFORME" in status_raw.upper():
                    bg, fg, label = COR_FAIL_BG, COR_FAIL_FG, "✕ Não Conforme"
                elif "⚠" in status_raw or "RESSALVA" in status_raw.upper() or "AUSENTE" in status_raw.upper():
                    bg, fg, label = COR_WARN_BG, COR_WARN_FG, "⚠ Com Ressalvas"
                else:
                    bg, fg, label = COR_BONE, COR_MUTED, "—"

                campo_header_tb = Table([[
                    Paragraph(f"<b>C{num}</b>", st_campo_num),
                    Paragraph(html_escape(nome.strip()), st_campo_titulo),
                    Paragraph(f"<font color='{fg.hexval()}'>{label}</font>",
                             ParagraphStyle("ci", fontSize=8, fontName="Helvetica-Bold",
                                           leading=10, alignment=TA_RIGHT))
                ]], colWidths=[15*mm, W - 15*mm - 30*mm, 30*mm])
                campo_header_tb.setStyle(TableStyle([
                    ("BACKGROUND", (0,0), (-1,-1), bg),
                    ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
                    ("LEFTPADDING", (0,0), (-1,-1), 8),
                    ("RIGHTPADDING", (0,0), (-1,-1), 8),
                    ("TOPPADDING", (0,0), (-1,-1), 6),
                    ("BOTTOMPADDING", (0,0), (-1,-1), 6),
                ]))

                # Campo body (do buffer)
                campo_body_items = []
                for ln in campo_buffer:
                    ln_clean = ln.strip()
                    if not ln_clean:
                        continue
                    # Bullets ✓ / ✗ / ⚠
                    bullet_m = _re.match(r'^[\s]*([✓✗❌✅⚠️⚠])\s+(.+)$', ln_clean)
                    if bullet_m:
                        sym = bullet_m.group(1)
                        rest = bullet_m.group(2)
                        # Cor do símbolo
                        if sym in ("✓", "✅"):
                            sym_color = COR_PASS_FG.hexval()
                        elif sym in ("✗", "❌"):
                            sym_color = COR_FAIL_FG.hexval()
                        else:
                            sym_color = COR_WARN_FG.hexval()
                        campo_body_items.append(Paragraph(
                            f"<font color='{sym_color}'><b>{sym}</b></font> {md_to_html(rest)}",
                            st_check
                        ))
                    elif ln_clean.startswith("**Sugestões"):
                        campo_body_items.append(Spacer(1, 2*mm))
                        campo_body_items.append(Paragraph(md_to_html(ln_clean), st_body_bold))
                    else:
                        campo_body_items.append(Paragraph(md_to_html(ln_clean), st_body))

                if campo_body_items:
                    # Header sempre junto com primeira linha do body
                    flowables.append(KeepTogether([campo_header_tb,
                                                    Spacer(1, 1*mm),
                                                    campo_body_items[0]]))
                    # Restante das linhas (podem quebrar página naturalmente)
                    for item in campo_body_items[1:]:
                        flowables.append(item)
                    flowables.append(Spacer(1, 4*mm))
                else:
                    flowables.append(campo_header_tb)
                    flowables.append(Spacer(1, 4*mm))

            for line in lines:
                line = line.rstrip()
                stripped = line.strip()
                if not stripped or _re.match(r'^[─━=-]+$', stripped):
                    continue

                # Detecta CAMPO header
                campo_m = _re.match(
                    r'^[\*#\s]{0,6}CAMPO\s+(\d+)\s*[—–\-]\s*([^:\*\n]+?)\**\s*:\s*\**\s*(.*?)\**\s*$',
                    stripped
                )
                if campo_m:
                    flush_campo()
                    current_campo = (campo_m.group(1), campo_m.group(2), campo_m.group(3))
                    campo_buffer = []
                    continue

                # Aviso legal
                if stripped.startswith(">") and ("aviso" in stripped.lower() or "⚠" in stripped):
                    # Pula — vamos renderizar separado
                    continue

                # SCORE/VEREDICTO no Passo 4 — pula (já mostrados no topo)
                if _re.match(r'^\*?\*?SCORE\s*:', stripped, _re.IGNORECASE):
                    continue
                if _re.match(r'^\*?\*?VEREDICTO\s*:', stripped, _re.IGNORECASE):
                    continue

                # Header ### TÍTULO
                if _re.match(r'^#{2,4}\s+', stripped):
                    titulo = _re.sub(r'^#+\s*', '', stripped).rstrip(":#").strip()
                    if current_campo:
                        campo_buffer.append(stripped)
                    else:
                        flowables.append(Spacer(1, 2*mm))
                        flowables.append(Paragraph(
                            f"<b>{md_to_html(titulo)}</b>",
                            ParagraphStyle("h", fontSize=10, fontName="Helvetica-Bold",
                                          textColor=COR_TEXTO, leading=13)
                        ))
                        flowables.append(Spacer(1, 1*mm))
                    continue

                # Item numerado (Passo 4)
                num_item_m = _re.match(
                    r'^(\d+)[.\)]\s*\*{0,2}([^*:]+)\*{0,2}\s*:\s*(.+)$', stripped
                )
                if num_item_m and not current_campo and passo_num == 4:
                    num = num_item_m.group(1)
                    titulo = num_item_m.group(2).strip()
                    desc = num_item_m.group(3).replace("**", "").strip()
                    item_tb = Table([[
                        Paragraph(f"<b>{num}</b>",
                                 ParagraphStyle("nn", fontSize=10, fontName="Helvetica-Bold",
                                               textColor=colors.white, leading=14, alignment=TA_CENTER)),
                        Paragraph(f"<b>{md_to_html(titulo)}:</b> {md_to_html(desc)}", st_body)
                    ]], colWidths=[10*mm, W - 10*mm - 8*mm])

                    # Cor do número baseada em criticidade
                    is_critical = bool(_re.search(r'cr[ií]tic|eliminat[óo]rio|reprovad', titulo + " " + desc, _re.IGNORECASE))
                    is_priority = bool(_re.search(r'prioridade|m[áa]xima|urgente', titulo, _re.IGNORECASE))
                    if is_critical:
                        num_bg, item_bg, item_border = COR_FAIL_FG, colors.HexColor("#FBEEEC"), COR_FAIL_FG
                    elif is_priority:
                        num_bg, item_bg, item_border = COR_WARN_FG, colors.HexColor("#FBF1E0"), COR_WARN_FG
                    else:
                        num_bg, item_bg, item_border = COR_MUTED, COR_BONE, COR_BORDA_F

                    item_tb.setStyle(TableStyle([
                        ("BACKGROUND", (0,0), (0,0), num_bg),
                        ("BACKGROUND", (1,0), (1,0), item_bg),
                        ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
                        ("ALIGN", (0,0), (0,0), "CENTER"),
                        ("LEFTPADDING", (0,0), (-1,-1), 6),
                        ("RIGHTPADDING", (0,0), (-1,-1), 6),
                        ("TOPPADDING", (0,0), (-1,-1), 6),
                        ("BOTTOMPADDING", (0,0), (-1,-1), 6),
                        ("LINEBEFORE", (1,0), (1,0), 1.5, item_border),
                    ]))
                    flowables.append(item_tb)
                    flowables.append(Spacer(1, 2*mm))
                    continue

                # Bullet
                bullet_m = _re.match(r'^[•·▸\-]\s+(.+)$', stripped)
                if bullet_m and not current_campo:
                    inner = bullet_m.group(1)
                    flowables.append(Paragraph(
                        f"<font color='{COR_PASS_FG.hexval()}'><b>✓</b></font> {md_to_html(inner)}",
                        st_check
                    ))
                    continue

                # Linha LABEL: valor
                label_m = _re.match(
                    r'^\*{0,2}([A-ZÁÀÂÃÉÊÍÓÔÕÚÇ][A-ZÁÀÂÃÉÊÍÓÔÕÚÇ\s\-]*?)\*{0,2}\s*:\s*\*{0,2}\s*(.*?)\*{0,2}\s*$',
                    stripped
                )
                if label_m and not current_campo:
                    label = label_m.group(1).strip()
                    value = label_m.group(2).strip()
                    flowables.append(Paragraph(
                        f"<font color='{COR_PASS_FG.hexval()}'><b>✓</b></font> "
                        f"<b>{html_escape(label)}:</b> {md_to_html(value)}",
                        st_check
                    ))
                    continue

                # Linha normal
                if current_campo:
                    campo_buffer.append(stripped)
                else:
                    flowables.append(Paragraph(md_to_html(stripped), st_body))

            flush_campo()
            return flowables

        # ── RENDER PASSOS ───────────────────────────────────────────────
        if passos:
            for p in passos:
                p_flowables = parse_passo_body(p["body"], p["num"])
                if p_flowables:
                    story.extend(render_passo_header(p["num"], p["titulo"], p_flowables))
        else:
            # Fallback: sem passos detectados, renderiza linha a linha
            for ln in relatorio.split("\n"):
                ln = ln.strip()
                if ln:
                    story.append(Paragraph(md_to_html(ln), st_body))

        # ── AVISO LEGAL ─────────────────────────────────────────────────
        story.append(Spacer(1, 4*mm))
        aviso_text = (
            "<b>Aviso legal:</b> Este relatório é gerado por inteligência artificial e tem caráter auxiliar. "
            "Não substitui a análise e aprovação de Responsável Técnico habilitado. "
            "A conformidade regulatória final é de responsabilidade exclusiva do RT inscrito no órgão competente."
        )
        aviso_tb = Table([[
            Paragraph("⚠", ParagraphStyle("av_icon", fontSize=14,
                                          textColor=COR_AVISO_FG, alignment=TA_CENTER)),
            Paragraph(aviso_text, st_aviso)
        ]], colWidths=[8*mm, W - 8*mm])
        aviso_tb.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,-1), COR_AVISO_BG),
            ("VALIGN", (0,0), (-1,-1), "TOP"),
            ("LEFTPADDING", (0,0), (-1,-1), 8),
            ("RIGHTPADDING", (0,0), (-1,-1), 10),
            ("TOPPADDING", (0,0), (-1,-1), 8),
            ("BOTTOMPADDING", (0,0), (-1,-1), 8),
            ("BOX", (0,0), (-1,-1), 0.8, COR_AVISO_BD),
        ]))
        story.append(aviso_tb)

        # ── ASSINATURA RT ───────────────────────────────────────────────
        story.append(Spacer(1, 8*mm))
        story.append(HRFlowable(width=W, thickness=0.5, color=COR_BORDA_F))
        story.append(Spacer(1, 4*mm))
        story.append(Paragraph("RESPONSÁVEL TÉCNICO", st_assinatura_label))
        story.append(Spacer(1, 3*mm))

        nome_show = nome_rt if nome_rt else "_______________________________________"
        crm_show = crm_rt if crm_rt else "_______________"
        assin_tb = Table([[
            Paragraph(f"Nome: {html_escape(nome_show)}", st_assinatura_value),
            Paragraph(f"Registro: {html_escape(crm_show)}", st_assinatura_value),
        ]], colWidths=[W*0.6, W*0.4])
        assin_tb.setStyle(TableStyle([
            ("VALIGN", (0,0), (-1,-1), "TOP"),
            ("LEFTPADDING", (0,0), (-1,-1), 0),
            ("RIGHTPADDING", (0,0), (-1,-1), 0),
        ]))
        story.append(assin_tb)
        story.append(Spacer(1, 2*mm))
        story.append(Paragraph(f"Data da análise: {data_str}", st_meta))

        # Build
        doc.build(story)
        filename = f"InspectIA_relatorio_{produto[:30].replace(' ','_')}_{_dt.datetime.now().strftime('%Y%m%d')}.pdf"
        from fastapi.responses import Response as _Resp
        return _Resp(content=buf.getvalue(), media_type="application/pdf",
                     headers={"Content-Disposition": f'attachment; filename="{filename}"',
                              "Access-Control-Allow-Origin": "*"})

    except ImportError as e:
        print(f"[PDF] ImportError: {e}", flush=True)
        return JSONResponse(
            {"error": f"reportlab não instalado: {str(e)[:200]}. Adicione ao requirements.txt."},
            status_code=500, headers={"Access-Control-Allow-Origin": "*"})
    except Exception as e:
        import traceback as _tb
        tb_str = _tb.format_exc()
        print(f"[PDF] Erro ao gerar relatório PDF: {type(e).__name__}: {e}", flush=True)
        print(f"[PDF] Traceback:\n{tb_str}", flush=True)
        return JSONResponse(
            {"error": f"{type(e).__name__}: {str(e)[:300]}"},
            status_code=500,
            headers={"Access-Control-Allow-Origin": "*"})


# ═══════════════════════════════════════════════════════════════════════════════
# M1_5 — HISTÓRICO DE VERSÕES POR PRODUTO (SUPABASE)
# Tabela: rotulo_versoes (produto_id, versao, campos_json, relatorio_json, created_at)
# ═══════════════════════════════════════════════════════════════════════════════

@app.post("/versoes/salvar")
async def salvar_versao(request: Request):
    """
    Salva uma versão do rótulo no histórico.
    Body: { produto_id, nome_produto, campos_json, relatorio_json, score, veredicto }
    """
    try:
        import datetime as _dt
        body = await request.json()
        produto_id   = body.get("produto_id") or body.get("nome_produto", "produto").lower().replace(" ", "_")[:40]
        nome_produto = body.get("nome_produto", "")
        campos_json  = body.get("campos_json", {})
        relatorio    = body.get("relatorio_json", "")
        score        = body.get("score")
        veredicto    = body.get("veredicto", "")

        # Conta versões existentes para numerar
        versoes_existentes = await _sb_get("rotulo_versoes",
                                           filters={"produto_id": produto_id}, limit=100)
        num_versao = len(versoes_existentes) + 1

        registro = {
            "produto_id":    produto_id,
            "nome_produto":  nome_produto,
            "versao":        num_versao,
            "campos_json":   json.dumps(campos_json, ensure_ascii=False),
            "relatorio_json": relatorio[:8000] if relatorio else "",
            "score":         score,
            "veredicto":     veredicto,
            "created_at":    _dt.datetime.now().isoformat(),
        }

        salvo = await _sb_insert("rotulo_versoes", registro)

        return JSONResponse({
            "ok": True,
            "versao": num_versao,
            "produto_id": produto_id,
            "persistido": salvo,
            "mensagem": f"Versão {num_versao} salva com sucesso."
        }, headers={"Access-Control-Allow-Origin": "*"})

    except Exception as e:
        return JSONResponse({"error": str(e)[:200]}, status_code=500,
                            headers={"Access-Control-Allow-Origin": "*"})


@app.get("/versoes/{produto_id}")
async def listar_versoes(produto_id: str):
    """Lista todas as versões de um produto, mais recente primeiro."""
    try:
        rows = await _sb_get("rotulo_versoes",
                             filters={"produto_id": produto_id}, limit=50)
        # Ordena por versão decrescente
        rows.sort(key=lambda x: x.get("versao", 0), reverse=True)
        return JSONResponse({
            "produto_id": produto_id,
            "total": len(rows),
            "versoes": [{
                "versao":       r.get("versao"),
                "nome_produto": r.get("nome_produto", ""),
                "score":        r.get("score"),
                "veredicto":    r.get("veredicto", ""),
                "created_at":   r.get("created_at", ""),
                "campos_json":  r.get("campos_json", "{}"),
            } for r in rows]
        }, headers={"Access-Control-Allow-Origin": "*"})
    except Exception as e:
        return JSONResponse({"error": str(e)[:200]}, status_code=500,
                            headers={"Access-Control-Allow-Origin": "*"})


@app.get("/versoes/{produto_id}/{versao}")
async def recuperar_versao(produto_id: str, versao: int):
    """Recupera campos_json completo de uma versão específica."""
    try:
        rows = await _sb_get("rotulo_versoes",
                             filters={"produto_id": produto_id}, limit=100)
        row = next((r for r in rows if r.get("versao") == versao), None)
        if not row:
            return JSONResponse({"error": "Versão não encontrada."},
                                status_code=404, headers={"Access-Control-Allow-Origin": "*"})
        campos = {}
        try:
            campos = json.loads(row.get("campos_json", "{}"))
        except Exception:
            pass
        return JSONResponse({
            "versao":       row.get("versao"),
            "nome_produto": row.get("nome_produto", ""),
            "score":        row.get("score"),
            "veredicto":    row.get("veredicto", ""),
            "created_at":   row.get("created_at", ""),
            "campos":       campos,
            "relatorio":    row.get("relatorio_json", ""),
        }, headers={"Access-Control-Allow-Origin": "*"})
    except Exception as e:
        return JSONResponse({"error": str(e)[:200]}, status_code=500,
                            headers={"Access-Control-Allow-Origin": "*"})


@app.get("/admin/stats")
async def admin_stats():
    """Painel de aprendizado: métricas de qualidade por categoria."""
    if not _cases_db:
        return JSONResponse({
            "total_cases": 0,
            "message": "Nenhuma validação com feedback ainda."
        }, headers={"Access-Control-Allow-Origin": "*"})

    from collections import Counter, defaultdict
    categorias = Counter(c["categoria"] for c in _cases_db if c.get("categoria"))
    feedbacks  = Counter(c["feedback"]  for c in _cases_db if c.get("feedback"))

    # Precisão por categoria
    precisao = {}
    by_cat = defaultdict(list)
    for c in _cases_db:
        if c.get("categoria") and c.get("feedback"):
            by_cat[c["categoria"]].append(c["feedback"])
    for cat, fbs in by_cat.items():
        corretos = sum(1 for f in fbs if f in ("correto", "parcialmente_correto"))
        precisao[cat] = round(corretos / len(fbs) * 100) if fbs else 0

    # Erros mais comuns
    erros_text = " ".join(c.get("erros_encontrados", "") for c in _cases_db)
    erros_comuns = Counter(w for w in erros_text.lower().split()
                           if len(w) > 4 and w not in {"campo","nenhum","ok","sim","não","para"})

    auto_stored   = sum(1 for x in _cases_db if x.get("auto_stored"))
    com_feedback  = sum(1 for x in _cases_db if x.get("feedback"))
    score_medio   = None
    scores = [x["score_agente"] for x in _cases_db if x.get("score_agente") is not None]
    if scores:
        score_medio = round(sum(scores) / len(scores), 1)

    return JSONResponse({
        "total_cases":        len(_cases_db),
        "auto_armazenados":   auto_stored,
        "com_feedback_rt":    com_feedback,
        "score_medio_agente": score_medio,
        "feedbacks":          dict(feedbacks),
        "precisao_pct":       precisao,
        "top_categorias":     dict(categorias.most_common(8)),
        "erros_mais_comuns":  dict(erros_comuns.most_common(10)),
        "fewshot_por_categoria": {
            cat: {
                "total": sum(1 for x in _cases_db if x.get("categoria") == cat),
                "com_feedback": sum(1 for x in _cases_db if x.get("categoria") == cat and x.get("feedback")),
                "auto_alta_confianca": sum(1 for x in _cases_db
                                           if x.get("categoria") == cat
                                           and (x.get("score_agente") or 0) >= 10),
            }
            for cat in set(x.get("categoria","") for x in _cases_db) if cat
        },
        "como_melhorar": {
            "proxima_melhoria": "Adicione feedback RT em 10+ casos para ativar exemplos mais precisos",
            "status_fewshot": (
                "Ativo — injetando exemplos" if len(_cases_db) >= 5
                else f"Aguardando {5 - len(_cases_db)} validações para ativar"
            ),
        }
    }, headers={"Access-Control-Allow-Origin": "*"})



# ══════════════════════════════════════════════════════════════════════════════
# CRIADOR DE RÓTULOS — Gera conteúdo + PDF/PNG profissional
# ══════════════════════════════════════════════════════════════════════════════

SP_CRIAR_ROTULO = """Você é um especialista em rotulagem de alimentos brasileira.
Sua tarefa é gerar TODOS os campos obrigatórios de um rótulo de alimento conforme a legislação vigente.

RETORNE SOMENTE JSON válido, sem markdown, sem texto fora do JSON.

Estrutura obrigatória:
{
  "denominacao": "Denominação de venda exata conforme RTIQ",
  "ingredientes": "Lista completa em ordem decrescente com INS e funções tecnológicas",
  "conteudo_liquido": "Conteúdo líquido formatado com símbolo INMETRO (ex: 200 g ℮)",
  "fabricante": "Razão Social\nCNPJ: XX.XXX.XXX/XXXX-XX\nEndereço completo\nCidade - UF CEP",
  "carimbo": "SIF/SIE/SIM + número (ex: SIF 3456)",
  "conservacao": "Instrução de conservação com temperatura específica",
  "lote": "Veja fundo da embalagem",
  "validade": "Veja fundo da embalagem",
  "alergenos": "CONTÉM: X, Y E DERIVADOS. / PODE CONTER: Z.",
  "transgenicos": "(Não contém transgênicos)" ou texto conforme aplicável,
  "lupa_necessaria": true/false,
  "lupa_motivo": "ALTO EM SÓDIO" / "ALTO EM GORDURAS SATURADAS" / "ALTO EM AÇÚCAR" ou null,
  "porcao": "100g (porção padrão conforme IN 75/2020)",
  "tabela_nutricional": {
    "porcao": "100g",
    "energia_kcal": "XXX kcal",
    "energia_kj": "XXX kJ",
    "carboidratos": "Xg",
    "acucares_totais": "Xg",
    "acucares_adicionados": "Xg",
    "proteinas": "Xg",
    "gorduras_totais": "Xg",
    "gorduras_saturadas": "Xg",
    "gorduras_trans": "0g",
    "fibra": "Xg",
    "sodio": "XXmg"
  },
  "gluten": "CONTÉM GLÚTEN" ou "NÃO CONTÉM GLÚTEN",
  "lactose": "CONTÉM LACTOSE" ou "NÃO CONTÉM LACTOSE" ou null,
  "observacoes_rt": "Orientações específicas ao RT sobre este produto",
  "legislacoes": ["IN 4/2000", "RDC 727/2022"]
}

REGRAS OBRIGATÓRIAS:
1. Denominação: use exatamente o nome técnico do RTIQ aplicável + espécie animal
2. Ingredientes: ordem decrescente, aditivos com INS e função tecnológica
3. Tabela nutricional: use valores TACO para a categoria se não informado
4. Alérgenos: 14 grupos da RDC 727/2022, CONTÉM para intencionais, PODE CONTER para cruzada
5. Lupa: obrigatória se sódio >= 600mg, gordura saturada >= 6g, ou açúcar adicionado >= 15g por 100g
6. Lote e validade: sempre "Veja fundo/tampa da embalagem" pois são impressos na linha
7. Conteúdo líquido: incluir símbolo ℮ (INMETRO) e unidade correta
"""

# ═══════════════════════════════════════════════════════════════════════════════
# NP10 — SPs ESPECÍFICOS POR CAMINHO/CATEGORIA
# ═══════════════════════════════════════════════════════════════════════════════

SP_CRIAR_POA = SP_CRIAR_ROTULO + """
CAMINHO: POA — PRODUTO DE ORIGEM ANIMAL
Normas: IN 22/2005 + Port. 1485/2025 (MAPA) + RTIQ específico da categoria.
OBRIGATÓRIO: carimbo SIF/SIE/SIM nunca pode ser null. Declarar espécie animal na denominação.
Citar RTIQ aplicável em legislacoes[]. CMS declarar espécie e %. Proteína de soja: declarar % em embutidos.
"""

SP_CRIAR_SUPLEMENTO = SP_CRIAR_ROTULO + """
CAMINHO: SUPLEMENTOS ALIMENTARES — RDC 243/2018 + RDC 786/2023 + IN 28/2018
carimbo: null (suplementos não têm SIF/SIE/SIM).
Adicione ao JSON: "notificacao_anvisa": "⚠️ NÚMERO FICTÍCIO — SUBSTITUIR PELO NÚMERO REAL ANTES DE USAR. Obter em: consultas.anvisa.gov.br" (NUNCA invente um número real — use EXATAMENTE esse texto),
"frase_nao_medicamento": "Este produto não é um medicamento" (obrigatório),
"cafeina_mg_porcao": mg de cafeína por porção se houver (RDC 786/2023).
Whey: mín. 10g proteína/porção. Cafeína: máx. 210mg/porção. Colágeno de peixe = alérgeno PEIXE.

⚠️ AVISO CRÍTICO — OBRIGATÓRIO SEGUIR:
NUNCA gere número de notificação ANVISA real ou verossímil (ex: "6.4832.0012.001-9", "123456789").
O número real é obtido EXCLUSIVAMENTE pelo fabricante no Portal ANVISA (consultas.anvisa.gov.br).
Usar número fictício em produto comercializado = INFRAÇÃO SANITÁRIA (Lei 6.437/1977) = embargo + multa.
O campo "notificacao_anvisa" DEVE conter exatamente o texto de aviso acima — nunca um número.
"""

SP_CRIAR_BEBIDA = SP_CRIAR_ROTULO + """
CAMINHO: BEBIDAS — Decreto 6.871/2009 + RDC 173/2006 + IN MAPA 37/2018 + RDC 273/2005
carimbo: null (exceto bebidas lácteas POA).
Adicione ao JSON: "tipo_bebida": "suco_integral"|"nectar"|"refresco"|"energetico"|"cerveja"|"vinho"|"agua",
"percentual_suco": % de suco (obrigatório para néctar e refresco),
"teor_alcoolico": % v/v (obrigatório para cerveja e vinho),
"cafeina_mg_porcao": mg por porção se energético.
Suco integral = 100% fruta, PROIBIDO aditivos conservantes. Néctar: mín. 30-50% suco por fruta.
Energético: frase advertência: "NÃO RECOMENDADO PARA CRIANÇAS, GESTANTES, IDOSOS E PORTADORES DE DOENÇAS CARDIOVASCULARES".
Porções IN 75/2020: suco/néctar=200mL, cerveja=330mL, vinho=150mL.
"""

SP_CRIAR_PANIFICACAO = SP_CRIAR_ROTULO + """
CAMINHO: PANIFICAÇÃO — RDC 90/2000 + RDC 263/2005 + RDC 711/2022 + RDC 712/2022
carimbo: null.
Adicione ao JSON: "tipo_produto": "pao"|"biscoito_salgado"|"biscoito_doce"|"massa"|"torrada"|"wafer",
"percentual_integral": % de farinha integral (obrigatório se "integral" na denominação — mín. 50%).
Farinha de trigo enriquecida: declarar ferro e ácido fólico. Lecitina de soja = alérgeno SOJA.
Qualquer trigo = CONTÉM GLÚTEN. Porções: pão=50g, biscoito=30g, macarrão cru=80g.
"""

SP_CRIAR_CHOCOLATE = SP_CRIAR_ROTULO + """
CAMINHO: CHOCOLATES E CACAU — RDC 264/2005
carimbo: null.
Adicione ao JSON: "tipo_chocolate": "ao_leite"|"meio_amargo"|"amargo"|"branco"|"cobertura",
"percentual_cacau": % de sólidos de cacau.
Mínimos: ao leite=25%, meio amargo=40%, amargo=50%, branco=20% manteiga de cacau.
Cobertura (<25% cacau): PROIBIDO usar "chocolate" na denominação.
Lecitina de soja = SOJA. Leite em pó = LEITE. Porção=30g.
"""

SP_CRIAR_CONDIMENTO = SP_CRIAR_ROTULO + """
CAMINHO: CONDIMENTOS, MOLHOS E TEMPEROS — RDC 276/2005
carimbo: null.
Adicione ao JSON: "tipo_condimento": "maionese"|"ketchup"|"mostarda"|"shoyu"|"molho_ingles"|"tempero_composto"|"vinagre",
"percentual_lipidios": % lipídios (maionese mín. 50% — abaixo disso = "Molho tipo maionese"),
"percentual_tomate": % extrato tomate (ketchup mín. 6%).
Shoyu: alérgenos SOJA + TRIGO obrigatórios. Maionese: alérgeno OVO obrigatório.
Porções: maionese/ketchup=15g, shoyu=10mL.
"""

SP_CRIAR_VEGETAL = SP_CRIAR_ROTULO + """
CAMINHO: VEGETAIS PROCESSADOS E CONSERVAS — RDC 272/2005 + RDC 714/2022
carimbo: null.
Adicione ao JSON: "tipo_vegetal": "palmito"|"cogumelo"|"azeitona"|"milho_conserva"|"ervilha_conserva"|"outros",
"especie_vegetal": espécie obrigatória (ex: "Pupunha", "Shitake", "Preta"),
"calibre": calibre da azeitona (obrigatório — ex: "Extra Grande"),
"peso_drenado": peso drenado em gramas (obrigatório — INMETRO Port. 157/2002).
Conteúdo líquido: declarar peso total E peso drenado. Palmito e cogumelo: espécie na denominação.
"""

SP_CRIAR_ORGANICO = SP_CRIAR_ROTULO + """
CAMINHO: ALIMENTOS ORGÂNICOS — Lei 10.831/2003 + Decreto 6.323/2007
carimbo: null se não-POA, SIF/SIE/SIM se POA orgânico.
Adicione ao JSON: "certificadora": "nome do OAC credenciado pelo MAPA",
"numero_certificado": "número do certificado",
"simbolo_sisorg": true (obrigatório no painel principal).
"Orgânico" sem certificação = fraude. Aditivos sintéticos: maioria proibida em orgânicos.
"""

SP_CRIAR_POR_CAMINHO = {
    "poa": SP_CRIAR_POA, "laticinios": SP_CRIAR_POA, "embutidos": SP_CRIAR_POA,
    "presunto": SP_CRIAR_POA, "bacon": SP_CRIAR_POA, "hamburguer": SP_CRIAR_POA,
    "salame": SP_CRIAR_POA, "charque": SP_CRIAR_POA, "aves": SP_CRIAR_POA,
    "pescado": SP_CRIAR_POA, "mel": SP_CRIAR_POA, "ovos": SP_CRIAR_POA,
    "carnes": SP_CRIAR_POA,
    "suplemento": SP_CRIAR_SUPLEMENTO, "whey": SP_CRIAR_SUPLEMENTO,
    "bebida": SP_CRIAR_BEBIDA, "suco": SP_CRIAR_BEBIDA,
    "cerveja": SP_CRIAR_BEBIDA, "vinho": SP_CRIAR_BEBIDA,
    "panificacao": SP_CRIAR_PANIFICACAO, "pao": SP_CRIAR_PANIFICACAO,
    "biscoito": SP_CRIAR_PANIFICACAO, "massa": SP_CRIAR_PANIFICACAO,
    "chocolate": SP_CRIAR_CHOCOLATE, "cacau": SP_CRIAR_CHOCOLATE,
    "condimento": SP_CRIAR_CONDIMENTO, "molho": SP_CRIAR_CONDIMENTO,
    "maionese": SP_CRIAR_CONDIMENTO,
    "vegetal": SP_CRIAR_VEGETAL, "conserva": SP_CRIAR_VEGETAL,
    "palmito": SP_CRIAR_VEGETAL,
    "organico": SP_CRIAR_ORGANICO,
}

def get_sp_criar(caminho: str) -> str:
    """Retorna o SP correto para o caminho/categoria selecionado."""
    return SP_CRIAR_POR_CAMINHO.get((caminho or "").lower(), SP_CRIAR_ROTULO)

import tempfile as _tempfile, uuid as _uuid

def _gerar_pdf_label(campos: dict, formato: str, tema: str) -> tuple[bytes, bytes]:
    """
    Gera PDF profissional do rótulo e converte para PNG.
    Retorna (pdf_bytes, png_bytes).
    Formatos: retangular (210x100mm), quadrado (100x100mm), circular (90x90mm)
    Temas: moderno, classico, premium
    """
    from reportlab.pdfgen import canvas as rl_canvas
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.platypus import Paragraph, Table, TableStyle
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import io as _io

    # ── Dimensões por formato ──────────────────────────────────────────────────
    dims = {
        "retangular": (210*mm, 100*mm),
        "quadrado":   (100*mm, 100*mm),
        "circular":   (90*mm,  90*mm),
    }
    W, H = dims.get(formato, dims["retangular"])

    # ── Paletas de tema ────────────────────────────────────────────────────────
    temas = {
        "moderno":  {"bg": colors.HexColor("#111827"), "acc": colors.HexColor("#16a34a"),
                     "txt": colors.white, "light": colors.HexColor("#f9fafb"),
                     "dark": colors.HexColor("#1f2937"), "border": colors.HexColor("#374151")},
        "classico": {"bg": colors.HexColor("#1e3a5f"), "acc": colors.HexColor("#c9922a"),
                     "txt": colors.white, "light": colors.HexColor("#f8f5f0"),
                     "dark": colors.HexColor("#1e3a5f"), "border": colors.HexColor("#2d5a8e")},
        "premium":  {"bg": colors.HexColor("#1a0a2e"), "acc": colors.HexColor("#c9a227"),
                     "txt": colors.white, "light": colors.HexColor("#fdf9f0"),
                     "dark": colors.HexColor("#2d1457"), "border": colors.HexColor("#4a1a7a")},
    }
    T = temas.get(tema, temas["moderno"])

    buf = _io.BytesIO()
    c = rl_canvas.Canvas(buf, pagesize=(W, H))

    # ── Fundo ──────────────────────────────────────────────────────────────────
    c.setFillColor(colors.white)
    c.rect(0, 0, W, H, fill=1, stroke=0)

    # ── Funções auxiliares ─────────────────────────────────────────────────────
    def rect_fill(x, y, w, h, color):
        c.setFillColor(color); c.rect(x, y, w, h, fill=1, stroke=0)

    def text(txt, x, y, size=7, bold=False, color=None, align="left", max_w=None):
        c.setFont("Helvetica-Bold" if bold else "Helvetica", size)
        c.setFillColor(color or colors.HexColor("#111827"))
        if align == "center":
            c.drawCentredString(x, y, str(txt)[:80])
        elif align == "right":
            c.drawRightString(x, y, str(txt)[:80])
        else:
            txt_str = str(txt)
            if max_w:
                # Truncate to fit
                while len(txt_str) > 3 and c.stringWidth(txt_str, "Helvetica-Bold" if bold else "Helvetica", size) > max_w:
                    txt_str = txt_str[:-4] + "..."
            c.drawString(x, y, txt_str)

    def wrap_text(txt, x, y, max_w, size=6.5, line_h=8, color=None, max_lines=6):
        """Simple word wrap."""
        c.setFont("Helvetica", size)
        c.setFillColor(color or colors.HexColor("#374151"))
        words = str(txt).split()
        lines, line = [], ""
        for w in words:
            test = (line + " " + w).strip()
            if c.stringWidth(test, "Helvetica", size) <= max_w:
                line = test
            else:
                if line: lines.append(line)
                line = w
        if line: lines.append(line)
        for i, ln in enumerate(lines[:max_lines]):
            c.drawString(x, y - i*line_h, ln)
        return len(lines[:max_lines]) * line_h

    # ══════════════════════════════════════════════════════════════════════════
    # LAYOUT RETANGULAR (210 x 100mm)
    # ══════════════════════════════════════════════════════════════════════════
    if formato == "retangular":
        margin = 3*mm
        # Cabeçalho escuro
        rect_fill(0, H - 22*mm, W, 22*mm, T["bg"])
        # Nome produto grande
        prod = (campos.get("denominacao") or "Produto")[:45]
        c.setFont("Helvetica-Bold", 14)
        c.setFillColor(T["txt"])
        c.drawString(margin, H - 14*mm, prod)
        # Carimbo no canto direito
        carimbo = campos.get("carimbo", "SIF 000")
        rect_fill(W - 28*mm, H - 19*mm, 25*mm, 16*mm, colors.white)
        c.setFont("Helvetica-Bold", 7)
        c.setFillColor(T["bg"])
        c.drawCentredString(W - 15.5*mm, H - 12*mm, carimbo)
        c.setFont("Helvetica", 5.5)
        c.drawCentredString(W - 15.5*mm, H - 16*mm, "INSPECIONADO")
        # Linha accent
        rect_fill(0, H - 24*mm, W, 2*mm, T["acc"])

        # ── Coluna esquerda (ingredientes + fabricante + conservação) ──────────
        col_w = W * 0.52
        y = H - 27*mm
        text("INGREDIENTES", margin, y, size=6, bold=True, color=T["dark"])
        y -= 2*mm
        wrap_text(campos.get("ingredientes", ""), margin, y, col_w - 2*margin, size=6, line_h=7.5, max_lines=5)
        y -= 42

        text("FABRICANTE", margin, y, size=6, bold=True, color=T["dark"])
        y -= 2*mm
        wrap_text(campos.get("fabricante", ""), margin, y, col_w - 2*margin, size=6, line_h=7.5, max_lines=3)
        y -= 26

        text("CONSERVAR: " + (campos.get("conservacao","") or "")[:60],
             margin, y, size=5.5, color=colors.HexColor("#6b7280"), max_w=col_w-2*margin)
        y -= 7
        text(campos.get("gluten","") + " · " + (campos.get("lactose","") or ""),
             margin, y, size=5.5, bold=True, color=colors.HexColor("#dc2626"), max_w=col_w-2*margin)
        y -= 7
        alergs = campos.get("alergenos","")[:100]
        text(alergs, margin, y, size=5.5, color=colors.HexColor("#374151"), max_w=col_w-2*margin)

        # Lote / validade / conteúdo
        y = 4*mm
        text(campos.get("conteudo_liquido",""), margin, y, size=7, bold=True, color=T["dark"])
        text("LOTE: Veja embalagem   VALIDADE: Veja embalagem",
             margin + 30*mm, y, size=5.5, color=colors.HexColor("#9ca3af"))

        # ── Divisória vertical ─────────────────────────────────────────────────
        c.setStrokeColor(colors.HexColor("#e5e7eb"))
        c.setLineWidth(0.5)
        c.line(col_w, H - 25*mm, col_w, 2*mm)

        # ── Tabela nutricional (coluna direita) ────────────────────────────────
        tx = col_w + 2*mm
        tw = W - col_w - margin
        tn = campos.get("tabela_nutricional", {})

        ty = H - 27*mm
        # Header tabela
        rect_fill(tx, ty - 5*mm, tw, 5*mm, T["dark"])
        c.setFont("Helvetica-Bold", 7); c.setFillColor(T["txt"])
        c.drawCentredString(tx + tw/2, ty - 3.5*mm, "Informação Nutricional")
        ty -= 5*mm
        rect_fill(tx, ty - 4*mm, tw, 4*mm, colors.HexColor("#f3f4f6"))
        c.setFont("Helvetica", 5.5); c.setFillColor(colors.HexColor("#374151"))
        c.drawString(tx + 1*mm, ty - 3*mm, f"Porção: {tn.get('porcao', '100g')} (1 porção)")
        ty -= 4*mm

        rows_nut = [
            ("Valor energético", f"{tn.get('energia_kcal','?')} = {tn.get('energia_kj','?')} kJ"),
            ("Carboidratos", tn.get("carboidratos","?")),
            ("   Acucares totais", tn.get("acucares_totais","?")),
            ("   Acucares adicionados", tn.get("acucares_adicionados","?")),
            ("Proteinas", tn.get("proteinas","?")),
            ("Gorduras totais", tn.get("gorduras_totais","?")),
            ("   Gord. saturadas", tn.get("gorduras_saturadas","?")),
            ("   Gord. trans", tn.get("gorduras_trans","0g")),
            ("Fibra alimentar", tn.get("fibra","?")),
            ("Sodio", tn.get("sodio","?")),
        ]
        for i, (nm, vl) in enumerate(rows_nut):
            row_bg = colors.HexColor("#ffffff") if i%2==0 else colors.HexColor("#f9fafb")
            rect_fill(tx, ty - 5*mm, tw, 5*mm, row_bg)
            c.setStrokeColor(colors.HexColor("#e5e7eb")); c.setLineWidth(0.3)
            c.line(tx, ty - 5*mm, tx + tw, ty - 5*mm)
            c.setFont("Helvetica", 5.5); c.setFillColor(colors.HexColor("#374151"))
            c.drawString(tx + 1*mm, ty - 3.5*mm, nm[:28])
            c.setFont("Helvetica-Bold", 5.5); c.setFillColor(colors.HexColor("#111827"))
            c.drawRightString(tx + tw - 1*mm, ty - 3.5*mm, str(vl)[:12])
            ty -= 5*mm

        # Lupa se necessário
        if campos.get("lupa_necessaria") and campos.get("lupa_motivo"):
            rect_fill(tx, ty - 6*mm, tw, 6*mm, colors.HexColor("#111827"))
            c.setFont("Helvetica-Bold", 7); c.setFillColor(colors.white)
            c.drawCentredString(tx + tw/2, ty - 4*mm, campos["lupa_motivo"][:25])
            ty -= 6*mm

        # Borda do rótulo inteiro
        c.setStrokeColor(colors.HexColor("#374151")); c.setLineWidth(1)
        c.rect(0.5, 0.5, W - 1, H - 1, fill=0, stroke=1)

    # ══════════════════════════════════════════════════════════════════════════
    # LAYOUT QUADRADO (100 x 100mm)
    # ══════════════════════════════════════════════════════════════════════════
    elif formato == "quadrado":
        m = 3*mm
        rect_fill(0, H - 20*mm, W, 20*mm, T["bg"])
        prod = (campos.get("denominacao") or "Produto")[:30]
        c.setFont("Helvetica-Bold", 11); c.setFillColor(T["txt"])
        c.drawCentredString(W/2, H - 13*mm, prod)
        c.setFont("Helvetica", 7); c.setFillColor(colors.HexColor("#9ca3af"))
        c.drawCentredString(W/2, H - 17*mm, campos.get("carimbo","SIF 000"))
        rect_fill(0, H - 22*mm, W, 2*mm, T["acc"])

        # Conteúdo líquido proeminente
        c.setFont("Helvetica-Bold", 14); c.setFillColor(T["dark"])
        c.drawCentredString(W/2, H - 28*mm, campos.get("conteudo_liquido",""))

        # Ingredientes
        y = H - 34*mm
        text("INGREDIENTES", m, y, size=5.5, bold=True, color=T["dark"])
        y -= 2*mm
        wrap_text(campos.get("ingredientes",""), m, y, W - 2*m, size=5.5, line_h=7, max_lines=4)
        y -= 34

        # Fabricante em uma linha
        fab = (campos.get("fabricante","") or "").split("\n")[0][:40]
        text("Fab: " + fab, m, y, size=5.5, color=colors.HexColor("#6b7280"), max_w=W-2*m)
        y -= 7

        # Alérgenos
        text(campos.get("gluten","") + " · " + (campos.get("lactose","") or ""),
             m, y, size=5.5, bold=True, color=colors.HexColor("#dc2626"), max_w=W-2*m)
        y -= 7
        text((campos.get("alergenos",""))[:80], m, y, size=5, color=colors.HexColor("#374151"), max_w=W-2*m)

        # Tabela nutricional compacta no rodapé
        tn = campos.get("tabela_nutricional", {})
        ty = 22*mm
        rect_fill(m, ty, W - 2*m, 4*mm, T["dark"])
        c.setFont("Helvetica-Bold", 6); c.setFillColor(T["txt"])
        c.drawCentredString(W/2, ty + 2.5*mm, f"Informacao Nutricional | Porcao: {tn.get('porcao','100g')}")
        ty += 4*mm
        cols_nut = [
            ("Energia", f"{tn.get('energia_kcal','?')}"),
            ("Prot.", tn.get("proteinas","?")),
            ("Carb.", tn.get("carboidratos","?")),
            ("Gord.", tn.get("gorduras_totais","?")),
            ("Sodio", tn.get("sodio","?")),
        ]
        cw = (W - 2*m) / len(cols_nut)
        for i, (nm, vl) in enumerate(cols_nut):
            cx = m + i*cw
            bg = colors.HexColor("#f9fafb") if i%2==0 else colors.white
            rect_fill(cx, m, cw, ty - m, bg)
            c.setFont("Helvetica", 5); c.setFillColor(colors.HexColor("#6b7280"))
            c.drawCentredString(cx + cw/2, ty - 4*mm, nm)
            c.setFont("Helvetica-Bold", 6); c.setFillColor(colors.HexColor("#111827"))
            c.drawCentredString(cx + cw/2, ty - 9*mm, str(vl)[:8])

        c.setStrokeColor(colors.HexColor("#374151")); c.setLineWidth(1)
        c.rect(0.5, 0.5, W-1, H-1, fill=0, stroke=1)

    # ══════════════════════════════════════════════════════════════════════════
    # LAYOUT CIRCULAR (90 x 90mm — renderizado em quadrado, circular via borda)
    # ══════════════════════════════════════════════════════════════════════════
    else:
        cx, cy, r = W/2, H/2, W/2 - 1*mm
        # Círculo de fundo
        c.setFillColor(T["bg"]); c.circle(cx, cy, r, fill=1, stroke=0)
        # Círculo accent externo
        c.setStrokeColor(T["acc"]); c.setLineWidth(3)
        c.circle(cx, cy, r - 1, fill=0, stroke=1)
        # Arco interno branco para conteúdo
        c.setFillColor(colors.white); c.circle(cx, cy, r - 8*mm, fill=1, stroke=0)

        # Produto no topo do círculo
        prod = (campos.get("denominacao") or "Produto")[:22]
        c.setFont("Helvetica-Bold", 9); c.setFillColor(T["txt"])
        c.drawCentredString(cx, H - 14*mm, prod)
        c.setFont("Helvetica", 6); c.setFillColor(T["acc"])
        c.drawCentredString(cx, H - 18*mm, campos.get("carimbo","SIF 000"))

        # Conteúdo líquido central
        c.setFont("Helvetica-Bold", 18); c.setFillColor(T["dark"])
        c.drawCentredString(cx, cy + 6*mm, campos.get("conteudo_liquido",""))

        # Fabricante
        fab = (campos.get("fabricante","") or "").split("\n")[0][:28]
        c.setFont("Helvetica", 6); c.setFillColor(colors.HexColor("#6b7280"))
        c.drawCentredString(cx, cy - 2*mm, fab)

        # Glúten/lactose
        c.setFont("Helvetica-Bold", 6); c.setFillColor(colors.HexColor("#dc2626"))
        c.drawCentredString(cx, cy - 8*mm, campos.get("gluten",""))

        # Conservação
        conserv = (campos.get("conservacao","") or "")[:35]
        c.setFont("Helvetica", 5.5); c.setFillColor(colors.HexColor("#374151"))
        c.drawCentredString(cx, 14*mm, conserv)
        c.setFont("Helvetica", 5); c.setFillColor(colors.HexColor("#9ca3af"))
        c.drawCentredString(cx, 10*mm, "LOTE e VALIDADE: veja embalagem")

    c.save()
    pdf_bytes = buf.getvalue()

    # ── Converte para PNG usando PyMuPDF ──────────────────────────────────────
    png_bytes = b""
    try:
        import fitz as _fitz
        doc = _fitz.open(stream=pdf_bytes, filetype="pdf")
        page = doc[0]
        mat = _fitz.Matrix(4, 4)  # 4x zoom = alta resolução
        pix = page.get_pixmap(matrix=mat, alpha=False)
        png_bytes = pix.tobytes("png")
    except Exception:
        pass

    return pdf_bytes, png_bytes


@app.post("/criar")
async def criar_rotulo(request: Request):
    """
    Gera rótulo completo: conteúdo via Claude + validação automática.
    Retorna SSE stream com campos JSON e depois validação streaming.
    """
    form = await request.form()
    produto     = form.get("produto", "")
    categoria   = form.get("categoria", "")
    caminho     = form.get("caminho", categoria)
    especie     = form.get("especie", "")
    orgao       = form.get("orgao", "SIF")
    num_reg     = form.get("num_registro", "")
    peso        = form.get("peso", "")
    fabricante  = form.get("fabricante", "")
    endereco    = form.get("endereco", "")
    cnpj        = form.get("cnpj", "")
    ingredientes= form.get("ingredientes", "")
    nutricional = form.get("nutricional", "")
    obs         = form.get("obs", "")
    formato     = form.get("formato", "retangular")
    tema        = form.get("tema", "moderno")
    # Campos extras por caminho (NP10)
    cafeina_mg         = form.get("cafeina_mg", "")
    notificacao_anvisa = form.get("notificacao_anvisa", "")
    tipo_bebida        = form.get("tipo_bebida", "")
    percentual_suco    = form.get("percentual_suco", "")
    teor_alcoolico     = form.get("teor_alcoolico", "")
    tipo_produto       = form.get("tipo_produto", "")
    percentual_integral= form.get("percentual_integral", "")
    tipo_chocolate     = form.get("tipo_chocolate", "")
    percentual_cacau   = form.get("percentual_cacau", "")
    tipo_condimento    = form.get("tipo_condimento", "")
    percentual_lipidios= form.get("percentual_lipidios", "")
    tipo_vegetal       = form.get("tipo_vegetal", "")
    especie_vegetal    = form.get("especie_vegetal", "")
    peso_drenado       = form.get("peso_drenado", "")
    certificadora      = form.get("certificadora", "")
    numero_certificado = form.get("numero_certificado", "")

    if not produto or not categoria:
        return JSONResponse({"error": "Produto e categoria são obrigatórios"},
                            headers={"Access-Control-Allow-Origin": "*"})

    # NP10 — seleciona SP correto pelo caminho
    sp_ativo = get_sp_criar(caminho)

    user_msg = f"""Gere o rótulo para:
CAMINHO: {caminho.upper()}
PRODUTO: {produto}
CATEGORIA: {categoria}
ESPÉCIE: {especie or "não informado"}
ÓRGÃO DE INSPEÇÃO: {orgao} {num_reg}
PESO/VOLUME: {peso}
FABRICANTE: {fabricante}
ENDEREÇO: {endereco}
CNPJ: {cnpj}
INGREDIENTES FORNECIDOS: {ingredientes or "gerar com base na categoria"}
INFO NUTRICIONAL: {nutricional or "calcular pela tabela TACO"}
OBSERVAÇÕES: {obs}
{f"CAFEÍNA POR PORÇÃO: {cafeina_mg}mg" if cafeina_mg else ""}
{f"NOTIFICAÇÃO ANVISA: {notificacao_anvisa}" if notificacao_anvisa else ""}
{f"TIPO DE BEBIDA: {tipo_bebida}" if tipo_bebida else ""}
{f"% SUCO: {percentual_suco}" if percentual_suco else ""}
{f"TEOR ALCOÓLICO: {teor_alcoolico}% v/v" if teor_alcoolico else ""}
{f"TIPO DE PRODUTO: {tipo_produto}" if tipo_produto else ""}
{f"% FARINHA INTEGRAL: {percentual_integral}" if percentual_integral else ""}
{f"TIPO DE CHOCOLATE: {tipo_chocolate}" if tipo_chocolate else ""}
{f"% SÓLIDOS DE CACAU: {percentual_cacau}" if percentual_cacau else ""}
{f"TIPO DE CONDIMENTO: {tipo_condimento}" if tipo_condimento else ""}
{f"% LIPÍDIOS: {percentual_lipidios}" if percentual_lipidios else ""}
{f"TIPO DE VEGETAL: {tipo_vegetal}" if tipo_vegetal else ""}
{f"ESPÉCIE/CALIBRE: {especie_vegetal}" if especie_vegetal else ""}
{f"PESO DRENADO: {peso_drenado}g" if peso_drenado else ""}
{f"CERTIFICADORA ORGÂNICA: {certificadora}" if certificadora else ""}
{f"Nº CERTIFICADO: {numero_certificado}" if numero_certificado else ""}

Retorne SOMENTE o JSON conforme especificado."""

    async def stream_criar():
        headers_cors = {"Access-Control-Allow-Origin": "*",
                        "Content-Type": "text/event-stream",
                        "Cache-Control": "no-cache"}
        try:
            # ── Fase 1: Gerar campos via Claude ──────────────────────────────
            resp = httpx.AsyncClient(timeout=60.0)
            async with resp as client:
                r = await client.post(
                    "https://api.anthropic.com/v1/messages",
                    headers={"x-api-key": os.environ.get("ANTHROPIC_API_KEY",""),
                             "anthropic-version": "2023-06-01",
                             "content-type": "application/json"},
                    json={"model": ANTHROPIC_MODEL,
                          "max_tokens": 2000,
                          "system": sp_ativo,
                          "messages": [{"role": "user", "content": user_msg}]}
                )
                data = r.json()
                raw_json = data["content"][0]["text"]

            # Limpa markdown se veio
            raw_json = raw_json.strip()
            if raw_json.startswith("```"):
                raw_json = raw_json.split("```")[1]
                if raw_json.startswith("json"):
                    raw_json = raw_json[4:]
            raw_json = raw_json.strip().rstrip("```")

            campos = json.loads(raw_json)

            # Emite campos para o frontend renderizar preview
            yield f"data: {json.dumps({'tipo': 'campos', 'dados': json.dumps(campos)})}\n\n"

            # ── Fase 2: Gerar PDF + PNG ───────────────────────────────────────
            try:
                pdf_bytes, png_bytes = _gerar_pdf_label(campos, formato, tema)
                pdf_b64 = base64.b64encode(pdf_bytes).decode()
                png_b64 = base64.b64encode(png_bytes).decode() if png_bytes else ""
                yield f"data: {json.dumps({'tipo': 'arquivos', 'pdf': pdf_b64, 'png': png_b64, 'produto': produto})}\n\n"
            except Exception as e:
                yield f"data: {json.dumps({'tipo': 'aviso', 'msg': 'Preview gerado; PDF indisponível: ' + str(e)[:60]})}\n\n"

            # ── Fase 3: Validação automática do rótulo gerado ─────────────────
            yield f"data: {json.dumps({'tipo': 'validacao_inicio'})}\n\n"

            kb_text = await get_kb_for_categories(detect_categories(categoria + " " + produto))
            kb_section = f"\n\nCONTEXTO LEGISLATIVO:\n{kb_text}" if kb_text else ""

            val_system = SP_VALIDACAO.replace("{kb_section}", kb_section)
            val_msg = f"""Valide este rótulo que acabou de ser GERADO automaticamente.
Produto: {campos.get('denominacao', produto)}
Campos gerados: {json.dumps(campos, ensure_ascii=False)[:3000]}"""

            async with httpx.AsyncClient(timeout=120.0) as client2:
                async with client2.stream("POST",
                    "https://api.anthropic.com/v1/messages",
                    headers={"x-api-key": os.environ.get("ANTHROPIC_API_KEY",""),
                             "anthropic-version": "2023-06-01",
                             "content-type": "application/json"},
                    json={"model": ANTHROPIC_MODEL,
                          "max_tokens": 16384, "stream": True,
                          "system": val_system,
                          "messages": [{"role": "user", "content": val_msg}]}
                ) as resp2:
                    async for line in resp2.aiter_lines():
                        if line.startswith("data:"):
                            ev = json.loads(line[5:].strip())
                            if ev.get("type") == "content_block_delta":
                                chunk = ev.get("delta",{}).get("text","")
                                if chunk:
                                    yield f"data: {json.dumps({'tipo': 'validacao_texto', 'text': chunk})}\n\n"

            yield "data: [DONE]\n\n"

        except Exception as e:
            yield f"data: {json.dumps({'tipo': 'erro', 'msg': str(e)[:200]})}\n\n"
            yield "data: [DONE]\n\n"

    return StreamingResponse(stream_criar(),
        media_type="text/event-stream",
        headers={"Access-Control-Allow-Origin": "*",
                 "Cache-Control": "no-cache",
                 "X-Accel-Buffering": "no"})



# ══════════════════════════════════════════════════════════════════════════════
# EXTRAÇÃO DE RECEITA — Fase 2 do criador de rótulos
# ══════════════════════════════════════════════════════════════════════════════

SP_EXTRAIR_RECEITA = """Voce e um especialista em rotulagem de alimentos brasileira.
Analise o documento/imagem fornecido (ficha tecnica, receita ou formulacao do produto) e extraia TODAS as informacoes relevantes para um rotulo.

IMPORTANTE: Se o documento NAO for uma receita, ficha tecnica ou formulacao de produto alimenticio (ex: e um guia, manual, legislacao, apresentacao), retorne todos os campos como null e coloque "documento_invalido": true.

RETORNE SOMENTE JSON valido, sem markdown, sem texto fora do JSON.

INSTRUCOES DE EXTRACAO:
- ingredientes: procure por "ingredientes:", "composicao:", "formula:", "formulacao:", lista de materia-prima. Ordene decrescente por quantidade.
- conservacao: procure "conservar", "armazenar", "manter refrigerado", "temperatura", "validade", "vida util". Inclua temperatura se encontrada.
- nutricional: procure tabela nutricional, "informacao nutricional", valores por 100g ou por porcao.
- peso: procure "peso liquido", "peso liq.", "g", "kg", "mL", "L" proximo ao nome do produto.
- produto: nome ou denominacao do produto alimenticio.
- Se encontrar informacoes parciais, extraia o que houver e liste o resto em campos_nao_encontrados.

{
  "documento_invalido": false,
  "produto": "nome do produto ou null",
  "categoria": "embutidos|laticinios|pescado|carnes|ovos|mel|prato_pronto_poa|outro|null",
  "especie": "especie animal ou null",
  "ingredientes": "lista completa em ordem decrescente ou null",
  "peso": "peso/volume ou null",
  "conservacao": "instrucao de conservacao com temperatura ou null",
  "nutricional": {
    "porcao": "porcao encontrada ou null",
    "energia_kcal": "valor ou null",
    "carboidratos": "valor ou null",
    "proteinas": "valor ou null",
    "gorduras_totais": "valor ou null",
    "gorduras_saturadas": "valor ou null",
    "gorduras_trans": "0g",
    "fibra": "valor ou null",
    "sodio": "valor ou null",
    "acucares_totais": "valor ou null",
    "acucares_adicionados": "valor ou null"
  },
  "alergenos_identificados": [],
  "orgao": "SIF|SIE|SIM|null",
  "observacoes": "outras informacoes relevantes ou null",
  "campos_nao_encontrados": ["lista dos campos nao encontrados"],
  "confianca": "alta|media|baixa"
}

Use null para campos nao encontrados. Nunca invente valores."""


@app.post("/extrair-receita")
async def extrair_receita(
    arquivo: UploadFile = File(...),
):
    """Le documento de receita/ficha tecnica e extrai campos para pre-preencher o formulario."""
    if not ANTHROPIC_API_KEY:
        return JSONResponse({"error": "API nao configurada"}, status_code=400,
                            headers={"Access-Control-Allow-Origin": "*"})

    contents = await arquivo.read()
    content_type = (arquivo.content_type or "").lower()
    filename = (arquivo.filename or "").lower()

    try:
        is_pdf = content_type == "application/pdf" or filename.endswith(".pdf")
        is_word = filename.endswith(".docx") or filename.endswith(".doc")
        msg_content = []

        if is_pdf:
            pages = await pdf_to_images_b64(contents)
            if pages and pages[0].get("is_error"):
                return JSONResponse({"error": pages[0]["error"]}, status_code=400,
                                    headers={"Access-Control-Allow-Origin": "*"})
            if pages and pages[0].get("is_pdf"):
                msg_content = [
                    {"type": "document", "source": {"type": "base64",
                     "media_type": "application/pdf", "data": pages[0]["b64"]}},
                    {"type": "text", "text": "Extraia todas as informacoes desta receita/ficha tecnica."}
                ]
            elif pages:
                raw = base64.b64decode(pages[0]["b64"])
                proc, _ = preprocess_image(raw, pages[0]["mime"])
                msg_content = [
                    {"type": "image", "source": {"type": "base64",
                     "media_type": "image/jpeg", "data": base64.b64encode(proc).decode()}},
                    {"type": "text", "text": "Extraia todas as informacoes desta receita/ficha tecnica."}
                ]
        elif is_word:
            try:
                import zipfile, io as _io, re as _re
                with zipfile.ZipFile(_io.BytesIO(contents)) as z:
                    xml = z.read("word/document.xml").decode("utf-8", errors="ignore")
                text_content = " ".join(_re.findall(r"<w:t[^>]*>([^<]+)</w:t>", xml))[:8000]
                msg_content = [{"type": "text",
                    "text": "Extraia informacoes desta receita:\n\n" + text_content}]
            except Exception:
                return JSONResponse(
                    {"error": "Nao foi possivel ler o Word. Tente PDF ou imagem."},
                    status_code=400, headers={"Access-Control-Allow-Origin": "*"})
        else:
            # Imagem
            ext = filename.rsplit(".", 1)[-1] if "." in filename else "jpg"
            mime_map = {"jpg": "image/jpeg", "jpeg": "image/jpeg",
                        "png": "image/png", "webp": "image/webp"}
            mime = mime_map.get(ext, "image/jpeg")
            proc, mime = preprocess_image(contents, mime)
            msg_content = [
                {"type": "image", "source": {"type": "base64",
                 "media_type": mime, "data": base64.b64encode(proc).decode()}},
                {"type": "text", "text": "Extraia todas as informacoes desta receita/ficha tecnica."}
            ]

        async with httpx.AsyncClient(timeout=60.0) as client:
            r = await client.post(
                "https://api.anthropic.com/v1/messages",
                headers={"x-api-key": ANTHROPIC_API_KEY,
                         "anthropic-version": "2023-06-01",
                         "content-type": "application/json"},
                json={"model": ANTHROPIC_MODEL, "max_tokens": 2000,
                      "system": SP_EXTRAIR_RECEITA,
                      "messages": [{"role": "user", "content": msg_content}]}
            )
            raw = r.json()["content"][0]["text"].strip()

        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
        campos = json.loads(raw.strip().rstrip("`"))
        return JSONResponse(campos, headers={"Access-Control-Allow-Origin": "*"})

    except Exception as e:
        return JSONResponse({"error": str(e)[:200]}, status_code=500,
                            headers={"Access-Control-Allow-Origin": "*"})


@app.post("/perfil-empresa")
async def salvar_perfil_empresa(request: Request):
    """Salva perfil da empresa no Supabase."""
    try:
        body = await request.json()
        import datetime as _dt_pf
        perfil = {
            "perfil_id": body.get("perfil_id", "default"),
            "razao_social": body.get("razao_social", ""),
            "cnpj": body.get("cnpj", ""),
            "endereco": body.get("endereco", ""),
            "cidade_uf": body.get("cidade_uf", ""),
            "cep": body.get("cep", ""),
            "orgao_padrao": body.get("orgao_padrao", "SIF"),
            "logo_b64": body.get("logo_b64", ""),
            "updated_at": _dt_pf.datetime.now().isoformat(),
        }
        salvo = await _sb_upsert("perfis_empresa", perfil)
        return JSONResponse({"status": "ok", "supabase": salvo},
                            headers={"Access-Control-Allow-Origin": "*"})
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500,
                            headers={"Access-Control-Allow-Origin": "*"})


@app.get("/perfil-empresa/{perfil_id}")
async def carregar_perfil_empresa(perfil_id: str = "default"):
    """Carrega perfil da empresa do Supabase."""
    try:
        rows = await _sb_get("perfis_empresa", {"perfil_id": perfil_id}, limit=1)
        return JSONResponse(rows[0] if rows else {},
                            headers={"Access-Control-Allow-Origin": "*"})
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500,
                            headers={"Access-Control-Allow-Origin": "*"})



# ══════════════════════════════════════════════════════════════════════════════
# FASE 4 — GERADOR DE DESIGN PROFISSIONAL DE RÓTULO
# Claude atua como designer sênior e gera SVG criativo com logo integrada
# ══════════════════════════════════════════════════════════════════════════════

def _extrair_cores_logo(logo_b64: str) -> dict:
    """Extrai cores dominantes da logo para usar no design."""
    try:
        from PIL import Image as _PIL
        import io as _io, collections as _col
        img = _PIL.open(_io.BytesIO(base64.b64decode(logo_b64))).convert("RGB")
        img.thumbnail((80, 80))
        pixels = list(img.getdata())
        cnt = _col.Counter()
        for r, g, b in pixels:
            bri = (r + g + b) / 3
            if 30 < bri < 220:
                cnt[((r//40)*40, (g//40)*40, (b//40)*40)] += 1
        top = cnt.most_common(3)
        primaria = top[0][0] if top else (30, 30, 30)
        return {
            "primaria": "#{:02x}{:02x}{:02x}".format(*primaria),
            "r": primaria[0], "g": primaria[1], "b": primaria[2],
        }
    except Exception:
        return {"primaria": "#1a1a2e", "r": 26, "g": 26, "b": 46}


# ═══════════════════════════════════════════════════════════════════════════════
# D2 — TEMPLATE VISUAL: MEL E PRODUTOS APÍCOLAS
# ═══════════════════════════════════════════════════════════════════════════════
DESIGN_TEMPLATE_MEL = """
═══════════════════════════════════════════════════════════════
PARTE 7 — TEMPLATE VISUAL DA CATEGORIA: MEL E APÍCOLA
═══════════════════════════════════════════════════════════════

▶ PALETA OBRIGATÓRIA
  Cor primária:    #D97706 (âmbar dourado)
  Cor secundária:  #F59E0B (mel claro, gradiente)
  Cor de acento:   #92400E (marrom terra, textos e bordas)
  Fundo zona técnica: #FFFBEB (creme levemente amarelado)
  Texto principal: #1C1917 (quase preto — boa leitura sobre fundo claro)
  Destaque claim:  #FFFFFF sobre fundo âmbar

▶ PADRÃO DE FUNDO — FAVOS HEXAGONAIS (obrigatório na zona emocional)
  Gere um padrão de hexágonos regulares no fundo da zona emocional usando <polygon> SVG.
  Hexágono padrão: 6 pontos, raio≈18px, stroke=#F59E0B, strokeWidth=0.8, fill=none (ou fill=#F59E0B opacity=0.08)
  Distribua em grelha offset (linhas pares deslocadas +15px no x) para criar padrão real de favo.
  Área de cobertura: toda a zona emocional (40% esquerda do rótulo retangular).
  O padrão deve ser sutil — fundo, não destaque.

▶ ELEMENTO GRÁFICO PRINCIPAL — GOTA DE MEL
  Posição: zona emocional, centralizada abaixo da logo.
  Forma: path SVG de gota invertida — arredondada em baixo, ponta em cima.
  Exemplo de path aproximado: M cx,y-40 C cx+30,y-40 cx+30,y+10 cx,y+30 C cx-30,y+10 cx-30,y-40 cx,y-40 Z
  Fill: gradiente radial de #F59E0B (centro) para #D97706 (borda).
  Sombra suave: feDropShadow dx=0 dy=3 stdDeviation=4 flood-color=#92400E opacity=0.25.
  Tamanho: ≈ 60×80px — presença visual forte mas não dominante.

▶ ELEMENTO SECUNDÁRIO — ABELHA ESTILIZADA (opcional, se espaço permitir)
  Abelha simplificada com 3 elements SVG: corpo oval (#1C1917), asas (ellipse translúcido #BFDBFE), listras (rect amarelas).
  Posição: canto superior direito da zona emocional ou acima da gota.
  Tamanho: ≈ 30×20px — discreto.

▶ TIPOGRAFIA
  Nome do produto: font-family="Georgia, 'Times New Roman', serif" — elegante, artesanal
  Sub-denominação: font-family="Arial, Helvetica, sans-serif" weight=400 — legível
  Claim: font-family="Georgia, serif" font-style="italic" — romantismo natural
  Dados técnicos (ingredientes, tabela): font-family="Arial, Helvetica, sans-serif" size=8-9px

▶ CLAIM VISUAL — usar um destes:
  "Mel puro de abelhas brasileiras"
  "Naturalmente doce, genuinamente brasileiro"
  "Do apiário para sua mesa"
  Posição: abaixo da gota de mel, centralizado, italic, cor #92400E, 11px.

▶ CONTEÚDO LÍQUIDO — destaque visual obrigatório
  Caixa com borda arredondada (rx=6), fill=#D97706, texto branco bold 16px.
  Exemplo: rect + text centralizado mostrando "500g" ou "1kg".

▶ ZONA TÉCNICA — diferencial apícola
  Fundo #FFFBEB (creme mel) — não branco puro.
  Separador entre ingredientes e tabela: linha de hexágonos pequenos (≈8px) em #F59E0B.
  Header da tabela nutricional: fill=#D97706, texto branco.
  Carimbo oval: stroke=#92400E, texto #92400E.

▶ SELOS / BADGES (se espaço)
  Badge "100% NATURAL" — círculo pequeno, fill=#D97706, texto branco 7px bold.
  Badge "SEM CONSERVANTES" — mesmo estilo.
  Posição: zona emocional, canto inferior.
"""

# ═══════════════════════════════════════════════════════════════════════════════
# D3 — TEMPLATE VISUAL: EMBUTIDOS E CARNES (LINGUIÇA, SALAME, PRESUNTO)
# ═══════════════════════════════════════════════════════════════════════════════
DESIGN_TEMPLATE_EMBUTIDOS = """
═══════════════════════════════════════════════════════════════
PARTE 7 — TEMPLATE VISUAL DA CATEGORIA: EMBUTIDOS E CARNES
═══════════════════════════════════════════════════════════════

▶ PALETA OBRIGATÓRIA
  Cor primária:    #991B1B (vermelho cárneo intenso)
  Cor secundária:  #7F1D1D (vinho escuro — gradiente e sombras)
  Cor de acento:   #B45309 (âmbar/marrom — detalhes e bordas)
  Fundo zona técnica: #FAFAFA (quase branco)
  Fundo zona emocional: gradiente linear de #991B1B (topo) → #7F1D1D (base)
  Texto sobre fundo escuro: #FFFFFF
  Texto técnico: #1C1917

▶ TEXTURA DE FUNDO — MADEIRA RÚSTICA (zona emocional)
  Simule textura de madeira com linhas diagonais finas em SVG:
  Gere 15-20 linhas diagonais (45°) com stroke=#7F1D1D opacity=0.15, strokeWidth=1.
  Espaçamento entre linhas: ≈ 18px.
  Sobreposição: um retângulo com fill=#991B1B opacity=0.85 sobre as linhas — cria efeito de madeira vista através de verniz escuro.

▶ ELEMENTO GRÁFICO PRINCIPAL — SILHUETA DE ANIMAL OU CORTE TRANSVERSAL
  OPÇÃO A (linguiça/salame): círculo com padrão interno de "corte transversal" — circles concêntricos em tons de vermelho/marrom simulando textura de embutido fatiado.
    Círculo externo: r=45px fill=#7F1D1D
    Círculo meio: r=35px fill=#991B1B
    Círculo interno: r=20px fill=#B45309
    Pontos irregulares (simulando gordura): circles r=3-5px fill=#FED7AA opacity=0.6 distribuídos.
  OPÇÃO B (presunto/pernil): silhueta simplificada de pernil usando path orgânico, fill=#7F1D1D opacity=0.4.
  Posição: zona emocional, centralizada. Tamanho: ≈ 90×90px.

▶ TIPOGRAFIA — BOLD E INDUSTRIAL
  Nome do produto: font-family="Arial Black, Impact, sans-serif" font-weight="900"
    Letras em caixa alta, letra-espaçamento ligeiramente expandido (letter-spacing=1).
  Sub-denominação: font-family="Arial, Helvetica, sans-serif" weight=700 uppercase.
  Claim: font-family="Arial, Helvetica, sans-serif" weight=400 italic branco.
  Dados técnicos: font-family="Arial, Helvetica, sans-serif" size=8-9px.

▶ CLAIM VISUAL — usar um destes:
  "Selecionado campo a campo"
  "Tradição no sabor desde sempre"
  "Feito com carnes selecionadas"
  "Receita de família, qualidade de indústria"
  Posição: abaixo da denominação, branco, italic, 10-11px.

▶ FAIXA DE ACENTO HORIZONTAL
  Linha/faixa horizontal de 4px, fill=#B45309, atravessando toda a largura da zona emocional.
  Posição: entre o nome do produto e o claim — separa e organiza a hierarquia.

▶ CONTEÚDO LÍQUIDO — destaque visual obrigatório
  Caixa no canto inferior da zona emocional.
  fill=#B45309, bordas retas (rx=2), texto branco bold 14px.

▶ CARIMBO OVAL — integração visual forte
  Para embutidos, o carimbo é elemento de credibilidade — dar destaque.
  Tamanho ligeiramente maior que o padrão: rx=35, ry=22.
  Borda dupla: stroke externo=#991B1B 3px + stroke interno=#991B1B 1px, separados por 3px.
  fill=white. Texto em #991B1B bold.
  Posição: zona emocional, canto inferior — não na zona técnica.

▶ ZONA TÉCNICA
  Header tabela nutricional: fill=#991B1B, texto branco.
  Linhas alternadas: #FFFFFF e #FFF5F5 (levemente rosado).
  Separador ingredientes/tabela: linha 1px #991B1B opacity=0.3.
"""

# ═══════════════════════════════════════════════════════════════════════════════
# D4 — TEMPLATE VISUAL: LATICÍNIOS (QUEIJO, IOGURTE, REQUEIJÃO)
# ═══════════════════════════════════════════════════════════════════════════════
DESIGN_TEMPLATE_LATICINIOS = """
═══════════════════════════════════════════════════════════════
PARTE 7 — TEMPLATE VISUAL DA CATEGORIA: LATICÍNIOS
═══════════════════════════════════════════════════════════════

▶ PALETA OBRIGATÓRIA
  Cor primária:    #1E40AF (azul confiança — laticínios BR)
  Cor secundária:  #DBEAFE (azul muito claro — fundos e gradientes)
  Cor de acento:   #10B981 (verde frescor — claims e badges)
  Fundo zona emocional: gradiente linear #1E40AF (topo) → #1D4ED8 (base)
  Fundo zona técnica: #F8FAFF (azul quase branco)
  Texto sobre fundo azul: #FFFFFF
  Texto técnico: #1E3A5F

▶ ELEMENTO GRÁFICO PRINCIPAL — GOTA DE LEITE
  Gota de leite caindo: forma clássica, ponta para cima, arredondada embaixo.
  Path: similar à gota de mel mas proporcões mais alongadas.
  Fill: branco (#FFFFFF) com opacidade 0.9.
  Reflexo interno: elipse pequena no canto superior esquerdo, fill=branco opacity=0.6 — efeito 3D.
  Sombra: feDropShadow dx=0 dy=4 stdDeviation=6 flood-color=#1E40AF opacity=0.3.
  Tamanho: ≈ 55×75px. Posição: zona emocional, centralizada.

▶ ELEMENTO SECUNDÁRIO — PASTAGEM / NATUREZA
  3-4 formas orgânicas simples (colinas suaves) no rodapé da zona emocional.
  Paths suaves (curvas Bezier), fill=#DBEAFE opacity=0.2 — sugestão de campo verde/azul.
  Efeito: horizonte natural, produto ligado à natureza.

▶ TIPOGRAFIA — CLEAN E CONFIÁVEL
  Nome do produto: font-family="'Montserrat', Arial, sans-serif" weight=700
    Letras title-case (não tudo maiúsculo) — transmite modernidade e cuidado.
  Sub-denominação: font-family="Arial, Helvetica, sans-serif" weight=400 branco opacity=0.9.
  Claim: font-family="'Montserrat', Arial, sans-serif" weight=300 italic branco.
  Dados técnicos: font-family="Arial, Helvetica, sans-serif" size=8-9px #1E3A5F.

▶ CLAIM VISUAL — usar um destes por tipo:
  Queijo: "Cremoso. Natural. Perfeito."
  Iogurte: "Probióticos naturais para o seu dia"
  Leite: "Do campo para sua família"
  Requeijão: "Cremosidade que abraça"
  Manteiga/ghee: "Puro sabor da natureza"
  Posição: abaixo da denominação, branco, 300 weight, italic, 10px.

▶ LINHA DIVISÓRIA — ONDAS
  Em vez de linha reta, use um path de onda suave como divisória entre zona emocional e técnica:
  Path de senoide suave com 3-4 ondas, fill=#DBEAFE, altura ≈ 12px.
  Isso cria transição orgânica que remete a líquido/leite.

▶ CONTEÚDO LÍQUIDO
  Pill/badge arredondado (rx=12), fill=#10B981 (verde), texto branco bold 14px.
  Posição: zona emocional, canto inferior esquerdo.

▶ ZONA TÉCNICA
  Header tabela nutricional: fill=#1E40AF, texto branco.
  Linhas alternadas: #FFFFFF e #F0F4FF (azul muito suave).
  Badge "FONTE DE CÁLCIO" (se aplicável): pequeno, fill=#10B981, texto branco 7px.
  Carimbo oval: stroke=#1E40AF, texto #1E40AF.

▶ BADGES DE ATRIBUTO (zona emocional, canto inferior)
  "SEM CONSERVANTES" — pill pequeno, fill=branco opacity=0.2, texto branco 7px.
  "RICO EM PROTEÍNAS" — idem.
"""

# ═══════════════════════════════════════════════════════════════════════════════
# D5 — TEMPLATE VISUAL: PESCADO (TILÁPIA, SALMÃO, CAMARÃO, ATUM)
# ═══════════════════════════════════════════════════════════════════════════════
DESIGN_TEMPLATE_PESCADO = """
═══════════════════════════════════════════════════════════════
PARTE 7 — TEMPLATE VISUAL DA CATEGORIA: PESCADO
═══════════════════════════════════════════════════════════════

▶ PALETA OBRIGATÓRIA
  Cor primária:    #0369A1 (azul oceano profundo)
  Cor secundária:  #7DD3FC (azul água claro)
  Cor de acento:   #0EA5E9 (ciano vibrante — detalhes e ação)
  Cor quente:      #F97316 (laranja coral — contraste e apetência)
  Fundo zona emocional: gradiente linear #0369A1 (topo) → #0C4A6E (base)
  Fundo zona técnica: #F0F9FF (azul céu pálido)
  Texto sobre fundo escuro: #FFFFFF
  Texto técnico: #0C4A6E

▶ PADRÃO DE FUNDO — ONDAS (zona emocional)
  Gere 4-5 ondas SVG (paths senoidais suaves) no fundo da zona emocional.
  Ondas em profundidade: as mais altas são #7DD3FC opacity=0.1, as mais baixas #0EA5E9 opacity=0.06.
  Cada onda desloca ≈ 20px verticalmente da anterior.
  Ondas criam sensação de movimento e frescor aquático.

▶ ELEMENTO GRÁFICO PRINCIPAL — PEIXE ESTILIZADO
  Peixe simplificado com 2 elements SVG:
    Corpo: ellipse alongada, fill branco opacity=0.15, largura≈80px altura≈35px.
    Cauda: path triangular, fill branco opacity=0.12.
    Olho: circle r=4, fill=#F97316.
    Escamas hint: 3-4 arcos (path de semicírculo) sobre o corpo, stroke=branco opacity=0.1.
  Tamanho total: ≈ 95×45px. Posição: zona emocional, centralizada.
  Efeito: elegante, não cartunesco.

▶ ELEMENTO ALTERNATIVO (para atum em lata / conserva)
  Lata estilizada: rect arredondado com linhas horizontais de textura metálica.
  Fill: gradiente linear de #7DD3FC → #0369A1.

▶ TIPOGRAFIA — CLEAN E MODERNO
  Nome do produto: font-family="'Oswald', Arial Narrow, sans-serif" — condensado, impacto vertical.
  Se Oswald não disponível: Arial Narrow ou simplesmente scaleX(0.85) em Arial.
  Sub-denominação: font-family="Arial, sans-serif" weight=300 branco opacity=0.85.
  Claim: font-family="Arial, sans-serif" weight=300 italic branco opacity=0.9.
  Dados técnicos: font-family="Arial, sans-serif" size=8-9px #0C4A6E.

▶ CLAIM VISUAL — usar um destes por tipo:
  Peixe fresco/tilápia: "Fresco da água para sua mesa"
  Salmão: "Proteína nobre do oceano"
  Camarão: "Do mar para sua receita"
  Atum em conserva: "Selecionado no ápice do frescor"
  Geral: "Sabor que vem do mar"
  Posição: abaixo do peixe estilizado, branco italic 10px.

▶ FAIXA DE ACENTO — CORAL
  Linha/faixa de 3px fill=#F97316 atravessando a zona emocional horizontalmente.
  Posição: entre a denominação e o claim.
  Cria contraste quente que remete ao produto (carne de peixe, camarão).

▶ CONTEÚDO LÍQUIDO
  Badge ciano: rx=4, fill=#0EA5E9, texto branco bold 14px.
  Ou: caixa com borda #F97316, fill transparente, texto #F97316 bold.

▶ ZONA TÉCNICA
  Header tabela nutricional: fill=#0369A1, texto branco.
  Linhas alternadas: #FFFFFF e #E0F2FE (azul água pálido).
  Badge "RICO EM ÔMEGA-3" (se aplicável): fill=#0EA5E9, texto branco 7px.
  Badge "PROTEÍNA MAGRA": fill=#10B981, texto branco 7px.
  Carimbo oval: stroke=#0369A1, texto #0369A1.

▶ PARA CONSERVAS (atum, sardinha em lata):
  Adicionar elemento "conserva" — tampa de lata estilizada no topo da zona emocional.
  Linhas metálicas sutis no fundo da zona técnica: stroke=#7DD3FC opacity=0.1.
"""

# Mapeamento categoria → template
DESIGN_TEMPLATES = {
    # Mel e apícola
    "mel":              DESIGN_TEMPLATE_MEL,
    "mel_qualidade":    DESIGN_TEMPLATE_MEL,
    "apicola_derivados":DESIGN_TEMPLATE_MEL,

    # Embutidos e carnes
    "embutidos":        DESIGN_TEMPLATE_EMBUTIDOS,
    "salame":           DESIGN_TEMPLATE_EMBUTIDOS,
    "presunto":         DESIGN_TEMPLATE_EMBUTIDOS,
    "hamburguer":       DESIGN_TEMPLATE_EMBUTIDOS,
    "bacon":            DESIGN_TEMPLATE_EMBUTIDOS,
    "carne_moida":      DESIGN_TEMPLATE_EMBUTIDOS,
    "charque":          DESIGN_TEMPLATE_EMBUTIDOS,
    "fiambre":          DESIGN_TEMPLATE_EMBUTIDOS,
    "apresuntado":      DESIGN_TEMPLATE_EMBUTIDOS,
    "carne_maturada":   DESIGN_TEMPLATE_EMBUTIDOS,
    "carnes_temperadas":DESIGN_TEMPLATE_EMBUTIDOS,
    "jerked_beef":      DESIGN_TEMPLATE_EMBUTIDOS,
    "carne_sol":        DESIGN_TEMPLATE_EMBUTIDOS,
    "frango_inteiro":   DESIGN_TEMPLATE_EMBUTIDOS,
    "almondega_kibe":   DESIGN_TEMPLATE_EMBUTIDOS,
    "pato_pronto_poa":  DESIGN_TEMPLATE_EMBUTIDOS,
    "corned_beef":      DESIGN_TEMPLATE_EMBUTIDOS,

    # Laticínios
    "laticinios_geral":       DESIGN_TEMPLATE_LATICINIOS,
    "queijo_coalho_manteiga": DESIGN_TEMPLATE_LATICINIOS,
    "mussarela":              DESIGN_TEMPLATE_LATICINIOS,
    "leite_uht":              DESIGN_TEMPLATE_LATICINIOS,
    "leite_pasteurizado":     DESIGN_TEMPLATE_LATICINIOS,
    "doce_de_leite":          DESIGN_TEMPLATE_LATICINIOS,
    "leite_fermentado":       DESIGN_TEMPLATE_LATICINIOS,
    "requeijao":              DESIGN_TEMPLATE_LATICINIOS,
    "leite_em_po":            DESIGN_TEMPLATE_LATICINIOS,
    "leite_condensado":       DESIGN_TEMPLATE_LATICINIOS,
    "bebida_lactea":          DESIGN_TEMPLATE_LATICINIOS,
    "soro_leite":             DESIGN_TEMPLATE_LATICINIOS,
    "leite_aromatizado":      DESIGN_TEMPLATE_LATICINIOS,
    "composto_lacteo":        DESIGN_TEMPLATE_LATICINIOS,
    "nata":                   DESIGN_TEMPLATE_LATICINIOS,
    "queijo_provolone":       DESIGN_TEMPLATE_LATICINIOS,
    "queijo_parmesao":        DESIGN_TEMPLATE_LATICINIOS,
    "queijo_prato":           DESIGN_TEMPLATE_LATICINIOS,
    "queijo_processado":      DESIGN_TEMPLATE_LATICINIOS,

    # Pescado
    "pescado_fresco":       DESIGN_TEMPLATE_PESCADO,
    "pescado_congelado":    DESIGN_TEMPLATE_PESCADO,
    "pescado_salgado":      DESIGN_TEMPLATE_PESCADO,
    "camarao":              DESIGN_TEMPLATE_PESCADO,
    "lagosta":              DESIGN_TEMPLATE_PESCADO,
    "moluscos_cefalopodes": DESIGN_TEMPLATE_PESCADO,
    "conserva_sardinha":    DESIGN_TEMPLATE_PESCADO,
    "conserva_sardinhas":   DESIGN_TEMPLATE_PESCADO,
    "conserva_atum":        DESIGN_TEMPLATE_PESCADO,
    "conserva_atuns":       DESIGN_TEMPLATE_PESCADO,
    "conserva_peixe":       DESIGN_TEMPLATE_PESCADO,
    "conserva_peixes":      DESIGN_TEMPLATE_PESCADO,
}

def get_design_template(categoria: str) -> str:
    """Retorna o template de design para a categoria detectada."""
    return DESIGN_TEMPLATES.get(categoria, "")


SP_DESIGN_ROTULO = """Você é um diretor de arte sênior de uma agência de branding especializada em embalagens de alimentos brasileiros. Seus trabalhos estão nas prateleiras de supermercados como Carrefour, Pão de Açúcar e Assaí. Você já criou identidades para marcas como Sadia, Perdigão, Vigor e Parati.

RETORNE APENAS O CÓDIGO SVG. Nenhuma palavra antes ou depois. Nenhum markdown. Nenhum comentário fora do SVG.

═══════════════════════════════════════════════════════════════
PARTE 1 — PRINCÍPIOS DE DESIGN DE EMBALAGEM (OBRIGATÓRIO LER)
═══════════════════════════════════════════════════════════════

▶ HIERARQUIA VISUAL — 3 NÍVEIS IMUTÁVEIS
Todo rótulo profissional tem exatamente 3 camadas de leitura:
  NÍVEL 1 — APELO (0,3 seg): o consumidor SENTE antes de ler
    → Logo da marca + cor dominante + elemento visual emocional
    → Ocupa 35-45% da área visível — geralmente faixa lateral ou superior
    → Tipografia: bold, grande (32-48px), contraste máximo
  NÍVEL 2 — IDENTIDADE (1-2 seg): o consumidor LÊ o nome
    → Denominação do produto + peso/volume + claims principais
    → Ocupa 25-35% — próximo ao nível 1
    → Tipografia: semi-bold, 18-28px
  NÍVEL 3 — INFORMAÇÃO (5+ seg): o consumidor VERIFICA detalhes
    → Ingredientes + tabela nutricional + alérgenos + fabricante
    → Ocupa 30-40% — geralmente no verso/painel técnico
    → Tipografia: regular, 8-11px

▶ REGRA DOS TERÇOS (aplicar em todos os formatos)
Divida o rótulo em 3 colunas e 3 linhas imaginárias.
Os elementos principais (logo, nome do produto) ficam NAS INTERSEÇÕES dessas linhas,
nunca perfeitamente centrados e nunca nos cantos extremos.
Isso cria tensão visual positiva — o olho "viaja" pelo rótulo.

▶ ZONA EMOCIONAL vs ZONA TÉCNICA
Qualquer rótulo tem duas zonas:
  ZONA EMOCIONAL (face principal): vende. Contém logo, produto, claim, conteúdo líquido, lupa.
  ZONA TÉCNICA (faces secundárias): informa. Contém ingredientes, tabela, alérgenos, fabricante.
Para SVG de face única, dividir horizontalmente: 40% emocional (esquerda) / 60% técnica (direita).

▶ PALETA DE CORES — REGRAS PARA ALIMENTOS
Máximo 3 cores base + branco/preto. Nunca mais.
  Cor primária = cor da marca (extraída da logo ou fornecida)
  Cor secundária = 20-30% mais clara que a primária (para gradientes e fundos)
  Cor de acento = complementar ou análoga para calls-to-action (conteúdo líquido, claims)
  Backgrounds: nunca cinza puro (#808080) — use off-whites (#FAFAFA, #F5F0EB) ou cores pastéis
  Texto sobre fundo colorido: sempre branco (#FFF) ou preto (#111) — nunca cinza

▶ TIPOGRAFIA — PERSONALIDADE DO PRODUTO
Cada categoria de produto tem uma voz tipográfica:
  Carnes/Embutidos: Bebas Neue, Anton, Black Han Sans — bold, industrial, apetitoso
  Laticínios/Mel: Montserrat, Raleway, Lato — limpo, fresco, confiável
  Pescado: Oswald, Source Sans — moderno, saudável, premium
  Premium/Artesanal: Playfair Display, Cormorant — elegante, tradicional
  Natural/Orgânico: Nunito, Poppins — suave, amigável, saudável
  Genérico/Industrial: Inter, Roboto — neutro, eficiente

Hierarquia de peso: produto=800 (Black), sub-denominação=600 (SemiBold), texto=400 (Regular)

▶ CLAIM VISUAL — GERADO POR IA (D7 — OBRIGATÓRIO)
NÃO use claims genéricos. Gere um claim ÚNICO e ESPECÍFICO para este produto
baseado nos dados reais fornecidos: denominação, ingredientes, categoria e carimbo.

ALGORITMO DE GERAÇÃO DO CLAIM:
  1. Leia os INGREDIENTES fornecidos — extraia o diferencial real do produto
  2. Leia a DENOMINAÇÃO — identifique o tipo exato (ex: "Linguiça Toscana Suína Frescal")
  3. Leia o CARIMBO — SIF indica alcance nacional (mais prestígio), SIE/SIM indica regional
  4. Combine em uma frase de 3-6 palavras que seja verdadeira e apelativa

REGRAS DO CLAIM PERFEITO:
  ✅ Específico ao produto — "Defumada na lenha de algaroba" > "Tradição no sabor"
  ✅ Baseado em ingrediente real — se tem alho, "Com alho selecionado do campo"
  ✅ Curto — máximo 6 palavras, preferencialmente 3-4
  ✅ Sem exageros legais — não usar "melhor", "número 1", "superior"
  ✅ Emocional mas verdadeiro — conecta produto à origem ou ao momento de consumo
  ❌ Nunca genérico — "Qualidade garantida", "Sabor incomparável", "Produto de qualidade"
  ❌ Nunca claim de saúde sem base — "Saudável", "Nutritivo", "Rico em proteínas" (só se ingredientes confirmarem)

EXEMPLOS DE CLAIMS BEM GERADOS (use como referência de qualidade, não copie):
  Linguiça toscana suína (ingredientes: carne suína, alho, pimenta-do-reino):
    → "Temperada com pimenta e alho do campo"
  Mel silvestre (ingredientes: mel silvestre puro):
    → "Direto das flores do cerrado"
  Queijo mussarela (ingredientes: leite pasteurizado integral):
    → "Do leite fresco para sua mesa"
  Filé de tilápia congelado (ingredientes: tilápia, água, sal):
    → "Pescado e congelado no mesmo dia"
  Presunto cozido fatiado (SIF, ingredientes: pernil suíno, sal, conservantes):
    → "Pernil selecionado, fatiado na medida"
  Manteiga de garrafa (ingredientes: creme de leite):
    → "Cremosa como sempre foi"
  Salame italiano (ingredientes: carne suína, carne bovina, pimenta, vinho):
    → "Curado com vinho e especiarias"

POSICIONAMENTO DO CLAIM NO SVG:
  • Imediatamente abaixo da denominação do produto
  • Fonte: italic, peso 300-400, tamanho 10-12px
  • Cor: branco com opacity 0.85-0.9 (sobre fundo escuro) OU cor de acento (sobre fundo claro)
  • Nunca em caixa alta — sempre title-case ou minúsculo

▶ ELEMENTO GRÁFICO DE APOIO
Adicione 1 elemento gráfico sutil relacionado ao produto. Use formas SVG simples:
  Carnes: silhueta de animal (path simples), pingo de gordura, chama
  Laticínios: gota de leite, forma de queijo, ondas de creme
  Pescado: escamas (padrão hexagonal), onda do mar, peixe estilizado
  Mel: favos hexagonais no fundo, abelha estilizada, gota de mel
  Ovos: círculos concêntricos, formato oval suave

═══════════════════════════════════════════════════════════════
PARTE 2 — ESPECIFICAÇÕES TÉCNICAS POR FORMATO
═══════════════════════════════════════════════════════════════

▶ FORMATO RETANGULAR (viewBox="0 0 900 420")
┌─────────────────────────────────────────────────────────────┐
│ ZONA EMOCIONAL (0-320px) │ ZONA TÉCNICA (320-880px)         │
│                           │                                  │
│ • Fundo: cor primária     │ • Fundo: branco ou off-white     │
│ • Logo: y=30, máx 140px   │ • Ingredientes: fonte 9px        │
│ • Produto: y=180, 32-40px │ • Tabela nutricional: y=120      │
│ • Claim: y=240, 11px      │ • Rodapé: alérgenos+carimbo      │
│ • Peso: y=280, bold 18px  │                                  │
└─────────────────────────────────────────────────────────────┘
Divisória: linha vertical em x=320, cor de acento, espessura 3px
Header tabela: fundo cor primária, texto branco, altura 24px
Linhas tabela: alternando #FFFFFF e #F7F7F7, altura 20px cada

▶ FORMATO CIRCULAR (viewBox="0 0 600 600")
Centro: cx=300, cy=300. Raio externo: 292px.
  Anel externo (r=292 a r=265): cor primária
  Área branca (r=265 a r=0): fundo do produto
  Faixa superior (arco de 60° a 300°, y < 150): cor primária, logo centrada
  Produto: cx=300, y=280, 28-34px bold
  Claim: cx=300, y=310, 10px italic
  Tabela nutricional: quadrado centralizado, y=330 a y=500, fundo #F5F5F5
  Carimbo: cx=490, cy=490, oval pequeno

▶ FORMATO QUADRADO (viewBox="0 0 600 600")
Dividir em quadrantes:
  Superior esquerdo (0-300, 0-280): zona emocional com logo + produto
  Superior direito (300-600, 0-280): elemento gráfico + claim + peso
  Inferior (0-600, 280-600): zona técnica completa (ingredientes + tabela + rodapé)
Linha divisória horizontal: y=280, cor de acento, espessura 2px

═══════════════════════════════════════════════════════════════
PARTE 3 — TABELA NUTRICIONAL — ELEMENTO GRÁFICO PROFISSIONAL (D9)
═══════════════════════════════════════════════════════════════

A tabela nutricional é um ELEMENTO GRÁFICO, não apenas texto. Deve ter a mesma
qualidade visual do resto do rótulo. Siga estas especificações exatas:

▶ ESTRUTURA SVG DA TABELA (implementar com <rect> + <text>, não apenas texto)

BLOCO CONTAINER:
  <rect> com fill="#FFFFFF", stroke="#DDDDDD", stroke-width="0.5", rx="3"
  Padding interno: 6px em todos os lados
  Largura: adaptar ao espaço disponível na zona técnica

HEADER — "Informação Nutricional":
  <rect> altura=22px, fill=cor_primária_da_marca, rx="3 3 0 0" (arredondado só em cima)
  <text> centralizado, fill="#FFFFFF", font-weight="700", font-size="9px"
  font-family="Arial, sans-serif"

LINHA DE PORÇÃO — "Porção Xg (X porções por embalagem)":
  <rect> altura=18px, fill="#F0F0F0"
  <text> fill="#444444", font-size="8px", font-style="italic"
  Separador abaixo: <line> stroke="#CCCCCC" stroke-width="1px"

LINHA DE COLUNAS (apenas se houver %VD):
  <rect> altura=14px, fill="#E8E8E8"
  <text> "Por porção" alinhado à direita, font-size="7px", fill="#666666"
  <text> "%VD*" mais à direita, font-size="7px", fill="#666666"

LINHAS DE NUTRIENTES (10 linhas obrigatórias, ordem RDC 429/2020):
  Altura de cada linha: 17px — espaçamento generoso, não apertado
  Alternância de fundo:
    Linhas pares:   fill="#FFFFFF"
    Linhas ímpares: fill="#F7F8FA" (azul-acinzentado muito suave — mais sofisticado que cinza puro)
  Separador entre linhas: <line> stroke="#EBEBEB" stroke-width="0.5"

  Para CADA linha, usar este padrão SVG:
    <rect x="{tx}" y="{ty}" width="{tw}" height="17" fill="{bg}"/>
    <line x1="{tx}" y1="{ty+17}" x2="{tx+tw}" y2="{ty+17}" stroke="#EBEBEB" stroke-width="0.5"/>
    <!-- Nome do nutriente (esquerda) -->
    <text x="{tx+6}" y="{ty+12}" font-family="Arial, sans-serif" font-size="8"
          font-weight="{700 se linha principal, 400 se sub-item}" fill="#222222">{nome}</text>
    <!-- Valor (direita) — FONTE MONOSPACED para alinhamento perfeito -->
    <text x="{tx+tw-6}" y="{ty+12}" font-family="'Courier New', Courier, monospace"
          font-size="8" font-weight="700" fill="#111111" text-anchor="end">{valor}</text>

  HIERARQUIA VISUAL dos nutrientes:
    Linha principal (Valor energético, Carboidratos, Proteínas, Gorduras totais, Fibra, Sódio):
      font-weight="700", fill="#111111"
    Sub-item (Açúcares totais, Açúcares adicionados, Gorduras saturadas, Gorduras trans):
      font-weight="400", fill="#444444"
      Recuo de 10px no texto (x = tx + 16 ao invés de tx + 6)
      Prefixar com "  " (espaço) para indicar hierarquia

ORDEM EXATA DOS NUTRIENTES (RDC 429/2020):
  1. Valor energético         → {energia_kcal} kcal = {energia_kj} kJ   [PRINCIPAL]
  2. Carboidratos             → {carboidratos}                            [PRINCIPAL]
  3.   Açúcares totais        → calcular se ausente                       [sub-item]
  4.   Açúcares adicionados   → calcular se ausente                       [sub-item]
  5. Proteínas                → {proteinas}                               [PRINCIPAL]
  6. Gorduras totais          → {gorduras_totais}                         [PRINCIPAL]
  7.   Gorduras saturadas     → {gorduras_saturadas}                      [sub-item]
  8.   Gorduras trans         → {gorduras_trans}                          [sub-item]
  9. Fibra alimentar          → {fibra}                                   [PRINCIPAL]
  10. Sódio                   → {sodio}                                   [PRINCIPAL]

RODAPÉ DA TABELA:
  <rect> altura=16px, fill="#F5F5F5"
  <text> "* % Valores Diários com base em dieta de 2.000 kcal ou 8.400 kJ."
  font-size="6.5px", fill="#888888", font-style="italic"
  Segunda linha (se espaço): "** VD não estabelecido."

▶ LUPA DE ADVERTÊNCIA FRONTAL (RDC 429/2020 — obrigatória se aplicável)
  Condições: sódio ≥ 600mg/100g OU gordura saturada ≥ 6g/100g OU açúcar adicionado ≥ 15g/100g
  Posição: zona EMOCIONAL (frente do rótulo), não na tabela
  Formato: retângulo com lupa estilizada + texto de advertência
    <rect> fill="#000000", rx="4"
    Ícone lupa: círculo stroke="#FFFFFF" + linha diagonal — ambos em branco
    <text> "ALTO EM {SÓDIO / GORDURAS SATURADAS / AÇÚCAR}" fill="#FFFFFF" font-weight="700" font-size="9px"
  Se múltiplas advertências: empilhar verticalmente, uma abaixo da outra

▶ DIMENSÕES RECOMENDADAS DA TABELA POR FORMATO:
  Retangular (zona técnica 320-880px): tabela de x=330 a x=620, altura total ≈ 220px
  Quadrado (zona inferior): tabela de x=10 a x=290, y=290 a y=560
  Circular: tabela centralizada, largura ≈ 180px, y=330 a y=490

═══════════════════════════════════════════════════════════════
PARTE 4 — CARIMBO DE INSPEÇÃO — PADRÃO OFICIAL
═══════════════════════════════════════════════════════════════
Carimbo OVAL obrigatório (não redondo, não retangular):
  Formato: <ellipse> com proporção 3:2 (largura:altura)
  Borda dupla: anel externo stroke 3px + anel interno stroke 1px, cor escura
  Texto interno (3 linhas, centralizado):
    Linha 1: "INSPECIONADO" — bold, 7px
    Linha 2: sigla do órgão (SIF / SIE/RJ / SIM) — bold, 8px
    Linha 3: número do registro — regular, 7px
  Cores: fundo branco, texto e bordas na cor primária escura ou #1a1a2e
  Posicionamento: canto inferior direito da zona técnica

═══════════════════════════════════════════════════════════════
PARTE 5 — ELEMENTOS LEGAIS OBRIGATÓRIOS (NUNCA OMITIR)
═══════════════════════════════════════════════════════════════
Todos estes elementos DEVEM aparecer no SVG gerado:
  ✓ Logo da empresa: <image href="data:image/png;base64,{logo_b64}" width="mín.130" height="mín.70"/>
  ✓ Denominação: {denominacao}
  ✓ Claim emocional gerado: [criar baseado na categoria]
  ✓ Conteúdo líquido: {conteudo_liquido} — fonte bold, destaque visual
  ✓ Fabricante: {fabricante} — fonte 8px, pode ser comprimido em 2 linhas
  ✓ Conservação: {conservacao} — fonte 8px
  ✓ Glúten: {gluten} — fonte 8px bold, cor destacada
  ✓ Lactose: {lactose} — fonte 8px bold (se aplicável)
  ✓ Alérgenos: {alergenos} — fonte 8px, background levemente amarelado (#FFFBEA)
  ✓ Transgênicos: {transgenicos} — fonte 7px
  ✓ Carimbo oval: {carimbo}
  ✓ Lote e validade: "LOTE E VALIDADE: veja embalagem" — fonte 7.5px
  ✓ Tabela nutricional completa (todos os 10 nutrientes)
  ✓ Lupa se necessário (sódio ≥600mg ou gordura sat. ≥6g ou açúcar adicionado ≥15g)

═══════════════════════════════════════════════════════════════
PARTE 6 — QUALIDADE DO SVG (PADRÃO TÉCNICO)
═══════════════════════════════════════════════════════════════
  • Fontes: usar Google Fonts via @import no <style> — ver PARTE 8 abaixo (OBRIGATÓRIO)
  • Todos os textos: clip dentro da área do rótulo com clipPath se necessário
  • Nenhum texto pode ultrapassar a borda do rótulo
  • Gradientes: usar <linearGradient> ou <radialGradient> para profundidade
  • Sombras: <filter> com feDropShadow nos elementos principais (logo, nome do produto)
  • Borda do rótulo: rect ou path com bordas arredondadas rx="8" e stroke de 1.5px
  • Resolução: o SVG deve ser legível quando renderizado em 800x400px ou maior

═══════════════════════════════════════════════════════════════
PARTE 8 — FONTES PREMIUM VIA GOOGLE FONTS (D6 — OBRIGATÓRIO)
═══════════════════════════════════════════════════════════════

▶ ESTRUTURA OBRIGATÓRIA DO SVG — SEMPRE COMEÇAR ASSIM:
O SVG DEVE ter um bloco <defs><style> logo após a abertura <svg ...>, importando
as fontes Google correspondentes à categoria do produto.

TEMPLATE BASE (copie e adapte para a categoria):
<svg viewBox="0 0 900 420" xmlns="http://www.w3.org/2000/svg">
  <defs>
    <style>
      @import url('https://fonts.googleapis.com/css2?family=FONTE_PRINCIPAL:wght@400;600;700;900&amp;family=FONTE_SECUNDARIA:wght@300;400;600&amp;display=swap');
    </style>
    <!-- gradients, filters, clipPaths aqui -->
  </defs>
  <!-- resto do SVG -->
</svg>

▶ FONTES POR CATEGORIA (escolha baseado no template da PARTE 7):

EMBUTIDOS / CARNES:
  @import url('https://fonts.googleapis.com/css2?family=Bebas+Neue&amp;family=Oswald:wght@400;600;700&amp;display=swap');
  Nome produto:    font-family="'Bebas Neue', Arial Black, sans-serif"
  Sub/claim:       font-family="'Oswald', Arial Narrow, sans-serif"
  Dados técnicos:  font-family="'Oswald', Arial, sans-serif" font-weight="300"

LATICÍNIOS:
  @import url('https://fonts.googleapis.com/css2?family=Montserrat:wght@300;400;600;700&amp;display=swap');
  Nome produto:    font-family="'Montserrat', Arial, sans-serif" font-weight="700"
  Sub/claim:       font-family="'Montserrat', Arial, sans-serif" font-weight="300"
  Dados técnicos:  font-family="'Montserrat', Arial, sans-serif" font-weight="400"

MEL / ORGÂNICOS / NATURAIS:
  @import url('https://fonts.googleapis.com/css2?family=Nunito:wght@300;400;600;700&amp;family=Playfair+Display:ital,wght@0,700;1,400&amp;display=swap');
  Nome produto:    font-family="'Playfair Display', Georgia, serif" font-weight="700"
  Claim (italic):  font-family="'Playfair Display', Georgia, serif" font-style="italic"
  Ingredientes:    font-family="'Nunito', Arial, sans-serif" font-weight="400"

PESCADO:
  @import url('https://fonts.googleapis.com/css2?family=Oswald:wght@300;400;600;700&amp;family=Nunito:wght@300;400&amp;display=swap');
  Nome produto:    font-family="'Oswald', Arial Narrow, sans-serif" font-weight="700"
  Claim:           font-family="'Oswald', Arial Narrow, sans-serif" font-weight="300"
  Dados técnicos:  font-family="'Nunito', Arial, sans-serif" font-weight="400"

PREMIUM / ARTESANAL (qualquer categoria):
  @import url('https://fonts.googleapis.com/css2?family=Playfair+Display:ital,wght@0,400;0,700;1,400&amp;family=Montserrat:wght@300;400;600&amp;display=swap');
  Nome produto:    font-family="'Playfair Display', Georgia, serif" font-weight="700"
  Claim:           font-family="'Playfair Display', Georgia, serif" font-style="italic"
  Dados técnicos:  font-family="'Montserrat', Arial, sans-serif" font-weight="300"

▶ REGRAS CRÍTICAS DE IMPLEMENTAÇÃO:
  1. O @import DEVE usar &amp; para separar parâmetros (não & puro — causa erro XML no SVG)
  2. O bloco <style> DEVE ser filho direto de <defs>, que DEVE ser filho direto de <svg>
  3. Sempre declare fallback após a fonte Google: font-family="'Bebas Neue', Arial Black, sans-serif"
  4. Pesos disponíveis dependem do @import — só use pesos declarados no URL
  5. Nunca use font-family="Bebas Neue" sem aspas — use sempre font-family="'Bebas Neue', fallback"

▶ EXEMPLO FUNCIONAL COMPLETO (embutidos):
<svg viewBox="0 0 900 420" xmlns="http://www.w3.org/2000/svg">
  <defs>
    <style>@import url('https://fonts.googleapis.com/css2?family=Bebas+Neue&amp;family=Oswald:wght@300;400;700&amp;display=swap');</style>
    <linearGradient id="grad1" x1="0%" y1="0%" x2="0%" y2="100%">
      <stop offset="0%" style="stop-color:#991B1B"/>
      <stop offset="100%" style="stop-color:#7F1D1D"/>
    </linearGradient>
  </defs>
  <rect width="320" height="420" fill="url(#grad1)"/>
  <text x="160" y="180" text-anchor="middle"
        font-family="'Bebas Neue', Arial Black, sans-serif"
        font-size="42" fill="white">LINGUIÇA TOSCANA</text>
  <text x="160" y="210" text-anchor="middle"
        font-family="'Oswald', Arial Narrow, sans-serif"
        font-weight="300" font-style="italic"
        font-size="13" fill="white" opacity="0.85">Selecionada campo a campo</text>
</svg>

{evite_rules}

{design_template}

═══════════════════════════════════════════════════════════════
DADOS DO RÓTULO A GERAR
═══════════════════════════════════════════════════════════════
Formato: {formato}
Denominação: {denominacao}
Carimbo: {carimbo}
Fabricante: {fabricante}
Ingredientes: {ingredientes}
Conservação: {conservacao}
Conteúdo líquido: {conteudo_liquido}
Alérgenos: {alergenos}
Transgênicos: {transgenicos}
Glúten: {gluten}
Lactose: {lactose}
Cor primária da marca: {cor_primaria}
Porção: {porcao}
Energia: {energia_kcal} kcal
Proteínas: {proteinas}
Carboidratos: {carboidratos}
Gorduras totais: {gorduras_totais}
Gorduras saturadas: {gorduras_saturadas}
Gorduras trans: {gorduras_trans}
Fibra: {fibra}
Sódio: {sodio}

Agora gere o SVG profissional completo. APENAS o código SVG, começando com <svg e terminando com </svg>.
"""

# ═══════════════════════════════════════════════════════════════════════════════
# D10 — ADAPTAÇÃO MULTIPANEL AUTOMÁTICA
# Distribuição inteligente de conteúdo por face da embalagem
# ═══════════════════════════════════════════════════════════════════════════════

SP_PAINEL = {
    "frente": """
PAINEL: FRENTE DA EMBALAGEM (Face Principal — o consumidor vê primeiro)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

FUNÇÃO: VENDER. Esta face decide a compra. Cada elemento deve servir ao apelo.

CONTEÚDO OBRIGATÓRIO (nesta ordem de hierarquia visual):
  1. Logo da empresa — posição privilegiada, topo ou terço superior
  2. Denominação do produto — maior elemento tipográfico, máxima legibilidade
  3. Claim emocional gerado por IA (ver regras D7) — imediatamente abaixo da denominação
  4. Elemento gráfico de categoria (D2-D5) — central ou lateral, forte presença visual
  5. Conteúdo líquido / peso — destaque visual, badge ou caixa
  6. Lupa de advertência (se aplicável) — canto inferior, obrigatória se critérios atingidos
  7. Carimbo SIF/SIE/SIM — canto inferior oposto à lupa, presença institucional

CONTEÚDO PROIBIDO NA FRENTE:
  ✗ Lista de ingredientes completa
  ✗ Tabela nutricional completa
  ✗ CNPJ e endereço do fabricante
  ✗ Instruções de conservação detalhadas
  ✗ Declaração de transgênicos

LAYOUT viewBox="0 0 400 600" (orientação retrato — face de embalagem):
  Zona emocional (0-600px altura total): toda a face é zona emocional
  Logo: x=20, y=20, máx 120×60px
  Denominação: centralizado, y=200-250, font-size=36-48px bold
  Claim: centralizado, y=270, font-size=12px italic
  Elemento gráfico: centralizado, y=300-420, 120-150px
  Peso/volume: badge no canto inferior direito, x=300, y=540
  Carimbo: canto inferior esquerdo, cx=60, cy=555
  Lupa (se necessário): faixa preta na base, y=560-600
""",

    "verso": """
PAINEL: VERSO DA EMBALAGEM (Face Técnica — o consumidor consulta após compra)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

FUNÇÃO: INFORMAR. Esta face cumpre a legislação e constrói confiança.

CONTEÚDO OBRIGATÓRIO (ordem de disposição vertical):
  1. Tabela nutricional completa (PARTE 3 — elemento gráfico profissional D9)
  2. Lista de ingredientes completa (precedida de "Ingredientes:")
  3. Identificação do fabricante (Razão Social + CNPJ + endereço completo)
  4. Declaração de alérgenos (fundo #FFFBEA, texto bold, destaque visual)
  5. Glúten e lactose (logo abaixo dos alérgenos)
  6. Declaração de transgênicos
  7. SAC / contato do consumidor (se disponível)
  8. Lote e validade: "LOTE E VALIDADE: veja embalagem" — rodapé

LAYOUT viewBox="0 0 400 600" (mesmo tamanho da frente):
  Fundo: branco (#FFFFFF) ou off-white (#FAFAFA) — zona técnica pura
  Margem: 14px em todos os lados
  Tabela nutricional: x=14, y=14, largura=372px
  Ingredientes: abaixo da tabela, y dinâmico, font-size=8px
  Fabricante: abaixo dos ingredientes, font-size=8px, cor #444
  Alérgenos: faixa com fundo #FFFBEA, borda #F59E0B 1px, padding 8px
  Rodapé (lote/validade + transgênicos): y=570-590, font-size=7px

ESTÉTICA DO VERSO:
  Minimalista — sem elementos decorativos excessivos
  Linha fina da cor primária (1px) como header separando da margem superior
  Tabela nutricional usa cor primária da marca no header (mantém identidade)
  Texto técnico: Arial 8px #333333 — máxima legibilidade
""",

    "lateral": """
PAINEL: LATERAL DA EMBALAGEM (Face Secundária — visível na prateleira de lado)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

FUNÇÃO: IDENTIFICAR na prateleira + cumprir obrigações secundárias.

CONTEÚDO OBRIGATÓRIO:
  1. Denominação do produto (repetida, menor que na frente)
  2. Instruções de conservação completas (temperatura, prazo pós-abertura)
  3. Declaração de alérgenos (obrigatória em todas as faces — RDC 727/2022)
  4. Glúten e lactose
  5. Código de barras (placeholder visual — retângulo com linhas verticais)

LAYOUT viewBox="0 0 180 600" (orientação retrato estreita — lateral de caixa):
  Fundo: cor primária da marca (espelho da frente — identidade na prateleira)
  Denominação: rotacionada 90° ou horizontal, topo, branco bold, font-size=14px
  Conservação: texto vertical ou horizontal pequeno, branco opacity=0.9, font-size=8px
  Alérgenos: faixa horizontal no meio, fundo branco/amarelo, texto escuro, font-size=7px
  Código de barras (placeholder):
    Bloco branco rx=3 na base
    30-40 linhas verticais de largura variável (1-3px), altura 50px, fill=#000
    Número fictício abaixo: "7890000000000", font-size=6px
  Linha decorativa longitudinal: faixa de 3px da cor de acento ao longo da lateral

DICA: laterais estreitas pedem tipografia compacta e elementos essenciais apenas.
Não tente encaixar a tabela nutricional na lateral — ela vai no verso.
""",
}

def _montar_prompt_painel(painel_key: str, campos: dict, cores: dict, logo_b64_small: str, design_template: str, evite_rules: str) -> list:
    """Monta msg_content para um painel específico da embalagem."""
    tn = campos.get("tabela_nutricional") or {}

    instrucao_painel = SP_PAINEL[painel_key]

    prompt = (SP_DESIGN_ROTULO
        .replace("{logo_b64}", logo_b64_small)
        .replace("{cor_primaria}", cores["primaria"])
        .replace("{formato}", "MULTIPANEL — " + painel_key.upper())
        .replace("{denominacao}", campos.get("denominacao", ""))
        .replace("{carimbo}", campos.get("carimbo", ""))
        .replace("{fabricante}", (campos.get("fabricante") or "").replace("\n", " | "))
        .replace("{ingredientes}", campos.get("ingredientes", ""))
        .replace("{conservacao}", campos.get("conservacao", ""))
        .replace("{conteudo_liquido}", campos.get("conteudo_liquido", ""))
        .replace("{alergenos}", campos.get("alergenos", ""))
        .replace("{transgenicos}", campos.get("transgenicos", ""))
        .replace("{porcao}", tn.get("porcao", "100g"))
        .replace("{energia_kcal}", tn.get("energia_kcal", ""))
        .replace("{proteinas}", tn.get("proteinas", ""))
        .replace("{carboidratos}", tn.get("carboidratos", ""))
        .replace("{gorduras_totais}", tn.get("gorduras_totais", ""))
        .replace("{gorduras_saturadas}", tn.get("gorduras_saturadas", ""))
        .replace("{gorduras_trans}", tn.get("gorduras_trans", "0g"))
        .replace("{fibra}", tn.get("fibra", ""))
        .replace("{sodio}", tn.get("sodio", ""))
        .replace("{gluten}", campos.get("gluten", ""))
        .replace("{lactose}", campos.get("lactose", ""))
        .replace("{evite_rules}", evite_rules or "")
        .replace("{design_template}", design_template)
    ) + f"\n\n{instrucao_painel}\n\nGere APENAS o SVG deste painel. Comece com <svg e termine com </svg>."

    if logo_b64_small:
        return [
            {"type": "image", "source": {"type": "base64", "media_type": "image/png", "data": logo_b64_small}},
            {"type": "text", "text": prompt + "\n\nA imagem acima é a LOGO DA EMPRESA. Use-a no SVG da frente."}
        ]
    return [{"type": "text", "text": prompt}]


async def _gerar_painel(painel_key: str, msg_content: list) -> dict:
    """Gera um único painel SVG via API."""
    try:
        async with httpx.AsyncClient(timeout=90.0) as client:
            r = await client.post(
                "https://api.anthropic.com/v1/messages",
                headers={"x-api-key": ANTHROPIC_API_KEY,
                         "anthropic-version": "2023-06-01",
                         "content-type": "application/json"},
                json={"model": ANTHROPIC_MODEL,
                      "max_tokens": 8000,
                      "system": "Voce e um designer senior de embalagens. Gere apenas codigo SVG valido para o painel solicitado, nada mais.",
                      "messages": [{"role": "user", "content": msg_content}]}
            )
            svg_raw = _limpar_svg(r.json()["content"][0]["text"].strip())

        png_b64 = ""
        try:
            import fitz as _fitz
            doc = _fitz.open(stream=svg_raw.encode(), filetype="svg")
            pix = doc[0].get_pixmap(matrix=_fitz.Matrix(2, 2), alpha=False)
            png_b64 = base64.b64encode(pix.tobytes("png")).decode()
        except Exception:
            pass

        return {"painel": painel_key, "svg": svg_raw, "png": png_b64, "ok": True}
    except Exception as e:
        return {"painel": painel_key, "svg": "", "png": "", "ok": False, "error": str(e)[:100]}


@app.post("/gerar-design-multipanel")
async def gerar_design_multipanel(request: Request):
    """
    D10 — Gera frente + verso + lateral em paralelo.
    Body: mesmo formato do /gerar-design.
    Response: { "paineis": { "frente": {svg,png}, "verso": {svg,png}, "lateral": {svg,png} }, "cores": {...} }
    Painéis solicitados via body.get("paineis", ["frente","verso","lateral"]).
    """
    if not ANTHROPIC_API_KEY:
        return JSONResponse({"error": "API nao configurada"}, status_code=400,
                            headers={"Access-Control-Allow-Origin": "*"})
    try:
        body = await request.json()
        campos = body.get("campos", {})
        logo_b64 = body.get("logo_b64", "")
        paineis_solicitados = body.get("paineis", ["frente", "verso", "lateral"])
        # Valida painéis
        paineis_validos = [p for p in paineis_solicitados if p in SP_PAINEL]
        if not paineis_validos:
            paineis_validos = ["frente", "verso", "lateral"]

        cores = _extrair_cores_logo(logo_b64) if logo_b64 else {"primaria": "#1a1a2e"}
        categoria_design = campos.get("categoria", "outro")
        design_template = get_design_template(categoria_design)
        evite_rules = await _get_evite_rules(categoria_design, campos.get("orgao_sigla", ""))

        # Redimensiona logo uma vez
        logo_b64_small = ""
        if logo_b64:
            try:
                from PIL import Image as _PIL
                import io as _io
                logo_img = _PIL.open(_io.BytesIO(base64.b64decode(logo_b64)))
                logo_img.thumbnail((400, 400))
                buf = _io.BytesIO()
                logo_img.save(buf, "PNG")
                logo_b64_small = base64.b64encode(buf.getvalue()).decode()
            except Exception:
                logo_b64_small = logo_b64[:50000]

        # Monta prompts e dispara em paralelo
        tarefas = [
            _gerar_painel(
                painel_key,
                _montar_prompt_painel(painel_key, campos, cores, logo_b64_small, design_template, evite_rules)
            )
            for painel_key in paineis_validos
        ]
        resultados = await asyncio.gather(*tarefas)

        paineis_dict = {r["painel"]: r for r in resultados}

        return JSONResponse({
            "paineis": paineis_dict,
            "cores": cores,
            "paineis_gerados": paineis_validos,
            "labels": {
                "frente":  "🎯 Frente",
                "verso":   "📋 Verso",
                "lateral": "↔️ Lateral",
            }
        }, headers={"Access-Control-Allow-Origin": "*"})

    except Exception as e:
        return JSONResponse({"error": str(e)[:300]}, status_code=500,
                            headers={"Access-Control-Allow-Origin": "*"})
async def gerar_design_rotulo(request: Request):
    """
    Fase 4 — Gera design profissional do rotulo como SVG.
    Claude atua como designer senior, produz SVG criativo com logo integrada.
    Body JSON: { campos: {...}, logo_b64: "...", formato: "circular|retangular|quadrado" }
    """
    if not ANTHROPIC_API_KEY:
        return JSONResponse({"error": "API nao configurada"}, status_code=400,
                            headers={"Access-Control-Allow-Origin": "*"})
    try:
        body = await request.json()
        campos = body.get("campos", {})
        logo_b64 = body.get("logo_b64", "")
        formato = body.get("formato", "retangular")

        # Extrai cores da logo
        cores = _extrair_cores_logo(logo_b64) if logo_b64 else {"primaria": "#1a1a2e"}

        tn = campos.get("tabela_nutricional") or {}

        _logo_placeholder = logo_b64_small if logo_b64 else ""
        prompt = (SP_DESIGN_ROTULO
            .replace("{logo_b64}", _logo_placeholder)
            .replace("{cor_primaria}", cores["primaria"])
            .replace("{formato}", formato.upper())
            .replace("{denominacao}", campos.get("denominacao", ""))
            .replace("{carimbo}", campos.get("carimbo", ""))
            .replace("{fabricante}", (campos.get("fabricante") or "").replace("\n", " | "))
            .replace("{ingredientes}", campos.get("ingredientes", ""))
            .replace("{conservacao}", campos.get("conservacao", ""))
            .replace("{conteudo_liquido}", campos.get("conteudo_liquido", ""))
            .replace("{alergenos}", campos.get("alergenos", ""))
            .replace("{transgenicos}", campos.get("transgenicos", ""))
            .replace("{porcao}", tn.get("porcao", "100g"))
            .replace("{energia_kcal}", tn.get("energia_kcal", ""))
            .replace("{proteinas}", tn.get("proteinas", ""))
            .replace("{carboidratos}", tn.get("carboidratos", ""))
            .replace("{gorduras_totais}", tn.get("gorduras_totais", ""))
            .replace("{gorduras_saturadas}", tn.get("gorduras_saturadas", ""))
            .replace("{gorduras_trans}", tn.get("gorduras_trans", "0g"))
            .replace("{fibra}", tn.get("fibra", ""))
            .replace("{sodio}", tn.get("sodio", ""))
            .replace("{gluten}", campos.get("gluten", ""))
            .replace("{lactose}", campos.get("lactose", ""))
            .replace("{evite_rules}", "")
            .replace("{design_template}", design_template)
        )

        # Sistema 1: busca referências visuais da categoria
        categoria_design = campos.get("categoria", "outro")

        # D2-D5: injeta template visual da categoria
        design_template = get_design_template(categoria_design)

        refs = await _get_referencias_para_design(categoria_design, campos.get("orgao_sigla", ""))
        
        # Sistema 2: carrega regras EVITE acumuladas do feedback
        evite_rules = await _get_evite_rules(categoria_design, campos.get("orgao_sigla", ""))
        if evite_rules:
            prompt += evite_rules

        # Monta a mensagem com logo + referências se disponíveis
        if logo_b64:
            try:
                from PIL import Image as _PIL
                import io as _io
                logo_img = _PIL.open(_io.BytesIO(base64.b64decode(logo_b64)))
                logo_img.thumbnail((400, 400))
                buf = _io.BytesIO()
                logo_img.save(buf, "PNG")
                logo_b64_small = base64.b64encode(buf.getvalue()).decode()
            except Exception:
                logo_b64_small = logo_b64[:50000]

            msg_content = [
                {"type": "image", "source": {"type": "base64",
                 "media_type": "image/png", "data": logo_b64_small}},
                {"type": "text", "text": prompt + "\n\nA imagem acima é a LOGO DA EMPRESA. "
                 "Use-a como <image href='data:image/png;base64," + logo_b64_small + "' .../> no SVG. "
                 "Extraia as cores dominantes da logo para o esquema de cores do rotulo."}
            ]
        else:
            msg_content = [{"type": "text", "text": prompt}]

        # Sistema 1: injeta referências visuais como imagens no prompt
        if refs:
            ref_texts = []
            for i, ref_b64 in enumerate(refs[:2]):
                try:
                    ref_small = ref_b64[:40000]  # limite seguro
                    msg_content.insert(0, {"type": "image", "source": {
                        "type": "base64", "media_type": "image/png", "data": ref_small}})
                    ref_texts.append(f"Imagem {i+1}: rotulo de referencia aprovado para a categoria {categoria_design}")
                except Exception:
                    pass
            if ref_texts:
                ref_intro = {"type": "text", "text": "ROTULOS DE REFERENCIA APROVADOS (use como inspiracao de layout e qualidade):\n" + "\n".join(ref_texts)}
                msg_content.insert(0, ref_intro)

        async with httpx.AsyncClient(timeout=90.0) as client:
            r = await client.post(
                "https://api.anthropic.com/v1/messages",
                headers={"x-api-key": ANTHROPIC_API_KEY,
                         "anthropic-version": "2023-06-01",
                         "content-type": "application/json"},
                json={"model": ANTHROPIC_MODEL,
                      "max_tokens": 8000,
                      "system": "Voce e um designer senior de embalagens. Gere apenas codigo SVG valido, nada mais.",
                      "messages": [{"role": "user", "content": msg_content}]}
            )
            data = r.json()
            svg_raw = data["content"][0]["text"].strip()

        # Limpa markdown se vier
        if "```svg" in svg_raw:
            svg_raw = svg_raw.split("```svg")[1].split("```")[0].strip()
        elif "```xml" in svg_raw:
            svg_raw = svg_raw.split("```xml")[1].split("```")[0].strip()
        elif "```" in svg_raw:
            svg_raw = svg_raw.split("```")[1].split("```")[0].strip()

        # Garante que comeca com <svg
        if not svg_raw.startswith("<svg"):
            idx = svg_raw.find("<svg")
            if idx >= 0:
                svg_raw = svg_raw[idx:]

        # Converte SVG para PNG usando PyMuPDF
        png_b64 = ""
        try:
            import fitz as _fitz
            doc = _fitz.open(stream=svg_raw.encode(), filetype="svg")
            page = doc[0]
            mat = _fitz.Matrix(3, 3)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            png_b64 = base64.b64encode(pix.tobytes("png")).decode()
        except Exception:
            pass

        return JSONResponse({
            "svg": svg_raw,
            "png": png_b64,
            "cores": cores,
        }, headers={"Access-Control-Allow-Origin": "*"})

    except Exception as e:
        return JSONResponse({"error": str(e)[:300]}, status_code=500,
                            headers={"Access-Control-Allow-Origin": "*"})


# ═══════════════════════════════════════════════════════════════════════════════
# D8 — MÚLTIPLAS VARIAÇÕES DE DESIGN (3 opções para o RT escolher)
# Gera Bold/Colorida + Clássica/Tradicional + Clean/Minimalista em paralelo
# ═══════════════════════════════════════════════════════════════════════════════

ESTILOS_VARIACAO = {
    "bold": """
ESTILO DESTA VARIAÇÃO: BOLD / COLORIDA / IMPACTANTE
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Objetivo: máximo impacto visual na prateleira. O consumidor vê a 3 metros.

REGRAS ESPECÍFICAS DESTE ESTILO:
• Cor primária: use a versão mais SATURADA e INTENSA possível da paleta da categoria
• Nome do produto: maior que o normal — ocupa 40-50% da altura da zona emocional
• Fonte: bold máximo (font-weight=900 ou Black), sem italic
• Contraste: texto branco sobre fundo escuro saturado — nunca cinza, nunca pastéis
• Elemento gráfico: mais proeminente e centralizado — ocupa 35% da zona emocional
• Gradiente: linear forte de escuro→mais escuro (não suave, não pastel)
• Claim: em destaque — caixa com fundo de acento, texto bold
• Conteúdo líquido: badge grande, fill sólido, nunca outline
• Bordas internas: use divisórias espessas (3-4px) na cor de acento
• Sensação: produto premium de supermercado grande, Sadia/Perdigão tier
""",
    "classica": """
ESTILO DESTA VARIAÇÃO: CLÁSSICA / TRADICIONAL / ARTESANAL
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Objetivo: transmitir tradição, confiança e história da marca. Rótulo "de família".

REGRAS ESPECÍFICAS DESTE ESTILO:
• Paleta: tons mais escuros e envelhecidos — vermelho vinho (não vermelho), dourado (não amarelo)
• Background zona emocional: textura simulada (madeira, pergaminho) com linhas diagonais sutis
• Fonte nome: serifada sempre — Georgia, Times, Playfair Display
• Ornamentos: use elementos decorativos como linhas duplas (rect finos sobrepostos), bordas ornamentadas
• Faixa de título: retângulo com bordas duplas finas contendo o nome do produto — estilo "etiqueta de vinho"
• Escudo/brasão: elemento oval ou shield estilizado ao redor do carimbo — integra logo e carimbo
• Claim: em script/italic serifado, estilo caligráfico
• Conteúdo líquido: oval ou círculo com borda dupla, não badge moderno
• Sombras: mais pronunciadas nos elementos principais
• Sensação: produto artesanal premiado, queijaria mineira, defumados da serra
""",
    "clean": """
ESTILO DESTA VARIAÇÃO: CLEAN / MINIMALISTA / PREMIUM MODERNO
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Objetivo: sofisticação pelo que NÃO tem. Produto premium urbano.

REGRAS ESPECÍFICAS DESTE ESTILO:
• Fundo zona emocional: UMA cor sólida flat — sem gradientes, sem texturas, sem padrões
• Paleta máxima: 2 cores + branco. Nada mais.
• Espaço negativo: use generosamente — menos elementos, mais respiro
• Nome do produto: tamanho moderado (não gigante), font-weight=600 (não 900), letter-spacing=2-3
• Fonte: sans-serif geométrica — Arial com letter-spacing, simulando Futura/Helvetica Neue
• Linha divisória: única, fina (1px), cor de acento — substitui todos os outros separadores
• Elemento gráfico: minimalista — ícone geométrico simples (círculo, linha, ponto) ao invés de ilustração
• Claim: muito pequeno (9px), all-caps, letter-spacing=3, cor de acento
• Conteúdo líquido: apenas texto, sem badge — integrado à hierarquia tipográfica
• Carimbo: reduzido ao mínimo funcional, sem ornamentos
• Sensação: produto orgânico premium, marca de nicho, mercado gourmet
""",
}

def _limpar_svg(svg_raw: str) -> str:
    """Remove markdown e garante que começa com <svg."""
    for marker in ["```svg", "```xml", "```"]:
        if marker in svg_raw:
            parts = svg_raw.split(marker)
            svg_raw = parts[1].split("```")[0].strip() if len(parts) > 1 else svg_raw
            break
    if not svg_raw.startswith("<svg"):
        idx = svg_raw.find("<svg")
        if idx >= 0:
            svg_raw = svg_raw[idx:]
    return svg_raw

async def _gerar_uma_variacao(
    msg_content: list,
    estilo_key: str,
    estilo_prompt: str,
) -> dict:
    """Gera uma variação de design com o estilo especificado."""
    conteudo = list(msg_content)  # cópia
    # Injeta instrução de estilo no último elemento de texto
    for item in reversed(conteudo):
        if item.get("type") == "text":
            item["text"] = item["text"] + "\n\n" + estilo_prompt
            break

    try:
        async with httpx.AsyncClient(timeout=90.0) as client:
            r = await client.post(
                "https://api.anthropic.com/v1/messages",
                headers={"x-api-key": ANTHROPIC_API_KEY,
                         "anthropic-version": "2023-06-01",
                         "content-type": "application/json"},
                json={"model": ANTHROPIC_MODEL,
                      "max_tokens": 8000,
                      "system": "Voce e um designer senior de embalagens. Gere apenas codigo SVG valido, nada mais.",
                      "messages": [{"role": "user", "content": conteudo}]}
            )
            svg_raw = _limpar_svg(r.json()["content"][0]["text"].strip())

        png_b64 = ""
        try:
            import fitz as _fitz
            doc = _fitz.open(stream=svg_raw.encode(), filetype="svg")
            pix = doc[0].get_pixmap(matrix=_fitz.Matrix(2, 2), alpha=False)
            png_b64 = base64.b64encode(pix.tobytes("png")).decode()
        except Exception:
            pass

        return {"estilo": estilo_key, "svg": svg_raw, "png": png_b64, "ok": True}

    except Exception as e:
        return {"estilo": estilo_key, "svg": "", "png": "", "ok": False, "error": str(e)[:100]}


@app.post("/gerar-design-variacoes")
async def gerar_design_variacoes(request: Request):
    """
    D8 — Gera 3 variações de design em paralelo para o RT escolher.
    Bold/Colorida + Clássica/Tradicional + Clean/Minimalista.
    Body: mesmo formato do /gerar-design.
    Response: { "variacoes": [ {estilo, svg, png}, ... ], "cores": {...} }
    """
    if not ANTHROPIC_API_KEY:
        return JSONResponse({"error": "API nao configurada"}, status_code=400,
                            headers={"Access-Control-Allow-Origin": "*"})
    try:
        body = await request.json()
        campos = body.get("campos", {})
        logo_b64 = body.get("logo_b64", "")
        formato = body.get("formato", "retangular")

        cores = _extrair_cores_logo(logo_b64) if logo_b64 else {"primaria": "#1a1a2e"}
        tn = campos.get("tabela_nutricional") or {}
        categoria_design = campos.get("categoria", "outro")
        design_template = get_design_template(categoria_design)
        evite_rules = await _get_evite_rules(categoria_design, campos.get("orgao_sigla", ""))

        # Monta prompt base (igual ao /gerar-design)
        logo_b64_small = ""
        if logo_b64:
            try:
                from PIL import Image as _PIL
                import io as _io
                logo_img = _PIL.open(_io.BytesIO(base64.b64decode(logo_b64)))
                logo_img.thumbnail((400, 400))
                buf = _io.BytesIO()
                logo_img.save(buf, "PNG")
                logo_b64_small = base64.b64encode(buf.getvalue()).decode()
            except Exception:
                logo_b64_small = logo_b64[:50000]

        prompt_base = (SP_DESIGN_ROTULO
            .replace("{logo_b64}", logo_b64_small)
            .replace("{cor_primaria}", cores["primaria"])
            .replace("{formato}", formato.upper())
            .replace("{denominacao}", campos.get("denominacao", ""))
            .replace("{carimbo}", campos.get("carimbo", ""))
            .replace("{fabricante}", (campos.get("fabricante") or "").replace("\n", " | "))
            .replace("{ingredientes}", campos.get("ingredientes", ""))
            .replace("{conservacao}", campos.get("conservacao", ""))
            .replace("{conteudo_liquido}", campos.get("conteudo_liquido", ""))
            .replace("{alergenos}", campos.get("alergenos", ""))
            .replace("{transgenicos}", campos.get("transgenicos", ""))
            .replace("{porcao}", tn.get("porcao", "100g"))
            .replace("{energia_kcal}", tn.get("energia_kcal", ""))
            .replace("{proteinas}", tn.get("proteinas", ""))
            .replace("{carboidratos}", tn.get("carboidratos", ""))
            .replace("{gorduras_totais}", tn.get("gorduras_totais", ""))
            .replace("{gorduras_saturadas}", tn.get("gorduras_saturadas", ""))
            .replace("{gorduras_trans}", tn.get("gorduras_trans", "0g"))
            .replace("{fibra}", tn.get("fibra", ""))
            .replace("{sodio}", tn.get("sodio", ""))
            .replace("{gluten}", campos.get("gluten", ""))
            .replace("{lactose}", campos.get("lactose", ""))
            .replace("{evite_rules}", evite_rules or "")
            .replace("{design_template}", design_template)
        )

        # Monta msg_content base
        if logo_b64_small:
            msg_content_base = [
                {"type": "image", "source": {"type": "base64",
                 "media_type": "image/png", "data": logo_b64_small}},
                {"type": "text", "text": prompt_base +
                 "\n\nA imagem acima é a LOGO DA EMPRESA. Use-a no SVG. "
                 "Gere o SVG profissional completo. APENAS o código SVG."}
            ]
        else:
            msg_content_base = [{"type": "text", "text": prompt_base +
                "\n\nGere o SVG profissional completo. APENAS o código SVG."}]

        # Dispara as 3 variações em paralelo
        tarefas = [
            _gerar_uma_variacao(msg_content_base, estilo, prompt)
            for estilo, prompt in ESTILOS_VARIACAO.items()
        ]
        variacoes = await asyncio.gather(*tarefas)

        return JSONResponse({
            "variacoes": list(variacoes),
            "cores": cores,
            "labels": {
                "bold":     "🎨 Bold / Colorida",
                "classica": "🏛️ Clássica / Tradicional",
                "clean":    "✨ Clean / Minimalista",
            }
        }, headers={"Access-Control-Allow-Origin": "*"})

    except Exception as e:
        return JSONResponse({"error": str(e)[:300]}, status_code=500,
                            headers={"Access-Control-Allow-Origin": "*"})



# ══════════════════════════════════════════════════════════════════════════════
# SISTEMA 1 — Banco de referências visuais de design
# SISTEMA 2 — Feedback loop de design
# ══════════════════════════════════════════════════════════════════════════════

# Cache em memória para referências e restrições acumuladas
_referencias_cache: dict = {}   # {categoria_orgao: [lista de imagens b64]}
_evite_cache: dict = {}         # {categoria_orgao: [lista de restrições]}


@app.post("/referencias-design")
async def salvar_referencia_design(request: Request):
    """Upload de rótulo aprovado como referência visual."""
    try:
        body = await request.json()
        ref = {
            "categoria":  body.get("categoria", "outro"),
            "orgao":      body.get("orgao", "SIF"),
            "estilo":     body.get("estilo", "moderno"),
            "imagem_b64": body.get("imagem_b64", ""),
            "descricao":  body.get("descricao", ""),
            "aprovado":   True,
        }
        if not ref["imagem_b64"]:
            return JSONResponse({"error": "imagem_b64 obrigatoria"}, status_code=400,
                                headers={"Access-Control-Allow-Origin": "*"})

        salvo = await _sb_upsert("referencias_design", ref)

        # Invalida cache para essa categoria
        key = f"{ref['categoria']}_{ref['orgao']}"
        _referencias_cache.pop(key, None)

        return JSONResponse({"status": "ok", "supabase": salvo},
                            headers={"Access-Control-Allow-Origin": "*"})
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500,
                            headers={"Access-Control-Allow-Origin": "*"})


@app.get("/referencias-design/{categoria}")
async def listar_referencias(categoria: str, orgao: str = ""):
    """Lista referências disponíveis para uma categoria."""
    try:
        filtros = {"categoria": categoria, "aprovado": True}
        if orgao:
            filtros["orgao"] = orgao
        rows = await _sb_get("referencias_design", filtros, limit=10)
        # Retorna sem imagem_b64 para não explodir a resposta
        return JSONResponse(
            [{"id": r.get("id"), "categoria": r.get("categoria"),
              "orgao": r.get("orgao"), "estilo": r.get("estilo"),
              "descricao": r.get("descricao")} for r in (rows or [])],
            headers={"Access-Control-Allow-Origin": "*"})
    except Exception as e:
        return JSONResponse([], headers={"Access-Control-Allow-Origin": "*"})


async def _get_referencias_para_design(categoria: str, orgao: str) -> list[str]:
    """Busca até 2 imagens de referência para injetar no prompt de design."""
    key = f"{categoria}_{orgao}"
    if key in _referencias_cache:
        return _referencias_cache[key]
    try:
        filtros = {"categoria": categoria, "aprovado": True}
        if orgao:
            filtros["orgao"] = orgao
        rows = await _sb_get("referencias_design", filtros, limit=2)
        imgs = [r["imagem_b64"] for r in (rows or []) if r.get("imagem_b64")]
        _referencias_cache[key] = imgs
        return imgs
    except Exception:
        return []


@app.post("/feedback-design")
async def salvar_feedback_design(request: Request):
    """Salva avaliação de design. Restrições ruins viram regras EVITE."""
    try:
        body = await request.json()
        categoria = body.get("categoria", "outro")
        orgao     = body.get("orgao", "SIF")
        rating    = body.get("rating", "")      # "perfeito" | "bom" | "ruim"
        problema  = body.get("problema", "")    # texto livre quando ruim

        fb = {
            "categoria": categoria,
            "orgao":     orgao,
            "rating":    rating,
            "problema":  problema,
            "timestamp": __import__("datetime").datetime.now().isoformat(),
        }
        await _sb_upsert("design_feedback", fb)

        # Se ruim com descrição, adiciona ao cache de restrições
        if rating == "ruim" and problema.strip():
            key = f"{categoria}_{orgao}"
            if key not in _evite_cache:
                _evite_cache[key] = []
            _evite_cache[key].append(problema.strip()[:120])
            # Mantém no máximo 10 restrições por categoria
            _evite_cache[key] = _evite_cache[key][-10:]

        return JSONResponse({"status": "ok"},
                            headers={"Access-Control-Allow-Origin": "*"})
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500,
                            headers={"Access-Control-Allow-Origin": "*"})


@app.get("/feedback-design/stats")
async def stats_feedback_design():
    """Retorna restrições acumuladas por categoria (para debug/admin)."""
    return JSONResponse(_evite_cache,
                        headers={"Access-Control-Allow-Origin": "*"})


async def _get_evite_rules(categoria: str, orgao: str) -> str:
    """Monta bloco EVITE para injetar no prompt com base nos feedbacks."""
    # Tenta carregar do Supabase se cache vazio
    key = f"{categoria}_{orgao}"
    if key not in _evite_cache:
        try:
            rows = await _sb_get("design_feedback",
                                  {"categoria": categoria, "orgao": orgao, "rating": "ruim"},
                                  limit=15)
            problemas = [r["problema"] for r in (rows or [])
                         if r.get("problema") and r["problema"].strip()]
            _evite_cache[key] = problemas[-10:]
        except Exception:
            _evite_cache[key] = []

    regras = _evite_cache.get(key, [])
    if not regras:
        return ""
    return "\n\nRESTRIÇÕES BASEADAS EM FEEDBACKS ANTERIORES — OBRIGATÓRIO RESPEITAR:\n" + \
           "\n".join(f"- EVITE: {r}" for r in regras)


@app.get("/monitor/alertas-ativos")
async def alertas_rotulos_impactados():
    """Rótulos possivelmente impactados pelas últimas alertas do DOU."""
    try:
        alertas = await _sb_get("monitor_alertas", {}, limit=20)
        alertas_recentes = [a for a in (alertas or []) if not a.get("descartado")][:5]
        rotulos = await _sb_get("validacoes", {}, limit=100)
        impactos = []
        for alerta in alertas_recentes:
            norma     = alerta.get("norma", "")
            categoria = alerta.get("categoria", "")
            titulo    = alerta.get("titulo", "")
            data_pub  = alerta.get("data_publicacao", "")
            rotulos_impactados = []
            if rotulos and categoria:
                for r in rotulos:
                    cats    = r.get("categorias") or []
                    produto = r.get("produto", "") or ""
                    hit = (
                        (isinstance(cats, list) and any(categoria.lower() in ci.lower() for ci in cats))
                        or categoria.lower() in produto.lower()
                    )
                    if hit:
                        rotulos_impactados.append({
                            "produto":          produto,
                            "data_validacao":   r.get("criado_em", ""),
                            "case_id":          r.get("case_id", ""),
                        })
            impactos.append({
                "norma": norma, "titulo": titulo, "data_publicacao": data_pub,
                "categoria": categoria,
                "rotulos_possivelmente_impactados": rotulos_impactados[:10],
                "total_impactados": len(rotulos_impactados),
            })
        return JSONResponse({"alertas": impactos, "total": len(impactos)},
                            headers={"Access-Control-Allow-Origin": "*"})
    except Exception as e:
        return JSONResponse({"alertas": [], "error": str(e)[:200]},
                            headers={"Access-Control-Allow-Origin": "*"})


# ═══════════════════════════════════════════════════════════════════════════════
# M1_8 — TRIAL 14 DIAS + EMAILS DE ONBOARDING
# Env: RESEND_API_KEY ou SENDGRID_API_KEY
# ═══════════════════════════════════════════════════════════════════════════════

_RESEND_KEY = os.environ.get("RESEND_API_KEY", "")
_EMAIL_FROM = os.environ.get("EMAIL_FROM", "noreply@validarotulo.com.br")

_EMAILS_ONBOARDING = {
    "d0": {
        "subject": "🎉 Bem-vindo ao ValidaRótulo IA!",
        "html": """<h2>Olá {nome}! Seu trial de 14 dias começou.</h2>
<p>Você tem acesso completo ao ValidaRótulo IA por 14 dias. Aqui está o que você pode fazer:</p>
<ul>
  <li>✅ <b>Validar rótulos</b> — envie a arte e receba relatório com 14 campos avaliados</li>
  <li>✏️ <b>Criar rótulos</b> — gere campos legais + design profissional com IA</li>
  <li>📄 <b>Baixar relatório PDF</b> — documento formal com campo de assinatura do RT</li>
</ul>
<p><a href="{frontend_url}" style="background:#7c3aed;color:#fff;padding:12px 24px;border-radius:8px;text-decoration:none;font-weight:bold">Começar agora →</a></p>
<p style="color:#9ca3af;font-size:12px">Trial expira em 14 dias. Dúvidas? Responda este email.</p>"""
    },
    "d3": {
        "subject": "💡 Dica: como tirar o melhor do ValidaRótulo IA",
        "html": """<h2>Olá {nome}, aqui vai uma dica rápida!</h2>
<p>Para validações mais precisas, envie <b>todas as faces da embalagem</b> — frente, verso e laterais juntas.</p>
<p>O agente analisa cada face e cruza informações entre elas (ex: ingredientes no verso vs alérgenos na frente).</p>
<p>Você também pode digitar uma observação como <i>"linguiça suína toscana, SIE-SP"</i> para ajudar na detecção.</p>
<p><a href="{frontend_url}" style="background:#7c3aed;color:#fff;padding:12px 24px;border-radius:8px;text-decoration:none;font-weight:bold">Validar agora →</a></p>"""
    },
    "d10": {
        "subject": "⏰ Seu trial expira em 4 dias",
        "html": """<h2>Olá {nome}, seu trial termina em breve.</h2>
<p>Você tem <b>4 dias restantes</b> de acesso completo ao ValidaRótulo IA.</p>
<p>Continue validando rótulos com total tranquilidade — e quando o trial terminar, escolha o plano certo para você:</p>
<ul>
  <li><b>Starter R$ 97/mês</b> — 30 validações + PDF + Criador completo</li>
  <li><b>Pro R$ 247/mês</b> — ilimitado + 5 usuários + suporte</li>
</ul>
<p>Use o cupom <b>EARLY30</b> para 30% off por 6 meses.</p>
<p><a href="{frontend_url}/planos.html" style="background:#7c3aed;color:#fff;padding:12px 24px;border-radius:8px;text-decoration:none;font-weight:bold">Ver planos →</a></p>"""
    },
    "d14": {
        "subject": "🔒 Seu trial terminou — continue sem interrupção",
        "html": """<h2>Olá {nome}, seu trial de 14 dias terminou.</h2>
<p>Para continuar validando rótulos, escolha um plano:</p>
<ul>
  <li><b>Starter R$ 97/mês</b> — 30 validações + PDF + Criador</li>
  <li><b>Pro R$ 247/mês</b> — ilimitado + multi-usuário</li>
</ul>
<p>Cupom <b>EARLY30</b> — 30% off nos primeiros 6 meses (expira em breve).</p>
<p><a href="{frontend_url}/planos.html" style="background:#7c3aed;color:#fff;padding:12px 24px;border-radius:8px;text-decoration:none;font-weight:bold">Assinar agora →</a></p>"""
    },
}

async def _enviar_email_onboarding(email: str, nome: str, etapa: str) -> bool:
    """Envia email de onboarding via Resend. Falha silenciosamente se não configurado."""
    if not _RESEND_KEY:
        return False
    tpl = _EMAILS_ONBOARDING.get(etapa)
    if not tpl:
        return False
    try:
        html = tpl["html"].replace("{nome}", nome or "RT").replace("{frontend_url}", _FRONTEND_URL)
        async with httpx.AsyncClient(timeout=10.0) as client:
            r = await client.post(
                "https://api.resend.com/emails",
                headers={"Authorization": f"Bearer {_RESEND_KEY}",
                         "Content-Type": "application/json"},
                json={"from": _EMAIL_FROM, "to": [email],
                      "subject": tpl["subject"], "html": html}
            )
            return r.status_code in (200, 201)
    except Exception:
        return False


# ─── Task #17 — Email de relatório após validação ──────────────────────────
async def _enviar_email_relatorio(
    email: str,
    nome: str,
    produto: str,
    veredicto: str,
    score: float | None,
    relatorio_md: str,
    case_id: str = "",
) -> bool:
    """
    Envia o relatório de validação por email após cada validação.
    Falha silenciosamente se Resend não configurado.
    """
    if not _RESEND_KEY or not email:
        return False

    score_str = f"{score}/14" if score is not None else "—"
    cor_veredicto = (
        "#166534" if veredicto and "APROVADO" in veredicto and "RESSALVAS" not in veredicto
        else "#854f0b" if veredicto and "RESSALVAS" in veredicto
        else "#a32d2d"
    )
    emoji_veredicto = (
        "✅" if veredicto and "APROVADO" in veredicto and "RESSALVAS" not in veredicto
        else "⚠️" if veredicto and "RESSALVAS" in veredicto
        else "❌"
    )

    # Converte markdown básico para HTML (negrito, listas, cabeçalhos)
    import re as _re
    rel_html = relatorio_md[:6000]  # limita tamanho
    rel_html = _re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', rel_html)
    rel_html = _re.sub(r'^### (.+)$', r'<h4 style="margin:12px 0 4px;color:#111">\1</h4>', rel_html, flags=_re.MULTILINE)
    rel_html = _re.sub(r'^## (.+)$', r'<h3 style="margin:14px 0 6px;color:#111">\1</h3>', rel_html, flags=_re.MULTILINE)
    rel_html = _re.sub(r'^# (.+)$', r'<h2 style="margin:16px 0 8px;color:#111">\1</h2>', rel_html, flags=_re.MULTILINE)
    rel_html = _re.sub(r'^• (.+)$', r'<li style="margin:3px 0">\1</li>', rel_html, flags=_re.MULTILINE)
    rel_html = _re.sub(r'^- (.+)$', r'<li style="margin:3px 0">\1</li>', rel_html, flags=_re.MULTILINE)
    rel_html = rel_html.replace('\n\n', '</p><p style="margin:8px 0">')
    rel_html = rel_html.replace('\n', '<br>')

    html = f"""
    <div style="font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',sans-serif;max-width:600px;margin:0 auto;padding:24px;background:#fff">
      <div style="background:#060A0F;padding:16px 20px;border-radius:10px 10px 0 0;display:flex;align-items:center;gap:10px">
        <span style="font-size:18px;font-weight:700;color:#F2EDD6">Inspect<span style="color:#D4860A">-IA</span></span>
        <span style="font-size:12px;color:#9ca3af">Validação de Rótulos</span>
      </div>

      <div style="background:#f9fafb;border:1px solid #e5e7eb;border-top:none;padding:20px 24px;border-radius:0 0 10px 10px">
        <p style="font-size:14px;color:#374151;margin:0 0 16px">Olá, <strong>{nome or 'RT'}</strong>. Segue o relatório da sua validação:</p>

        <div style="background:#fff;border:1px solid #e5e7eb;border-radius:8px;padding:14px 16px;margin-bottom:16px">
          <div style="font-size:12px;color:#6b7280;margin-bottom:4px">Produto</div>
          <div style="font-size:15px;font-weight:600;color:#111">{produto or '—'}</div>
          <div style="display:flex;gap:16px;margin-top:10px">
            <div>
              <div style="font-size:11px;color:#6b7280">Veredicto</div>
              <div style="font-size:14px;font-weight:700;color:{cor_veredicto}">{emoji_veredicto} {veredicto or '—'}</div>
            </div>
            <div>
              <div style="font-size:11px;color:#6b7280">Score</div>
              <div style="font-size:14px;font-weight:700;color:#111">{score_str}</div>
            </div>
          </div>
        </div>

        <div style="background:#fff;border:1px solid #e5e7eb;border-radius:8px;padding:14px 16px;margin-bottom:20px;font-size:13px;color:#374151;line-height:1.7">
          <p style="margin:8px 0">{rel_html}</p>
        </div>

        <p style="font-size:11px;color:#9ca3af;margin:0;text-align:center">
          Este relatório foi gerado pelo Inspect-IA e tem caráter auxiliar. Não substitui análise de RT habilitado.<br>
          <a href="{_FRONTEND_URL}" style="color:#1A7A4A">Acessar plataforma</a>
        </p>
      </div>
    </div>
    """

    try:
        async with httpx.AsyncClient(timeout=10.0) as client:
            r = await client.post(
                "https://api.resend.com/emails",
                headers={"Authorization": f"Bearer {_RESEND_KEY}",
                         "Content-Type": "application/json"},
                json={
                    "from": _EMAIL_FROM,
                    "to": [email],
                    "subject": f"{emoji_veredicto} Relatório ValidaRótulo — {produto or 'Produto'} ({score_str})",
                    "html": html,
                }
            )
            return r.status_code in (200, 201)
    except Exception:
        return False


@app.post("/billing/enviar-emails-trial")
async def enviar_emails_trial(request: Request):
    """
    Cron job diário — verifica usuários nos dias D3, D10, D14 do trial e envia emails.
    Chamar diariamente via cron-job.org: POST /billing/enviar-emails-trial
    """
    if not _SUPABASE_ON:
        return JSONResponse({"ok": False, "msg": "Supabase não configurado"},
                            headers={"Access-Control-Allow-Origin": "*"})
    try:
        import datetime as _dt_trial
        hoje = _dt_trial.datetime.now()

        rows = await _sb_get("assinaturas", {"status": "trial"}, limit=500)
        enviados = []

        for row in rows:
            trial_ends = row.get("trial_ends_at", "")
            if not trial_ends:
                continue
            try:
                ends_dt = _dt_trial.datetime.fromisoformat(trial_ends)
                dias_restantes = (ends_dt - hoje).days
                email = row.get("email", "")
                nome  = row.get("nome", "")
                uid   = row.get("user_id", "")

                # Busca email/nome do user se não tiver na tabela
                if not email and uid and _SUPABASE_ON:
                    user_rows = await _sb_get("perfis_empresa", {"perfil_id": uid}, limit=1)
                    if user_rows:
                        email = user_rows[0].get("email", "")

                etapa = None
                if dias_restantes == 11:   etapa = "d3"
                elif dias_restantes == 4:  etapa = "d10"
                elif dias_restantes <= 0:  etapa = "d14"

                if etapa and email:
                    ok = await _enviar_email_onboarding(email, nome, etapa)
                    if ok:
                        enviados.append({"email": email, "etapa": etapa})
            except Exception:
                continue

        return JSONResponse({"ok": True, "enviados": len(enviados), "detalhes": enviados},
                            headers={"Access-Control-Allow-Origin": "*"})
    except Exception as e:
        return JSONResponse({"error": str(e)[:200]},
                            headers={"Access-Control-Allow-Origin": "*"})


# ═══════════════════════════════════════════════════════════════════════════════
# M1_9 — RATE LIMITING POR PLANO
# Limites: trial/free=5/mês, starter=30/mês, pro=ilimitado
# Contador: tabela uso_validacoes (user_id, ano_mes, contagem)
# ═══════════════════════════════════════════════════════════════════════════════

_LIMITES_PLANO = {
    "trial":   5,
    "free":    5,
    "starter": 30,
    "pro":     999999,
    "":        5,  # sem plano = trial
}

async def _get_plano_usuario(user_id: str) -> str:
    """Retorna o plano atual do usuário ('trial', 'starter', 'pro')."""
    if not _SUPABASE_ON or not user_id:
        return "trial"
    try:
        rows = await _sb_get("assinaturas", {"user_id": user_id}, limit=1)
        if not rows:
            return "trial"
        row = rows[0]
        status = row.get("status", "trial")
        plano  = row.get("plano", "trial")
        # Trial expirado → downgrade para free
        if status == "trial":
            import datetime as _dt_pl
            trial_ends = row.get("trial_ends_at", "")
            if trial_ends:
                try:
                    ends = _dt_pl.datetime.fromisoformat(trial_ends)
                    if _dt_pl.datetime.now() > ends:
                        return "free"
                except Exception:
                    pass
        if status in ("cancelado", "pagamento_falhou"):
            return "free"
        return plano if plano in _LIMITES_PLANO else "trial"
    except Exception:
        return "trial"

async def _checar_e_incrementar_uso(user_id: str) -> dict:
    """
    Verifica se o usuário ainda tem validações disponíveis e incrementa o contador.
    Retorna: { ok: bool, usado: int, limite: int, plano: str }
    """
    import datetime as _dt_uso
    if not user_id:
        return {"ok": True, "usado": 0, "limite": 999999, "plano": "anonimo"}

    plano  = await _get_plano_usuario(user_id)
    limite = _LIMITES_PLANO.get(plano, 5)
    ano_mes = _dt_uso.datetime.now().strftime("%Y-%m")

    if not _SUPABASE_ON:
        return {"ok": True, "usado": 0, "limite": limite, "plano": plano}

    try:
        rows = await _sb_get("uso_validacoes",
                             {"user_id": user_id, "ano_mes": ano_mes}, limit=1)
        contagem = rows[0].get("contagem", 0) if rows else 0

        if contagem >= limite:
            return {"ok": False, "usado": contagem, "limite": limite, "plano": plano,
                    "msg": f"Limite de {limite} validações/mês atingido. "
                           f"{'Faça upgrade para continuar.' if plano != 'pro' else 'Entre em contato.'}"}

        # Incrementa
        nova = contagem + 1
        asyncio.ensure_future(_sb_upsert("uso_validacoes", {
            "user_id": user_id,
            "ano_mes": ano_mes,
            "contagem": nova,
            "plano": plano,
            "updated_at": _dt_uso.datetime.now().isoformat(),
        }))
        return {"ok": True, "usado": nova, "limite": limite, "plano": plano}

    except Exception:
        return {"ok": True, "usado": 0, "limite": limite, "plano": plano}


@app.get("/billing/uso/{user_id}")
async def consultar_uso(user_id: str):
    """Retorna uso atual do usuário no mês corrente."""
    import datetime as _dt_cons
    ano_mes = _dt_cons.datetime.now().strftime("%Y-%m")
    plano   = await _get_plano_usuario(user_id)
    limite  = _LIMITES_PLANO.get(plano, 5)
    contagem = 0
    if _SUPABASE_ON:
        try:
            rows = await _sb_get("uso_validacoes",
                                 {"user_id": user_id, "ano_mes": ano_mes}, limit=1)
            contagem = rows[0].get("contagem", 0) if rows else 0
        except Exception:
            pass
    return JSONResponse({
        "user_id":  user_id,
        "plano":    plano,
        "usado":    contagem,
        "limite":   limite,
        "restante": max(0, limite - contagem),
        "ano_mes":  ano_mes,
        "percentual": round(contagem / limite * 100) if limite < 999999 else 0,
    }, headers={"Access-Control-Allow-Origin": "*"})


# ─────────────────────────────────────────────────────────────────────────────
# AUTH — Supabase Auth (email + senha)
# ─────────────────────────────────────────────────────────────────────────────
_SUPA_AUTH = os.environ.get("SUPABASE_URL", "").rstrip("/") + "/auth/v1"
_SUPA_ANON = os.environ.get("SUPABASE_KEY", "")

async def _auth_post(path: str, payload: dict, token: str = "") -> dict:
    """Chama endpoint da Supabase Auth API."""
    headers = {
        "apikey": _SUPA_ANON,
        "Content-Type": "application/json",
    }
    if token:
        headers["Authorization"] = f"Bearer {token}"
    try:
        async with httpx.AsyncClient(timeout=10.0) as cl:
            r = await cl.post(f"{_SUPA_AUTH}{path}", json=payload, headers=headers)
            return r.json()
    except Exception as e:
        return {"error": {"message": str(e)}}

async def _auth_get_user(token: str) -> dict:
    """Retorna dados do usuário pelo access_token."""
    headers = {"apikey": _SUPA_ANON, "Authorization": f"Bearer {token}"}
    try:
        async with httpx.AsyncClient(timeout=8.0) as cl:
            r = await cl.get(f"{_SUPA_AUTH}/user", headers=headers)
            return r.json()
    except Exception as e:
        return {"error": str(e)}


@app.post("/auth/signup")
async def auth_signup(request: Request):
    """Cria nova conta com email + senha."""
    body = await request.json()
    email = (body.get("email") or "").strip().lower()
    password = body.get("password") or ""
    nome = (body.get("nome") or "").strip()

    if not email or not password:
        return JSONResponse({"error": "Email e senha são obrigatórios."},
                            status_code=400, headers={"Access-Control-Allow-Origin": "*"})
    if len(password) < 6:
        return JSONResponse({"error": "Senha deve ter pelo menos 6 caracteres."},
                            status_code=400, headers={"Access-Control-Allow-Origin": "*"})

    result = await _auth_post("/signup", {
        "email": email,
        "password": password,
        "data": {"nome": nome}
    })

    if result.get("error") or "error_code" in result:
        msg = result.get("error", {}).get("message") or result.get("msg") or "Erro ao criar conta."
        if "already" in msg.lower() or "exists" in msg.lower():
            msg = "Este email já está cadastrado."
        return JSONResponse({"error": msg}, status_code=400,
                            headers={"Access-Control-Allow-Origin": "*"})

    # M1_8 — Configura trial de 14 dias no Supabase + envia email D0
    user_id = result.get("user", {}).get("id") or result.get("id", "")
    if user_id and _SUPABASE_ON:
        import datetime as _dt_reg
        trial_ends = (_dt_reg.datetime.now() + _dt_reg.timedelta(days=14)).isoformat()
        asyncio.ensure_future(_sb_upsert("assinaturas", {
            "user_id":       user_id,
            "plano":         "trial",
            "status":        "trial",
            "trial_ends_at": trial_ends,
            "stripe_sub_id": "",
            "stripe_cust_id":"",
            "updated_at":    _dt_reg.datetime.now().isoformat(),
        }))
        asyncio.ensure_future(_enviar_email_onboarding(email, nome, "d0"))

    return JSONResponse({
        "ok": True,
        "message": "Conta criada. Verifique seu email para confirmar o cadastro.",
        "user": {"email": email, "nome": nome}
    }, headers={"Access-Control-Allow-Origin": "*"})


@app.post("/auth/login")
async def auth_login(request: Request):
    """Login com email + senha. Retorna access_token."""
    body = await request.json()
    email = (body.get("email") or "").strip().lower()
    password = body.get("password") or ""

    if not email or not password:
        return JSONResponse({"error": "Email e senha são obrigatórios."},
                            status_code=400, headers={"Access-Control-Allow-Origin": "*"})

    result = await _auth_post("/token?grant_type=password", {
        "email": email,
        "password": password
    })

    if result.get("error") or "error_code" in result or not result.get("access_token"):
        # Extração robusta da mensagem (Supabase varia: error pode ser dict OU string,
        # e em versões diferentes usa msg / error_description / error_code)
        err_obj = result.get("error")
        if isinstance(err_obj, dict):
            raw_msg = err_obj.get("message", "") or err_obj.get("description", "")
        elif isinstance(err_obj, str):
            raw_msg = err_obj
        else:
            raw_msg = ""
        raw_msg = (raw_msg
                   or result.get("msg", "")
                   or result.get("error_description", "")
                   or result.get("error_code", "")
                   or "")
        raw_lower = raw_msg.lower()

        # Email não confirmado (várias variações possíveis do Supabase / GoTrue)
        if ("not_confirmed" in raw_lower
                or "not confirmed" in raw_lower
                or "email_not_confirmed" in raw_lower
                or ("confirm" in raw_lower and "email" in raw_lower)):
            return JSONResponse({
                "error": "Confirme seu email antes de entrar. Verifique sua caixa de entrada (e a pasta de spam).",
                "code": "email_not_confirmed"
            }, status_code=403, headers={"Access-Control-Allow-Origin": "*"})

        # Rate limit
        if "rate" in raw_lower or "too many" in raw_lower or "limit" in raw_lower:
            return JSONResponse({
                "error": "Muitas tentativas seguidas. Aguarde alguns segundos e tente de novo.",
                "code": "rate_limited"
            }, status_code=429, headers={"Access-Control-Allow-Origin": "*"})

        # Credenciais inválidas (Supabase não diferencia email errado de senha errada por design)
        if ("invalid" in raw_lower or "credentials" in raw_lower
                or "invalid_grant" in raw_lower or "wrong" in raw_lower):
            return JSONResponse({
                "error": "Email ou senha incorretos. Se acabou de criar a conta, confirme o email primeiro.",
                "code": "invalid_credentials"
            }, status_code=401, headers={"Access-Control-Allow-Origin": "*"})

        # Genérico — devolve o que vier do Supabase pra ajudar diagnóstico
        return JSONResponse({
            "error": raw_msg or "Não foi possível entrar. Tente novamente em alguns minutos.",
            "code": "unknown"
        }, status_code=401, headers={"Access-Control-Allow-Origin": "*"})

    user_data = result.get("user", {})
    meta = user_data.get("user_metadata", {})
    return JSONResponse({
        "ok": True,
        "access_token": result["access_token"],
        "refresh_token": result.get("refresh_token", ""),
        "user": {
            "id": user_data.get("id", ""),
            "email": user_data.get("email", email),
            "nome": meta.get("nome") or meta.get("name") or "",
        }
    }, headers={"Access-Control-Allow-Origin": "*"})


@app.post("/auth/logout")
async def auth_logout(request: Request):
    """Invalida o token no Supabase."""
    auth_header = request.headers.get("Authorization", "")
    token = auth_header.replace("Bearer ", "").strip()
    if token:
        await _auth_post("/logout", {}, token=token)
    return JSONResponse({"ok": True}, headers={"Access-Control-Allow-Origin": "*"})


@app.get("/auth/me")
async def auth_me(request: Request):
    """Retorna dados do usuário logado."""
    auth_header = request.headers.get("Authorization", "")
    token = auth_header.replace("Bearer ", "").strip()
    if not token:
        return JSONResponse({"error": "Não autenticado."}, status_code=401,
                            headers={"Access-Control-Allow-Origin": "*"})
    user = await _auth_get_user(token)
    if user.get("error") or not user.get("id"):
        return JSONResponse({"error": "Token inválido ou expirado."}, status_code=401,
                            headers={"Access-Control-Allow-Origin": "*"})
    meta = user.get("user_metadata", {})
    return JSONResponse({
        "id": user["id"],
        "email": user.get("email", ""),
        "nome": meta.get("nome") or meta.get("name") or "",
    }, headers={"Access-Control-Allow-Origin": "*"})


@app.post("/auth/reset-password")
async def auth_reset_password(request: Request):
    """Envia email de reset de senha. O link aponta pra reset-password.html."""
    body = await request.json()
    email = (body.get("email") or "").strip().lower()
    if not email:
        return JSONResponse({"error": "Email obrigatório."},
                            status_code=400, headers={"Access-Control-Allow-Origin": "*"})
    # IMPORTANTE: redirect_to precisa estar nas Redirect URLs do Supabase
    # (já está coberto pelo wildcard https://cosmic-llama-9576a5.netlify.app/**)
    redirect_to = (body.get("redirect_url")
                   or "https://cosmic-llama-9576a5.netlify.app/reset-password.html")
    # Supabase moderno aceita redirect_to no body; também passamos via query como fallback
    safe_redirect = redirect_to.replace(" ", "%20")
    result = await _auth_post(f"/recover?redirect_to={safe_redirect}", {
        "email": email,
        "redirect_to": redirect_to,
    })
    # Sempre retornar sucesso (não vazar se email existe ou não)
    return JSONResponse({"ok": True, "message": "Se o email existir, você receberá um link de redefinição."},
                        headers={"Access-Control-Allow-Origin": "*"})


@app.post("/auth/update-password")
async def auth_update_password(request: Request):
    """Atualiza senha do usuário usando o access_token recebido via email de recovery."""
    body = await request.json()
    access_token = (body.get("access_token") or "").strip()
    new_password = body.get("password") or ""

    if not access_token:
        return JSONResponse({"error": "Token de recuperação ausente. Peça um novo link."},
                            status_code=400, headers={"Access-Control-Allow-Origin": "*"})
    if not new_password or len(new_password) < 6:
        return JSONResponse({"error": "Nova senha deve ter pelo menos 6 caracteres."},
                            status_code=400, headers={"Access-Control-Allow-Origin": "*"})

    # PUT /auth/v1/user com Bearer token + body {password: "..."}
    headers = {
        "apikey": _SUPA_ANON,
        "Content-Type": "application/json",
        "Authorization": f"Bearer {access_token}",
    }
    try:
        async with httpx.AsyncClient(timeout=10.0) as cl:
            r = await cl.put(f"{_SUPA_AUTH}/user",
                             json={"password": new_password},
                             headers=headers)
            try:
                result = r.json()
            except Exception:
                result = {}
            status = r.status_code
    except Exception as e:
        return JSONResponse({"error": "Erro de conexão. Tente novamente em alguns segundos."},
                            status_code=500, headers={"Access-Control-Allow-Origin": "*"})

    if status >= 400 or result.get("error") or "error_code" in result:
        err_obj = result.get("error")
        if isinstance(err_obj, dict):
            raw_msg = err_obj.get("message", "") or err_obj.get("description", "")
        elif isinstance(err_obj, str):
            raw_msg = err_obj
        else:
            raw_msg = ""
        raw_msg = (raw_msg
                   or result.get("msg", "")
                   or result.get("error_description", "")
                   or "")
        raw_lower = raw_msg.lower()

        if "expired" in raw_lower or ("invalid" in raw_lower and ("token" in raw_lower or "jwt" in raw_lower)):
            return JSONResponse({
                "error": "O link de recuperação expirou ou já foi usado. Peça um novo na tela de login.",
                "code": "token_expired"
            }, status_code=401, headers={"Access-Control-Allow-Origin": "*"})

        if "weak" in raw_lower or "short" in raw_lower or ("password" in raw_lower and "length" in raw_lower):
            return JSONResponse({
                "error": "Senha muito curta ou fraca. Use ao menos 6 caracteres.",
                "code": "weak_password"
            }, status_code=400, headers={"Access-Control-Allow-Origin": "*"})

        if "same" in raw_lower:
            return JSONResponse({
                "error": "A nova senha precisa ser diferente da anterior.",
                "code": "same_password"
            }, status_code=400, headers={"Access-Control-Allow-Origin": "*"})

        return JSONResponse({
            "error": raw_msg or "Não foi possível atualizar a senha.",
            "code": "unknown"
        }, status_code=400, headers={"Access-Control-Allow-Origin": "*"})

    return JSONResponse({
        "ok": True,
        "message": "Senha atualizada com sucesso. Faça login com a nova senha."
    }, headers={"Access-Control-Allow-Origin": "*"})

@app.on_event("startup")
async def startup_load():
    """No startup, carrega dados persistidos do Supabase em background."""
    asyncio.ensure_future(_startup_background())

async def _startup_background():
    """
    Carrega Supabase em background sem bloquear o health check.
    Tenta 3 vezes com backoff — Render free tier pode ser lento no cold start.
    T1 fix: também faz warmup do _kb_cache com as categorias mais usadas.
    """
    global _cases_db, _monitor_history, _kb_cache, _startup_ready, _startup_ts
    if not _SUPABASE_ON:
        return
    import asyncio as _aio
    import datetime as _dt_startup

    # ── 1. Recarrega _cases_db ──────────────────────────────────────────────
    for attempt in range(3):
        try:
            rows = await _aio.wait_for(load_cases_from_supabase(), timeout=15.0)
            if rows:
                _cases_db = rows
                break
        except Exception:
            if attempt < 2:
                await _aio.sleep(3 * (attempt + 1))

    # ── 2. Monitor de alertas ───────────────────────────────────────────────
    try:
        alerts = await _aio.wait_for(load_monitor_from_supabase(), timeout=8.0)
        if alerts:
            _monitor_history = alerts
    except Exception:
        pass

    # ── 3. T1 fix — Warmup do _kb_cache com normas mais usadas ────────────
    # Pré-carrega as categorias POA/NP mais frequentes para que as primeiras
    # validações pós-restart já tenham o conteúdo completo dos PDFs.
    TOP_WARMUP_CATEGORIES = [
        # POA — carnes e embutidos (mais validados)
        "embutidos", "salame", "presunto", "mortadela",
        # POA — laticínios
        "leite_uht", "iogurte", "queijo_mussarela",
        # POA — aves e pescado
        "aves_frango", "pescado_congelado",
        # Gerais ANVISA (aplicam a todos)
        "rotulagem_anvisa_rdc715", "rdc_429_rotulagem_nutri", "in_75_porcoes",
        "alergenos_rdc727", "cereais_rdc711",
        # Bebidas MAPA
        "bebidas_dec6871", "lei_8918_bebidas",
    ]
    warmed = 0
    for cat in TOP_WARMUP_CATEGORIES:
        if cat in _kb_cache:
            warmed += 1
            continue
        urls = MAPA_URLS.get(cat, [])
        if not urls:
            continue
        try:
            texts = await _aio.gather(*[
                _aio.wait_for(fetch_pdf_text(u, max_chars=3000), timeout=5.0)
                for u in urls[:2]  # max 2 URLs por categoria no warmup
            ], return_exceptions=True)
            result = "\n\n".join(t for t in texts if isinstance(t, str) and t)
            if result:
                _kb_cache[cat] = result
                warmed += 1
        except Exception:
            pass
        await _aio.sleep(0.3)  # throttle para não sobrecarregar no cold start

    # ── 4. Carrega kb_documents do Supabase (crawlados pelo GitHub Actions) ──
    # Esses documentos são indexados pelo crawler semanal e têm prioridade
    # sobre os fallbacks textuais para categorias cobertas pelo crawler.
    try:
        async with httpx.AsyncClient(timeout=10.0) as client:
            r = await client.get(
                f"{_SUPABASE_URL}/rest/v1/kb_documents",
                headers={"apikey": _SUPABASE_KEY,
                         "Authorization": f"Bearer {_SUPABASE_KEY}"},
                params={"select": "chave,conteudo,categoria,orgao",
                        "order": "atualizado_em.desc",
                        "limit": "500"}
            )
            if r.status_code == 200 and isinstance(r.json(), list):
                docs = r.json()
                for doc in docs:
                    chave = doc.get("chave")
                    conteudo = doc.get("conteudo", "")
                    if chave and conteudo and chave not in _kb_cache:
                        _kb_cache[chave] = conteudo
    except Exception:
        pass  # kb_documents é complemento — falha silenciosa

    _startup_ready = True
    _startup_ts = _dt_startup.datetime.now().isoformat()


# ═══════════════════════════════════════════════════════════════════════════════
# M1_7 — STRIPE: CHECKOUT + WEBHOOK + PORTAL DO CLIENTE
# Env vars: STRIPE_SECRET_KEY, STRIPE_WEBHOOK_SECRET
# Price IDs: STRIPE_PRICE_STARTER_MENSAL, STRIPE_PRICE_PRO_MENSAL,
#            STRIPE_PRICE_STARTER_ANUAL, STRIPE_PRICE_PRO_ANUAL
# ═══════════════════════════════════════════════════════════════════════════════

_STRIPE_KEY     = os.environ.get("STRIPE_SECRET_KEY", "")
_STRIPE_WHK_SEC = os.environ.get("STRIPE_WEBHOOK_SECRET", "")
_STRIPE_PRICES  = {
    "starter_mensal": os.environ.get("STRIPE_PRICE_STARTER_MENSAL", ""),
    "starter_anual":  os.environ.get("STRIPE_PRICE_STARTER_ANUAL", ""),
    "pro_mensal":     os.environ.get("STRIPE_PRICE_PRO_MENSAL", ""),
    "pro_anual":      os.environ.get("STRIPE_PRICE_PRO_ANUAL", ""),
}
_FRONTEND_URL = os.environ.get("FRONTEND_URL", "https://cosmic-llama-9576a5.netlify.app")


@app.post("/billing/checkout")
async def billing_checkout(request: Request):
    """
    Cria sessão de checkout Stripe.
    Body: { plano: "starter"|"pro", cobranca: "mensal"|"anual",
            email: str, user_id: str, cupom: str (opcional) }
    """
    if not _STRIPE_KEY:
        return JSONResponse({"error": "Stripe não configurado. Adicione STRIPE_SECRET_KEY no Render."},
                            status_code=503, headers={"Access-Control-Allow-Origin": "*"})
    try:
        body     = await request.json()
        plano    = body.get("plano", "starter")
        cobranca = body.get("cobranca", "mensal")
        email    = body.get("email", "")
        user_id  = body.get("user_id", "")
        cupom    = body.get("cupom", "").strip().upper()

        price_key = f"{plano}_{cobranca}"
        price_id  = _STRIPE_PRICES.get(price_key, "")
        if not price_id:
            return JSONResponse({"error": f"Price ID não configurado para {price_key}. "
                                          "Adicione STRIPE_PRICE_* no Render."},
                                status_code=400, headers={"Access-Control-Allow-Origin": "*"})

        payload: dict = {
            "mode": "subscription",
            "line_items": [{"price": price_id, "quantity": 1}],
            "success_url": f"{_FRONTEND_URL}/planos.html?checkout=sucesso&plano={plano}",
            "cancel_url":  f"{_FRONTEND_URL}/planos.html?checkout=cancelado",
            "metadata":    {"user_id": user_id, "plano": plano, "cobranca": cobranca},
        }
        if email:
            payload["customer_email"] = email
        if cupom:
            payload["discounts"] = [{"coupon": cupom}]

        async with httpx.AsyncClient(timeout=20.0) as client:
            r = await client.post(
                "https://api.stripe.com/v1/checkout/sessions",
                headers={"Authorization": f"Bearer {_STRIPE_KEY}",
                         "Content-Type": "application/x-www-form-urlencoded"},
                content="&".join(f"{k}={v}" for k, v in _flatten_stripe(payload).items())
            )
            data = r.json()

        if r.status_code != 200:
            return JSONResponse({"error": data.get("error", {}).get("message", "Erro Stripe")},
                                status_code=400, headers={"Access-Control-Allow-Origin": "*"})

        return JSONResponse({"url": data["url"], "session_id": data["id"]},
                            headers={"Access-Control-Allow-Origin": "*"})

    except Exception as e:
        return JSONResponse({"error": str(e)[:200]}, status_code=500,
                            headers={"Access-Control-Allow-Origin": "*"})


def _flatten_stripe(obj: dict, prefix: str = "") -> dict:
    """Converte dict aninhado para formato x-www-form-urlencoded do Stripe."""
    result = {}
    for k, v in obj.items():
        key = f"{prefix}[{k}]" if prefix else k
        if isinstance(v, dict):
            result.update(_flatten_stripe(v, key))
        elif isinstance(v, list):
            for i, item in enumerate(v):
                if isinstance(item, dict):
                    result.update(_flatten_stripe(item, f"{key}[{i}]"))
                else:
                    result[f"{key}[{i}]"] = str(item)
        elif isinstance(v, bool):
            result[key] = "true" if v else "false"
        elif v is not None:
            result[key] = str(v)
    return result


@app.post("/billing/webhook")
async def billing_webhook(request: Request):
    """
    Recebe eventos do Stripe e atualiza assinatura no Supabase.
    Configurar no Stripe Dashboard → Webhooks → endpoint URL + eventos:
      customer.subscription.created, .updated, .deleted
      checkout.session.completed
      invoice.payment_failed
    """
    payload = await request.body()
    sig     = request.headers.get("stripe-signature", "")

    # Verifica assinatura do webhook
    if _STRIPE_WHK_SEC:
        try:
            import hmac, hashlib, time as _time
            parts = {p.split("=")[0]: p.split("=")[1] for p in sig.split(",") if "=" in p}
            ts    = parts.get("t", "0")
            v1    = parts.get("v1", "")
            signed_payload = f"{ts}.{payload.decode()}"
            expected = hmac.new(_STRIPE_WHK_SEC.encode(),
                                signed_payload.encode(), hashlib.sha256).hexdigest()
            if not hmac.compare_digest(expected, v1):
                return JSONResponse({"error": "Assinatura inválida"}, status_code=400)
        except Exception:
            pass  # Em dev sem secret, continua

    try:
        event = json.loads(payload)
        etype = event.get("type", "")
        data  = event.get("data", {}).get("object", {})

        import datetime as _dt_wh

        if etype == "checkout.session.completed":
            user_id = data.get("metadata", {}).get("user_id", "")
            plano   = data.get("metadata", {}).get("plano", "starter")
            sub_id  = data.get("subscription", "")
            if user_id and _SUPABASE_ON:
                asyncio.ensure_future(_sb_upsert("assinaturas", {
                    "user_id":          user_id,
                    "plano":            plano,
                    "status":           "ativo",
                    "stripe_sub_id":    sub_id,
                    "stripe_cust_id":   data.get("customer", ""),
                    "updated_at":       _dt_wh.datetime.now().isoformat(),
                }))

        elif etype in ("customer.subscription.updated", "customer.subscription.deleted"):
            sub_id = data.get("id", "")
            status = "cancelado" if etype.endswith("deleted") else data.get("status", "ativo")
            if sub_id and _SUPABASE_ON:
                rows = await _sb_get("assinaturas", {"stripe_sub_id": sub_id}, limit=1)
                if rows:
                    asyncio.ensure_future(_sb_upsert("assinaturas", {
                        **rows[0],
                        "status":     status,
                        "updated_at": _dt_wh.datetime.now().isoformat(),
                    }))

        elif etype == "invoice.payment_failed":
            sub_id = data.get("subscription", "")
            if sub_id and _SUPABASE_ON:
                rows = await _sb_get("assinaturas", {"stripe_sub_id": sub_id}, limit=1)
                if rows:
                    asyncio.ensure_future(_sb_upsert("assinaturas", {
                        **rows[0],
                        "status":     "pagamento_falhou",
                        "updated_at": _dt_wh.datetime.now().isoformat(),
                    }))

        return JSONResponse({"received": True})

    except Exception as e:
        return JSONResponse({"error": str(e)[:200]}, status_code=400)


@app.post("/billing/portal")
async def billing_portal(request: Request):
    """
    Cria sessão do portal do cliente Stripe (gerenciar assinatura).
    Body: { user_id: str }
    """
    if not _STRIPE_KEY:
        return JSONResponse({"error": "Stripe não configurado."},
                            status_code=503, headers={"Access-Control-Allow-Origin": "*"})
    try:
        body    = await request.json()
        user_id = body.get("user_id", "")

        # Busca customer_id no Supabase
        cust_id = ""
        if _SUPABASE_ON and user_id:
            rows = await _sb_get("assinaturas", {"user_id": user_id}, limit=1)
            if rows:
                cust_id = rows[0].get("stripe_cust_id", "")

        if not cust_id:
            return JSONResponse({"error": "Assinatura não encontrada para este usuário."},
                                status_code=404, headers={"Access-Control-Allow-Origin": "*"})

        async with httpx.AsyncClient(timeout=15.0) as client:
            r = await client.post(
                "https://api.stripe.com/v1/billing_portal/sessions",
                headers={"Authorization": f"Bearer {_STRIPE_KEY}",
                         "Content-Type": "application/x-www-form-urlencoded"},
                content=f"customer={cust_id}&return_url={_FRONTEND_URL}/planos.html"
            )
            data = r.json()

        if r.status_code != 200:
            return JSONResponse({"error": data.get("error", {}).get("message", "Erro Stripe")},
                                status_code=400, headers={"Access-Control-Allow-Origin": "*"})

        return JSONResponse({"url": data["url"]},
                            headers={"Access-Control-Allow-Origin": "*"})

    except Exception as e:
        return JSONResponse({"error": str(e)[:200]}, status_code=500,
                            headers={"Access-Control-Allow-Origin": "*"})


# ═══════════════════════════════════════════════════════════════════════════════
# NP11 — TABELA NUTRICIONAL AVULSA (R$35/cálculo)
# Fluxo: ingredientes → preview gratuito → Stripe checkout → PDF download
# ═══════════════════════════════════════════════════════════════════════════════

_SP_TABELA_NUTRICIONAL = """Você é um nutricionista especializado em rotulagem alimentar brasileira.
Calcule a tabela nutricional completa conforme RDC 429/2020 + IN 75/2020 (ANVISA).

Dado os ingredientes com suas quantidades, retorne SOMENTE JSON válido:
{
  "produto": "Nome do produto",
  "porcao_g": 30,
  "porcao_descricao": "1 unidade (30g)",
  "porcoes_por_embalagem": 10,
  "energia_kcal": 150,
  "energia_kj": 628,
  "carboidratos_g": 20.5,
  "acucares_totais_g": 8.2,
  "acucares_adicionados_g": 5.0,
  "proteinas_g": 3.1,
  "gorduras_totais_g": 6.4,
  "gorduras_saturadas_g": 2.1,
  "gorduras_trans_g": 0.0,
  "gorduras_monoinsaturadas_g": 2.8,
  "gorduras_poliinsaturadas_g": 1.5,
  "fibra_alimentar_g": 1.2,
  "sodio_mg": 180,
  "vd_carboidratos_pct": 7,
  "vd_acucares_adicionados_pct": 10,
  "vd_proteinas_pct": 4,
  "vd_gorduras_totais_pct": 12,
  "vd_gorduras_saturadas_pct": 11,
  "vd_gorduras_trans_pct": null,
  "vd_fibra_pct": 5,
  "vd_sodio_pct": 8,
  "lupa_alto_sodio": false,
  "lupa_alto_gordura_saturada": false,
  "lupa_alto_acucar": false,
  "alergenos_contem": ["GLÚTEN", "LEITE"],
  "alergenos_pode_conter": [],
  "gluten": "CONTÉM GLÚTEN",
  "lactose": "CONTÉM LACTOSE",
  "observacoes": "Orientações técnicas sobre este produto",
  "legislacao_base": ["RDC 429/2020", "IN 75/2020"]
}

REGRAS CRÍTICAS:
- Use tabela TACO e IBGE como referência para composição dos ingredientes
- %VD baseado nos valores de referência da IN 75/2020
- Lupa: sódio ≥600mg/100g → lupa_alto_sodio=true | gordura saturada ≥6g/100g → lupa_alto_gordura_saturada=true | açúcar adicionado ≥15g/100g → lupa_alto_acucar=true
- gorduras_trans_g: NUNCA arredondar para 0 se >0.2g/porção
- Porção padrão: use IN 75/2020 para a categoria do produto
- %VD gorduras trans: sempre null (sem VD definido pela ANVISA)
- Calcule com precisão de 1 casa decimal para macros, inteiro para energia e sódio
"""

@app.post("/calcular-tabela")
async def calcular_tabela(request: Request):
    """
    NP11 — Calcula tabela nutricional a partir de ingredientes.
    Retorna preview gratuito. PDF requer pagamento via /tabela-checkout.
    Body: { produto, ingredientes_lista, porcao_g, peso_total_g, obs }
    """
    try:
        body = await request.json()
        produto         = body.get("produto", "")
        ingredientes    = body.get("ingredientes_lista", "")
        porcao_g        = body.get("porcao_g", "")
        peso_total_g    = body.get("peso_total_g", "")
        obs             = body.get("obs", "")

        if not ingredientes:
            return JSONResponse({"error": "Ingredientes são obrigatórios."},
                                status_code=400, headers={"Access-Control-Allow-Origin": "*"})

        user_msg = f"""Calcule a tabela nutricional para:
PRODUTO: {produto or "Produto alimentício"}
PORÇÃO: {porcao_g or "conforme categoria IN 75/2020"}g
PESO TOTAL DA EMBALAGEM: {peso_total_g or "não informado"}g
OBSERVAÇÕES: {obs or "nenhuma"}

INGREDIENTES (em ordem decrescente com quantidades):
{ingredientes}

Retorne SOMENTE o JSON conforme especificado."""

        async with httpx.AsyncClient(timeout=45.0) as client:
            r = await client.post(
                "https://api.anthropic.com/v1/messages",
                headers={"x-api-key": os.environ.get("ANTHROPIC_API_KEY", ""),
                         "anthropic-version": "2023-06-01",
                         "content-type": "application/json"},
                json={"model": ANTHROPIC_MODEL,
                      "max_tokens": 1500,
                      "system": _SP_TABELA_NUTRICIONAL,
                      "messages": [{"role": "user", "content": user_msg}]}
            )
            data = r.json()

        raw = data["content"][0]["text"].strip()
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"): raw = raw[4:]
        raw = raw.strip().rstrip("```")
        tabela = json.loads(raw)

        return JSONResponse({"ok": True, "tabela": tabela},
                            headers={"Access-Control-Allow-Origin": "*"})

    except json.JSONDecodeError:
        return JSONResponse({"error": "Erro ao processar resposta. Tente novamente."},
                            status_code=500, headers={"Access-Control-Allow-Origin": "*"})
    except Exception as e:
        return JSONResponse({"error": str(e)[:200]}, status_code=500,
                            headers={"Access-Control-Allow-Origin": "*"})


@app.post("/tabela-checkout")
async def tabela_checkout(request: Request):
    """
    NP11 — Cria sessão Stripe Checkout para pagamento da tabela nutricional avulsa.
    NP11 fix: persiste tabela_json no Supabase ANTES do checkout.
    Body: { tabela_json, produto, email }
    """
    if not _STRIPE_KEY:
        return JSONResponse({"error": "Stripe não configurado."},
                            status_code=503, headers={"Access-Control-Allow-Origin": "*"})
    try:
        body        = await request.json()
        tabela_json = body.get("tabela_json", "{}")
        produto     = body.get("produto", "Tabela Nutricional")
        email       = body.get("email", "")

        price_id = os.environ.get("STRIPE_PRICE_TABELA", "")
        if not price_id:
            return JSONResponse({"error": "Configure STRIPE_PRICE_TABELA no Render."},
                                status_code=503, headers={"Access-Control-Allow-Origin": "*"})

        import urllib.parse

        # NP11 fix — Persiste tabela no Supabase ANTES de criar a sessão Stripe
        # Assim o usuário pode baixar mesmo após fechar a aba
        tabela_key = f"tabela_{__import__('hashlib').md5((tabela_json + produto).encode()).hexdigest()[:16]}"
        if _SUPABASE_ON:
            try:
                await _sb_upsert("tabelas_checkout", {
                    "chave": tabela_key,
                    "tabela_json": tabela_json[:8000],
                    "produto": produto[:80],
                    "email": email[:120],
                    "status": "pendente",
                    "created_at": __import__("datetime").datetime.now().isoformat(),
                })
            except Exception:
                pass  # falha silenciosa — checkout prossegue

        payload = {
            "mode": "payment",
            "line_items[0][price]": price_id,
            "line_items[0][quantity]": "1",
            # NP11 fix: success_url usa tabela_key (curto e estável) em vez de JSON encodado
            "success_url": f"{_FRONTEND_URL}/tabela-nutricional.html?status=sucesso&tabela_key={tabela_key}",
            "cancel_url":  f"{_FRONTEND_URL}/tabela-nutricional.html?status=cancelado",
            "metadata[produto]": produto[:80],
            "metadata[tipo]": "tabela_avulsa",
            "metadata[tabela_key]": tabela_key,
        }
        if email:
            payload["customer_email"] = email

        async with httpx.AsyncClient(timeout=20.0) as client:
            r = await client.post(
                "https://api.stripe.com/v1/checkout/sessions",
                headers={"Authorization": f"Bearer {_STRIPE_KEY}",
                         "Content-Type": "application/x-www-form-urlencoded"},
                content="&".join(f"{k}={urllib.parse.quote(str(v))}" for k, v in payload.items())
            )
            data = r.json()

        if r.status_code != 200:
            return JSONResponse({"error": data.get("error", {}).get("message", "Erro Stripe")},
                                status_code=400, headers={"Access-Control-Allow-Origin": "*"})

        return JSONResponse({"url": data["url"], "session_id": data["id"], "tabela_key": tabela_key},
                            headers={"Access-Control-Allow-Origin": "*"})

    except Exception as e:
        return JSONResponse({"error": str(e)[:200]}, status_code=500,
                            headers={"Access-Control-Allow-Origin": "*"})


@app.get("/tabela-redownload/{tabela_key}")
async def tabela_redownload(tabela_key: str):
    """
    NP11 fix — Permite redownload da tabela mesmo após fechar a aba.
    Busca tabela_json no Supabase pela chave e retorna os dados para o frontend gerar o PDF.
    """
    if not _SUPABASE_ON:
        return JSONResponse({"error": "Supabase não configurado."},
                            status_code=503, headers={"Access-Control-Allow-Origin": "*"})
    try:
        rows = await _sb_get("tabelas_checkout", {"chave": tabela_key}, limit=1)
        if not rows:
            return JSONResponse({"error": "Tabela não encontrada. Link pode ter expirado."},
                                status_code=404, headers={"Access-Control-Allow-Origin": "*"})
        row = rows[0]
        return JSONResponse({
            "ok": True,
            "tabela_json": row.get("tabela_json", "{}"),
            "produto": row.get("produto", ""),
            "email": row.get("email", ""),
        }, headers={"Access-Control-Allow-Origin": "*"})
    except Exception as e:
        return JSONResponse({"error": str(e)[:200]}, status_code=500,
                            headers={"Access-Control-Allow-Origin": "*"})


@app.post("/tabela-pdf")
async def tabela_pdf(request: Request):
    """
    NP11 — Gera PDF da tabela nutricional após pagamento confirmado.
    Body: { tabela_json, produto, nome_rt, crm_rt }
    """
    try:
        body        = await request.json()
        tabela      = body.get("tabela_json") or {}
        if isinstance(tabela, str):
            tabela = json.loads(tabela)
        produto     = body.get("produto", tabela.get("produto", "Produto"))
        nome_rt     = body.get("nome_rt", "")
        crm_rt      = body.get("crm_rt", "")

        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib import colors
            from reportlab.lib.units import cm
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            import io as _io, datetime as _dt

            buf = _io.BytesIO()
            doc = SimpleDocTemplate(buf, pagesize=A4,
                                    topMargin=2*cm, bottomMargin=2*cm,
                                    leftMargin=2*cm, rightMargin=2*cm)
            styles = getSampleStyleSheet()
            verde  = colors.HexColor("#1A7A4A")
            ambar  = colors.HexColor("#D4860A")
            escuro = colors.HexColor("#060A0F")

            elems = []

            # Cabeçalho
            titulo_style = ParagraphStyle("titulo", parent=styles["Heading1"],
                                          fontSize=18, textColor=escuro, spaceAfter=4)
            sub_style    = ParagraphStyle("sub", parent=styles["Normal"],
                                          fontSize=10, textColor=verde, spaceAfter=16)
            body_style   = ParagraphStyle("body", parent=styles["Normal"],
                                          fontSize=9, textColor=escuro)
            small_style  = ParagraphStyle("small", parent=styles["Normal"],
                                          fontSize=8, textColor=colors.HexColor("#6b7280"))

            elems.append(Paragraph("Tabela Nutricional", titulo_style))
            elems.append(Paragraph(f"Produto: <b>{produto}</b> · Gerado por Inspect-IA", sub_style))

            # Linha amarela decorativa
            from reportlab.platypus import HRFlowable
            elems.append(HRFlowable(width="100%", thickness=2, color=ambar))
            elems.append(Spacer(1, 0.3*cm))

            # Info de porção
            porcao    = tabela.get("porcao_g", "—")
            porcao_d  = tabela.get("porcao_descricao", f"{porcao}g")
            porcoes   = tabela.get("porcoes_por_embalagem", "—")
            elems.append(Paragraph(f"<b>Porção:</b> {porcao_d} | <b>Porções por embalagem:</b> {porcoes}", body_style))
            elems.append(Spacer(1, 0.3*cm))

            def vd(val):
                return f"{val}%" if val is not None else "**"

            # Tabela nutricional
            dados = [
                ["Informação Nutricional", f"Por porção ({porcao}g)", "% VD*"],
                ["Valor energético", f"{tabela.get('energia_kcal','—')} kcal = {tabela.get('energia_kj','—')} kJ", "—"],
                ["Carboidratos", f"{tabela.get('carboidratos_g','—')}g", vd(tabela.get('vd_carboidratos_pct'))],
                ["    Açúcares totais", f"{tabela.get('acucares_totais_g','—')}g", "—"],
                ["    Açúcares adicionados", f"{tabela.get('acucares_adicionados_g','—')}g", vd(tabela.get('vd_acucares_adicionados_pct'))],
                ["Proteínas", f"{tabela.get('proteinas_g','—')}g", vd(tabela.get('vd_proteinas_pct'))],
                ["Gorduras totais", f"{tabela.get('gorduras_totais_g','—')}g", vd(tabela.get('vd_gorduras_totais_pct'))],
                ["    Gorduras saturadas", f"{tabela.get('gorduras_saturadas_g','—')}g", vd(tabela.get('vd_gorduras_saturadas_pct'))],
                ["    Gorduras trans", f"{tabela.get('gorduras_trans_g','—')}g", "**"],
                ["Fibra alimentar", f"{tabela.get('fibra_alimentar_g','—')}g", vd(tabela.get('vd_fibra_pct'))],
                ["Sódio", f"{tabela.get('sodio_mg','—')}mg", vd(tabela.get('vd_sodio_pct'))],
            ]

            col_w = [9*cm, 4.5*cm, 2.5*cm]
            t = Table(dados, colWidths=col_w)
            t.setStyle(TableStyle([
                ("BACKGROUND",  (0,0), (-1,0), verde),
                ("TEXTCOLOR",   (0,0), (-1,0), colors.white),
                ("FONTNAME",    (0,0), (-1,0), "Helvetica-Bold"),
                ("FONTSIZE",    (0,0), (-1,0), 9),
                ("ALIGN",       (1,0), (-1,-1), "CENTER"),
                ("FONTNAME",    (0,1), (-1,-1), "Helvetica"),
                ("FONTSIZE",    (0,1), (-1,-1), 8),
                ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, colors.HexColor("#f0fdf4")]),
                ("GRID",        (0,0), (-1,-1), 0.3, colors.HexColor("#d1d5db")),
                ("LEFTPADDING", (0,0), (-1,-1), 6),
                ("RIGHTPADDING",(0,0), (-1,-1), 6),
                ("TOPPADDING",  (0,0), (-1,-1), 4),
                ("BOTTOMPADDING",(0,0), (-1,-1), 4),
            ]))
            elems.append(t)
            elems.append(Spacer(1, 0.3*cm))

            # Lutas
            lutas = []
            if tabela.get("lupa_alto_sodio"):     lutas.append("⚠ ALTO EM SÓDIO")
            if tabela.get("lupa_alto_gordura_saturada"): lutas.append("⚠ ALTO EM GORDURAS SATURADAS")
            if tabela.get("lupa_alto_acucar"):    lutas.append("⚠ ALTO EM AÇÚCAR")
            if lutas:
                elems.append(Paragraph(f"<b>Lupas obrigatórias:</b> {' | '.join(lutas)}", body_style))
                elems.append(Spacer(1, 0.2*cm))

            # Alérgenos
            contem     = ", ".join(tabela.get("alergenos_contem", []))
            pode_conter= ", ".join(tabela.get("alergenos_pode_conter", []))
            if contem:
                elems.append(Paragraph(f"<b>CONTÉM:</b> {contem} E DERIVADOS.", body_style))
            if pode_conter:
                elems.append(Paragraph(f"<b>PODE CONTER:</b> {pode_conter}.", body_style))
            elems.append(Paragraph(tabela.get("gluten",""), body_style))
            elems.append(Spacer(1, 0.2*cm))

            # Rodapé técnico
            elems.append(Paragraph("* % Valores Diários com base em uma dieta de 2.000 kcal ou 8.400 kJ. Seus valores diários podem ser maiores ou menores dependendo de suas necessidades energéticas.", small_style))
            elems.append(Paragraph("** Gorduras trans não possuem Valor Diário de referência estabelecido.", small_style))
            elems.append(Spacer(1, 0.5*cm))
            elems.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#e5e7eb")))
            elems.append(Spacer(1, 0.3*cm))

            # Obs técnicas
            obs_txt = tabela.get("observacoes", "")
            if obs_txt:
                elems.append(Paragraph(f"<b>Observações técnicas:</b> {obs_txt}", small_style))
                elems.append(Spacer(1, 0.3*cm))

            # Legislação
            legs = ", ".join(tabela.get("legislacao_base", ["RDC 429/2020", "IN 75/2020"]))
            elems.append(Paragraph(f"<b>Legislação aplicada:</b> {legs}", small_style))
            elems.append(Spacer(1, 0.5*cm))

            # Assinatura RT
            if nome_rt or crm_rt:
                elems.append(HRFlowable(width="50%", thickness=0.5, color=escuro))
                elems.append(Paragraph(f"{nome_rt}", body_style))
                elems.append(Paragraph(f"{crm_rt}", body_style))
                elems.append(Paragraph("Responsável Técnico", small_style))
                elems.append(Spacer(1, 0.3*cm))

            # Gerado por
            elems.append(Paragraph(
                f"Documento gerado por Inspect-IA — inspect-ia.com.br | {_dt.datetime.now().strftime('%d/%m/%Y %H:%M')}",
                small_style))
            elems.append(Paragraph("⚠ Este documento é auxiliar. A responsabilidade técnica permanece com o RT habilitado.", small_style))

            doc.build(elems)
            pdf_bytes = buf.getvalue()

        except ImportError:
            return JSONResponse({"error": "reportlab não instalado."},
                                status_code=500, headers={"Access-Control-Allow-Origin": "*"})

        from fastapi.responses import Response as _Resp
        return _Resp(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f'attachment; filename="tabela_nutricional_{produto[:30].replace(" ","_")}.pdf"',
                "Access-Control-Allow-Origin": "*"
            }
        )

    except Exception as e:
        return JSONResponse({"error": str(e)[:200]}, status_code=500,
                            headers={"Access-Control-Allow-Origin": "*"})


@app.get("/debug/feedback-status")
async def feedback_status():
    """
    Diagnóstico do ciclo de feedback — quantos casos em memória vs Supabase.
    Útil para confirmar que o aprendizado está persistindo entre restarts.
    """
    supabase_total = 0
    supabase_com_feedback = 0
    if _SUPABASE_ON:
        try:
            rows_total = await _sb_get("validacoes", limit=1)
            # Conta total via header (Supabase retorna count no header com Prefer: count=exact)
            async with httpx.AsyncClient(timeout=8.0) as client:
                r = await client.get(
                    f"{_SUPABASE_URL}/rest/v1/validacoes",
                    headers={"apikey": _SUPABASE_KEY,
                             "Authorization": f"Bearer {_SUPABASE_KEY}",
                             "Prefer": "count=exact"},
                    params={"select": "case_id", "limit": "1"}
                )
                ct = r.headers.get("content-range", "0/0")
                supabase_total = int(ct.split("/")[-1]) if "/" in ct else 0
                # Com feedback
                r2 = await client.get(
                    f"{_SUPABASE_URL}/rest/v1/validacoes",
                    headers={"apikey": _SUPABASE_KEY,
                             "Authorization": f"Bearer {_SUPABASE_KEY}",
                             "Prefer": "count=exact"},
                    params={"select": "case_id", "feedback": "not.is.null", "limit": "1"}
                )
                ct2 = r2.headers.get("content-range", "0/0")
                supabase_com_feedback = int(ct2.split("/")[-1]) if "/" in ct2 else 0
        except Exception:
            pass

    em_memoria = len(_cases_db)
    com_feedback_mem = sum(1 for c in _cases_db if c.get("feedback"))

    return JSONResponse({
        "status": "ok",
        "memoria": {
            "total_casos": em_memoria,
            "com_feedback_rt": com_feedback_mem,
            "sem_feedback": em_memoria - com_feedback_mem,
        },
        "supabase": {
            "conectado": _SUPABASE_ON,
            "total_casos": supabase_total,
            "com_feedback_rt": supabase_com_feedback,
        },
        "saude": "ok" if supabase_total > 0 or not _SUPABASE_ON else "aviso: Supabase vazio",
        "nota": "Casos em memória devem ser iguais ou menores que Supabase (máx 500 em RAM)."
    }, headers={"Access-Control-Allow-Origin": "*"})


# ═══════════════════════════════════════════════════════════════════════════════
# VALIDAÇÃO EM LOTE — Task #9
# Processa N rótulos sequencialmente, emite progresso via SSE
# Retorna relatório consolidado ao final
# ═══════════════════════════════════════════════════════════════════════════════

@app.post("/validar-lote")
async def validar_lote(
    imagens: list[UploadFile] = File(...),
    obs_geral: str = Form(default=""),
    seg_tipo: str = Form(default=""),
    seg_np_categoria: str = Form(default=""),
    seg_estado: str = Form(default=""),
    seg_categoria: str = Form(default=""),
    user_id: str = Form(default=""),
):
    """
    Valida N rótulos em sequência.
    Emite eventos SSE:
      { tipo: 'inicio',    total: N }
      { tipo: 'progresso', idx: i, nome: str, status: 'processando' }
      { tipo: 'resultado', idx: i, nome: str, score: float, veredicto: str, resumo: str }
      { tipo: 'erro',      idx: i, nome: str, msg: str }
      { tipo: 'concluido', resultados: [...], resumo_geral: str }
    """
    if not ANTHROPIC_API_KEY:
        return JSONResponse({"error": "ANTHROPIC_API_KEY não configurada"}, status_code=400,
                            headers={"Access-Control-Allow-Origin": "*"})

    if len(imagens) < 1:
        return JSONResponse({"error": "Envie ao menos 1 imagem."}, status_code=400,
                            headers={"Access-Control-Allow-Origin": "*"})

    if len(imagens) > 20:
        return JSONResponse({"error": "Máximo de 20 rótulos por lote."}, status_code=400,
                            headers={"Access-Control-Allow-Origin": "*"})

    # Rate limiting por plano
    if user_id:
        uso = await _checar_e_incrementar_uso(user_id, quantidade=len(imagens))
        if not uso["ok"]:
            return JSONResponse({
                "error": uso.get("msg", "Limite de validações atingido."),
                "limite_atingido": True,
            }, status_code=429, headers={"Access-Control-Allow-Origin": "*"})

    # Lê todas as imagens antes de iniciar o stream
    imagens_data = []
    for img in imagens:
        raw = await img.read()
        mime = img.content_type or "image/jpeg"
        b64 = __import__("base64").b64encode(raw).decode()
        imagens_data.append({"nome": img.filename or f"rotulo_{len(imagens_data)+1}", "b64": b64, "mime": mime})

    async def gerar_lote():
        import json as _j, base64 as _b64, asyncio as _aio

        total = len(imagens_data)
        yield f"data: {_j.dumps({'tipo': 'inicio', 'total': total})}\n\n"

        resultados = []

        for idx, img_data in enumerate(imagens_data):
            nome = img_data["nome"]
            yield f"data: {_j.dumps({'tipo': 'progresso', 'idx': idx+1, 'total': total, 'nome': nome, 'status': 'processando'})}\n\n"

            try:
                # Gera relatório completo para este rótulo
                relatorio_completo = ""
                async with httpx.AsyncClient(timeout=120.0) as client:
                    payload = {
                        "model": ANTHROPIC_MODEL,
                        "max_tokens": 2000,
                        "system": SP_VALIDACAO.replace("{kb_section}", ""),
                        "messages": [{
                            "role": "user",
                            "content": [
                                {"type": "image", "source": {
                                    "type": "base64",
                                    "media_type": img_data["mime"],
                                    "data": img_data["b64"]
                                }},
                                {"type": "text", "text": (
                                    f"Valide este rótulo.\n"
                                    f"{('Tipo: ' + seg_tipo) if seg_tipo else ''}\n"
                                    f"{('Categoria: ' + seg_np_categoria) if seg_np_categoria else ''}\n"
                                    f"{('Obs: ' + obs_geral) if obs_geral else ''}"
                                ).strip()}
                            ]
                        }],
                        "stream": True,
                    }
                    headers_api = {
                        "x-api-key": ANTHROPIC_API_KEY,
                        "anthropic-version": "2023-06-01",
                        "content-type": "application/json",
                    }
                    async with client.stream("POST", "https://api.anthropic.com/v1/messages",
                                             json=payload, headers=headers_api) as resp:
                        if resp.status_code != 200:
                            raise Exception(f"API error {resp.status_code}")
                        async for line in resp.aiter_lines():
                            if not line.startswith("data: "): continue
                            raw = line[6:].strip()
                            if raw == "[DONE]": break
                            try:
                                ev = _j.loads(raw)
                                if ev.get("type") == "content_block_delta":
                                    delta = ev.get("delta", {})
                                    if delta.get("type") == "text_delta":
                                        relatorio_completo += delta.get("text", "")
                            except Exception:
                                continue

                # Extrai score e veredicto
                score = extrair_score(relatorio_completo)
                veredicto = extrair_veredicto(relatorio_completo)

                # Extrai resumo executivo (correções prioritárias)
                import re as _re
                resumo_match = _re.search(
                    r"CORREÇÕES PRIORITÁRIAS.*?\n(.+?)(?=\n###|\Z)",
                    relatorio_completo, _re.DOTALL | _re.IGNORECASE
                )
                resumo = resumo_match.group(1)[:400].strip() if resumo_match else ""

                # Conta não conformidades
                nao_conformes = len(_re.findall(r"NÃO CONFORME", relatorio_completo, _re.IGNORECASE))
                ressalvas = len(_re.findall(r"COM RESSALVAS", relatorio_completo, _re.IGNORECASE))

                resultado = {
                    "idx": idx + 1,
                    "nome": nome,
                    "score": score,
                    "veredicto": veredicto,
                    "nao_conformes": nao_conformes,
                    "ressalvas": ressalvas,
                    "resumo": resumo,
                    "relatorio_completo": relatorio_completo,
                }
                resultados.append(resultado)

                yield f"data: {_j.dumps({'tipo': 'resultado', **{k: v for k, v in resultado.items() if k != 'relatorio_completo'}})}\n\n"

                # Pausa breve entre rótulos para não sobrecarregar a API
                if idx < total - 1:
                    await _aio.sleep(0.5)

            except Exception as e:
                erro = {"idx": idx + 1, "nome": nome, "score": None, "veredicto": "ERRO",
                        "nao_conformes": 0, "ressalvas": 0, "resumo": "", "relatorio_completo": ""}
                resultados.append(erro)
                yield f"data: {_j.dumps({'tipo': 'erro', 'idx': idx+1, 'nome': nome, 'msg': str(e)[:100]})}\n\n"

        # Resumo geral do lote
        aprovados = sum(1 for r in resultados if "APROVADO" in (r.get("veredicto") or "") and "RESSALVAS" not in (r.get("veredicto") or ""))
        ressalvas_c = sum(1 for r in resultados if "RESSALVAS" in (r.get("veredicto") or ""))
        reprovados = sum(1 for r in resultados if "REPROVADO" in (r.get("veredicto") or ""))
        scores_validos = [r["score"] for r in resultados if r.get("score") is not None]
        media_score = round(sum(scores_validos) / len(scores_validos), 1) if scores_validos else None

        yield f"data: {_j.dumps({'tipo': 'concluido', 'total': total, 'aprovados': aprovados, 'ressalvas': ressalvas_c, 'reprovados': reprovados, 'media_score': media_score, 'resultados': [{k: v for k, v in r.items() if k != 'relatorio_completo'} for r in resultados]})}\n\n"

    return StreamingResponse(
        gerar_lote(),
        media_type="text/event-stream",
        headers={
            "Access-Control-Allow-Origin": "*",
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
        }
    )


# ═══════════════════════════════════════════════════════════════════════════════
# COMPARATIVO DE VERSÕES — Task #12
# ═══════════════════════════════════════════════════════════════════════════════

@app.get("/comparativo/{case_id}")
async def get_comparativo(case_id: str):
    """
    Retorna comparativo entre a validação atual (case_id) e a validação anterior
    do mesmo produto. Usado para mostrar o que melhorou após correções.
    """
    # Encontra a validação atual
    atual = next((c for c in _cases_db if c.get("case_id") == case_id), None)
    if not atual:
        # Tenta Supabase
        if _SUPABASE_ON:
            try:
                rows = await _sb_get("validacoes", limit=1)
                async with httpx.AsyncClient(timeout=8.0) as client:
                    r = await client.get(
                        f"{_SUPABASE_URL}/rest/v1/validacoes",
                        headers={"apikey": _SUPABASE_KEY,
                                 "Authorization": f"Bearer {_SUPABASE_KEY}"},
                        params={"select": "*", "case_id": f"eq.{case_id}", "limit": "1"}
                    )
                    rows = r.json() if r.status_code == 200 and isinstance(r.json(), list) else []
                    atual = rows[0] if rows else None
            except Exception:
                pass
        if not atual:
            return JSONResponse({"error": "Validação não encontrada."},
                                status_code=404, headers={"Access-Control-Allow-Origin": "*"})

    produto_atual = (atual.get("produto") or "").lower().strip()

    # Busca validações anteriores do mesmo produto (exclui a atual)
    candidatos = [
        c for c in _cases_db
        if c.get("case_id") != case_id
        and c.get("produto")
        and _similaridade_produto(c.get("produto", ""), produto_atual) >= 0.6
    ]

    if not candidatos:
        return JSONResponse({
            "tem_comparativo": False,
            "mensagem": "Primeira validação deste produto — sem versão anterior para comparar."
        }, headers={"Access-Control-Allow-Origin": "*"})

    # Pega a mais recente
    candidatos.sort(key=lambda c: c.get("timestamp") or "", reverse=True)
    anterior = candidatos[0]

    # Extrai status dos campos das duas versões
    rel_atual    = atual.get("relatorio_completo", "")
    rel_anterior = anterior.get("relatorio_completo", "")

    campos_atual    = extrair_campos_status(rel_atual)    if rel_atual    else {}
    campos_anterior = extrair_campos_status(rel_anterior) if rel_anterior else {}

    # Gera o diff
    diff = []
    for c in range(1, 15):
        status_ant = campos_anterior.get(c, "NAO_DETECTADO")
        status_now = campos_atual.get(c, "NAO_DETECTADO")
        nome = CAMPOS_NOME.get(c, f"Campo {c}")

        if status_ant == "NAO_DETECTADO" and status_now == "NAO_DETECTADO":
            continue  # campo não detectado em nenhuma versão

        tendencia = "igual"
        if status_ant != status_now:
            # Hierarquia: CONFORME > COM_RESSALVAS > NAO_CONFORME
            ordem = {"CONFORME": 3, "COM_RESSALVAS": 2, "NAO_CONFORME": 1, "AUSENTE": 1, "NAO_DETECTADO": 0}
            if ordem.get(status_now, 0) > ordem.get(status_ant, 0):
                tendencia = "melhora"
            else:
                tendencia = "piora"

        diff.append({
            "campo":      c,
            "nome":       nome,
            "anterior":   status_ant,
            "atual":      status_now,
            "tendencia":  tendencia,
        })

    melhorias = [d for d in diff if d["tendencia"] == "melhora"]
    piorou    = [d for d in diff if d["tendencia"] == "piora"]
    iguais    = [d for d in diff if d["tendencia"] == "igual"
                 and d["atual"] not in ("NAO_DETECTADO",)]

    score_anterior = anterior.get("score_agente")
    score_atual    = atual.get("score_agente")
    delta_score    = round(score_atual - score_anterior, 1) if score_atual and score_anterior else None

    return JSONResponse({
        "tem_comparativo": True,
        "produto":         atual.get("produto"),
        "score_anterior":  score_anterior,
        "score_atual":     score_atual,
        "delta_score":     delta_score,
        "veredicto_anterior": anterior.get("veredicto"),
        "veredicto_atual":    atual.get("veredicto"),
        "timestamp_anterior": anterior.get("timestamp") or anterior.get("created_at"),
        "melhorias":    melhorias,
        "piorou":       piorou,
        "iguais_conformes": [d for d in iguais if d["atual"] == "CONFORME"],
        "iguais_erros":     [d for d in iguais if d["atual"] in ("NAO_CONFORME","AUSENTE")],
        "diff_completo": diff,
    }, headers={"Access-Control-Allow-Origin": "*"})


def _similaridade_produto(a: str, b: str) -> float:
    """Similaridade simples entre nomes de produto (Jaccard de palavras)."""
    a_words = set((a or "").lower().split())
    b_words = set((b or "").lower().split())
    if not a_words or not b_words:
        return 0.0
    inter = a_words & b_words
    union = a_words | b_words
    return len(inter) / len(union)


# ═══════════════════════════════════════════════════════════════════════════════
# HISTÓRICO DE VALIDAÇÕES — Task #11
# ═══════════════════════════════════════════════════════════════════════════════

@app.get("/historico")
async def get_historico(
    user_id: str = "",
    busca: str = "",
    veredicto: str = "",
    categoria: str = "",
    score_min: str = "",
    score_max: str = "",
    limite: int = 50,
    offset: int = 0,
):
    """
    Retorna histórico de validações com busca e filtros.
    Parâmetros:
      user_id   — filtra pelo usuário (opcional)
      busca     — texto livre sobre produto
      veredicto — APROVADO | APROVADO COM RESSALVAS | REPROVADO
      categoria — categoria do produto
      score_min / score_max — faixa de score
      limite    — paginação (max 100)
      offset    — paginação
    """
    try:
        limite = min(int(limite), 100)

        # Busca na memória local primeiro (inclui casos recentes ainda não persistidos)
        resultados = list(_cases_db)

        # Aplica filtros na memória
        if busca:
            b = busca.lower()
            resultados = [c for c in resultados
                         if b in (c.get("produto") or "").lower()
                         or b in (c.get("categoria") or "").lower()]
        if veredicto:
            resultados = [c for c in resultados
                         if veredicto.upper() in (c.get("veredicto") or "").upper()]
        if categoria:
            resultados = [c for c in resultados
                         if categoria.lower() in (c.get("categoria") or "").lower()
                         or categoria.lower() in (c.get("caminho_np") or "").lower()]
        if score_min:
            try:
                sm = float(score_min)
                resultados = [c for c in resultados if (c.get("score_agente") or 0) >= sm]
            except ValueError:
                pass
        if score_max:
            try:
                sx = float(score_max)
                resultados = [c for c in resultados if (c.get("score_agente") or 14) <= sx]
            except ValueError:
                pass

        # Ordena por timestamp desc
        resultados.sort(key=lambda c: c.get("timestamp") or "", reverse=True)
        total = len(resultados)
        pagina = resultados[offset:offset + limite]

        # Se Supabase disponível e poucos resultados em memória, busca lá também
        supabase_extra = []
        if _SUPABASE_ON and total < limite:
            try:
                params = {
                    "select": "case_id,produto,categoria,caminho_np,feedback,score_agente,score_real,erros_encontrados,ressalvas_auto,timestamp,created_at",
                    "order": "created_at.desc",
                    "limit": str(limite),
                    "offset": str(offset),
                }
                if busca:
                    params["produto"] = f"ilike.%{busca}%"
                async with httpx.AsyncClient(timeout=8.0) as client:
                    r = await client.get(
                        f"{_SUPABASE_URL}/rest/v1/validacoes",
                        headers={"apikey": _SUPABASE_KEY,
                                 "Authorization": f"Bearer {_SUPABASE_KEY}"},
                        params=params
                    )
                    if r.status_code == 200:
                        sb_rows = r.json() if isinstance(r.json(), list) else []
                        # Mescla sem duplicatas (case_id)
                        mem_ids = {c.get("case_id") for c in pagina if c.get("case_id")}
                        supabase_extra = [row for row in sb_rows
                                         if row.get("case_id") not in mem_ids]
            except Exception:
                pass

        merged = pagina + supabase_extra
        merged.sort(key=lambda c: c.get("created_at") or c.get("timestamp") or "", reverse=True)

        # Formata para o frontend
        items = []
        for c in merged[:limite]:
            score = c.get("score_agente") or c.get("score_real")
            veredicto_c = c.get("veredicto") or (
                "APROVADO" if score and score >= 13 else
                "APROVADO COM RESSALVAS" if score and score >= 8 else
                "REPROVADO" if score is not None else None
            )
            items.append({
                "case_id":    c.get("case_id", ""),
                "produto":    c.get("produto", "Produto não identificado"),
                "categoria":  c.get("categoria") or c.get("caminho_np") or "",
                "score":      score,
                "veredicto":  veredicto_c,
                "feedback_rt": c.get("feedback"),
                "erros":      c.get("erros_encontrados") or c.get("erros_auto") or "",
                "ressalvas":  c.get("ressalvas_auto") or "",
                "timestamp":  c.get("created_at") or c.get("timestamp") or "",
                "com_feedback": bool(c.get("feedback")),
            })

        return JSONResponse({
            "ok": True,
            "total": total,
            "offset": offset,
            "limite": limite,
            "items": items,
        }, headers={"Access-Control-Allow-Origin": "*"})

    except Exception as e:
        return JSONResponse({"error": str(e)[:200]}, status_code=500,
                            headers={"Access-Control-Allow-Origin": "*"})


@app.post("/kb/crawl-full")
async def kb_crawl_full(request: Request):
    """KB Crawler v5 — fontes corretas rodando no Render."""
    if not _SUPABASE_ON:
        return JSONResponse({"ok": False, "msg": "Supabase não configurado"},
                            headers={"Access-Control-Allow-Origin": "*"})
    asyncio.ensure_future(_crawl_v5_background())
    return JSONResponse({"ok": True, "msg": "Crawl v5 iniciado. Verifique /kb/crawl-status em 5 minutos."},
                        headers={"Access-Control-Allow-Origin": "*"})


_kb_crawl_status = {"rodando": False, "ultimo_resultado": None}


@app.get("/kb/crawl-status")
async def kb_crawl_status():
    return JSONResponse(_kb_crawl_status, headers={"Access-Control-Allow-Origin": "*"})


_CRAWL_V5_URLS = [
    # ── ANVISA via bvsms.saude.gov.br ─────────────────────────────────────────
    # PDFs com texto selecionável e HTMLs diretamente acessíveis — confirmados funcionando.
    # anvisalegis individual foi descartado (retorna 191c JS shell via httpx).
    # Index do anvisalegis mantido pois é página HTML estática (não usa AJAX).
    ("https://anvisalegis.datalegis.net/action/ActionDatalegis.php?acao=abrirTematica&cod_modulo=135&cod_menu=1686",
     "ANVISA","rotulagem","anvisalegis_rotulagem_index"),
    # ── Rotulagem obrigatória — normas base ───────────────────────────────────
    ("https://bvsms.saude.gov.br/bvs/saudelegis/anvisa/2020/RDC_429_2020_.pdf",
     "ANVISA","rotulagem","rdc_429_2020"),
    ("https://bvsms.saude.gov.br/bvs/saudelegis/anvisa/2020/IN_75_2020_.pdf",
     "ANVISA","rotulagem","in_75_2020"),
    ("https://bvsms.saude.gov.br/bvs/saudelegis/anvisa/2022/RDC_727_2022_.pdf",
     "ANVISA","rotulagem","rdc_727_2022"),
    ("https://bvsms.saude.gov.br/bvs/saudelegis/anvisa/2022/RDC_715_2022_.pdf",
     "ANVISA","rotulagem","rdc_715_2022_lactose"),
    # rdc_259_2002 foi revogada pela rdc_727_2022 (já seeded com conteúdo completo)
    # ── Suplementos (chave padronizada com seed SQL) ─────────────────────────
    ("https://anvisalegis.datalegis.net/action/ActionDatalegis.php?acao=abrirTextoAto&link=S&tipo=RDC&numeroAto=00000243&seqAto=000&valorAno=2018&orgao=RDC/DC/ANVISA/MS&cod_modulo=310&cod_menu=8542",
     "ANVISA","suplementos","rdc_243_2018_suplementos"),
    # ── Categorias de alimentos — versões vigentes ────────────────────────────
    # Normas 2005 AINDA VIGENTES (não revogadas em 2022):
    ("https://anvisalegis.datalegis.net/action/ActionDatalegis.php?acao=abrirTextoAto&link=S&tipo=RDC&numeroAto=00000268&seqAto=000&valorAno=2005&orgao=RDC/DC/ANVISA/MS&cod_modulo=310&cod_menu=9882",
     "ANVISA","alimentos","rdc_268_2005_proteina_vegetal"),
    ("https://bvsms.saude.gov.br/bvs/saudelegis/anvisa/2005/rdc0263_22_09_2005.html",
     "ANVISA","alimentos","rdc_263_2005_cereais_farinhas"),
    ("https://bvsms.saude.gov.br/bvs/saudelegis/anvisa/2005/rdc0272_22_09_2005.html",
     "ANVISA","alimentos","rdc_272_2005_fermentos_leveduras"),
    ("https://bvsms.saude.gov.br/bvs/saudelegis/anvisa/2005/rdc0278_22_09_2005.html",
     "ANVISA","alimentos","rdc_278_2005_vinagre"),
    ("https://bvsms.saude.gov.br/bvs/saudelegis/anvisa/2005/res0274_22_09_2005.html",
     "ANVISA","bebidas","rdc_274_2005_bebidas_nao_alcoolicas"),
    # Normas 2005 REVOGADAS → substituídas pelas 2022 (já incluídas acima):
    # rdc_264/266/273/276/277/2005 → rdc_723/713/719/716/2022 (chaves no seed)
    # rdc_259/2002 → rdc_727/2022 (chave no seed)
    # ── Aditivos aromas e contaminantes (normas 2022 vigentes) ───────────────
    ("https://anvisalegis.datalegis.net/action/ActionDatalegis.php?acao=abrirTextoAto&link=S&tipo=RDC&numeroAto=00000725&seqAto=000&valorAno=2022&orgao=RDC/DC/ANVISA/MS&cod_modulo=310&cod_menu=8542",
     "ANVISA","aditivos","rdc_725_2022_aromas"),
    ("https://anvisalegis.datalegis.net/action/ActionDatalegis.php?acao=abrirTextoAto&tipo=RDC&numeroAto=00000722&seqAto=002&valorAno=2022&orgao=RDC/DC/ANVISA/MS&cod_menu=9434&cod_modulo=310",
     "ANVISA","alimentos","rdc_722_2022_contaminantes"),
    # ── Normas 2022 que substituem as de 2005 revogadas ──────────────────────
    # RDC 723/2022 = substitui RDC 264/2005 (chocolate) + RDC 266/2005 (sorvete)
    ("https://anvisalegis.datalegis.net/action/ActionDatalegis.php?acao=abrirTextoAto&tipo=RDC&numeroAto=00000723&seqAto=002&valorAno=2022&orgao=RDC/DC/ANVISA/MS&cod_menu=1696&cod_modulo=134",
     "ANVISA","alimentos","rdc_723_2022_acucar_chocolate"),
    # RDC 713/2022 = substitui RDC 266/2005 (sorvetes/gelados)
    ("https://anvisalegis.datalegis.net/action/ActionDatalegis.php?acao=abrirTextoAto&tipo=RDC&numeroAto=00000713&seqAto=002&valorAno=2022&orgao=RDC/DC/ANVISA/MS&cod_menu=1696&cod_modulo=134",
     "ANVISA","alimentos","rdc_713_2022_gelados_sorvetes"),
    # RDC 719/2022 = substitui RDC 273/2005 (misturas/prontos)
    ("https://anvisalegis.datalegis.net/action/ActionDatalegis.php?acao=abrirTextoAto&link=S&tipo=RDC&numeroAto=00000719&seqAto=002&valorAno=2022&orgao=RDC/DC/ANVISA/MS&cod_modulo=310&cod_menu=8542",
     "ANVISA","alimentos","rdc_719_2022_misturas_prontos"),
    # RDC 716/2022 = substitui RDC 276/2005 (condimentos) + RDC 277/2005 (café/chá)
    ("https://anvisalegis.datalegis.net/action/ActionDatalegis.php?acao=abrirTextoAto&link=S&tipo=RDC&numeroAto=00000716&seqAto=002&valorAno=2022&orgao=RDC/DC/ANVISA/MS&cod_modulo=310&cod_menu=8542",
     "ANVISA","alimentos","rdc_716_2022_condimentos_cafe"),
    # RDC 267/2003 = alegações funcionais — verificada: AINDA VIGENTE (não revogada)
    ("https://anvisalegis.datalegis.net/action/ActionDatalegis.php?acao=abrirTextoAto&link=S&tipo=RDC&numeroAto=00000267&seqAto=000&valorAno=2003&orgao=RDC/DC/ANVISA/MS&cod_modulo=310&cod_menu=8542",
     "ANVISA","funcional","rdc_267_2003_funcional"),
    # ── MAPA POA — página de legislação HTML ─────────────────────────────────
    ("https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao",
     "MAPA","poa","mapa_poa_legislacao_index"),
    # ── MAPA Vegetal — RTIQs HTML ─────────────────────────────────────────────
    ("https://www.gov.br/agricultura/pt-br/assuntos/defesa-agropecuaria/suasa/regulamentos-tecnicos-de-identidade-e-qualidade-de-produtos-de-origem-vegetal",
     "MAPA","qualidade_vegetal","mapa_rtiq_vegetal_index"),
    # ── MAPA SISBI ────────────────────────────────────────────────────────────
    ("https://www.gov.br/agricultura/pt-br/assuntos/defesa-agropecuaria/suasa/sisbi-1",
     "MAPA","poa","mapa_sisbi_poa"),
    # ── MAPA bebidas e vegetal ────────────────────────────────────────────────
    ("https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-vegetal/legislacao-programas-nacionais-e-seguranca-dos-alimentos-1/legislacao/bebidas",
     "MAPA","bebidas","mapa_bebidas_legislacao"),
    ("https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-vegetal/legislacao-programas-nacionais-e-seguranca-dos-alimentos-1/legislacao/legislacaoPOV",
     "MAPA","qualidade_vegetal","mapa_pov_legislacao"),
    # ── DOU/in.gov.br (funciona no Render) ────────────────────────────────────
    ("https://www.in.gov.br/en/web/dou/-/instrucao-normativa-n-36-de-20-de-setembro-de-2018-42584174",
     "MAPA","bebidas","in_36_2018"),
    ("https://www.in.gov.br/en/web/dou/-/instrucao-normativa-n-49-de-23-de-outubro-de-2019-223577337",
     "MAPA","qualidade_vegetal","in_49_2019"),
    ("https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-vegetal/legislacao-programas-nacionais-e-seguranca-dos-alimentos-1/legislacao/bebidas/IN412019Kombucha.pdf",
     "MAPA","bebidas","in_41_2019_kombucha"),
    ("https://www.in.gov.br/en/web/dou/-/portaria-mapa-n-521-de-1-de-dezembro-de-2022-447310581",
     "MAPA","qualidade_vegetal","portaria_521_2022"),
    ("https://www.in.gov.br/en/web/dou/-/portaria-mapa-n-586-de-16-de-maio-de-2023-486234511",
     "MAPA","bebidas","portaria_586_2023"),
    ("https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-vegetal/legislacao-programas-nacionais-e-seguranca-dos-alimentos-1/legislacao/bebidas/IN652019Cerveja.pdf",
     "MAPA","bebidas","in_65_2019_cerveja"),
    ("https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN172020SisbiRotulagemPOA.pdf",
     "MAPA","poa","in_17_2020_sisbi"),
    ("https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN92019RotulagemPOA.pdf",
     "MAPA","poa","in_9_2019_rotulagem_poa"),
    ("https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-animal/legislacao/IN312020SisbiPOA.pdf",
     "MAPA","poa","in_31_2020_sisbi_poa"),
    ("https://www.gov.br/agricultura/pt-br/assuntos/inspecao/produtos-vegetal/legislacao-programas-nacionais-e-seguranca-dos-alimentos-1/legislacao/bebidas/IN602019Bebidas.pdf",
     "MAPA","bebidas","in_60_2019_bebidas"),
    # ── CONAR PDF (funciona) ──────────────────────────────────────────────────
    ("https://www.conar.org.br/pdf/codigo-conar-2021.pdf",
     "CONAR","publicidade","codigo_conar_2021"),
    # ── SIE estaduais HTML ────────────────────────────────────────────────────
    ("https://www.adagri.ce.gov.br/legislacao-sie/",
     "ADAGRI","sie_ce","adagri_ce_sie"),
    ("https://goias.gov.br/agricultura/legislacoes-do-sim-servico-de-inspecao-municipal/",
     "AGRODEFESA","sie_go","agrodefesa_go"),
    ("https://ima.mg.gov.br/agroindustria/produtos-de-origem-animal",
     "IMA","sie_mg","ima_mg_poa"),
]


async def _crawl_v5_background():
    """Crawl v5: HTML-first, sem PDFs escaneados, delay respeitoso."""
    global _kb_crawl_status
    _kb_crawl_status["rodando"] = True
    import re as _re

    salvos = 0; mantidos = 0; falhos = 0; detalhes = {}

    async with httpx.AsyncClient(timeout=20.0, follow_redirects=True, headers={
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) Chrome/124.0 Safari/537.36",
        "Accept": "text/html,application/xhtml+xml,*/*;q=0.8",
        "Accept-Language": "pt-BR,pt;q=0.9",
    }) as client:
        for url, orgao, categoria, chave in _CRAWL_V5_URLS:
            try:
                r = await asyncio.wait_for(client.get(url), timeout=18.0)
                if r.status_code != 200:
                    falhos += 1; detalhes[chave] = f"HTTP {r.status_code}"
                    await asyncio.sleep(0.5); continue

                ctype = r.headers.get("content-type", "")
                if "pdf" in ctype or url.lower().endswith(".pdf"):
                    try:
                        from pypdf import PdfReader
                        reader = PdfReader(__import__("io").BytesIO(r.content))
                        texto = "".join(pg.extract_text() or "" for pg in reader.pages)[:7000].strip()
                    except Exception:
                        texto = ""
                else:
                    raw = r.text
                    raw = _re.sub(r"<script[^>]*>.*?</script>", " ", raw, flags=_re.DOTALL|_re.I)
                    raw = _re.sub(r"<style[^>]*>.*?</style>", " ", raw, flags=_re.DOTALL|_re.I)
                    raw = _re.sub(r"<[^>]+>", " ", raw)
                    texto = _re.sub(r"\s{2,}", " ", raw).strip()[:7000]

                if len(texto) < 80:
                    falhos += 1; detalhes[chave] = "vazio"
                    await asyncio.sleep(0.5); continue

                rows = await _sb_get("kb_documents", {"chave": chave}, limit=1)
                existente = rows[0].get("tamanho_chars", 0) if rows else 0
                if existente >= len(texto):
                    mantidos += 1; detalhes[chave] = f"mantido ({existente}c)"
                    await asyncio.sleep(0.3); continue

                ok = await _sb_upsert("kb_documents", {
                    "chave": chave, "titulo": chave.replace("_"," ").title(),
                    "fonte": url, "orgao": orgao, "categoria": categoria,
                    "conteudo": texto, "tamanho_chars": len(texto),
                    "atualizado_em": __import__("datetime").datetime.now().isoformat(),
                }, conflict_col="chave")
                if ok:
                    salvos += 1; detalhes[chave] = f"salvo ({len(texto)}c)"; _kb_cache[chave] = texto
                else:
                    falhos += 1; detalhes[chave] = "erro_supabase"

            except asyncio.TimeoutError:
                falhos += 1; detalhes[chave] = "timeout"
            except Exception as e:
                falhos += 1; detalhes[chave] = f"erro:{str(e)[:40]}"
            await asyncio.sleep(0.6)

    _kb_crawl_status = {
        "rodando": False,
        "ultimo_resultado": {
            "salvos": salvos, "mantidos": mantidos, "falhos": falhos,
            "total_urls": len(_CRAWL_V5_URLS),
            "timestamp": __import__("datetime").datetime.now().isoformat(),
            "detalhes": detalhes,
        }
    }


async def discover_anvisa_kb():
    """
    T2 fix — Tenta descobrir e indexar normas ANVISA via padrão direto de URL.
    antigo.anvisa.gov.br/documents/33916/392655/ contém todas as RDCs de alimentos.
    Testa URLs de RDCs conhecidas que ainda não estão no cache.
    Execute mensalmente ou sempre que suspeitar de normas novas.
    """
    # RDCs de alimentos conhecidas no antigo.anvisa.gov.br — padrão direto de PDF
    ANVISA_DIRECT_URLS = {
        # Rotulagem geral (base)
        "rdc_259_rotulagem_geral":   "https://antigo.anvisa.gov.br/documents/33916/392655/RDC+259+2002.pdf",
        "rdc_360_info_nutri":        "https://antigo.anvisa.gov.br/documents/33916/392655/RDC360_2003.pdf",
        # Categorias de alimentos
        "rdc_263_cereais":           "https://antigo.anvisa.gov.br/documents/33916/392655/RDC+263_2005.pdf",
        "rdc_265_amido":             "https://antigo.anvisa.gov.br/documents/33916/392655/RDC+265_2005.pdf",
        "rdc_267_cogumelos":         "https://antigo.anvisa.gov.br/documents/33916/392655/RDC+267_2005.pdf",
        "rdc_269_proteinas":         "https://antigo.anvisa.gov.br/documents/33916/392655/RDC+269_2005.pdf",
        "rdc_271_acucar":            "https://antigo.anvisa.gov.br/documents/33916/392655/RDC+271_2005.pdf",
        "rdc_272_vegetais":          "https://antigo.anvisa.gov.br/documents/33916/392655/RDC+272_2005.pdf",
        "rdc_273_alim_enriquecidos": "https://antigo.anvisa.gov.br/documents/33916/392655/RDC+273_2005.pdf",
        "rdc_274_bebidas_nao_alc":   "https://antigo.anvisa.gov.br/documents/33916/392655/RDC+274_2005.pdf",
        "rdc_275_boas_praticas":     "https://antigo.anvisa.gov.br/documents/33916/392655/RDC+275_2005.pdf",
        "rdc_278_vinagre":           "https://antigo.anvisa.gov.br/documents/33916/392655/RDC+278_2005.pdf",
        # Aditivos e aromas
        "rdc_2_aditivos":            "https://antigo.anvisa.gov.br/documents/33916/392655/RDC+02+2007+Aromas.pdf",
        "resolucao_386_aditivos":    "https://antigo.anvisa.gov.br/documents/33916/392655/RESOLUCAO_386_1999_.pdf",
        # Contaminantes e limites
        "rdc_42_contaminantes":      "https://antigo.anvisa.gov.br/documents/33916/392655/RDC42_2013_contaminantes.pdf",
        "rdc_722_moldes_2022":       "https://antigo.anvisa.gov.br/documents/33916/392655/RDC_722_2022.pdf",
    }

    results = {}
    discovered = 0
    already_cached = 0
    failed = 0

    async with httpx.AsyncClient(timeout=10.0) as client:
        for key, url in ANVISA_DIRECT_URLS.items():
            if key in _kb_cache and _kb_cache[key]:
                already_cached += 1
                results[key] = f"cache ({len(_kb_cache[key])} chars)"
                continue
            try:
                text = await asyncio.wait_for(fetch_pdf_text(url, max_chars=4000), timeout=8.0)
                if text and len(text) > 100:
                    _kb_cache[key] = text
                    # Adiciona à MAPA_URLS para uso futuro
                    if key not in MAPA_URLS:
                        MAPA_URLS[key] = [url]
                    discovered += 1
                    results[key] = f"descoberto ({len(text)} chars)"
                else:
                    failed += 1
                    results[key] = "vazio/inacessivel"
            except Exception as e:
                failed += 1
                results[key] = f"erro: {str(e)[:50]}"
            await asyncio.sleep(0.5)  # throttle

    return JSONResponse({
        "ok": True,
        "descobertos": discovered,
        "ja_em_cache": already_cached,
        "falhou": failed,
        "total_testados": len(ANVISA_DIRECT_URLS),
        "nota": "AnvisaLegis tem falha de banco Oracle (ORA-28001) — usando URLs diretas do antigo.anvisa.gov.br",
        "detalhes": results,
    }, headers={"Access-Control-Allow-Origin": "*"})


@app.get("/sentry/test")
def sentry_test():
    """
    Endpoint de teste do Sentry — gera um erro intencional para confirmar que
    o monitoramento está funcionando. Usar apenas para validação inicial.
    """
    try:
        import sentry_sdk
        sentry_sdk.capture_message("Sentry test — ValidaRótulo IA funcionando", level="info")
        return JSONResponse({"ok": True, "msg": "Evento enviado ao Sentry. Verifique o dashboard."})
    except Exception:
        return JSONResponse({"ok": False, "msg": "Sentry não configurado — adicione SENTRY_DSN no Render."})


@app.get("/health/readiness")
def health_readiness():
    """
    T1 fix — Endpoint de readiness: mostra o estado real do sistema pós-startup.
    Útil para diagnosticar se o warmup foi concluído e quantas normas estão em cache.
    """
    return JSONResponse({
        "startup_completo": _startup_ready,
        "startup_ts": _startup_ts,
        "casos_em_memoria": len(_cases_db),
        "kb_cache_entradas": len(_kb_cache),
        "kb_cache_keys": list(_kb_cache.keys())[:20],
        "supabase_on": _SUPABASE_ON,
        "resend_on": bool(_RESEND_KEY),
        "stripe_on": bool(_STRIPE_KEY),
        "anthropic_on": bool(ANTHROPIC_API_KEY),
    }, headers={"Access-Control-Allow-Origin": "*"})


@app.get("/")
def health():
    return {
        "status": "ok",
        "service": "ValidaRótulo IA v6",
        "supabase": "conectado" if _SUPABASE_ON else "não configurado (in-memory)",
        "cases_em_memoria": len(_cases_db),
        "kb_categories": len(MAPA_URLS),
        "kb_cached": list(_kb_cache.keys()),
        "endpoints": ["/validar", "/eval", "/feedback", "/admin/stats", "/monitor/check"],
    }


@app.get("/kb/preload")
async def preload_kb():
    """Pré-carrega todas as legislações em cache."""
    results = {}
    for category in MAPA_URLS:
        text = await get_kb_for_categories([category])
        results[category] = f"{len(text)} chars" if text else "erro/vazio"
    return {"status": "ok", "total": len(MAPA_URLS), "loaded": results}


@app.get("/kb/status")
def kb_status():
    """Mostra status atual do cache da KB."""
    return {
        "total_categories": len(MAPA_URLS),
        "cached": len(_kb_cache),
        "cached_categories": list(_kb_cache.keys()),
        "pending": [k for k in MAPA_URLS if k not in _kb_cache],
    }


@app.get("/kb/refresh-vegetal")
async def refresh_kb_vegetal():
    """
    Baixa e atualiza todas as normas não-POA (MAPA Vegetal + ANVISA Alimentos).
    Acessa diretamente planalto.gov.br e in.gov.br — sem dependência de JavaScript.
    Execute este endpoint sempre que suspeitar que a KB está desatualizada.
    Retorna relatório de quais normas foram carregadas com sucesso.
    """
    # URLs prioritárias para não-POA vegetal
    VEGETAL_KB_KEYS = [
        # Bebidas MAPA
        "lei_8918_bebidas", "dec_6871_bebidas_tipos", "lei_7678_vinho",
        "in_mapa_49_2018_suco", "port_mapa_123_2021", "in_mapa_41_2019_kombucha",
        "lei_13648_2018_polpa",
        # Qualidade vegetal
        "lei_10831_organicos", "dec_6323_organicos", "rdc_714_frutas_hortalicas",
        "lei_9972_classificacao",
        # ANVISA Alimentos
        "rdc_429_rotulagem_nutri", "in_75_porcoes", "rdc_266_sorvetes",
        "rdc_268_proteina_vegetal", "rdc_243_alimentos_funcionais",
        "rdc_277_cafe_cha", "rdc_264_chocolate", "rdc_270_oleos",
        "rdc_276_condimentos", "rdc_716_oleos_2022", "rdc_713_acucar_2022",
        # Gerais
        "cereais_rdc711", "cereais_integrais_rdc712", "suplementos_rdc243",
        "dec_lei_986_1969", "bebidas_dec6871",
    ]

    from typing import Optional
    results = {}
    ok_count = 0
    fail_count = 0

    async with httpx.AsyncClient(timeout=15.0) as client:
        for key in VEGETAL_KB_KEYS:
            urls = MAPA_URLS.get(key, [])
            if not urls:
                results[key] = "sem URL mapeada"
                fail_count += 1
                continue
            text = ""
            for url in urls:
                try:
                    raw = await fetch_pdf_text(url, max_chars=3000)
                    if raw and len(raw) > 100:
                        text = raw
                        break
                except Exception as e:
                    pass
            if text:
                _kb_cache[key] = text
                results[key] = f"✅ {len(text)} chars"
                ok_count += 1
            else:
                results[key] = "❌ não carregado"
                fail_count += 1

    return {
        "status": "ok",
        "ok": ok_count,
        "falhou": fail_count,
        "total": len(VEGETAL_KB_KEYS),
        "detalhes": results,
        "nota": "Execute /kb/refresh-vegetal periodicamente para manter KB atualizada."
    }


# ═══════════════════════════════════════════════════════════════════════════════
# EXCEPTION HANDLER (garante CORS mesmo em crashes)
# ═══════════════════════════════════════════════════════════════════════════════
# CURADOR — Endpoints de aprovação/rejeição de feedbacks (Fase C — Parte 2)
# ═══════════════════════════════════════════════════════════════════════════════
#
# Fluxo:
#  1. Curador GitHub Action (curador_diario_v2.py) roda 8h Brasília
#  2. Manda email pra Giovanna com botões "Aprovar" / "Rejeitar"
#  3. Cada botão tem token único (tabela curador_tokens)
#  4. Click no link → endpoints abaixo
#  5. Aprovação: 2 etapas (preview → confirm) — Decisão B
#  6. Reversão: manual no Supabase (Decisão C — sem endpoint admin)
# ═══════════════════════════════════════════════════════════════════════════════

NOMES_CAMPOS_CURADOR = {
    1:  "Denominação de Venda",
    2:  "Lista de Ingredientes",
    3:  "Conteúdo Líquido",
    4:  "Identificação do Fabricante",
    5:  "Declaração de Glúten",
    6:  "Declaração de Lactose",
    7:  "Instruções de Conservação",
    8:  "Carimbo de Inspeção",
    9:  "Tabela Nutricional",
    10: "Lupa Frontal",
    11: "Alérgenos",
    12: "Transgênicos",
    13: "Lote e Validade",
    14: "Porção Padrão",
}

# Cache em memória de docs enriquecidos pelo Sonnet
# Garante que preview e confirmar usem o MESMO doc (Sonnet é não-determinístico)
# TTL: 1h. Estrutura: { token: (doc, timestamp) }
_CURADOR_DOC_CACHE: dict = {}
_CURADOR_CACHE_TTL_SEC = 3600  # 1 hora


def _curador_cache_get(token: str) -> dict | None:
    """Retorna doc cacheado se ainda válido, senão None."""
    import time as _time
    entry = _CURADOR_DOC_CACHE.get(token)
    if not entry:
        return None
    doc, ts = entry
    if _time.time() - ts > _CURADOR_CACHE_TTL_SEC:
        _CURADOR_DOC_CACHE.pop(token, None)
        return None
    return doc


def _curador_cache_set(token: str, doc: dict) -> None:
    """Salva doc no cache. Limpa entradas expiradas pra não vazar memória."""
    import time as _time
    now = _time.time()
    # Limpeza oportunista (1% chance por chamada — evita varredura toda vez)
    if len(_CURADOR_DOC_CACHE) > 50:
        expired = [t for t, (_, ts) in _CURADOR_DOC_CACHE.items()
                   if now - ts > _CURADOR_CACHE_TTL_SEC]
        for t in expired:
            _CURADOR_DOC_CACHE.pop(t, None)
    _CURADOR_DOC_CACHE[token] = (doc, now)


async def _curador_buscar_token(token: str) -> dict | None:
    """Busca token na tabela curador_tokens. Retorna None se não existe."""
    if not _SUPABASE_ON or not token:
        return None
    url = f"{_SUPABASE_URL}/rest/v1/curador_tokens?token=eq.{token}&limit=1"
    try:
        async with httpx.AsyncClient(timeout=10.0) as client:
            r = await client.get(
                url,
                headers={
                    "apikey": _SUPABASE_KEY,
                    "Authorization": f"Bearer {_SUPABASE_KEY}",
                    "Accept": "application/json",
                },
            )
            if r.status_code == 200:
                rows = r.json() or []
                return rows[0] if rows else None
    except Exception:
        pass
    return None


async def _curador_marcar_token_usado(token: str) -> bool:
    """Marca token como usado (used_at = now)."""
    if not _SUPABASE_ON or not token:
        return False
    import datetime as _dt
    url = f"{_SUPABASE_URL}/rest/v1/curador_tokens?token=eq.{token}"
    try:
        async with httpx.AsyncClient(timeout=10.0) as client:
            r = await client.patch(
                url,
                headers={
                    "apikey": _SUPABASE_KEY,
                    "Authorization": f"Bearer {_SUPABASE_KEY}",
                    "Content-Type": "application/json",
                    "Prefer": "return=minimal",
                },
                json={"used_at": _dt.datetime.now(_dt.timezone.utc).isoformat()},
            )
            return r.status_code in (200, 201, 204)
    except Exception:
        return False


async def _curador_buscar_feedback(case_id: str) -> dict | None:
    """Busca feedback original na validacoes pelo case_id."""
    if not _SUPABASE_ON or not case_id:
        return None
    url = f"{_SUPABASE_URL}/rest/v1/validacoes?case_id=eq.{case_id}&limit=1"
    try:
        async with httpx.AsyncClient(timeout=10.0) as client:
            r = await client.get(
                url,
                headers={
                    "apikey": _SUPABASE_KEY,
                    "Authorization": f"Bearer {_SUPABASE_KEY}",
                    "Accept": "application/json",
                },
            )
            if r.status_code == 200:
                rows = r.json() or []
                return rows[0] if rows else None
    except Exception:
        pass
    return None


async def _curador_atualizar_status_validacao(case_id: str, novo_status: str,
                                               aprovado_por: str = "",
                                               doc_kb_gerado: str = "") -> bool:
    """Atualiza status_curador na linha da validacoes."""
    if not _SUPABASE_ON or not case_id:
        return False
    import datetime as _dt
    payload = {"status_curador": novo_status}
    now_iso = _dt.datetime.now(_dt.timezone.utc).isoformat()
    if novo_status == "aprovado":
        payload["aprovado_por"] = aprovado_por
        payload["aprovado_em"] = now_iso
        if doc_kb_gerado:
            payload["doc_kb_gerado"] = doc_kb_gerado
    elif novo_status == "rejeitado":
        payload["rejeitado_em"] = now_iso

    url = f"{_SUPABASE_URL}/rest/v1/validacoes?case_id=eq.{case_id}"
    try:
        async with httpx.AsyncClient(timeout=10.0) as client:
            r = await client.patch(
                url,
                headers={
                    "apikey": _SUPABASE_KEY,
                    "Authorization": f"Bearer {_SUPABASE_KEY}",
                    "Content-Type": "application/json",
                    "Prefer": "return=minimal",
                },
                json=payload,
            )
            return r.status_code in (200, 201, 204)
    except Exception:
        return False


async def _curador_enriquecer_correcao(feedback: dict, campo_num: int,
                                        comentario_rt: str) -> dict:
    """
    Skill B com guarda-rails: pede pro Sonnet 4.5 transformar a correção do RT
    em um doc estruturado pra KB. Não permite invenção de fatos.

    Retorna: { "titulo": "...", "conteudo": "...", "tags": [...], "norma_principal": "..." }
    """
    if not ANTHROPIC_API_KEY:
        return {
            "titulo": f"Correção C{campo_num} — {NOMES_CAMPOS_CURADOR.get(campo_num, '')}",
            "conteudo": comentario_rt,
            "tags": [feedback.get("categoria", ""), f"campo_{campo_num}", "feedback_aprovado"],
            "norma_principal": "",
        }

    nome_campo = NOMES_CAMPOS_CURADOR.get(campo_num, f"Campo {campo_num}")
    categoria = feedback.get("categoria", "geral") or "geral"
    produto_exemplo = feedback.get("produto", "")

    system_prompt = """Você é um curador técnico especializado em rotulagem de alimentos brasileira.

Sua missão: transformar uma correção de RT (Responsável Técnico) em um documento conciso \
e estruturado para a base de conhecimento de uma IA de validação de rótulos.

REGRAS CRÍTICAS DE GUARDA-RAILS:
1. NÃO invente fatos. Use APENAS o que o RT escreveu na correção.
2. NÃO invente artigos/parágrafos de normas. Se o RT cita "RDC 727 Art. 12 §1°", use exatamente essa referência.
3. Se o RT for vago ("isso está errado"), o doc final deve ser igualmente conservador.
4. Você PODE reformatar pra ficar técnico/legível, mas NÃO PODE adicionar interpretações novas.
5. NÃO invente tipos de produto ou cenários que o RT não mencionou.

ESTRUTURA DO DOC GERADO:
- Título: curto e descritivo (≤ 80 chars)
- Conteúdo: 2-4 parágrafos. Começa com a regra/exceção, depois exemplo do RT (se houver), \
depois implicação prática. Cita norma exata.
- Tags: 3-5 palavras-chave para retrieval

Responda APENAS em JSON válido, sem markdown. Schema:
{
  "titulo": "string ≤ 80 chars",
  "conteudo": "string em markdown, 2-4 parágrafos",
  "tags": ["tag1", "tag2", "tag3"],
  "norma_principal": "string — ex: 'RDC 727/2022 Art. 12 §1°' ou '' se não houver"
}"""

    user_msg = f"""Categoria do produto: {categoria}
Campo da rotulagem: C{campo_num} — {nome_campo}
Produto de exemplo (se útil pra contexto): {produto_exemplo or '(não relevante)'}

CORREÇÃO LITERAL DO RT:
\"\"\"{comentario_rt}\"\"\"

Gere o doc estruturado JSON conforme as regras."""

    try:
        async with httpx.AsyncClient(timeout=30.0) as client:
            r = await client.post(
                "https://api.anthropic.com/v1/messages",
                headers={
                    "x-api-key": ANTHROPIC_API_KEY,
                    "anthropic-version": "2023-06-01",
                    "content-type": "application/json",
                },
                json={
                    "model": "claude-sonnet-4-5",
                    "max_tokens": 1500,
                    "system": system_prompt,
                    "messages": [{"role": "user", "content": user_msg}],
                },
            )
            if r.status_code != 200:
                raise Exception(f"Anthropic API: {r.status_code} {r.text[:200]}")

            content = r.json().get("content", [])
            if not content:
                raise Exception("Resposta vazia do Sonnet")

            texto = content[0].get("text", "").strip()
            # Remove cercas markdown se aparecerem
            if texto.startswith("```"):
                texto = texto.split("```")[1]
                if texto.startswith("json"):
                    texto = texto[4:]
                texto = texto.strip()

            doc = json.loads(texto)

            # Garantir campos mínimos
            doc.setdefault("titulo", f"Correção C{campo_num} — {nome_campo}")
            doc.setdefault("conteudo", comentario_rt)
            doc.setdefault("tags", [categoria, f"campo_{campo_num}"])
            doc.setdefault("norma_principal", "")

            # Sempre adiciona tag de origem
            tags = doc.get("tags", [])
            if "feedback_aprovado" not in tags:
                tags.append("feedback_aprovado")
            doc["tags"] = tags

            return doc

    except Exception as e:
        if SENTRY_DSN:
            try:
                import sentry_sdk
                sentry_sdk.capture_message(
                    f"Curador: falha ao enriquecer correção via Sonnet | erro={e}",
                    level="warning"
                )
            except Exception:
                pass
        # Fallback: usa correção literal do RT
        return {
            "titulo": f"Correção C{campo_num} — {nome_campo}",
            "conteudo": comentario_rt,
            "tags": [categoria, f"campo_{campo_num}", "feedback_aprovado"],
            "norma_principal": "",
        }


async def _curador_salvar_doc_kb(doc: dict, case_id: str, campo_num: int,
                                  rt_senior_email: str) -> tuple[str, str]:
    """
    Salva o doc enriquecido em kb_documents.
    Retorna (chave_gerada, erro). Se sucesso: ("chave_xyz", "").
    Se falha: ("", "mensagem de erro detalhada").
    """
    if not _SUPABASE_ON:
        return ("", "Supabase não configurado")

    import datetime as _dt
    chave = f"feedback_aprovado_{case_id[:16]}_c{campo_num}_{_dt.datetime.now(_dt.timezone.utc).strftime('%Y%m%d%H%M')}"

    # Categoria principal: prioridade pra categoria do produto original (override)
    # Senão, usa primeira tag
    categoria_override = (doc.get("_categoria_override") or "").strip()
    tags_list = doc.get("tags", [])
    if isinstance(tags_list, str):
        tags_list = [t.strip() for t in tags_list.split(",") if t.strip()]
    if categoria_override:
        categoria_doc = categoria_override
    else:
        categoria_doc = tags_list[0] if tags_list else "geral"

    # Como kb_documents NÃO tem coluna `tags`, embutimos tags + norma no conteúdo
    # via markdown (footer técnico). Retrieval funciona via campo `categoria` + texto.
    conteudo_base = doc.get("conteudo") or ""
    norma_principal = (doc.get("norma_principal") or "").strip()
    tags_str = ", ".join(tags_list) if tags_list else ""
    footer_partes = []
    if norma_principal:
        footer_partes.append(f"**Norma principal:** {norma_principal}")
    if tags_str:
        footer_partes.append(f"**Tags:** {tags_str}")
    footer_partes.append(f"**Origem:** Feedback aprovado (case {case_id[:16]}) por {rt_senior_email}")
    conteudo_completo = conteudo_base.rstrip() + "\n\n---\n\n" + "\n".join(footer_partes)

    payload = {
        "chave": chave,
        "titulo": (doc.get("titulo") or "")[:200],
        "conteudo": conteudo_completo,
        "categoria": categoria_doc,
        "fonte": f"feedback_aprovado_case_{case_id[:16]}",
        "orgao": "MAPA/ANVISA",
        "tamanho_chars": len(conteudo_completo),
        "atualizado_em": _dt.datetime.now(_dt.timezone.utc).isoformat(),
        # Colunas adicionadas no ALTER TABLE para rastreabilidade
        "origem": "feedback_aprovado",
        "case_id_origem": case_id,
        "aprovado_por": rt_senior_email,
        "aprovado_em": _dt.datetime.now(_dt.timezone.utc).isoformat(),
    }

    print(f"[CURADOR] Salvando doc na KB: chave={chave} categoria={categoria_doc} chars={len(conteudo_completo)}", flush=True)

    url = f"{_SUPABASE_URL}/rest/v1/kb_documents?on_conflict=chave"
    try:
        async with httpx.AsyncClient(timeout=15.0) as client:
            r = await client.post(
                url,
                headers={
                    "apikey": _SUPABASE_KEY,
                    "Authorization": f"Bearer {_SUPABASE_KEY}",
                    "Content-Type": "application/json",
                    "Prefer": "resolution=merge-duplicates,return=minimal",
                },
                json=payload,
            )
            if r.status_code in (200, 201, 204):
                print(f"[CURADOR] ✓ Doc salvo: {chave}", flush=True)
                return (chave, "")

            # Falha — extrai mensagem do PostgREST pra UI
            err_msg = f"Supabase {r.status_code}"
            try:
                err_body = r.json()
                if isinstance(err_body, dict):
                    err_msg = err_body.get("message") or err_body.get("hint") or err_msg
                    if err_body.get("details"):
                        err_msg += f" — {err_body['details']}"
            except Exception:
                err_msg = f"Supabase {r.status_code}: {r.text[:200]}"

            print(f"[CURADOR] ✗ Falha ao salvar: {err_msg}", flush=True)

            # Log opcional no Sentry (com try/except robusto)
            try:
                if globals().get("SENTRY_DSN"):
                    import sentry_sdk
                    sentry_sdk.capture_message(
                        f"Curador: falha save KB | status={r.status_code} body={r.text[:300]}",
                        level="error"
                    )
            except Exception:
                pass

            return ("", err_msg)
    except Exception as e:
        err_msg = f"Exceção ao salvar: {type(e).__name__}: {str(e)[:200]}"
        print(f"[CURADOR] ✗ {err_msg}", flush=True)
        try:
            if globals().get("SENTRY_DSN"):
                import sentry_sdk
                sentry_sdk.capture_exception(e)
        except Exception:
            pass
        return ("", err_msg)


def _render_curador_html(title: str, body: str, status_class: str = "ok") -> str:
    """
    Página HTML simples de resposta do curador.
    status_class: 'ok' | 'error' | 'preview'
    """
    color = {
        "ok":      "#166534",
        "error":   "#992f2a",
        "preview": "#9c5a0e",
        "warn":    "#9c5a0e",
    }.get(status_class, "#166534")

    bg_soft = {
        "ok":      "#E8F0EA",
        "error":   "#F0DDDB",
        "preview": "#F5EBD8",
        "warn":    "#F5EBD8",
    }.get(status_class, "#E8F0EA")

    return f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{title} · InspectIA Curadoria</title>
<style>
  body {{ margin: 0; font-family: 'Helvetica Neue', Arial, sans-serif; background: #f0ede3; color: #0f2a20; padding: 40px 16px; }}
  .container {{ max-width: 640px; margin: 0 auto; }}
  .header {{ background: #0F2A20; color: #fff; padding: 18px 24px; border-radius: 10px 10px 0 0; font-size: 14px; font-weight: 600; }}
  .header small {{ display: block; color: #a8a4a0; font-size: 11px; margin-top: 4px; font-weight: 400; }}
  .card {{ background: #fff; border: 1px solid #e2ded2; border-top: none; border-radius: 0 0 10px 10px; padding: 28px 24px; }}
  h1 {{ color: {color}; font-size: 22px; margin: 0 0 16px; }}
  .status-bar {{ background: {bg_soft}; border-left: 4px solid {color}; padding: 12px 16px; border-radius: 6px; margin-bottom: 20px; font-size: 13px; }}
  .doc-preview {{ background: #fbfaf6; border: 1px solid #e2ded2; border-radius: 8px; padding: 18px 20px; margin: 16px 0; }}
  .doc-preview h3 {{ font-size: 11px; color: #6b6757; text-transform: uppercase; letter-spacing: 1px; margin: 0 0 8px; font-weight: 600; }}
  .doc-preview .titulo {{ font-size: 15px; font-weight: 600; color: #091a14; margin-bottom: 10px; }}
  .doc-preview .conteudo {{ font-size: 13px; color: #0f2a20; line-height: 1.6; white-space: pre-wrap; }}
  .doc-preview .tags {{ margin-top: 12px; }}
  .doc-preview .tag {{ display: inline-block; background: #e8f0ea; color: #166534; font-size: 10px; padding: 3px 8px; border-radius: 10px; margin-right: 4px; }}
  .actions {{ margin-top: 24px; display: flex; gap: 10px; flex-wrap: wrap; }}
  .btn {{ display: inline-block; padding: 11px 22px; border-radius: 8px; font-size: 13px; font-weight: 600; cursor: pointer; text-decoration: none; border: none; font-family: inherit; }}
  .btn-primary {{ background: #166534; color: #fff; }}
  .btn-primary:hover {{ background: #0E4A26; }}
  .btn-secondary {{ background: #fff; color: #0f2a20; border: 1px solid #d9d2be; }}
  .btn-secondary:hover {{ background: #fbfaf6; }}
  .footer {{ text-align: center; color: #6b6757; font-size: 11px; margin-top: 24px; }}
</style>
</head>
<body>
<div class="container">
  <div class="header">📋 InspectIA · Curadoria<small>Revisão de feedback de RT</small></div>
  <div class="card">
    <h1>{title}</h1>
    {body}
  </div>
  <div class="footer">InspectIA · Validador de Rótulos com IA · {datetime.now(timezone.utc).astimezone(timezone(timedelta(hours=-3))).strftime("%d/%m/%Y %H:%M")} BRT</div>
</div>
</body>
</html>"""


@app.get("/curador/relatorio/{case_id}")
async def curador_ver_relatorio(case_id: str):
    """
    Página HTML que mostra o relatório completo de uma validação.
    Linkada do email diário do curador — permite Giovanna ver o contexto
    completo da análise sem precisar reabrir a plataforma.
    """
    if not _SUPABASE_ON:
        return HTMLResponse(_render_curador_html(
            "Indisponível",
            "<div class='status-bar'>Sistema offline temporariamente.</div>",
            "fail"
        ), status_code=503)

    import html as _html_mod
    import re as _re

    # Busca o caso
    url = f"{_SUPABASE_URL}/rest/v1/validacoes?case_id=eq.{case_id}&limit=1"
    try:
        async with httpx.AsyncClient(timeout=10.0) as cl:
            r = await cl.get(url, headers={
                "apikey": _SUPABASE_KEY,
                "Authorization": f"Bearer {_SUPABASE_KEY}",
            })
            data = r.json() if r.status_code == 200 else []
    except Exception:
        data = []

    if not data:
        return HTMLResponse(_render_curador_html(
            "Caso não encontrado",
            f"<div class='status-bar'>Validação <code>{case_id[:24]}</code> não encontrada.</div>",
            "fail"
        ), status_code=404)

    caso = data[0]
    produto = (caso.get("produto") or "—").strip() or "—"
    categoria = (caso.get("categoria") or caso.get("caminho_np") or "—").strip() or "—"
    orgao = (caso.get("orgao") or "—").strip() or "—"
    score_agente = caso.get("score_agente")
    score_real = caso.get("score_real")
    imagem_url = (caso.get("imagem_url") or "").strip()
    relatorio = (caso.get("relatorio_completo") or "").strip()
    feedback_rt = (caso.get("feedback") or "").strip()
    rt_comment = (caso.get("rt_comment") or "").strip()
    created_at = (caso.get("created_at") or caso.get("timestamp") or "").strip()

    if not relatorio:
        relatorio_html = "<em style='color:#6b6757;'>Relatório completo não foi salvo para esta validação.</em>"
    else:
        # Conversão markdown-like → HTML simples
        # Preserva quebras de linha; renderiza **bold**, # headers
        rel_esc = _html_mod.escape(relatorio)
        # Headers (## ou ###)
        rel_esc = _re.sub(r'^### (.+)$', r'<h3 style="color:#0F2A20;margin:18px 0 8px;font-size:14px;">\1</h3>', rel_esc, flags=_re.MULTILINE)
        rel_esc = _re.sub(r'^## (.+)$', r'<h2 style="color:#0F2A20;margin:22px 0 10px;font-size:16px;border-bottom:1px solid #e2ded2;padding-bottom:6px;">\1</h2>', rel_esc, flags=_re.MULTILINE)
        rel_esc = _re.sub(r'^# (.+)$', r'<h1 style="color:#091A14;margin:24px 0 12px;font-size:20px;">\1</h1>', rel_esc, flags=_re.MULTILINE)
        # CAMPO X — destaque visual
        rel_esc = _re.sub(r'^(CAMPO \d+[^\n]*)$', r'<div style="background:#F0EDE3;padding:8px 12px;margin:14px 0 6px;border-left:3px solid #166534;font-weight:600;color:#091A14;">\1</div>', rel_esc, flags=_re.MULTILINE)
        # Bold **texto**
        rel_esc = _re.sub(r'\*\*([^*\n]+)\*\*', r'<strong>\1</strong>', rel_esc)
        # Status colorido
        rel_esc = rel_esc.replace("CONFORME", '<span style="color:#166534;font-weight:600;">CONFORME</span>') \
                         .replace("NÃO <span style=\"color:#166534;font-weight:600;\">CONFORME</span>", '<span style="color:#992F2A;font-weight:600;">NÃO CONFORME</span>') \
                         .replace("COM RESSALVAS", '<span style="color:#9C5A0E;font-weight:600;">COM RESSALVAS</span>')
        # Newlines → <br>
        rel_esc = rel_esc.replace("\n", "<br>\n")
        relatorio_html = f'<div style="font-family:Geist,system-ui,sans-serif;font-size:13px;line-height:1.7;color:#0F2A20;">{rel_esc}</div>'

    # Score box
    score_box = ""
    if score_agente is not None:
        score_box = f"<strong>Score IA:</strong> {score_agente}/14"
        if score_real is not None and score_real != score_agente:
            score_box += f" → <strong>Score após RT:</strong> {score_real}/14"

    # Imagem
    img_block = ""
    if imagem_url:
        img_block = f'''<div style="margin:16px 0;text-align:center;">
            <a href="{imagem_url}" target="_blank" rel="noopener">
                <img src="{imagem_url}" alt="Rótulo" style="max-width:100%;max-height:420px;border:1px solid #E2DED2;border-radius:8px;box-shadow:0 1px 3px rgba(15,42,32,0.08);">
            </a>
            <div style="font-size:11px;color:#6b6757;margin-top:6px;">Clique para abrir em tamanho real</div>
        </div>'''
    else:
        img_block = '<div style="background:#F0EDE3;border:1px dashed #D9D2BE;border-radius:8px;padding:32px;text-align:center;color:#6b6757;font-size:12px;margin:16px 0;">Foto do rótulo não disponível para esta validação</div>'

    # Feedback do RT (se houver)
    fb_block = ""
    if feedback_rt or rt_comment:
        fb_text = rt_comment or feedback_rt
        fb_block = f'''<div style="background:#F5EBD8;border-left:3px solid #9C5A0E;padding:14px 18px;margin:20px 0;border-radius:4px;">
            <div style="font-size:11px;color:#6b6757;text-transform:uppercase;letter-spacing:1px;font-weight:600;margin-bottom:6px;">💬 Comentário geral do RT</div>
            <div style="font-size:13px;color:#0f2a20;line-height:1.6;">{_html_mod.escape(fb_text)}</div>
        </div>'''

    body_html = f'''
<div style="background:#fff;border:1px solid #e2ded2;border-radius:10px;padding:24px;max-width:760px;margin:0 auto;">
    <div style="font-size:11px;color:#6b6757;text-transform:uppercase;letter-spacing:1px;font-weight:600;">Relatório completo · case {case_id[:24]}</div>
    <h1 style="font-size:22px;color:#091a14;margin:8px 0 4px;">{_html_mod.escape(produto)}</h1>
    <div style="font-size:13px;color:#6b6757;">{_html_mod.escape(categoria)} · {_html_mod.escape(orgao)}</div>
    <div style="font-size:13px;color:#0f2a20;margin-top:8px;">{score_box}</div>

    {img_block}

    {fb_block}

    <hr style="border:none;border-top:1px solid #e2ded2;margin:24px 0;">

    <div style="font-size:11px;color:#6b6757;text-transform:uppercase;letter-spacing:1px;font-weight:600;margin-bottom:12px;">📋 Análise técnica completa</div>
    {relatorio_html}
</div>
'''

    return HTMLResponse(_render_curador_html(
        f"Relatório · {produto[:60]}",
        body_html,
        "ok"
    ))


@app.get("/curador/aprovar/{token}")
async def curador_aprovar(token: str, confirmar: str = ""):
    """
    Aprovação em 2 etapas (Decisão B):
      1. Click do email → mostra preview do doc enriquecido
      2. User clica "Confirmar e salvar na KB" → salva
    """
    from datetime import datetime as _dtcls, timezone as _tz, timedelta as _td

    # 1. Valida token
    tok = await _curador_buscar_token(token)
    if not tok:
        return HTMLResponse(_render_curador_html(
            "Token inválido",
            "<div class='status-bar'>Este link de aprovação não existe ou já foi processado.</div>"
            "<p style='font-size:13px;color:#6b6757'>Se você acabou de clicar em um email do InspectIA e está vendo esta mensagem, "
            "provavelmente o link já foi usado ou foi aprovado/rejeitado em outra sessão. Tudo certo, não há ação adicional necessária.</p>",
            "error"
        ), status_code=404)

    if tok.get("acao") != "aprovar":
        return HTMLResponse(_render_curador_html(
            "Token de tipo errado",
            f"<div class='status-bar'>Este token é para ação '{tok.get('acao')}', não 'aprovar'.</div>",
            "error"
        ), status_code=400)

    if tok.get("used_at"):
        return HTMLResponse(_render_curador_html(
            "Token já usado",
            f"<div class='status-bar'>Esta correção já foi processada em {tok.get('used_at', '?')}.</div>",
            "warn"
        ), status_code=410)

    # Verifica expiração
    expires_at = tok.get("expires_at")
    if expires_at:
        try:
            exp_dt = _dtcls.fromisoformat(expires_at.replace("Z", "+00:00"))
            if exp_dt < _dtcls.now(_tz.utc):
                return HTMLResponse(_render_curador_html(
                    "Token expirado",
                    f"<div class='status-bar'>Este link expirou em {expires_at[:10]}. "
                    "O feedback voltará na próxima rodada do curador (manhã seguinte).</div>",
                    "warn"
                ), status_code=410)
        except Exception:
            pass

    case_id = tok.get("case_id", "")
    campo_num = tok.get("campo_num", 0)
    rt_senior_email = tok.get("rt_senior_email", "")

    # 2. Busca feedback original
    fb = await _curador_buscar_feedback(case_id)
    if not fb:
        return HTMLResponse(_render_curador_html(
            "Feedback não encontrado",
            f"<div class='status-bar'>O feedback original (case_id={case_id[:24]}) não existe mais na base.</div>",
            "error"
        ), status_code=404)

    # Extrai comentário do RT pra esse campo específico
    campos = fb.get("campos") or {}
    if isinstance(campos, str):
        try:
            campos = json.loads(campos)
        except Exception:
            campos = {}
    campo_dados = campos.get(str(campo_num)) or campos.get(campo_num) or {}
    comentario_rt = (campo_dados.get("comentario") or "").strip()

    if not comentario_rt:
        return HTMLResponse(_render_curador_html(
            "Sem correção para aprovar",
            f"<div class='status-bar'>O Campo {campo_num} não tem comentário do RT — nada para virar conhecimento.</div>",
            "warn"
        ), status_code=400)

    nome_campo = NOMES_CAMPOS_CURADOR.get(campo_num, f"Campo {campo_num}")

    # ─── ETAPA 1: PREVIEW (sem ?confirmar=1) ───
    if confirmar != "1":
        # Tenta usar cache primeiro (caso usuário recarregue a página)
        doc = _curador_cache_get(token)
        if not doc:
            # Gera doc enriquecido pela primeira vez
            doc = await _curador_enriquecer_correcao(fb, campo_num, comentario_rt)
            # Cacheia pra próxima etapa (confirmar) usar EXATAMENTE este doc
            _curador_cache_set(token, doc)

        tags_html = "".join(f'<span class="tag">{html.escape(t)}</span>' for t in doc.get("tags", []))
        norma = doc.get("norma_principal", "") or "—"
        produto = fb.get("produto", "—")

        body = f"""
        <div class='status-bar'>Revisão antes de salvar na KB · Campo {campo_num} — {html.escape(nome_campo)}</div>

        <p style='font-size:13px;color:#6b6757;margin-bottom:18px;line-height:1.6'>
          Você está prestes a salvar este aprendizado na base de conhecimento da IA.<br>
          Confira o conteúdo abaixo. Se algo estiver impreciso, clique em <strong>Cancelar</strong> e edite manualmente o feedback no Supabase.
        </p>

        <div class='doc-preview'>
          <h3>📄 Doc que será salvo na KB</h3>
          <div class='titulo'>{html.escape(doc.get("titulo", ""))}</div>
          <div class='conteudo'>{html.escape(doc.get("conteudo", ""))}</div>
          <div class='tags'>{tags_html}</div>
        </div>

        <div class='doc-preview' style='background:#f5ebd8;border-color:#ddc18b'>
          <h3>👤 Correção literal do RT</h3>
          <div style='font-size:13px;color:#0f2a20;line-height:1.6;font-style:italic'>"{html.escape(comentario_rt)}"</div>
        </div>

        <div style='font-size:11px;color:#6b6757;margin-top:12px'>
          <div><strong>Produto de origem:</strong> {html.escape(str(produto))}</div>
          <div><strong>Norma principal identificada:</strong> {html.escape(norma)}</div>
          <div><strong>case_id:</strong> <code style='font-size:10px'>{html.escape(case_id[:24])}</code></div>
        </div>

        <div class='actions'>
          <a href='/curador/aprovar/{token}?confirmar=1' class='btn btn-primary'>✓ Confirmar e salvar na KB</a>
          <a href='/curador/rejeitar/{token.replace(token, "")}' onclick='return false' class='btn btn-secondary' style='opacity:0.5;cursor:not-allowed' title='Use o link de Rejeitar do email'>Cancelar</a>
        </div>
        <p style='font-size:11px;color:#a8a4a0;margin-top:14px'>
          Para rejeitar esta correção em vez de aprovar, volte ao email e clique em "✗ Rejeitar".
        </p>
        """

        return HTMLResponse(_render_curador_html(
            f"Aprovar Campo {campo_num}",
            body,
            "preview"
        ))

    # ─── ETAPA 2: CONFIRMAÇÃO (com ?confirmar=1) ───

    # Re-verifica que token ainda não foi usado (race condition)
    tok_check = await _curador_buscar_token(token)
    if not tok_check or tok_check.get("used_at"):
        return HTMLResponse(_render_curador_html(
            "Token já foi usado",
            "<div class='status-bar'>Outra aba ou click já processou esta correção.</div>",
            "warn"
        ), status_code=410)

    # Recupera o MESMO doc enriquecido que o user viu no preview
    # (Sonnet é não-determinístico, regerar daria texto diferente)
    doc = _curador_cache_get(token)
    if not doc:
        # Cache expirou (>1h entre preview e confirmar) ou usuário pulou preview
        # Re-gera mas avisa no log
        print(f"[CURADOR] WARN: cache vazio na confirmação, regerando doc (token={token[:8]}...)", flush=True)
        doc = await _curador_enriquecer_correcao(fb, campo_num, comentario_rt)

    # Override: categoria do doc na KB usa categoria do PRODUTO original
    # (não a primeira tag — tag pode ser tema da correção, ex: "CNPJ")
    categoria_produto = (fb.get("categoria") or "").strip()
    if categoria_produto:
        # Mantém tags originais, só forçar a categoria principal
        doc["_categoria_override"] = categoria_produto

    # Salva na KB — retorna (chave, erro)
    chave_doc, erro_save = await _curador_salvar_doc_kb(doc, case_id, campo_num, rt_senior_email)

    if not chave_doc:
        # Mostra erro detalhado pro RT senior poder reportar
        return HTMLResponse(_render_curador_html(
            "Erro ao salvar",
            f"<div class='status-bar'>Não foi possível salvar o doc na KB.</div>"
            f"<p style='font-size:13px;color:#0f2a20;margin:14px 0;line-height:1.6'>"
            f"<strong>Erro técnico:</strong> {html.escape(erro_save or 'sem detalhes')}</p>"
            f"<p style='font-size:12px;color:#6b6757'>"
            f"O token <strong>NÃO</strong> foi consumido — você pode tentar novamente clicando no link do email. "
            f"Se persistir, mande este erro para o Hugo (hugodacach@gmail.com).</p>",
            "error"
        ), status_code=500)

    # Marca token como usado
    await _curador_marcar_token_usado(token)

    # Limpa cache do doc (já gravou na KB, não precisa mais)
    _CURADOR_DOC_CACHE.pop(token, None)

    # Atualiza status da validação
    await _curador_atualizar_status_validacao(
        case_id=case_id,
        novo_status="aprovado",
        aprovado_por=rt_senior_email,
        doc_kb_gerado=chave_doc,
    )

    # Limpa cache da KB pra próxima validação carregar o doc novo
    try:
        global _KB_CACHE_BY_CATEGORY, _KB_CACHE_TIMESTAMP
        _KB_CACHE_BY_CATEGORY = {}
        _KB_CACHE_TIMESTAMP = 0
    except Exception:
        pass

    body = f"""
    <div class='status-bar'>✓ Correção aprovada e salva na base de conhecimento</div>
    <p style='font-size:13px;color:#0f2a20;line-height:1.7'>
      Obrigada, <strong>Giovanna</strong>! O aprendizado foi incorporado.
    </p>
    <div class='doc-preview'>
      <h3>📚 Doc salvo na KB</h3>
      <div class='titulo'>{html.escape(doc.get("titulo", ""))}</div>
      <div style='font-size:11px;color:#6b6757;margin-top:8px'>
        Chave: <code>{html.escape(chave_doc)}</code>
      </div>
    </div>
    <p style='font-size:12px;color:#6b6757;margin-top:18px;line-height:1.6'>
      Próximas validações de produtos da categoria <strong>{html.escape(fb.get('categoria') or 'similar')}</strong> automaticamente
      consultarão este conhecimento. A IA vai aprender com sua correção.
    </p>
    <p style='font-size:11px;color:#a8a4a0;margin-top:18px'>
      Para reverter (raro), o admin pode deletar a linha com chave <code>{html.escape(chave_doc)}</code> diretamente na tabela <code>kb_documents</code> do Supabase.
    </p>
    """

    return HTMLResponse(_render_curador_html(
        f"Aprovado · Campo {campo_num}",
        body,
        "ok"
    ))


@app.get("/curador/rejeitar/{token}")
async def curador_rejeitar(token: str):
    """
    Rejeição em 1 click (mais simples — descartar não precisa preview).
    """
    from datetime import datetime as _dtcls, timezone as _tz

    tok = await _curador_buscar_token(token)
    if not tok:
        return HTMLResponse(_render_curador_html(
            "Token inválido",
            "<div class='status-bar'>Este link de rejeição não existe ou já foi processado.</div>",
            "error"
        ), status_code=404)

    if tok.get("acao") != "rejeitar":
        return HTMLResponse(_render_curador_html(
            "Token de tipo errado",
            f"<div class='status-bar'>Este token é para ação '{tok.get('acao')}', não 'rejeitar'.</div>",
            "error"
        ), status_code=400)

    if tok.get("used_at"):
        return HTMLResponse(_render_curador_html(
            "Já processado",
            f"<div class='status-bar'>Esta correção já foi rejeitada em {tok.get('used_at', '?')}.</div>",
            "warn"
        ), status_code=410)

    # Verifica expiração
    expires_at = tok.get("expires_at")
    if expires_at:
        try:
            exp_dt = _dtcls.fromisoformat(expires_at.replace("Z", "+00:00"))
            if exp_dt < _dtcls.now(_tz.utc):
                return HTMLResponse(_render_curador_html(
                    "Token expirado",
                    f"<div class='status-bar'>Este link expirou em {expires_at[:10]}.</div>",
                    "warn"
                ), status_code=410)
        except Exception:
            pass

    case_id = tok.get("case_id", "")
    campo_num = tok.get("campo_num", 0)

    # Marca token usado
    await _curador_marcar_token_usado(token)

    # Atualiza status da validação
    await _curador_atualizar_status_validacao(case_id, "rejeitado")

    nome_campo = NOMES_CAMPOS_CURADOR.get(campo_num, f"Campo {campo_num}")

    body = f"""
    <div class='status-bar'>✗ Correção rejeitada</div>
    <p style='font-size:13px;color:#0f2a20;line-height:1.7'>
      A correção do Campo {campo_num} ({html.escape(nome_campo)}) foi descartada.
      <strong>Nada será adicionado à KB.</strong>
    </p>
    <p style='font-size:12px;color:#6b6757;margin-top:14px;line-height:1.6'>
      O agente continuará validando rótulos como antes. Se houver outras correções neste mesmo feedback,
      você pode aprovar/rejeitar separadamente pelo email original.
    </p>
    <p style='font-size:11px;color:#a8a4a0;margin-top:18px'>
      case_id: <code>{html.escape(case_id[:24])}</code>
    </p>
    """

    return HTMLResponse(_render_curador_html(
        f"Rejeitado · Campo {campo_num}",
        body,
        "ok"
    ))


# ═══════════════════════════════════════════════════════════════════════════════
from fastapi import Request as _Request

@app.exception_handler(Exception)
async def global_exception_handler(request: _Request, exc: Exception):
    return JSONResponse(
        status_code=500,
        content={"error": str(exc)},
        headers={"Access-Control-Allow-Origin": "*"},
    )
